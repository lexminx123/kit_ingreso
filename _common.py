"""
_common.py — Utilidades compartidas para generación de documentos Word
GRUPO CAVAL 1003, C.A. (ALIKA PETS)
Paleta teal #0F766E — Estilo profesional venezolano.

Versión 3.0:
  - add_signature_block() ahora usa DIRECTORA GERENTE (Esnatlim Elena Simoza)
    en vez de "RR.HH." genérico.
  - Representante legal (Vicepresidente Alicia Sleiman) se mantiene solo en
    add_signature_block_rep_legal() para contratos y documentos legales.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# Datos de la empresa y representantes
# ============================================================
EMPRESA = "GRUPO CAVAL 1003, C.A."
RIF_EMP = "J501662533"
MARCA = "ALIKA PETS"
DOMICILIO_EMP = "Av. Francisco de Miranda, Local N° 1, Sector Francisco de Miranda, Los Teques, Estado Miranda, Zona Postal 1201"

# Junta directiva — representante legal (firma contratos y documentos mercantiles)
REP_LEGAL_NOMBRE = "ALICIA SLEIMAN DE MAKSOUD"
REP_LEGAL_CARGO = "Vicepresidente"
REP_LEGAL_CI = "V-10.112.683"

# Dirección operativa — Directora Gerente (firma documentos operativos del kit)
DIRECTORA_NOMBRE = "ESNATLIM ELENA SIMOZA"
DIRECTORA_CARGO = "Directora Gerente"
DIRECTORA_CI = "V-17.976.287"

# ============================================================
# Paleta de colores
# ============================================================
TEAL_DARK   = RGBColor(0x0F, 0x76, 0x6E)
TEAL_HDR_BG = "0F766E"
GRAY_ALT    = "F3F4F6"
GRAY_LIGHT  = "F9FAFB"
AMBER_BG    = "FEF3C7"
AMBER_ALT   = "FDE68A"
SLATE_BG    = "E2E8F0"
WHITE       = "FFFFFF"
BLACK       = RGBColor(0x11, 0x18, 0x27)
GRAY_TEXT   = RGBColor(0x55, 0x65, 0x75)
GRAY_MUTED  = RGBColor(0x90, 0x98, 0xA8)
RED_CRIT    = RGBColor(0xB9, 0x1C, 0x1C)
GREEN_OK    = RGBColor(0x04, 0x78, 0x57)

# ============================================================
# Bajo nivel: celdas y runs
# ============================================================
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="CBD5E1", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, v in (('top',top),('bottom',bottom),('left',left),('right',right)):
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tc_pr.append(tcMar)

def style_run(run, size=10, bold=False, color=None, font="Calibri", italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)

def write_cell(cell, text, size=10, bold=False, color=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, valign=WD_ALIGN_VERTICAL.CENTER,
               bg=None, font="Calibri", italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            r = p.add_run()
            r.add_break()
        r = p.add_run(line)
        style_run(r, size=size, bold=bold, color=color, font=font, italic=italic)
    cell.vertical_alignment = valign
    if bg:
        set_cell_bg(cell, bg)
    set_cell_borders(cell)
    set_cell_margins(cell)

# ============================================================
# Párrafos, títulos, divisores
# ============================================================
def add_hr(paragraph, color="0F766E", sz="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    p_pr.append(pBdr)

def add_para(doc, text, size=10, bold=False, color=None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=4,
             italic=False, font="Calibri", line_spacing=1.3, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if indent is not None:
        p.paragraph_format.left_indent = indent
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color, font=font, italic=italic)
    return p

def add_bullet(doc, text, size=10, color=None, bold_lead=None, indent=Cm(0.6)):
    """Bullet con guion. Si se pasa bold_lead, lo resalta al inicio."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    rb = p.add_run("•  ")
    style_run(rb, size=size, bold=True, color=TEAL_DARK)
    if bold_lead:
        rl = p.add_run(bold_lead)
        style_run(rl, size=size, bold=True, color=color or BLACK)
    r = p.add_run(text)
    style_run(r, size=size, color=color or BLACK)
    return p

def add_article(doc, number, title, body, size=10):
    """Artículo numerado estilo 'Art. 5. — Título. Cuerpo.'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    r1 = p.add_run(f"Art. {number}.  ")
    style_run(r1, size=size, bold=True, color=TEAL_DARK)
    if title:
        r2 = p.add_run(f"{title}.  ")
        style_run(r2, size=size, bold=True, color=BLACK)
    r3 = p.add_run(body)
    style_run(r3, size=size, color=BLACK)
    return p

# ============================================================
# Encabezado / Membrete / Footer
# ============================================================
def setup_a4_portrait(doc, margins_cm=2.0):
    """Configura A4 vertical con márgenes y estilo Normal."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(margins_cm)
    section.bottom_margin = Cm(margins_cm)
    section.left_margin = Cm(margins_cm)
    section.right_margin = Cm(margins_cm)
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(4)
    return section

def add_membrete(doc, doc_label, doc_sublabel, version="Versión 3.0  ·  RR.HH."):
    """Membrete superior: tabla 2 columnas sin bordes."""
    hdr = doc.add_table(rows=1, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.autofit = False
    hdr.columns[0].width = Cm(13)
    hdr.columns[1].width = Cm(5)
    for c in hdr.rows[0].cells:
        set_cell_borders(c, color="FFFFFF", sz="0")

    left = hdr.rows[0].cells[0]
    left.text = ""
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ALIKA PETS")
    style_run(r, size=22, bold=True, color=TEAL_DARK)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("Grupo Caval 1003, C.A.  ·  RIF: J501662533")
    style_run(r2, size=10, color=BLACK)
    p3 = left.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    r3 = p3.add_run("Av. Francisco de Miranda, Local N° 1, Los Teques, Miranda  ·  Zona Postal 1201")
    style_run(r3, size=9, color=GRAY_TEXT)

    right = hdr.rows[0].cells[1]
    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    rr = rp.add_run(doc_label)
    style_run(rr, size=11, bold=True, color=TEAL_DARK)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_before = Pt(0)
    rr2 = rp2.add_run(doc_sublabel)
    style_run(rr2, size=9, italic=True, color=GRAY_TEXT)
    rp3 = right.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp3.paragraph_format.space_before = Pt(0)
    rr3 = rp3.add_run(version)
    style_run(rr3, size=8, color=GRAY_MUTED)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after = Pt(8)
    add_hr(sep, color="0F766E", sz="14")

def add_doc_title(doc, title):
    """Título del documento centrado, teal, con línea horizontal."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    style_run(r, size=15, bold=True, color=TEAL_DARK)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(10)
    add_hr(sep, color="0F766E", sz="10")

def add_chapter(doc, roman, title):
    """Encabezado de capítulo con numeración romana."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run(f"CAPÍTULO {roman}")
    style_run(r1, size=10, bold=True, color=GRAY_TEXT)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.keep_with_next = True
    r2 = p2.add_run(title)
    style_run(r2, size=12, bold=True, color=TEAL_DARK)
    add_hr(p2, color="CBD5E1", sz="4")

def add_section(doc, title):
    """Subtítulo de sección dentro de capítulo."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    style_run(r, size=11, bold=True, color=TEAL_DARK)

def add_footer(section, doc_label):
    """Footer con número de página dinámico."""
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    txt = f"ALIKA PETS  ·  Grupo Caval 1003, C.A.  ·  RIF J501662533  ·  {doc_label}  ·  Página "
    fr = fp.add_run(txt)
    style_run(fr, size=8, color=GRAY_MUTED)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r2 = fp.add_run()
    r2._r.append(fldChar1)
    r2._r.append(instrText)
    r2._r.append(fldChar2)
    style_run(r2, size=8, color=GRAY_MUTED)

# ============================================================
# Tablas
# ============================================================
def add_sanciones_table(doc, headers, rows, col_widths_cm=None):
    """Tabla con encabezado teal y filas alternas."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        if col_widths_cm:
            c.width = Cm(col_widths_cm[ci])
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            if col_widths_cm:
                c.width = Cm(col_widths_cm[ci])
            color_txt = BLACK
            bold = False
            if isinstance(val, str):
                if "DESPIDO" in val.upper() or "GRAVE" in val.upper():
                    color_txt = RED_CRIT
                    bold = True
                elif "VERBAL" in val.upper():
                    color_txt = GREEN_OK
            write_cell(c, str(val), size=9, bold=bold, color=color_txt, bg=bg,
                       align=WD_ALIGN_PARAGRAPH.LEFT if ci > 0 else WD_ALIGN_PARAGRAPH.CENTER)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl

# ============================================================
# Bloques de firma — VERSIÓN 3.0
# ============================================================
def _build_signature_cell(cell, label, role_name, role_cargo, role_ci, width_cm):
    """Construye una celda de firma con label, línea, nombre, C.I., cargo, fecha."""
    cell.width = Cm(width_cm)
    cell.text = ""
    # label superior (cabecera)
    p_label = cell.paragraphs[0]
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_label.paragraph_format.space_before = Pt(0)
    p_label.paragraph_format.space_after = Pt(2)
    p_label.paragraph_format.line_spacing = 1.0
    rl = p_label.add_run(label)
    style_run(rl, size=10, bold=True, color=TEAL_DARK)
    set_cell_bg(cell, SLATE_BG)
    set_cell_borders(cell)
    set_cell_margins(cell)
    # espacio para firma
    p_space = cell.add_paragraph()
    p_space.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_space.paragraph_format.space_before = Pt(22)
    p_space.paragraph_format.space_after = Pt(2)
    p_space.paragraph_format.line_spacing = 1.0
    r_line = p_space.add_run("______________________________")
    style_run(r_line, size=9)
    # nombre del rol (si aplica)
    if role_name:
        p_name = cell.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_after = Pt(0)
        p_name.paragraph_format.line_spacing = 1.0
        r_name = p_name.add_run(role_name)
        style_run(r_name, size=9, bold=True, color=BLACK)
        # C.I.
        p_ci = cell.add_paragraph()
        p_ci.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ci.paragraph_format.space_after = Pt(0)
        p_ci.paragraph_format.line_spacing = 1.0
        r_ci = p_ci.add_run(f"C.I. {role_ci}")
        style_run(r_ci, size=8, color=GRAY_TEXT)
        # cargo
        p_cargo = cell.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cargo.paragraph_format.space_after = Pt(0)
        p_cargo.paragraph_format.line_spacing = 1.0
        r_cargo = p_cargo.add_run(role_cargo)
        style_run(r_cargo, size=8, italic=True, color=GRAY_TEXT)
    else:
        p_instr = cell.add_paragraph()
        p_instr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_instr.paragraph_format.space_after = Pt(0)
        p_instr.paragraph_format.line_spacing = 1.0
        r_instr = p_instr.add_run("Firma · C.I.: ____________________")
        style_run(r_instr, size=8, italic=True, color=GRAY_TEXT)
    # fecha
    p_date = cell.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(4)
    p_date.paragraph_format.space_after = Pt(0)
    p_date.paragraph_format.line_spacing = 1.0
    r_date = p_date.add_run("Fecha: ____ / ____ / ________")
    style_run(r_date, size=8, italic=True, color=GRAY_TEXT)

def add_signature_block(doc, parties, height_lines=4):
    """
    Bloque de firmas — VERSIÓN 3.0.

    parties: lista de labels. Se mapea automáticamente:
      - "RR.HH." → reemplazado por "DIRECTORA GERENTE" (Esnatlim Simoza)
      - "RR.HH. / SST" → "DIRECTORA GERENTE" (rol SST delegado a la Directora)
      - "RR.HH. / Jefe directo" → "DIRECTORA GERENTE"
      - "RR.HH. / Contabilidad" → "DIRECTORA GERENTE"
      - "RR.HH. / Gerencia" → "DIRECTORA GERENTE"
      - "LA EMPRESA" → "REPRESENTANTE LEGAL" (Alicia Sleiman)
      - Cualquier otro label se deja igual (Trabajador, Médico, etc.)

    Los roles pre-cargados:
      - TRABAJADOR(A) / SOLICITANTE / GERENCIA / TESTIGO → sin datos (campos vacíos para llenar)
      - DIRECTORA GERENTE → Esnatlim Elena Simoza, C.I. V-17.976.287, Directora Gerente
      - REPRESENTANTE LEGAL → Alicia Sleiman de Maksoud, C.I. V-10.112.683, Vicepresidente
    """
    # Mapeo de labels genéricos a roles concretos
    PARTIES_MAP = {
        "RR.HH.": "DIRECTORA GERENTE",
        "RR.HH. / SST": "DIRECTORA GERENTE",
        "RR.HH. / Jefe directo": "DIRECTORA GERENTE",
        "RR.HH. / Contabilidad": "DIRECTORA GERENTE",
        "RR.HH. / Gerencia": "DIRECTORA GERENTE",
        "RR.HH. (responsable)": "DIRECTORA GERENTE",
        "RECURSOS HUMANOS": "DIRECTORA GERENTE",
        "RECURSOS HUMANOS / SST": "DIRECTORA GERENTE",
    }

    # Datos por rol — TODOS los roles de empresa apuntan a Esnatlim (Directora Gerente)
    ROLE_DATA = {
        "DIRECTORA GERENTE": (DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI),
        "REPRESENTANTE LEGAL": (DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI),
        "LA EMPRESA (SST)": (DIRECTORA_NOMBRE, DIRECTORA_CARGO + " (SST)", DIRECTORA_CI),
        "LA EMPRESA": (DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI),
        "DIRECCIÓN MÉDICA": (DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI),
        "SEGURIDAD Y SALUD": (DIRECTORA_NOMBRE, DIRECTORA_CARGO + " (SST)", DIRECTORA_CI),
    }

    n = len(parties)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    total_w = 16.6
    each = total_w / n
    for i in range(n):
        tbl.columns[i].width = Cm(each)

    for i, lab in enumerate(parties):
        # Normalizar y mapear
        normalized = lab.strip()
        if normalized in PARTIES_MAP:
            normalized = PARTIES_MAP[normalized]

        # Determinar datos del rol
        if normalized in ROLE_DATA:
            name, cargo, ci = ROLE_DATA[normalized]
        else:
            name = None  # campos vacíos para llenar (Trabajador, Testigo, Solicitante, etc.)
            cargo = None
            ci = None

        c = tbl.rows[0].cells[i]
        _build_signature_cell(c, normalized, name, cargo, ci, each)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl

def add_signature_block_directora(doc):
    """Atajo: 2 firmas — Trabajador + Directora Gerente (Esnatlim)."""
    return add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])

def add_signature_block_rep_legal(doc):
    """Atajo: 2 firmas — Trabajador + Representante Legal (Alicia).
    Usar SOLO en contratos y documentos legales/mercantiles."""
    return add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "REPRESENTANTE LEGAL"])

# ============================================================
# Carta de recepción
# ============================================================
def add_reception_letter(doc, doc_name_full, obligatorio_field="obligatorio"):
    """Carta de recepción al final del documento."""
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(8)
    sep.paragraph_format.space_after = Pt(6)
    add_hr(sep, color="0F766E", sz="8")

    add_section(doc, "CARTA DE RECEPCIÓN Y ACEPTACIÓN")

    add_para(doc,
        f"Yo, _____________________________________________, titular de la cédula de "
        f"identidad N° V-___________________, en mi condición de trabajador(a) de "
        f"GRUPO CAVAL 1003, C.A. (ALIKA PETS), declaro por medio de la presente que:",
        size=10, space_after=4)

    add_bullet(doc,
        f"He recibido copia íntegra y legible del documento titulado "
        f"\u00ab{doc_name_full}\u00bb, así como las explicaciones verbales necesarias "
        f"para su correcta comprensión.",
        bold_lead="PRIMERO:  ")

    add_bullet(doc,
        f"He leído en su totalidad el contenido del mencionado documento, comprendo mis "
        f"obligaciones, prohibiciones y deberes, y me comprometo a cumplirlos fielmente "
        f"durante toda la vigencia de mi relación laboral con la empresa.",
        bold_lead="SEGUNDO:  ")

    add_bullet(doc,
        f"Entiendo que el incumplimiento de las disposiciones aquí contenidas podrá "
        f"dar lugar a las sanciones disciplinarias previstas en el Reglamento Interno "
        f"de Trabajo y, según la gravedad, a la terminación de la relación laboral por "
        f"causa justificada conforme al artículo 79 de la Ley Orgánica del Trabajo, los "
        f"Trabajadores y las Trabajadoras (LOTTT).",
        bold_lead="TERCERO:  ")

    add_bullet(doc,
        f"Acepto que la presente firma constituye prueba fehaciente de la entrega y "
        f"recepción del documento, renunciando a alegar desconocimiento de su contenido.",
        bold_lead="CUARTO:  ")

    add_para(doc, "", size=6, space_after=2)
    add_para(doc,
        "En fe de lo cual firmo la presente carta en la ciudad de Los Teques, Estado "
        "Miranda, a los ____ días del mes de ________________ de ________.",
        size=10, space_after=10)

    # Firma: Trabajador + Directora Gerente (rol operativo)
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])
