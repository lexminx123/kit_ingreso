"""
fix_funciones_completar.py — Completa las funciones de los 3 cargos curados,
añadiendo las funciones que un abogado laboralista NO puede dejar fuera
porque son necesarias para la operación, la normativa LOPCYMAT/LOTTT
o la protección legal de la empresa.

Las funciones que envió la empresa son CORRECTAS pero INCOMPLETAS.
Como abogado, Esnatlim no puede dejar cabos sueltos, así que
complementamos cada cargo con:

  Dog Groomer — Funciones faltantes:
    + Recepción de mascotas y entrega al cliente
    + Registro del servicio en sistema
    + Cobro del servicio
    + Reportar al veterinario hallazgos anómalos (ya estaba)
    + Cumplir protocolos de bioseguridad
    + Atender instrucciones de la dueña
    + Manejo de residuos peligrosos
    + Control de insumos (champús, etc.)

  Encargado de Tienda — Funciones faltantes:
    + Control de mermas y devoluciones
    + Arqueo de caja al cierre
    + Reportar incidentes de seguridad
    + Manejo de proveedores y órdenes de compra
    + Cumplir normativas de higiene de alimentos
    + Cumplir política de precios y promociones
    + Conciliación de pagos móviles/transferencias
    + Velar por imagen corporativa

  Auxiliar Veterinario — Funciones faltantes:
    + Recepción y control de citas
    + Cobro de servicios clínicos (ya lo mencionó)
    + Mantenimiento básico de equipos (autoclave, centrífuga)
    + Control de inventario de medicamentos
    + Manejo de cadáveres y residuos biológicos
    + Reporte de enfermedades zoonóticas
    + Atender instrucciones del MV
    + Cumplir LOPCYMAT
"""
import os, sys
sys.path.insert(0, "/home/z/my-project/output")
from _common import *
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

OUT_DIR = "/home/z/my-project/output/03_DESCRIPCION_DE_CARGOS"

# ============================================================
# CARGOS CON FUNCIONES COMPLEMENTADAS
# ============================================================
CARGOS = {
    "groomer": {
        "archivo": "03f_Funciones_Dog_Groomer.docx",
        "cargo": "DOG GROOMER (PELUQUERO CANINO)",
        "cargo_corto": "Dog Groomer",
        "area": "Peluquería Canina",
        "nivel": "Técnico especializado",
        "tipo_contrato": "Tiempo indeterminado (LOTTT art. 65) — período de prueba 30 días",
        "reporta_a": "Encargado(a) de Tienda / Gerente General",
        "supervisa_a": "Nadie directamente",
        "objetivo":
            "Realizar servicios de peluquería y estética canina y felina con técnicas "
            "profesionales, garantizando el bienestar de las mascotas, la satisfacción de "
            "los clientes y el cumplimiento de los protocolos de higiene, bioseguridad y "
            "manejo amigable de animales establecidos por la empresa.",
        "funciones": {
            "CORTES Y PELUQUERÍA": [
                "Cortes de raza y cortes higiénicos.",
                "Técnicas de tijera, máquina y rastrillo según tipo de pelo y raza.",
                "Asesorar al cliente sobre el corte más adecuado.",
            ],
            "BAÑO Y SECADO": [
                "Bañar con champús y acondicionadores según tipo de piel y pelaje.",
                "Aplicar tratamientos antipulgas y antiparasitarios tópicos.",
                "Secar con toalla y secador profesional, cepillar y desenredar.",
                "Limpiar oídos y cortar uñas.",
            ],
            "HIGIENE COMPLEMENTARIA": [
                "Perfumes, laqueado y terminaciones estéticas.",
            ],
            "EVALUACIÓN Y MANEJO ANIMAL": [
                "Evaluar estado general del animal al ingreso.",
                "Reportar al veterinario cualquier hallazgo anómalo.",
                "Aplicar contención ética y manejo con refuerzo positivo.",
                "Decidir si un animal no está en condiciones de ser atendido.",
            ],
            "LIMPIEZA Y MANTENIMIENTO": [
                "Desinfectar bañera, mesa, jaulas, tijeras y máquinas entre pacientes.",
                "Segregar residuos conforme COVENIN 2747-93.",
                "Mantener equipos en buen estado y reportar fallas.",
            ],
            "ATENCIÓN AL CLIENTE Y GESTIÓN DEL SERVICIO": [
                "Recibir a la mascota del cliente y verificar sus datos.",
                "Registrar el servicio en el sistema POS y agenda de citas.",
                "Cobrar el servicio al cliente (efectivo, tarjeta, transferencia o pago móvil).",
                "Entregar la mascota al cliente al finalizar el servicio y explicar cuidados posteriores.",
                "Atender llamadas y consultas sobre servicios de peluquería.",
                "Programar y confirmar citas futuras.",
            ],
            "CONTROL DE INSUMOS": [
                "Reportar niveles bajos de champús, acondicionadores, perfumes y otros insumos.",
                "Cuidar el uso de productos y reportar mermas.",
                "Mantener orden y limpieza del área de almacenamiento de insumos.",
            ],
            "BIOSEGURIDAD Y LOPCYMAT": [
                "Cumplir los protocolos de bioseguridad NT-01-2008.",
                "Usar obligatoriamente el EPP asignado (bata, guantes, mascarilla, gafas, botas).",
                "Mantener vacuna antirrábica pre-exposición vigente.",
                "Reportar inmediatamente cualquier incidente, accidente o casi-accidente (mordeduras, arañazos, caídas, cortes con herramientas).",
                "Participar en las capacitaciones de seguridad y salud laboral que programe la empresa.",
            ],
            "POLÍTICAS Y CONFIDENCIALIDAD": [
                "Cumplir el Reglamento Interno, el Código de Conducta y las políticas de la empresa.",
                "Mantener confidencialidad de datos de clientes, mascotas y la empresa.",
                "Obtener autorización previa y por escrito de la Dirección para publicar en redes sociales cualquier contenido que involucre la peluquería, mascotas, procedimientos o personal de la empresa.",
                "No contactar clientes de la empresa para ofrecer servicios externos durante la relación laboral y por 12 meses después.",
            ],
        },
        "competencias_tecnicas": [
            "Formación o certificación en peluquería canina / dog grooming.",
            "Experiencia mínima de 1 año como dog groomer (preferible).",
            "Conocimiento de técnicas de corte por raza (poodle, schnauzer, shih tzu, etc.).",
            "Manejo de herramientas: cuchillas, tijeras, máquinas eléctricas, secadores.",
            "Conocimiento de productos de higiene y cosmética veterinaria.",
            "Capacitación en manejo y contención de animales.",
            "Conocimiento básico de operación de POS y agenda.",
        ],
        "competencias_blandas": [
            "Paciencia y empatía con animales.",
            "Capacidad de lectura del lenguaje corporal canino/felino.",
            "Atención al detalle estético.",
            "Comunicación cordial con propietarios.",
            "Cuidado y respeto por el bienestar animal.",
        ],
        "requisitos": [
            "Cédula de identidad venezolana vigente.",
            "Certificación o portafolio que acredite experiencia en dog grooming.",
            "Vacuna antirrábica pre-exposición vigente.",
            "Sin antecedentes penales.",
            "Disponibilidad para fines de semana.",
        ],
        "lugar_trabajo": "Área de peluquería canina de ALIKA PETS, Av. Francisco de Miranda, Los Teques.",
        "jornada": "Tiempo completo, 40 horas semanales, con turnos rotativos incluyendo fines de semana.",
        "epp": "Bata/gabacha impermeable, guantes de nitrilo, mascarilla quirúrgica o con filtro, gafas protectoras, botas cerradas antideslizantes e impermeables, tapones auditivos (uso prolongado de secadores).",
        "riesgos_principales": "Mordeduras/arañazos, zoonosis cutáneas (tiña, sarna), cortes con herramientas, riesgo eléctrico en zona húmeda, posturas forzadas, asma/alergias por pelaje, dermatitis por humedad.",
    },
    "encargada_tienda": {
        "archivo": "03c_Funciones_Encargado_Tienda.docx",
        "cargo": "ENCARGADO(A) DE TIENDA",
        "cargo_corto": "Encargado(a) de Tienda",
        "area": "Tienda de Mascotas",
        "nivel": "Coordinación / Supervisión",
        "tipo_contrato": "Tiempo determinado (6 meses + prórroga) — período de prueba 30 días",
        "reporta_a": "Gerente General",
        "supervisa_a": "Asistentes de Ventas, personal de limpieza de tienda",
        "objetivo":
            "Supervisar las operaciones diarias de la tienda de mascotas, garantizando la "
            "correcta atención a clientes, el manejo adecuado de caja, el control de "
            "inventarios, la coordinación con proveedores y la implementación de "
            "estrategias comerciales que aumenten las ventas y fomenten la fidelización.",
        "funciones": {
            "ATENCIÓN AL CLIENTE": [
                "Saludar y recibir al cliente.",
                "Asesorar en alimentos, accesorios, higiene y antiparasitarios.",
                "Informar sobre promociones, programas de fidelización, devoluciones y garantías.",
                "Atender y escalar quejas o reclamos al Gerente General.",
                "Velar por la imagen corporativa de la tienda (uniforme, presentación personal, trato cordial).",
            ],
            "MANEJO DE CAJA": [
                "Abrir y cerrar caja diariamente (monto inicial, ingresos, egresos, arqueo).",
                "Cobrar con efectivo, punto de venta, transferencia o tarjeta.",
                "Conciliar pagos móviles y transferencias bancarias con los reportes del sistema POS.",
                "Emitir facturas y comprobantes.",
                "Realizar depósitos bancarios diarios y conciliación.",
                "Registrar operaciones en el sistema POS.",
                "Reportar diferencias o faltantes al Gerente General en un plazo máximo de 24 horas.",
            ],
            "REGISTRO Y CONTROL DE INVENTARIO": [
                "Mantener kárdex actualizado (entradas, salidas, mermas).",
                "Realizar conteos parciales semanales e inventario general mensual.",
                "Verificar fechas de vencimiento y rotación FIFO (First In, First Out).",
                "Coordinar recepción de mercancía con proveedores y verificar cantidades y estado.",
                "Reponer góndolas, exhibidores y neveras (cadena de frío).",
                "Registrar entradas y salidas de mercancía en el sistema POS.",
                "Comunicar niveles bajos de inventario al Gerente General.",
                "Participar en inventarios físicos mensuales.",
                "Controlar y reportar mermas, daños, productos vencidos y devoluciones a proveedores.",
                "Generar órdenes de compra y enviarlas a los proveedores según los niveles mínimos de stock.",
            ],
            "CARGA Y DESCARGA": [
                "Recibir mercancía de proveedores y verificar cantidades.",
                "Trasladar cajas a bodega y góndolas.",
                "Estibar, organizar y rotular productos conforme a las normas de almacenamiento.",
                "Verificar que los productos perecederos (alimentos refrigerados, mascotas vivas si aplica) se almacenen conforme a las normas sanitarias.",
            ],
            "MANTENIMIENTO DEL ÁREA": [
                "Velar por limpieza y orden de góndolas, neveras, mostrador y bodega.",
                "Revisar funcionamiento de equipos (neveras, balanzas, POS, iluminación).",
                "Revisar fechas de vencimiento periódicamente.",
                "Reportar fallas al Gerente General en un plazo máximo de 24 horas.",
            ],
            "POLÍTICAS Y REGISTROS": [
                "Cumplir y hacer cumplir las políticas internas de la empresa.",
                "Aplicar el procedimiento de reporte de incidentes conforme a la LOPCYMAT.",
                "Velar por la exhibición de precios y señalización.",
                "Mantener confidencialidad de datos de clientes, proveedores, márgenes y estrategias comerciales.",
                "Cumplir las normas de higiene y manipulación de alimentos para mascotas.",
                "Cumplir la política de precios, promociones y descuentos fijada por la Dirección.",
                "Obtener autorización previa y por escrito de la Dirección para publicar en redes sociales cualquier contenido que involucre la tienda, productos, clientes o personal.",
            ],
            "SEGURIDAD Y LOPCYMAT": [
                "Cumplir y hacer cumplir las normas de seguridad de la tienda.",
                "Reportar incidentes de seguridad (robo, hurto, asalto, sospechas) al Gerente General.",
                "Mantener el área de trabajo en condiciones seguras para clientes y personal.",
                "Cumplir el protocolo de apertura y cierre del establecimiento.",
                "Participar en las capacitaciones de seguridad y salud laboral.",
            ],
        },
        "competencias_tecnicas": [
            "Experiencia previa en atención al cliente y manejo de caja (mínimo 1 año).",
            "Manejo de sistemas de punto de venta (POS) y kárdex.",
            "Conocimientos de control de inventarios (FIFO, mermas, rotación, órdenes de compra).",
            "Conocimiento de productos para mascotas (alimentos, accesorios, higiene, antiparasitarios).",
            "Capacidad para coordinar con proveedores y conciliar depósitos bancarios.",
            "Conocimientos básicos de normativas LOPCYMAT, LOTTT.",
        ],
        "competencias_blandas": [
            "Orientación comercial y al cliente.",
            "Honestidad y transparencia en manejo de efectivo.",
            "Capacidad de organización y orden.",
            "Comunicación efectiva y resolución de quejas.",
            "Trabajo en equipo y liderazgo básico.",
        ],
        "requisitos": [
            "Cédula de identidad venezolana vigente.",
            "Sin antecedentes penales.",
            "Disponibilidad para trabajar fines de semana y feriados.",
            "Experiencia mínima de 1 año en retail (preferible).",
        ],
        "lugar_trabajo": "Tienda de mascotas ALIKA PETS, Av. Francisco de Miranda, Los Teques.",
        "jornada": "Tiempo completo, 40 horas semanales, con turnos rotativos incluyendo fines de semana.",
        "epp": "Uniforme corporativo, calzado cerrado antideslizante; guantes de nitrilo para carga de mercancía y limpieza.",
        "riesgos_principales": "Trastornos musculoesqueléticos por bipedestación y carga de mercancía, caídas por cajas o suelos mojados, exposición a polvos y alérgenos, riesgo de asalto, exposición a productos de limpieza.",
    },
    "auxiliar_vet": {
        "archivo": "03e_Funciones_Auxiliar_Veterinario.docx",
        "cargo": "AUXILIAR VETERINARIO(A)",
        "cargo_corto": "Auxiliar Veterinario(a)",
        "area": "Clínica Veterinaria",
        "nivel": "Técnico / Asistencial",
        "tipo_contrato": "Tiempo determinado (3 meses + prórroga) — período de prueba 15 días",
        "reporta_a": "Médico(a) Veterinario(a) / Encargado(a) de Clínica",
        "supervisa_a": "Nadie directamente",
        "objetivo":
            "Asistir al médico veterinario en los diversos procedimientos y servicios que "
            "presta la clínica, brindar cuidado directo a los animales hospitalizados, "
            "interactuar con los dueños y mantener limpias y desinfectadas las instalaciones, "
            "garantizando un ambiente higiénico y seguro.",
        "funciones": {
            "CONSULTAS": [
                "Recibir, pesar y tomar signos vitales del paciente.",
                "Preparar y desinfectar consultorio entre pacientes.",
                "Asistir en contención y manejo del animal.",
                "Registrar signos vitales en historia clínica.",
            ],
            "TRATAMIENTOS": [
                "Administrar medicamentos según indicación médica.",
                "Aplicar antiparasitarios y fluidoterapia.",
                "Retirar catéteres y puntos.",
            ],
            "LABORATORIO": [
                "Tomar muestras de orina, heces y secreciones.",
                "Operar centrífuga, microscopio y analizadores.",
                "Realizar mantenimiento básico de equipos de laboratorio (limpieza, calibración, reporte de fallas).",
            ],
            "CIRUGÍA": [
                "Preparar paciente (ayuno, rasurado, vía venosa).",
                "Preparar quirófano (material estéril, instrumental, monitores).",
                "Asistir como instrumentista y circulante.",
                "Monitorear signos vitales transoperatorios.",
                "Cuidar recuperación postoperatoria.",
                "Limpieza del área e instrumentos luego del procedimiento quirúrgico.",
            ],
            "HOSPITALIZACIÓN": [
                "Alimentación, hidratación, paseo, higiene y medicación de hospitalizados.",
                "Rondas de control por turno y registro de signos vitales en historia clínica.",
                "Limpiar y desinfectar jaulas y comederos.",
                "Reportar al MV cualquier cambio en el estado del paciente hospitalizado.",
            ],
            "ATENCIÓN AL CLIENTE": [
                "Recibir al cliente y a la mascota.",
                "Brindar información y/o asesoría al cliente en caso de ser necesario sobre los servicios.",
                "Hacerle entrega al cliente de los formatos correspondientes para ser llenados tales como Consentimiento informado, acta voluntaria, etc.",
                "Explicar indicaciones postoperatorias y de medicación.",
                "Cobrar servicios clínicos si el encargado no está disponible.",
                "Atender llamadas y consultas sobre los servicios de la clínica.",
                "Programar y confirmar citas de consulta, control, vacunación y desparasitación.",
            ],
            "CONTROL DE MEDICAMENTOS E INSUMOS": [
                "Verificar y reportar niveles de stock de medicamentos e insumos médicos al Encargado(a) de Clínica.",
                "Verificar y reportar fechas de vencimiento de medicamentos.",
                "Manejar y segregar residuos biológicos peligrosos (jeringas, gasas contaminadas, órganos) conforme a la norma COVENIN 2747-93.",
                "Manejar y disponer adecuadamente de cadáveres de animales conforme a las normas sanitarias.",
            ],
            "HIGIENE Y BIOSEGURIDAD": [
                "Cumplir protocolos NT-01-2008 y COVENIN 2747-93.",
                "Segregar residuos por color de bolsa (rojo=bioinfeccioso, negro=común, amarillo=cortopunzante, verde=reciclable).",
                "Esterilizar instrumental (autoclave).",
                "Realizar mantenimiento básico de la autoclave (limpieza, control de presión, reporte de fallas).",
                "Usar obligatoriamente el EPP asignado (bata, guantes, mascarilla, gafas, botas).",
                "Mantener vacuna antirrábica pre-exposición vigente.",
            ],
            "LOPCYMAT Y REPORTES": [
                "Reportar inmediatamente al Encargado(a) de Clínica o al Gerente cualquier incidente, accidente o casi-accidente (mordeduras, arañazos, cortes con agujas, exposición a fluidos, caídas).",
                "Participar en las capacitaciones de seguridad y salud laboral que programe la empresa.",
                "Atender y cumplir las instrucciones del Médico Veterinario(a) y del Encargado(a) de Clínica.",
            ],
            "POLÍTICAS Y CONFIDENCIALIDAD": [
                "Cumplir el Reglamento Interno, el Código de Conducta y las políticas de la empresa.",
                "Mantener confidencialidad de datos de clientes, mascotas, historias clínicas y estrategias comerciales.",
                "Obtener autorización previa y por escrito de la Dirección para publicar en redes sociales cualquier contenido que involucre la clínica, pacientes, procedimientos o personal.",
                "No contactar clientes de la empresa para ofrecer servicios externos durante la relación laboral y por 12 meses después.",
            ],
        },
        "competencias_tecnicas": [
            "Formación como Auxiliar Veterinario, Técnico Superior Universitario (TSU) en Veterinaria o afín (preferible).",
            "Experiencia previa en clínica veterinaria (mínimo 6 meses, preferible).",
            "Conocimiento de técnicas de sujeción y manejo de animales (cánidos, felinos).",
            "Conocimiento de instrumental quirúrgico, procedimientos de esterilización y operación de autoclave.",
            "Capacidad para tomar signos vitales y registrarlos en historia clínica.",
            "Manejo de laboratorio básico (centrífuga, microscopio, analizadores).",
            "Conocimiento básico de operación de POS y agenda de citas.",
        ],
        "competencias_blandas": [
            "Empatía y paciencia con animales.",
            "Capacidad de trabajo en equipo con el MV.",
            "Atención al detalle en el seguimiento de indicaciones médicas.",
            "Disposición para tareas de limpieza y mantenimiento.",
            "Comunicación cordial con propietarios de mascotas.",
        ],
        "requisitos": [
            "Cédula de identidad venezolana vigente.",
            "Formación técnica o certificación en auxiliar veterinario (preferible).",
            "Vacuna antirrábica pre-exposición vigente.",
            "Sin antecedentes penales.",
            "Disponibilidad para turnos rotativos.",
        ],
        "lugar_trabajo": "Clínica veterinaria ALIKA PETS (áreas de consulta, hospitalización, quirófano, laboratorio, recepción), Los Teques.",
        "jornada": "Tiempo completo, 40 horas semanales, con turnos rotativos.",
        "epp": "Bata impermeable, guantes de nitrilo, mascarilla quirúrgica, gafas protectoras (en limpieza), botas cerradas antideslizantes.",
        "riesgos_principales": "Mordeduras/arañazos por sujeción de animales, exposición a fluidos biológicos en limpieza de jaulas, exposición a zoonosis, pinchazos con cortopunzantes, exposición a desinfectantes químicos, trastornos musculoesqueléticos.",
    },
}

# ============================================================
# Funciones auxiliares (iguales a gen_actualizar_cargos_curados.py)
# ============================================================
def render_identificacion(doc, cfg):
    add_section(doc, "1. IDENTIFICACIÓN DEL CARGO")
    tbl = doc.add_table(rows=4, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(3.5), Cm(5.0), Cm(3.5), Cm(5.0)]):
        tbl.columns[i].width = w
    datos = [
        ("Título del cargo:", cfg["cargo"], "Nivel jerárquico:", cfg["nivel"]),
        ("Departamento / Área:", cfg["area"], "Tipo de contrato:", cfg["tipo_contrato"]),
        ("Reporta a:", cfg["reporta_a"], "Supervisa a:", cfg["supervisa_a"]),
        ("Cargo:", cfg["cargo_corto"], "", ""),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(datos):
        c = tbl.rows[ri].cells[0]; c.width = Cm(3.5)
        write_cell(c, l1, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[1]; c.width = Cm(5.0)
        write_cell(c, v1, size=9, bg=WHITE)
        if l2:
            c = tbl.rows[ri].cells[2]; c.width = Cm(3.5)
            write_cell(c, l2, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
            c = tbl.rows[ri].cells[3]; c.width = Cm(5.0)
            write_cell(c, v2, size=9, bg=WHITE)
        else:
            merged = tbl.rows[ri].cells[2].merge(tbl.rows[ri].cells[3])
            merged.text = ""
            set_cell_borders(merged, color="FFFFFF", sz="0")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

def render_objetivo(doc, cfg):
    add_section(doc, "2. OBJETIVO GENERAL DEL CARGO")
    add_para(doc, cfg["objetivo"], size=10, space_after=8)

def render_funciones(doc, cfg):
    add_section(doc, "3. FUNCIONES ESPECÍFICAS")
    add_para(doc,
        "Sin perjuicio de las demás funciones que le asigne la empresa, el(la) "
        f"{cfg['cargo_corto']} tendrá las siguientes funciones específicas, organizadas por "
        "área de actividad:",
        size=10, italic=True, color=GRAY_TEXT, space_after=6)
    for area, funciones in cfg["funciones"].items():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"▸ {area}")
        style_run(r, size=10, bold=True, color=TEAL_DARK)
        for fn in funciones:
            add_bullet(doc, fn, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

def render_competencias(doc, cfg):
    add_section(doc, "4. COMPETENCIAS Y REQUISITOS DEL CARGO")
    add_para(doc, "Competencias técnicas y profesionales:", size=10, bold=True, color=TEAL_DARK, space_after=4)
    for comp in cfg["competencias_tecnicas"]:
        add_bullet(doc, comp, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
    add_para(doc, "Competencias blandas / actitudinales:", size=10, bold=True, color=TEAL_DARK, space_after=4)
    for comp in cfg["competencias_blandas"]:
        add_bullet(doc, comp, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
    add_para(doc, "Requisitos académicos y legales:", size=10, bold=True, color=TEAL_DARK, space_after=4)
    for req in cfg["requisitos"]:
        add_bullet(doc, req, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

def render_condiciones(doc, cfg):
    add_section(doc, "5. CONDICIONES DE TRABAJO Y RIESGOS")
    add_para(doc,
        "El cargo se desempeña en condiciones que implican exposición a riesgos "
        "ocupacionales específicos, los cuales están detallados en la Notificación de "
        "Riesgos correspondiente, conforme al artículo 56 de la LOPCYMAT. Las condiciones "
        "generales son:",
        size=10, space_after=4)
    add_bullet(doc, cfg["lugar_trabajo"], bold_lead="Lugar de trabajo:  ")
    add_bullet(doc, cfg["jornada"], bold_lead="Jornada:  ")
    add_bullet(doc, cfg["epp"], bold_lead="EPP obligatorio:  ")
    add_bullet(doc, cfg["riesgos_principales"], bold_lead="Riesgos principales:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

def render_carta_recepcion(doc, cfg):
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(10)
    sep.paragraph_format.space_after = Pt(6)
    add_hr(sep, color="0F766E", sz="6")
    add_section(doc, "CARTA DE RECEPCIÓN DE LA DESCRIPCIÓN DE CARGO")
    add_para(doc,
        f"Yo, _____________________________________________, titular de la cédula de identidad "
        f"N° V-___________________, en mi condición de trabajador(a) de GRUPO CAVAL 1003, "
        f"C.A. (ALIKA PETS), designado(a) para el cargo de {cfg['cargo']}, declaro por medio "
        f"de la presente que:",
        size=10, space_after=6)
    add_bullet(doc,
        "He recibido copia íntegra y legible de la Descripción de Cargo correspondiente a "
        "mi posición, así como las explicaciones verbales necesarias para su correcta "
        "comprensión.", bold_lead="PRIMERO:  ")
    add_bullet(doc,
        "He leído en su totalidad el contenido del mencionado documento, comprendo las "
        "funciones, deberes y responsabilidades asignadas a mi cargo, y me comprometo a "
        "cumplirlas fielmente durante toda la vigencia de mi relación laboral con la "
        "empresa.", bold_lead="SEGUNDO:  ")
    add_bullet(doc,
        "Reconozco que las funciones aquí descritas me fueron suficientemente explicadas y "
        "notificadas por la empresa, conforme al artículo 26 de la Ley Orgánica del "
        "Trabajo, los Trabajadores y las Trabajadoras (LOTTT).", bold_lead="TERCERO:  ")
    add_bullet(doc,
        "Acepto que la presente firma constituye prueba fehaciente de la entrega y "
        "recepción del documento, renunciando a alegar desconocimiento de las funciones "
        "asignadas.", bold_lead="CUARTO:  ")
    add_para(doc, "", size=6, space_after=2)
    add_para(doc,
        "En fe de lo cual firmo la presente carta en la ciudad de Los Teques, Estado "
        "Miranda, a los ____ días del mes de ________________ de ________.",
        size=10, space_after=10)
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])

# ============================================================
# Build documento
# ============================================================
def build_descripcion_cargo(cfg):
    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=1.8)
    add_membrete(doc, "DESCRIPCIÓN DE CARGO", cfg["cargo_corto"], version="Versión 3.3  ·  RR.HH.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("DESCRIPCIÓN DE CARGO Y FUNCIONES")
    style_run(r, size=14, bold=True, color=TEAL_DARK)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(cfg["cargo"])
    style_run(r2, size=12, bold=True, color=TEAL_DARK)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(10)
    add_hr(sep, color="0F766E", sz="8")
    render_identificacion(doc, cfg)
    render_objetivo(doc, cfg)
    render_funciones(doc, cfg)
    render_competencias(doc, cfg)
    render_condiciones(doc, cfg)
    render_carta_recepcion(doc, cfg)
    add_footer(section, f"Descripción Cargo — {cfg['cargo_corto']} v3.3")
    out = os.path.join(OUT_DIR, cfg["archivo"])
    doc.save(out)
    return out

# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print("Complementando funciones de los 3 cargos curados...")
    for key, cfg in CARGOS.items():
        out = build_descripcion_cargo(cfg)
        size_kb = os.path.getsize(out) / 1024
        print(f"  ✓ {os.path.basename(out)}  ({size_kb:.1f} KB)")
    print("\n3 descripciones complementadas correctamente.")
