"""
fix_honorarios_no_discriminar_salario.py — Corrige los contratos de honorarios
profesionales eliminando la discriminación salarial (que era inadmisible).

En HP NO se discrimina salario base + cestaticket + bonos. Eso replica
estructura laboral y facilita la re-calificación por un juez.

Se reemplaza por:
  - Un ÚNICO concepto de honorarios profesionales mensuales
  - Pago en USD 250 mensuales (referenciado al BCV para pago en Bs)
  - Sin mención de cestaticket, bono de transporte, buen vivir, etc.
  - Forma de pago: mensual o quincenal según se acuerde
  - Retención ISLR según Art. 27
  - Facturación obligatoria
"""
import os, sys
sys.path.insert(0, "/home/z/my-project/output")
from _common import *
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# ============================================================
# 1. CONTRATO HP VETERINARIO — Reemplazar cláusula cuarta
# ============================================================
def fix_contrato_honorarios_vet():
    out = "/home/z/my-project/output/Contrato_Honorarios_Medico_Veterinario.docx"

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "CONTRATO DE HONORARIOS", "Profesional Veterinario(a)",
                 version="Versión 1.1  ·  Dirección")
    add_doc_title(doc, "CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Médico(a) Veterinario(a) — Honorarios Profesionales")
    style_run(r, size=11, italic=True, color=GRAY_TEXT)

    add_para(doc,
        "Entre los suscritos: GRUPO CAVAL 1003, C.A., sociedad mercantil de domicilio en "
        "Av. Francisco de Miranda, Local N° 1, Sector Francisco de Miranda, Los Teques, "
        "Estado Miranda, Zona Postal 1201, RIF N° J501662533, marca comercial ALIKA PETS, "
        "dedicada a la actividad de clínica veterinaria, tienda de mascotas y peluquería "
        "canina, en adelante «LA EMPRESA», representada en este acto por su Directora "
        "Gerente, ciudadana ESNATLIM ELENA SIMOZA, titular de la cédula de identidad N° "
        "V-17.976.287, por una parte; y por la otra, el(la) ciudadano(a) "
        "___________________________________________, venezolano(a), mayor de edad, titular "
        "de la cédula de identidad N° V-___________________, inscrito(a) en el Colegio de "
        "Médicos Veterinarios de Venezuela (CMV) bajo el N° __________________, con RIF "
        "personal N° __________________, en adelante «EL PROFESIONAL», quienes declaran "
        "ser mayores de edad y hábiles para contratar, han convenido en celebrar el "
        "presente CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES, el cual se regirá "
        "por las siguientes cláusulas:", size=10, space_after=8)

    # CLÁUSULA 1: OBJETO
    add_section(doc, "CLÁUSULA PRIMERA: OBJETO")
    add_para(doc,
        "EL PROFESIONAL se obliga a prestar a LA EMPRESA servicios profesionales de "
        "MEDICINA VETERINARIA, en su carácter de profesional liberal independiente, "
        "comprendiendo entre otras: consultas clínicas generales y especializadas; "
        "procedimientos quirúrgicos; atención de urgencias y emergencias; realización e "
        "interpretación de exámenes complementarios (radiografías, ecografías, "
        "laboratorios); supervisión de pacientes hospitalizados; aplicación de vacunas y "
        "pautas de desparasitación; redacción y firma de recetas, certificados e "
        "informes médicos; y demás servicios propios del ejercicio profesional veterinario.")
    add_para(doc,
        "Los servicios se prestarán CON PLENA AUTONOMÍA TÉCNICA. EL PROFESIONAL decidirá "
        "libremente el diagnóstico, tratamiento, procedimiento y medicación de cada "
        "paciente, conforme a su juicio clínico y a las normas del ejercicio profesional. "
        "LA EMPRESA reconoce esta autonomía técnica y NO impartirá instrucciones que "
        "afecten el criterio profesional del PROFESIONAL.", space_after=6)

    # CLÁUSULA 2: NATURALEZA
    add_section(doc, "CLÁUSULA SEGUNDA: NATURALEZA DEL CONTRATO")
    add_para(doc,
        "Las partes declaran expresamente que el presente contrato es de PRESTACIÓN DE "
        "SERVICIOS PROFESIONALES BAJO LA MODALIDAD DE HONORARIOS, conforme al artículo 40 "
        "de la Ley Orgánica del Trabajo, los Trabajadores y las Trabajadoras (LOTTT), a "
        "la Ley de Ejercicio de la Medicina Veterinaria y a la jurisprudencia del Tribunal "
        "Supremo de Justicia. En consecuencia:")
    add_bullet(doc,
        "EL PROFESIONAL actúa como profesional liberal independiente, sin relación de "
        "subordinación ni dependencia respecto a LA EMPRESA.",
        bold_lead="1. No subordinación:  ")
    add_bullet(doc,
        "EL PROFESIONAL organiza su tiempo y métodos de trabajo de manera autónoma, "
        "sujeto únicamente a los horarios de turnos acordados y a la entrega de resultados.",
        bold_lead="2. Autonomía técnica:  ")
    add_bullet(doc,
        "EL PROFESIONAL puede prestar servicios a otros consultorios, clínicas y "
        "particulares, en horarios distintos a los pactados con LA EMPRESA.",
        bold_lead="3. No exclusividad:  ")
    add_bullet(doc,
        "La remuneración se pagará en calidad de HONORARIOS PROFESIONALES, sin generar "
        "prestaciones sociales, utilidades, vacaciones, bono vacacional ni indemnizaciones "
        "laborales de ninguna naturaleza.",
        bold_lead="4. No salarial:  ")
    add_bullet(doc,
        "EL PROFESIONAL asume el riesgo técnico y civil de su actuación profesional, "
        "debiendo mantener vigente póliza de Seguro de Responsabilidad Civil Profesional.",
        bold_lead="5. Asume riesgo:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # CLÁUSULA 3: HORARIOS
    add_section(doc, "CLÁUSULA TERCERA: HORARIOS Y TURNOS")
    add_para(doc,
        "Las partes acuerdan que EL PROFESIONAL prestará sus servicios en turnos "
        "flexibles, conforme al siguiente esquema:")
    add_bullet(doc, "Días: de lunes a sábado (con un día de descanso acordado semanal).")
    add_bullet(doc, "Turno: _____ horas a _____ horas (con receso de 1 hora para almuerzo).")
    add_bullet(doc, "Guardias de emergencia: ____ veces al mes (según calendario rotativo).")
    add_para(doc,
        "EL PROFESIONAL podrá modificar la distribución semanal de sus turnos previa "
        "coordinación con LA EMPRESA, siempre que asegure la cobertura mínima pactada. "
        "Esta flexibilidad horaria es esencial para preservar la naturaleza no laboral "
        "del contrato.", space_after=6)

    # CLÁUSULA 4: HONORARIOS Y FORMA DE PAGO (CORREGIDA — sin discriminación salarial)
    add_section(doc, "CLÁUSULA CUARTA: HONORARIOS Y FORMA DE PAGO")
    add_para(doc,
        "Por los servicios profesionales prestados, LA EMPRESA pagará a EL PROFESIONAL "
        "la suma mensual única de USD 250,00 (DÓLARES DE LOS ESTADOS UNIDOS DE AMÉRICA "
        "DOSCIENTOS CINCUENTA CON 00/100), como HONORARIOS PROFESIONALES por la totalidad "
        "de los servicios objeto del presente contrato. Esta suma constituye la "
        "contraprestación íntegra y única por todos los servicios prestados, sin que "
        "proceda discriminación alguna en concepto de salario base, cestaticket, bono de "
        "alimentación, bono de transporte, bono vacacional, utilidades, prestaciones "
        "sociales o cualquier otro concepto de naturaleza salarial o laboral.",
        size=10, space_after=4)

    add_para(doc,
        "NATURALEZA DE LOS HONORARIOS: Las partes dejan expresamente constancia de que la "
        "suma mensual de USD 250,00 corresponde exclusivamente a HONORARIOS PROFESIONALES "
        "por la prestación de servicios profesionales veterinarios independientes. En "
        "consecuencia, NO constituye salario a ningún efecto legal, NO genera prestaciones "
        "sociales ni antigüedad, NO está sujeto a la LOTTT, y NO se calculará sobre ella "
        "vacaciones, utilidades, bono vacacional ni indemnizaciones de naturaleza laboral. "
        "Su único régimen aplicable es el del Código Civil, Código de Comercio y la Ley de "
        "Impuesto sobre la Renta (ISLR).",
        size=10, italic=True, color=GRAY_TEXT, space_after=4)

    add_para(doc,
        "MONEDA DE PAGO Y CONVERSIÓN CAMBIARIA: El monto de los honorarios se acuerda "
        "referenciado en moneda extranjera (USD). Si el pago efectivo se realiza en "
        "Bolívares (Bs.), la conversión se efectuará tomando como única referencia el "
        "tipo de cambio oficial publicado por el Banco Central de Venezuela (BCV) "
        "vigente para el día en que se realice el pago efectivo y oportuno, evitando así "
        "el impacto de la devaluación. En caso de pactarse el pago directo en USD en "
        "efectivo, LA EMPRESA entregará billetes aptos para la circulación.",
        size=10, space_after=4)

    add_para(doc,
        "FORMA DE PAGO: Mensual, dentro de los primeros cinco (5) días hábiles del mes "
        "siguiente al de prestación de los servicios, mediante transferencia bancaria a "
        "la cuenta del PROFESIONAL (Banco: _______________, Cuenta N°: _______________). "
        "El pago de honorarios estará sujeto a la retención del 3% por concepto de "
        "Impuesto sobre la Renta (ISLR) conforme al artículo 27 de la Ley de ISLR, "
        "retención que LA EMPRESA enterará al SENIAT dentro de los primeros 15 días del "
        "mes siguiente y entregará comprobante al PROFESIONAL.",
        size=10, space_after=4)

    add_para(doc,
        "FACTURACIÓN: EL PROFESIONAL se obliga a emitir factura o recibo por cada pago "
        "recibido, con sus datos de RIF, número de control, número de factura y "
        "descripción del servicio. Sin factura no procederá el pago. La factura debe "
        "ser emitida a nombre de GRUPO CAVAL 1003, C.A., RIF J501662533.",
        size=10, italic=True, color=GRAY_TEXT, space_after=6)

    # CLÁUSULA 5: OBLIGACIONES DEL PROFESIONAL
    add_section(doc, "CLÁUSULA QUINTA: OBLIGACIONES DEL PROFESIONAL")
    add_bullet(doc, "Prestar los servicios profesionales con diligencia, calidad y oportunidad.")
    add_bullet(doc, "Mantener vigente la inscripción en el CMV y la colegiatura profesional.")
    add_bullet(doc, "Mantener póliza de Seguro de Responsabilidad Civil Profesional vigente y entregar copia a LA EMPRESA.")
    add_bullet(doc, "Cumplir los protocolos clínicos, de bioseguridad (NT-01-2008) y manejo de sustancias controladas (SENAC).")
    add_bullet(doc, "Llevar y mantener actualizadas las historias clínicas de los pacientes atendidos, las cuales son propiedad de LA EMPRESA.")
    add_bullet(doc, "Redactar y firmar recetas, certificados, consentimientos informados e informes médicos.")
    add_bullet(doc, "Atender urgencias y emergencias según turnos asignados.")
    add_bullet(doc, "Mantener confidencialidad de la información de LA EMPRESA, clientes y pacientes (durante el contrato y por 5 años después).")
    add_bullet(doc, "Emitir facturas por cada pago y declarar el ISLR anualmente ante el SENIAT.")
    add_bullet(doc, "No contactar clientes de LA EMPRESA para ofrecer servicios externos durante la vigencia del contrato y por 12 meses después de su terminación (cláusula de no competencia parcial post-contractual).")
    add_bullet(doc, "OBTENER AUTORIZACIÓN PREVIA Y POR ESCRITO de LA EMPRESA para publicar en redes sociales o medios cualquier contenido que se origine en las instalaciones de la clínica, que involucre pacientes, mascotas, procedimientos, personal, uniformes o cualquier elemento identificable con LA EMPRESA. La publicación sin autorización previa constituye incumplimiento grave del presente contrato.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # CLÁUSULA 6: EQUIPOS
    add_section(doc, "CLÁUSULA SEXTA: EQUIPOS, HERRAMIENTAS Y BIENES")
    add_para(doc,
        "LA EMPRESA pondrá a disposición del PROFESIONAL todos los equipos médicos, "
        "instrumental clínico y quirúrgico, equipos de laboratorio, de radiología, "
        "monitores, materiales, insumos, medicamentos y demás bienes necesarios para la "
        "prestación de los servicios veterinarios. Estos bienes son propiedad exclusiva "
        "de LA EMPRESA y se entregan en comodato para el ejercicio de las funciones del "
        "PROFESIONAL.", size=10, space_after=4)
    add_para(doc,
        "EL PROFESIONAL se obliga a: (i) cuidar los equipos como un buen padre de "
        "familia; (ii) reportar de inmediato cualquier daño, deterioro, falla o pérdida; "
        "(iii) no sustraer, trasladar a terceros, ni utilizar los equipos para fines "
        "personales o externos; (iv) mantener los equipos en condiciones de higiene y "
        "bioseguridad conforme a la NT-01-2008; y (v) devolver todos los equipos al "
        "término del contrato, en el mismo estado en que los recibió, salvo el deterioro "
        "natural por su uso adecuado.", size=10, space_after=4)
    add_para(doc,
        "EQUIPOS PROPIOS DEL PROFESIONAL: EL PROFESIONAL podrá aportar para el ejercicio "
        "de sus funciones su propio estetoscopio, el cual será de su propiedad exclusiva. "
        "Dicho equipo deberá ser identificado e inventariado al inicio del contrato y al "
        "término será retirado por EL PROFESIONAL. Mientras el estetoscopio se encuentre "
        "en las instalaciones de LA EMPRESA, EL PROFESIONAL se obliga a mantenerlo en buen "
        "estado, desinfectado y bajo su responsabilidad. LA EMPRESA no se hace responsable "
        "por pérdida, robo o daño del estetoscopio del PROFESIONAL, salvo dolo o "
        "negligencia grave de la empresa.", size=10, italic=True, color=GRAY_TEXT, space_after=6)

    # CLÁUSULA 7: OBLIGACIONES DE LA EMPRESA
    add_section(doc, "CLÁUSULA SÉPTIMA: OBLIGACIONES DE LA EMPRESA")
    add_bullet(doc, "Pagar los honorarios en la forma y oportunidad pactadas en la Cláusula Cuarta.")
    add_bullet(doc, "Proveer el espacio físico, equipos, instrumental, medicamentos e insumos necesarios para la prestación de los servicios.")
    add_bullet(doc, "Garantizar el mantenimiento de las instalaciones y equipos.")
    add_bullet(doc, "Retener y enterar el 3% de ISLR al SENIAT, y entregar comprobante al PROFESIONAL.")
    add_bullet(doc, "Respetar la autonomía técnica del PROFESIONAL y no impartir instrucciones que afecten su criterio clínico.")
    add_bullet(doc, "Proporcionar acceso a las historias clínicas y registros de los pacientes.")
    add_bullet(doc, "Mantener el Sistema de Vigilancia Médica conforme a la NT-02-2008.")
    add_bullet(doc, "Autorizar o denegar por escrito, en un plazo máximo de 3 días hábiles, las solicitudes del PROFESIONAL para publicar contenido en redes sociales relacionado con la clínica, pacientes o procedimientos.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # CLÁUSULA 8: PROPIEDAD DE HISTORIAS CLÍNICAS
    add_section(doc, "CLÁUSULA OCTAVA: PROPIEDAD DE HISTORIAS CLÍNICAS")
    add_para(doc,
        "Las historias clínicas veterinarias, registros médicos, radiografías, resultados "
        "de laboratorio, fotografías clínicas de pacientes y demás documentación médica "
        "generada durante la prestación del servicio son PROPIEDAD EXCLUSIVA DE LA EMPRESA. "
        "EL PROFESIONAL no podrá sustraerlas, copiarlas, fotografiarlas, divulgarlas ni "
        "utilizarlas para fines personales o de terceros, ni durante ni después de la "
        "vigencia del presente contrato.")
    add_para(doc,
        "EL PROFESIONAL entregará a LA EMPRESA todas las historias clínicas y registros "
        "al término del contrato. La violación de esta cláusula generará responsabilidad "
        "civil por daños y perjuicios.", space_after=6)

    # CLÁUSULA 9: PUBLICACIONES Y REDES SOCIALES
    add_section(doc, "CLÁUSULA NOVENA: PUBLICACIONES Y REDES SOCIALES")
    add_para(doc,
        "Toda publicación, difusión o comunicación pública en redes sociales, medios "
        "tradicionales, plataformas digitales o cualquier canal de comunicación que "
        "incluya contenido originado en LA EMPRESA — incluyendo pero no limitado a: "
        "fotografías o videos de las instalaciones, consultorios, quirófano, "
        "hospitalización, laboratorio, recepción, áreas comunes; fotografías o videos de "
        "pacientes (mascotas), procedimientos médicos, quirúrgicos o de peluquería; "
        "casos clínicos, antes/después, testimonios de clientes; uniformes, logos o "
        "cualquier elemento identificable con LA EMPRESA; personal de la clínica en "
        "horario de trabajo o en las instalaciones — requiere AUTORIZACIÓN PREVIA, "
        "EXPRESA Y POR ESCRITO de la Dirección de LA EMPRESA.")
    add_para(doc,
        "LA EMPRESA dispondrá de un plazo de tres (3) días hábiles, contados a partir de "
        "la recepción de la solicitud de autorización, para pronunciarse por escrito. "
        "Transcurrido dicho plazo sin respuesta expresa, la solicitud se considerará "
        "DENEGADA. La publicación sin autorización previa constituye incumplimiento "
        "grave del contrato y facultará a LA EMPRESA a: (i) exigir la retirada "
        "inmediata del contenido; (ii) iniciar las acciones civiles y penales que "
        "correspondan por daños a la reputación y a la imagen corporativa; y (iii) "
        "resolver el contrato por causa grave conforme a la Cláusula Décima.",
        size=10, space_after=4)
    add_para(doc,
        "PROPIEDAD DE CONTENIDOS: Todo contenido creado, publicado o difundido en el "
        "marco de la prestación de servicios bajo este contrato, sea o no autorizado, "
        "que utilice imágenes, datos, casos, instalaciones, pacientes, uniformes, logos "
        "o cualquier elemento identificable con LA EMPRESA, será propiedad exclusiva de "
        "LA EMPRESA, conforme a la Ley sobre el Derecho de Autor (Art. 15 — obra por "
        "encargo). EL PROFESIONAL cede a LA EMPRESA todos los derechos patrimoniales "
        "sobre dichos contenidos, sin necesidad de contraprestación adicional, "
        "entendiéndose incluida dicha cesión en los honorarios pactados.",
        size=10, italic=True, color=GRAY_TEXT, space_after=6)

    # CLÁUSULA 10: DURACIÓN
    add_section(doc, "CLÁUSULA DÉCIMA: DURACIÓN")
    add_para(doc,
        "El presente contrato tendrá una duración de DOCE (12) MESES, contados a partir "
        "del ____ de ________________ de ______, hasta el ____ de ________________ de ______. "
        "Podrá prorrogarse por mutuo acuerdo mediante addendum suscrito por las partes con "
        "al menos 30 días de anticipación a su vencimiento.", space_after=6)

    # CLÁUSULA 11: TERMINACIÓN
    add_section(doc, "CLÁUSULA DÉCIMA PRIMERA: TERMINACIÓN")
    add_para(doc, "El contrato podrá terminar por:")
    add_bullet(doc, "Vencimiento del plazo pactado, sin necesidad de notificación.")
    add_bullet(doc, "Resolución por mutuo acuerdo, mediante acta suscrita por las partes.")
    add_bullet(doc, "Resolución unilateral por incumplimiento de cualquiera de las partes, previa notificación escrita con 15 días de anticipación.")
    add_bullet(doc, "Resolución inmediata por causa grave (mala praxis profesional, violación de confidencialidad, abandono de servicios, publicación no autorizada en redes sociales).")
    add_para(doc,
        "La terminación del contrato NO genera derecho a prestaciones sociales, "
        "indemnizaciones laborales, ni cualquier otro concepto de naturaleza salarial. "
        "Solo procederá el pago de los honorarios pendientes por servicios efectivamente "
        "prestados a la fecha de terminación.", space_after=6)

    # CLÁUSULA 12: RESPONSABILIDAD CIVIL
    add_section(doc, "CLÁUSULA DÉCIMA SEGUNDA: RESPONSABILIDAD CIVIL PROFESIONAL")
    add_para(doc,
        "EL PROFESIONAL responde civil y profesionalmente por los actos, omisiones y "
        "decisiones clínicas que adopte en el ejercicio de su profesión. LA EMPRESA no "
        "asume responsabilidad solidaria por la actividad profesional del PROFESIONAL, "
        "salvo que se demuestre culpa directa de la empresa (equipos defectuosos, "
        "instalaciones inseguras).")
    add_para(doc,
        "EL PROFESIONAL declara tener vigente póliza de Seguro de Responsabilidad Civil "
        "Profesional con cobertura no menor a USD 50.000,00, y entrega copia de la póliza "
        "a LA EMPRESA al momento de la suscripción del presente contrato.",
        space_after=6)

    # CLÁUSULA 13: CONFIDENCIALIDAD
    add_section(doc, "CLÁUSULA DÉCIMA TERCERA: CONFIDENCIALIDAD")
    add_para(doc,
        "EL PROFESIONAL se obliga a mantener en reserva toda la información de LA EMPRESA, "
        "sus clientes, pacientes, proveedores y estrategias comerciales, durante la "
        "vigencia del contrato y por CINCO (5) AÑOS después de su terminación. La "
        "violación de esta cláusula generará responsabilidad civil por daños y perjuicios.",
        space_after=6)

    # CLÁUSULA 14: LOPDP
    add_section(doc, "CLÁUSULA DÉCIMA CUARTA: PROTECCIÓN DE DATOS PERSONALES")
    add_para(doc,
        "Las partes se comprometen a cumplir la Ley Orgánica de Protección de Datos "
        "Personales (LOPDP) en el tratamiento de los datos personales de clientes, "
        "pacientes y personal. El profesional autoriza a LA EMPRESA el tratamiento de sus "
        "datos personales con fines administrativos y tributarios, conforme a la "
        "Autorización firmada por separado.", space_after=6)

    # CLÁUSULA 15: DOMICILIO Y JURISDICCIÓN
    add_section(doc, "CLÁUSULA DÉCIMA QUINTA: DOMICILIO Y JURISDICCIÓN")
    add_para(doc,
        "Para todos los efectos derivados del presente contrato, las partes eligen como "
        "domicilio procesal especial, con exclusión de cualquier otro, la ciudad de Los "
        "Teques, Estado Miranda, a cuyos tribunales declaran someterse. Las controversias "
        "derivadas del presente contrato se sustanciarán por la vía ordinaria civil, "
        "correspondiendo a los tribunales civiles competentes el conocimiento de las mismas. "
        "Se excluye expresamente la jurisdicción laboral por tratarse de un contrato de "
        "naturaleza civil y mercantil.", space_after=8)

    add_para(doc,
        "Se hacen dos (02) ejemplares de un mismo tenor y un mismo efecto, en la ciudad de "
        "Los Teques, a los ____ días del mes de ________________ de ________.",
        size=10, space_after=10)

    add_signature_block(doc, ["LA EMPRESA", "EL PROFESIONAL"])
    add_footer(section, "Contrato Honorarios Médico Veterinario v1.1")

    doc.save(out)
    return out


# ============================================================
# 2. CONTRATO HP DOG GROOMER — Reemplazar cláusula cuarta
# ============================================================
def fix_contrato_honorarios_groomer():
    out = "/home/z/my-project/output/Contrato_Honorarios_Dog_Groomer.docx"

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "CONTRATO DE HONORARIOS", "Peluquero Canino",
                 version="Versión 1.1  ·  Dirección")
    add_doc_title(doc, "CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Dog Groomer (Peluquero Canino) — Honorarios Profesionales")
    style_run(r, size=11, italic=True, color=GRAY_TEXT)

    add_para(doc,
        "Entre los suscritos: GRUPO CAVAL 1003, C.A., sociedad mercantil de domicilio en "
        "Av. Francisco de Miranda, Local N° 1, Sector Francisco de Miranda, Los Teques, "
        "Estado Miranda, Zona Postal 1201, RIF N° J501662533, marca comercial ALIKA PETS, "
        "dedicada a la actividad de clínica veterinaria, tienda de mascotas y peluquería "
        "canina, en adelante «LA EMPRESA», representada en este acto por su Directora "
        "Gerente, ciudadana ESNATLIM ELENA SIMOZA, titular de la cédula de identidad N° "
        "V-17.976.287, por una parte; y por la otra, el(la) ciudadano(a) "
        "___________________________________________, venezolano(a), mayor de edad, titular "
        "de la cédula de identidad N° V-___________________, con RIF personal N° "
        "__________________, con experiencia técnica certificada en peluquería canina y "
        "felina (anexar certificaciones o constancias), en adelante «EL PROFESIONAL», "
        "quienes declaran ser mayores de edad y hábiles para contratar, han convenido en "
        "celebrar el presente CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES, el cual "
        "se regirá por las siguientes cláusulas:", size=10, space_after=8)

    # CLÁUSULA 1: OBJETO
    add_section(doc, "CLÁUSULA PRIMERA: OBJETO")
    add_para(doc,
        "EL PROFESIONAL se obliga a prestar a LA EMPRESA servicios de PELUQUERÍA CANINA "
        "Y FELINA, en su carácter de técnico especializado independiente, comprendiendo: "
        "cortes de pelo según raza, tipo de pelaje o solicitud del cliente; baño, secado "
        "y cepillado profesional; corte y limado de uñas; limpieza del canal auditivo "
        "externo; evaluación inicial del estado del animal; manejo y contención amigable "
        "del animal; limpieza y desinfección del área de trabajo entre pacientes; y "
        "mantenimiento de las herramientas de peluquería.")
    add_para(doc,
        "Los servicios se prestarán CON PLENA AUTONOMÍA TÉCNICA. EL PROFESIONAL decide "
        "libremente las técnicas de corte, sujeción y manejo del animal, conforme a las "
        "buenas prácticas de peluquería canina y al bienestar animal. LA EMPRESA reconoce "
        "esta autonomía técnica.", space_after=6)

    # CLÁUSULA 2: NATURALEZA
    add_section(doc, "CLÁUSULA SEGUNDA: NATURALEZA DEL CONTRATO")
    add_para(doc,
        "Las partes declaran expresamente que el presente contrato es de PRESTACIÓN DE "
        "SERVICIOS BAJO LA MODALIDAD DE HONORARIOS PROFESIONALES, conforme al artículo 40 "
        "de la LOTTT y la jurisprudencia del Tribunal Supremo de Justicia. En consecuencia:")
    add_bullet(doc, "EL PROFESIONAL actúa como técnico independiente, sin relación de subordinación.",
               bold_lead="1. No subordinación:  ")
    add_bullet(doc, "EL PROFESIONAL organiza su trabajo de manera autónoma, sujeto a las citas acordadas.",
               bold_lead="2. Autonomía técnica:  ")
    add_bullet(doc, "EL PROFESIONAL puede atender a otros clientes y prestar servicios externos en horarios distintos a los pactados.",
               bold_lead="3. No exclusividad:  ")
    add_bullet(doc, "La remuneración se pagará en calidad de HONORARIOS, sin generar prestaciones sociales.",
               bold_lead="4. No salarial:  ")
    add_bullet(doc, "EL PROFESIONAL asume el riesgo de su actividad (mordeduras, cortes, lesiones a mascotas).",
               bold_lead="5. Asume riesgo:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # CLÁUSULA 3: HORARIOS
    add_section(doc, "CLÁUSULA TERCERA: HORARIOS Y TURNOS")
    add_para(doc,
        "Las partes acuerdan que EL PROFESIONAL prestará sus servicios con la siguiente "
        "flexibilidad horaria:")
    add_bullet(doc, "Días: de lunes a sábado (con un día de descanso acordado semanal).")
    add_bullet(doc, "Turno: _____ horas a _____ horas (con receso de 1 hora).")
    add_bullet(doc, "Citas asignadas por LA EMPRESA, previa coordinación con EL PROFESIONAL.")
    add_para(doc,
        "EL PROFESIONAL podrá aceptar o reprogramar citas, siempre que asegure la "
        "cobertura mínima pactada de ___ servicios/día. Esta flexibilidad es esencial "
        "para preservar la naturaleza no laboral del contrato.", space_after=6)

    # CLÁUSULA 4: HONORARIOS Y FORMA DE PAGO (CORREGIDA — sin discriminación salarial)
    add_section(doc, "CLÁUSULA CUARTA: HONORARIOS Y FORMA DE PAGO")
    add_para(doc,
        "Por los servicios prestados, LA EMPRESA pagará a EL PROFESIONAL la suma mensual "
        "única de USD 250,00 (DÓLARES DE LOS ESTADOS UNIDOS DE AMÉRICA DOSCIENTOS "
        "CINCUENTA CON 00/100), como HONORARIOS PROFESIONALES por la totalidad de los "
        "servicios objeto del presente contrato. Esta suma constituye la contraprestación "
        "íntegra y única por todos los servicios prestados, sin que proceda discriminación "
        "alguna en concepto de salario base, cestaticket, bono de alimentación, bono de "
        "transporte, bono vacacional, utilidades, prestaciones sociales o cualquier otro "
        "concepto de naturaleza salarial o laboral.",
        size=10, space_after=4)

    add_para(doc,
        "NATURALEZA DE LOS HONORARIOS: Las partes dejan expresamente constancia de que la "
        "suma mensual de USD 250,00 corresponde exclusivamente a HONORARIOS PROFESIONALES "
        "por la prestación de servicios técnicos de peluquería canina. En consecuencia, "
        "NO constituye salario a ningún efecto legal, NO genera prestaciones sociales "
        "ni antigüedad, NO está sujeto a la LOTTT, y NO se calculará sobre ella "
        "vacaciones, utilidades, bono vacacional ni indemnizaciones de naturaleza laboral. "
        "Su único régimen aplicable es el del Código Civil, Código de Comercio y la Ley "
        "de Impuesto sobre la Renta (ISLR).",
        size=10, italic=True, color=GRAY_TEXT, space_after=4)

    add_para(doc,
        "MONEDA DE PAGO Y CONVERSIÓN CAMBIARIA: El monto de los honorarios se acuerda "
        "referenciado en moneda extranjera (USD). Si el pago efectivo se realiza en "
        "Bolívares (Bs.), la conversión se efectuará tomando como única referencia el "
        "tipo de cambio oficial publicado por el Banco Central de Venezuela (BCV) "
        "vigente para el día en que se realice el pago efectivo y oportuno. En caso de "
        "pactarse el pago directo en USD en efectivo, LA EMPRESA entregará billetes "
        "aptos para la circulación.",
        size=10, space_after=4)

    add_para(doc,
        "FORMA DE PAGO: Mensual, dentro de los primeros cinco (5) días hábiles del mes "
        "siguiente al de prestación de los servicios, mediante transferencia bancaria a "
        "la cuenta del PROFESIONAL (Banco: _______________, Cuenta N°: _______________). "
        "El pago estará sujeto a la retención del 1% por concepto de ISLR (no profesional "
        "universitario, encomendado) conforme al artículo 27 de la Ley de ISLR, "
        "retención que LA EMPRESA enterará al SENIAT y entregará comprobante al PROFESIONAL.",
        size=10, space_after=4)

    add_para(doc,
        "FACTURACIÓN: EL PROFESIONAL se obliga a emitir recibo o factura por cada pago "
        "recibido, con sus datos de RIF, número de control y descripción del servicio. "
        "Sin factura no procederá el pago. La factura debe ser emitida a nombre de "
        "GRUPO CAVAL 1003, C.A., RIF J501662533.", size=10, italic=True, color=GRAY_TEXT, space_after=6)

    # CLÁUSULA 5: OBLIGACIONES DEL PROFESIONAL
    add_section(doc, "CLÁUSULA QUINTA: OBLIGACIONES DEL PROFESIONAL")
    add_bullet(doc, "Prestar los servicios con diligencia, calidad, oportunidad y respeto al bienestar animal.")
    add_bullet(doc, "Realizar cortes según raza, tipo de pelaje o solicitud del cliente.")
    add_bullet(doc, "Evaluar el estado del animal antes del servicio y reportar cualquier hallazgo al Médico Veterinario.")
    add_bullet(doc, "Suspender el servicio e informar al MV si el animal presenta signos de asfixia, fatiga extrema, síncope o agresividad inmanejable.")
    add_bullet(doc, "Aplicar técnicas de manejo y contención amigable, libres de crueldad.")
    add_bullet(doc, "Limpiar, desinfectar y ordenar la mesa de peluquería, bañera y jaulas entre cada paciente.")
    add_bullet(doc, "Mantener sus propias herramientas de corte (tijeras, cuchillas, máquinas) en buen estado.")
    add_bullet(doc, "Mantener confidencialidad de la información de LA EMPRESA y sus clientes (durante el contrato y por 3 años después).")
    add_bullet(doc, "Emitir facturas por cada pago y declarar ISLR anualmente.")
    add_bullet(doc, "No contactar clientes de LA EMPRESA para ofrecer servicios externos durante el contrato y por 12 meses después.")
    add_bullet(doc, "OBTENER AUTORIZACIÓN PREVIA Y POR ESCRITO de LA EMPRESA para publicar en redes sociales o medios cualquier contenido que se origine en las instalaciones de la peluquería, que involucre pacientes (mascotas), procedimientos, personal, uniformes o cualquier elemento identificable con LA EMPRESA. La publicación sin autorización previa constituye incumplimiento grave del presente contrato.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # CLÁUSULA 6: OBLIGACIONES DE LA EMPRESA
    add_section(doc, "CLÁUSULA SEXTA: OBLIGACIONES DE LA EMPRESA")
    add_bullet(doc, "Pagar los honorarios en la forma y oportunidad pactadas en la Cláusula Cuarta.")
    add_bullet(doc, "Proveer el espacio físico de peluquería, bañera, mesa, jaulas, secadora, productos (champús, acondicionadores) y servicios básicos.")
    add_bullet(doc, "Coordinar las citas y asignarlas previamente.")
    add_bullet(doc, "Mantener el área de peluquería limpia, segura y en condiciones adecuadas.")
    add_bullet(doc, "Retener y enterar el 1% de ISLR al SENIAT, entregar comprobante al PROFESIONAL.")
    add_bullet(doc, "Respetar la autonomía técnica del PROFESIONAL en los servicios de peluquería.")
    add_bullet(doc, "Autorizar o denegar por escrito, en un plazo máximo de 3 días hábiles, las solicitudes del PROFESIONAL para publicar contenido en redes sociales relacionado con la peluquería o pacientes.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # CLÁUSULA 7: HERRAMIENTAS Y EQUIPOS
    add_section(doc, "CLÁUSULA SÉPTIMA: HERRAMIENTAS Y EQUIPOS")
    add_para(doc,
        "LA EMPRESA pondrá a disposición del PROFESIONAL la bañera, mesa de peluquería, "
        "secadora, jaulas, champús, acondicionadores, toallas y demás insumos necesarios "
        "para la prestación de los servicios de peluquería canina. Estos bienes son "
        "propiedad exclusiva de LA EMPRESA y se entregan en comodato para el ejercicio de "
        "las funciones del PROFESIONAL.", size=10, space_after=4)
    add_para(doc,
        "EL PROFESIONAL se obliga a: (i) cuidar los equipos como un buen padre de familia; "
        "(ii) reportar de inmediato cualquier daño, deterioro, falla o pérdida a LA EMPRESA; "
        "(iii) no sustraer, trasladar a terceros, ni utilizar los equipos para fines "
        "personales o externos; (iv) mantener los equipos en condiciones de higiene y "
        "bioseguridad conforme a la NT-01-2008; y (v) devolver todos los equipos al término "
        "del contrato, en el mismo estado en que los recibió, salvo el deterioro natural "
        "por su uso adecuado.", size=10, space_after=4)
    add_para(doc,
        "HERRAMIENTAS PROPIAS DEL PROFESIONAL: EL PROFESIONAL utilizará sus propias "
        "herramientas de corte (tijeras, cuchillas, máquinas portátiles, peines, cepillos) "
        "que son de su propiedad exclusiva. Estas herramientas deberán ser identificadas e "
        "inventariadas al inicio del contrato y al término serán retiradas por EL "
        "PROFESIONAL. Mientras dichas herramientas se encuentren en las instalaciones de "
        "LA EMPRESA, EL PROFESIONAL se obliga a mantenerlas en buen estado, desinfectadas y "
        "bajo su responsabilidad. LA EMPRESA no se hace responsable por pérdida, robo o "
        "daño de las herramientas del PROFESIONAL, salvo dolo o negligencia grave de la "
        "empresa.", size=10, italic=True, color=GRAY_TEXT, space_after=6)

    # CLÁUSULA 8: PUBLICACIONES Y REDES SOCIALES
    add_section(doc, "CLÁUSULA OCTAVA: PUBLICACIONES Y REDES SOCIALES")
    add_para(doc,
        "Toda publicación, difusión o comunicación pública en redes sociales, medios "
        "tradicionales, plataformas digitales o cualquier canal de comunicación que "
        "incluya contenido originado en LA EMPRESA — incluyendo pero no limitado a: "
        "fotografías o videos de las instalaciones de peluquería, bañera, mesa, jaulas; "
        "fotografías o videos de pacientes (mascotas) antes/después del servicio; "
        "casos, testimonios de clientes; uniformes, logos o cualquier elemento "
        "identificable con LA EMPRESA; personal de la clínica en horario de trabajo o en "
        "las instalaciones — requiere AUTORIZACIÓN PREVIA, EXPRESA Y POR ESCRITO de la "
        "Dirección de LA EMPRESA.")
    add_para(doc,
        "LA EMPRESA dispondrá de un plazo de tres (3) días hábiles, contados a partir de "
        "la recepción de la solicitud de autorización, para pronunciarse por escrito. "
        "Transcurrido dicho plazo sin respuesta expresa, la solicitud se considerará "
        "DENEGADA. La publicación sin autorización previa constituye incumplimiento "
        "grave del contrato y facultará a LA EMPRESA a: (i) exigir la retirada "
        "inmediata del contenido; (ii) iniciar las acciones civiles y penales que "
        "correspondan por daños a la reputación y a la imagen corporativa; y (iii) "
        "resolver el contrato por causa grave.",
        size=10, space_after=4)
    add_para(doc,
        "PROPIEDAD DE CONTENIDOS: Todo contenido creado, publicado o difundido en el "
        "marco de la prestación de servicios bajo este contrato, sea o no autorizado, "
        "que utilice imágenes, datos, casos, instalaciones, pacientes, uniformes, logos "
        "o cualquier elemento identificable con LA EMPRESA, será propiedad exclusiva de "
        "LA EMPRESA, conforme a la Ley sobre el Derecho de Autor (Art. 15 — obra por "
        "encargo). EL PROFESIONAL cede a LA EMPRESA todos los derechos patrimoniales "
        "sobre dichos contenidos, sin necesidad de contraprestación adicional, "
        "entendiéndose incluida dicha cesión en los honorarios pactados.",
        size=10, italic=True, color=GRAY_TEXT, space_after=6)

    # CLÁUSULA 9: RESPONSABILIDAD CIVIL
    add_section(doc, "CLÁUSULA NOVENA: RESPONSABILIDAD CIVIL")
    add_para(doc,
        "EL PROFESIONAL responde civilmente por los daños causados a las mascotas durante "
        "el servicio, salvo que se demuestre que el daño fue causado por equipos "
        "defectuosos o instalaciones inseguras proporcionadas por LA EMPRESA. EL "
        "PROFESIONAL debe mantener póliza de Responsabilidad Civil (opcional pero "
        "recomendada) y reportar inmediatamente cualquier incidente.")
    add_para(doc,
        "En caso de mordedura o arañazo al PROFESIONAL durante el servicio, este será "
        "responsable de su atención médica, dado que se trata de un profesional "
        "independiente que asume el riesgo de su actividad. Se recomienda mantener "
        "vacuna antirrábica pre-exposición vigente.", space_after=6)

    # CLÁUSULA 10: DURACIÓN
    add_section(doc, "CLÁUSULA DÉCIMA: DURACIÓN")
    add_para(doc,
        "El presente contrato tendrá una duración de DOCE (12) MESES, contados a partir "
        "del ____ de ________________ de ______, hasta el ____ de ________________ de ______. "
        "Podrá prorrogarse por mutuo acuerdo mediante addendum suscrito por las partes con "
        "al menos 30 días de anticipación a su vencimiento.", space_after=6)

    # CLÁUSULA 11: TERMINACIÓN
    add_section(doc, "CLÁUSULA DÉCIMA PRIMERA: TERMINACIÓN")
    add_para(doc, "El contrato podrá terminar por:")
    add_bullet(doc, "Vencimiento del plazo pactado.")
    add_bullet(doc, "Resolución por mutuo acuerdo, mediante acta suscrita por las partes.")
    add_bullet(doc, "Resolución unilateral por incumplimiento, previa notificación escrita con 15 días de anticipación.")
    add_bullet(doc, "Resolución inmediata por causa grave (maltrato animal, robo, violación de confidencialidad, publicación no autorizada en redes sociales).")
    add_para(doc,
        "La terminación NO genera derecho a prestaciones sociales ni indemnizaciones "
        "laborales. Solo procederá el pago de honorarios pendientes por servicios "
        "efectivamente prestados.", space_after=6)

    # CLÁUSULA 12: CONFIDENCIALIDAD
    add_section(doc, "CLÁUSULA DÉCIMA SEGUNDA: CONFIDENCIALIDAD")
    add_para(doc,
        "EL PROFESIONAL se obliga a mantener en reserva toda la información de LA EMPRESA, "
        "sus clientes y pacientes, durante la vigencia del contrato y por TRES (3) AÑOS "
        "después de su terminación.", space_after=6)

    # CLÁUSULA 13: LOPDP
    add_section(doc, "CLÁUSULA DÉCIMA TERCERA: PROTECCIÓN DE DATOS PERSONALES")
    add_para(doc,
        "Las partes cumplirán la LOPDP. EL PROFESIONAL autoriza a LA EMPRESA el "
        "tratamiento de sus datos personales con fines administrativos y tributarios.",
        space_after=6)

    # CLÁUSULA 14: DOMICILIO Y JURISDICCIÓN
    add_section(doc, "CLÁUSULA DÉCIMA CUARTA: DOMICILIO Y JURISDICCIÓN")
    add_para(doc,
        "Para todos los efectos, las partes eligen como domicilio procesal especial la "
        "ciudad de Los Teques, Estado Miranda, a cuyos tribunales civiles declaran "
        "someterse. Se excluye expresamente la jurisdicción laboral.", space_after=8)

    add_para(doc,
        "Se hacen dos (02) ejemplares de un mismo tenor y un mismo efecto, en la ciudad de "
        "Los Teques, a los ____ días del mes de ________________ de ________.",
        size=10, space_after=10)

    add_signature_block(doc, ["LA EMPRESA", "EL PROFESIONAL"])
    add_footer(section, "Contrato Honorarios Dog Groomer v1.1")

    doc.save(out)
    return out


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print("Corrigiendo contratos HP (eliminando discriminación salarial)...")
    outs = [
        fix_contrato_honorarios_vet(),
        fix_contrato_honorarios_groomer(),
    ]
    for o in outs:
        size_kb = os.path.getsize(o) / 1024
        print(f"  ✓ {o}  ({size_kb:.1f} KB)")
    print("\nContratos HP corregidos correctamente.")
