"""
gen_actualizar_cargos_curados.py — Actualiza las descripciones de cargo con
las FUNCIONES CURADAS que envió la empresa (Esnatlim):
  - Dog Groomer
  - Encargada de Tienda
  - Auxiliar Veterinario

Reescribe solo las secciones 3 (Funciones Específicas) de cada cargo,
manteniendo el resto del documento (identificación, objetivo, competencias,
condiciones, carta de recepción).
"""
import os, sys
sys.path.insert(0, "/home/z/my-project/output")
from _common import *
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

OUT_DIR = "/home/z/my-project/output/03_DESCRIPCION_DE_CARGOS"

# ============================================================
# Datos de los 3 cargos a actualizar
# ============================================================
CARGOS = {
    "groomer": {
        "archivo": "03f_Funciones_Dog_Groomer.docx",
        "cargo": "DOG GROOMER (PELUQUERO CANINO)",
        "cargo_corto": "Dog Groomer",
        "area": "Peluquería Canina",
        "nivel": "Técnico especializado",
        "tipo_contrato": "Tiempo indeterminado (LOTTT art. 65) — período de prueba 30 días",
        "reporta_a": "Encargado(a) de Tienda / Gerente General",
        "supervisa_a": "Nadie directamente",
        "objetivo":
            "Realizar servicios de peluquería y estética canina y felina con técnicas "
            "profesionales, garantizando el bienestar de las mascotas, la satisfacción de "
            "los clientes y el cumplimiento de los protocolos de higiene, bioseguridad y "
            "manejo amigable de animales establecidos por la empresa.",
        "funciones": {
            "CORTES Y PELUQUERÍA": [
                "Cortes de raza y cortes higiénicos.",
                "Técnicas de tijera, máquina y rastrillo según tipo de pelo y raza.",
                "Asesorar al cliente sobre el corte más adecuado.",
            ],
            "BAÑO Y SECADO": [
                "Bañar con champús y acondicionadores según tipo de piel y pelaje.",
                "Aplicar tratamientos antipulgas y antiparasitarios tópicos.",
                "Secar con toalla y secador profesional, cepillar y desenredar.",
                "Limpiar oídos y cortar uñas.",
            ],
            "HIGIENE COMPLEMENTARIA": [
                "Perfumes, laqueado y terminaciones estéticas.",
            ],
            "EVALUACIÓN Y MANEJO ANIMAL": [
                "Evaluar estado general del animal al ingreso.",
                "Reportar al veterinario cualquier hallazgo anómalo.",
                "Aplicar contención ética y manejo con refuerzo positivo.",
                "Decidir si un animal no está en condiciones de ser atendido.",
            ],
            "LIMPIEZA Y MANTENIMIENTO": [
                "Desinfectar bañera, mesa, jaulas, tijeras y máquinas entre pacientes.",
                "Segregar residuos conforme COVENIN 2747-93.",
                "Mantener equipos en buen estado y reportar fallas.",
            ],
        },
        "competencias_tecnicas": [
            "Formación o certificación en peluquería canina / dog grooming.",
            "Experiencia mínima de 1 año como dog groomer (preferible).",
            "Conocimiento de técnicas de corte por raza (poodle, schnauzer, shih tzu, etc.).",
            "Manejo de herramientas: cuchillas, tijeras, máquinas eléctricas, secadores.",
            "Conocimiento de productos de higiene y cosmética veterinaria.",
            "Capacitación en manejo y contención de animales.",
        ],
        "competencias_blandas": [
            "Paciencia y empatía con animales.",
            "Capacidad de lectura del lenguaje corporal canino/felino.",
            "Atención al detalle estético.",
            "Comunicación cordial con propietarios.",
            "Cuidado y respeto por el bienestar animal.",
        ],
        "requisitos": [
            "Cédula de identidad venezolana vigente.",
            "Certificación o portafolio que acredite experiencia en dog grooming.",
            "Vacuna antirrábica pre-exposición vigente.",
            "Sin antecedentes penales.",
            "Disponibilidad para fines de semana.",
        ],
        "lugar_trabajo": "Área de peluquería canina de ALIKA PETS, Av. Francisco de Miranda, Los Teques.",
        "jornada": "Tiempo completo, 40 horas semanales, con turnos rotativos incluyendo fines de semana.",
        "epp": "Bata/gabacha impermeable, guantes de nitrilo, mascarilla quirúrgica o con filtro, gafas protectoras, botas cerradas antideslizantes e impermeables, tapones auditivos (uso prolongado de secadores).",
        "riesgos_principales": "Mordeduras/arañazos, zoonosis cutáneas (tiña, sarna), cortes con herramientas, riesgo eléctrico en zona húmeda, posturas forzadas, asma/alergias por pelaje, dermatitis por humedad.",
    },
    "encargada_tienda": {
        "archivo": "03c_Funciones_Encargado_Tienda.docx",
        "cargo": "ENCARGADO(A) DE TIENDA",
        "cargo_corto": "Encargado(a) de Tienda",
        "area": "Tienda de Mascotas",
        "nivel": "Coordinación / Supervisión",
        "tipo_contrato": "Tiempo determinado (6 meses + prórroga) — período de prueba 30 días",
        "reporta_a": "Gerente General",
        "supervisa_a": "Asistentes de Ventas, personal de limpieza de tienda",
        "objetivo":
            "Supervisar las operaciones diarias de la tienda de mascotas, garantizando la "
            "correcta atención a clientes, el manejo adecuado de caja, el control de "
            "inventarios, la coordinación con proveedores y la implementación de "
            "estrategias comerciales que aumenten las ventas y fomenten la fidelización.",
        "funciones": {
            "ATENCIÓN AL CLIENTE": [
                "Saludar y recibir al cliente.",
                "Asesorar en alimentos, accesorios, higiene y antiparasitarios.",
                "Informar sobre promociones, programas de fidelización, devoluciones y garantías.",
                "Atender y escalar quejas o reclamos.",
            ],
            "MANEJO DE CAJA": [
                "Abrir y cerrar caja diariamente (monto inicial, ingresos, egresos, arqueo).",
                "Cobrar con efectivo, punto de venta, transferencia o tarjeta.",
                "Emitir facturas y comprobantes.",
                "Depósitos bancarios diarios y conciliación.",
                "Registrar operaciones en el sistema POS.",
                "Reportar diferencias al Gerente.",
            ],
            "REGISTRO Y CONTROL DE INVENTARIO": [
                "Mantener kárdex actualizado (entradas, salidas, mermas).",
                "Conteos parciales semanales e inventario general mensual.",
                "Verificar fechas de vencimiento y rotación FIFO.",
                "Coordinar recepción de mercancía con proveedores.",
                "Reponer góndolas, exhibidores y neveras (cadena de frío).",
                "Registrar entradas y salidas de mercancía.",
                "Comunicar niveles bajos de inventario.",
                "Participar en inventarios físicos mensuales.",
            ],
            "CARGA Y DESCARGA": [
                "Recibir mercancía de proveedores y verificar cantidades.",
                "Trasladar cajas a bodega y góndolas.",
                "Estibar, organizar y rotular productos.",
            ],
            "MANTENIMIENTO DEL ÁREA": [
                "Velar por limpieza y orden de góndolas, neveras, mostrador y bodega.",
                "Revisar funcionamiento de equipos.",
                "Revisar fechas de vencimiento.",
                "Reportar fallas al Gerente.",
            ],
            "POLÍTICAS Y REGISTROS": [
                "Cumplir y hacer cumplir políticas internas.",
                "Aplicar procedimiento de reporte de incidentes.",
                "Velar por exhibición de precios y señalización.",
                "Mantener confidencialidad de datos.",
            ],
        },
        "competencias_tecnicas": [
            "Experiencia previa en atención al cliente y manejo de caja (mínimo 1 año).",
            "Manejo de sistemas de punto de venta (POS) y kárdex.",
            "Conocimientos de control de inventarios (FIFO, mermas, rotación).",
            "Conocimiento de productos para mascotas (alimentos, accesorios, higiene, antiparasitarios).",
            "Capacidad para coordinar con proveedores y conciliar depósitos bancarios.",
            "Conocimientos básicos de normativas LOPCYMAT, LOTTT.",
        ],
        "competencias_blandas": [
            "Orientación comercial y al cliente.",
            "Honestidad y transparencia en manejo de efectivo.",
            "Capacidad de organización y orden.",
            "Comunicación efectiva y resolución de quejas.",
            "Trabajo en equipo y liderazgo básico.",
        ],
        "requisitos": [
            "Cédula de identidad venezolana vigente.",
            "Sin antecedentes penales.",
            "Disponibilidad para trabajar fines de semana y feriados.",
            "Experiencia mínima de 1 año en retail (preferible).",
        ],
        "lugar_trabajo": "Tienda de mascotas ALIKA PETS, Av. Francisco de Miranda, Los Teques.",
        "jornada": "Tiempo completo, 40 horas semanales, con turnos rotativos incluyendo fines de semana.",
        "epp": "Uniforme corporativo, calzado cerrado antideslizante; guantes de nitrilo para carga de mercancía y limpieza.",
        "riesgos_principales": "Trastornos musculoesqueléticos por bipedestación y carga de mercancía, caídas por cajas o suelos mojados, exposición a polvos y alérgenos, riesgo de asalto, exposición a productos de limpieza.",
    },
    "auxiliar_vet": {
        "archivo": "03e_Funciones_Auxiliar_Veterinario.docx",
        "cargo": "AUXILIAR VETERINARIO(A)",
        "cargo_corto": "Auxiliar Veterinario(a)",
        "area": "Clínica Veterinaria",
        "nivel": "Técnico / Asistencial",
        "tipo_contrato": "Tiempo determinado (3 meses + prórroga) — período de prueba 15 días",
        "reporta_a": "Médico(a) Veterinario(a) / Encargado(a) de Clínica",
        "supervisa_a": "Nadie directamente",
        "objetivo":
            "Asistir al médico veterinario en los diversos procedimientos y servicios que "
            "presta la clínica, brindar cuidado directo a los animales hospitalizados, "
            "interactuar con los dueños y mantener limpias y desinfectadas las instalaciones, "
            "garantizando un ambiente higiénico y seguro.",
        "funciones": {
            "CONSULTAS": [
                "Recibir, pesar y tomar signos vitales del paciente.",
                "Preparar y desinfectar consultorio entre pacientes.",
                "Asistir en contención y manejo del animal.",
                "Registrar signos vitales en historia clínica.",
            ],
            "TRATAMIENTOS": [
                "Administrar medicamentos según indicación médica.",
                "Aplicar antiparasitarios y fluidoterapia.",
                "Retirar catéteres, puntos.",
            ],
            "LABORATORIO": [
                "Tomar muestras de orina, heces y secreciones.",
                "Operar centrífuga, microscopio y analizadores.",
            ],
            "CIRUGÍA": [
                "Preparar paciente (ayuno, rasurado, vía venosa).",
                "Preparar quirófano (material estéril, instrumental, monitores).",
                "Asistir como instrumentista y circulante.",
                "Monitorear signos vitales transoperatorios.",
                "Cuidar recuperación postoperatoria.",
                "Limpieza del área e instrumentos luego del procedimiento quirúrgico.",
            ],
            "HOSPITALIZACIÓN": [
                "Alimentación, hidratación, paseo, higiene y medicación de hospitalizados.",
                "Rondas de control por turno.",
                "Limpiar y desinfectar jaulas y comederos.",
            ],
            "ATENCIÓN AL CLIENTE": [
                "Recibir al cliente y a la mascota.",
                "Brindar información y/o asesoría al cliente en caso de ser necesario sobre los servicios.",
                "Hacerle entrega al cliente de los formatos correspondientes para ser llenados tales como Consentimiento informado, acta voluntaria, etc.",
                "Explicar indicaciones postoperatorias y de medicación.",
                "Cobrar servicios clínicos si el encargado no está disponible.",
            ],
            "HIGIENE Y BIOSEGURIDAD": [
                "Cumplir protocolos NT-01-2008 y COVENIN 2747-93.",
                "Segregar residuos por color de bolsa.",
                "Esterilizar instrumental (autoclave).",
            ],
        },
        "competencias_tecnicas": [
            "Formación como Auxiliar Veterinario, Técnico Superior Universitario (TSU) en Veterinaria o afín (preferible).",
            "Experiencia previa en clínica veterinaria (mínimo 6 meses, preferible).",
            "Conocimiento de técnicas de sujeción y manejo de animales (cánidos, felinos).",
            "Conocimiento de instrumental quirúrgico, procedimientos de esterilización y operación de autoclave.",
            "Capacidad para tomar signos vitales y registrarlos en historia clínica.",
            "Manejo de laboratorio básico (centrífuga, microscopio, analizadores).",
        ],
        "competencias_blandas": [
            "Empatía y paciencia con animales.",
            "Capacidad de trabajo en equipo con el MV.",
            "Atención al detalle en el seguimiento de indicaciones médicas.",
            "Disposición para tareas de limpieza y mantenimiento.",
            "Comunicación cordial con propietarios de mascotas.",
        ],
        "requisitos": [
            "Cédula de identidad venezolana vigente.",
            "Formación técnica o certificación en auxiliar veterinario (preferible).",
            "Vacuna antirrábica pre-exposición vigente.",
            "Sin antecedentes penales.",
            "Disponibilidad para turnos rotativos.",
        ],
        "lugar_trabajo": "Clínica veterinaria ALIKA PETS (áreas de consulta, hospitalización, quirófano, laboratorio, recepción), Los Teques.",
        "jornada": "Tiempo completo, 40 horas semanales, con turnos rotativos.",
        "epp": "Bata impermeable, guantes de nitrilo, mascarilla quirúrgica, gafas protectoras (en limpieza), botas cerradas antideslizantes.",
        "riesgos_principales": "Mordeduras/arañazos por sujeción de animales, exposición a fluidos biológicos en limpieza de jaulas, exposición a zoonosis, pinchazos con cortopunzantes, exposición a desinfectantes químicos, trastornos musculoesqueléticos.",
    },
}

# ============================================================
# Función: construir documento completo de descripción de cargo
# ============================================================
def render_identificacion(doc, cfg):
    add_section(doc, "1. IDENTIFICACIÓN DEL CARGO")
    tbl = doc.add_table(rows=4, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(3.5), Cm(5.0), Cm(3.5), Cm(5.0)]):
        tbl.columns[i].width = w
    datos = [
        ("Título del cargo:", cfg["cargo"], "Nivel jerárquico:", cfg["nivel"]),
        ("Departamento / Área:", cfg["area"], "Tipo de contrato:", cfg["tipo_contrato"]),
        ("Reporta a:", cfg["reporta_a"], "Supervisa a:", cfg["supervisa_a"]),
        ("Cargo:", cfg["cargo_corto"], "", ""),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(datos):
        c = tbl.rows[ri].cells[0]; c.width = Cm(3.5)
        write_cell(c, l1, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[1]; c.width = Cm(5.0)
        write_cell(c, v1, size=9, bg=WHITE)
        if l2:
            c = tbl.rows[ri].cells[2]; c.width = Cm(3.5)
            write_cell(c, l2, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
            c = tbl.rows[ri].cells[3]; c.width = Cm(5.0)
            write_cell(c, v2, size=9, bg=WHITE)
        else:
            # Fusionar celdas 2 y 3
            merged = tbl.rows[ri].cells[2].merge(tbl.rows[ri].cells[3])
            merged.text = ""
            set_cell_borders(merged, color="FFFFFF", sz="0")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

def render_objetivo(doc, cfg):
    add_section(doc, "2. OBJETIVO GENERAL DEL CARGO")
    add_para(doc, cfg["objetivo"], size=10, space_after=8)

def render_funciones(doc, cfg):
    add_section(doc, "3. FUNCIONES ESPECÍFICAS")
    add_para(doc,
        "Sin perjuicio de las demás funciones que le asigne la empresa, el(la) "
        f"{cfg['cargo_corto']} tendrá las siguientes funciones específicas, organizadas por "
        "área de actividad:",
        size=10, italic=True, color=GRAY_TEXT, space_after=6)

    for area, funciones in cfg["funciones"].items():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"▸ {area}")
        style_run(r, size=10, bold=True, color=TEAL_DARK)
        for fn in funciones:
            add_bullet(doc, fn, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

def render_competencias(doc, cfg):
    add_section(doc, "4. COMPETENCIAS Y REQUISITOS DEL CARGO")
    add_para(doc, "Competencias técnicas y profesionales:", size=10, bold=True, color=TEAL_DARK, space_after=4)
    for comp in cfg["competencias_tecnicas"]:
        add_bullet(doc, comp, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
    add_para(doc, "Competencias blandas / actitudinales:", size=10, bold=True, color=TEAL_DARK, space_after=4)
    for comp in cfg["competencias_blandas"]:
        add_bullet(doc, comp, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
    add_para(doc, "Requisitos académicos y legales:", size=10, bold=True, color=TEAL_DARK, space_after=4)
    for req in cfg["requisitos"]:
        add_bullet(doc, req, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

def render_condiciones(doc, cfg):
    add_section(doc, "5. CONDICIONES DE TRABAJO Y RIESGOS")
    add_para(doc,
        "El cargo se desempeña en condiciones que implican exposición a riesgos "
        "ocupacionales específicos, los cuales están detallados en la Notificación de "
        "Riesgos correspondiente, conforme al artículo 56 de la LOPCYMAT. Las condiciones "
        "generales son:",
        size=10, space_after=4)
    add_bullet(doc, cfg["lugar_trabajo"], bold_lead="Lugar de trabajo:  ")
    add_bullet(doc, cfg["jornada"], bold_lead="Jornada:  ")
    add_bullet(doc, cfg["epp"], bold_lead="EPP obligatorio:  ")
    add_bullet(doc, cfg["riesgos_principales"], bold_lead="Riesgos principales:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

def render_carta_recepcion(doc, cfg):
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(10)
    sep.paragraph_format.space_after = Pt(6)
    add_hr(sep, color="0F766E", sz="6")

    add_section(doc, "CARTA DE RECEPCIÓN DE LA DESCRIPCIÓN DE CARGO")

    add_para(doc,
        f"Yo, _____________________________________________, titular de la cédula de identidad "
        f"N° V-___________________, en mi condición de trabajador(a) de GRUPO CAVAL 1003, "
        f"C.A. (ALIKA PETS), designado(a) para el cargo de {cfg['cargo']}, declaro por medio "
        f"de la presente que:",
        size=10, space_after=6)

    add_bullet(doc,
        "He recibido copia íntegra y legible de la Descripción de Cargo correspondiente a "
        "mi posición, así como las explicaciones verbales necesarias para su correcta "
        "comprensión.", bold_lead="PRIMERO:  ")
    add_bullet(doc,
        "He leído en su totalidad el contenido del mencionado documento, comprendo las "
        "funciones, deberes y responsabilidades asignadas a mi cargo, y me comprometo a "
        "cumplirlas fielmente durante toda la vigencia de mi relación laboral con la "
        "empresa.", bold_lead="SEGUNDO:  ")
    add_bullet(doc,
        "Reconozco que las funciones aquí descritas me fueron suficientemente explicadas y "
        "notificadas por la empresa, conforme al artículo 26 de la Ley Orgánica del "
        "Trabajo, los Trabajadores y las Trabajadoras (LOTTT).", bold_lead="TERCERO:  ")
    add_bullet(doc,
        "Acepto que la presente firma constituye prueba fehaciente de la entrega y "
        "recepción del documento, renunciando a alegar desconocimiento de las funciones "
        "asignadas.", bold_lead="CUARTO:  ")

    add_para(doc, "", size=6, space_after=2)
    add_para(doc,
        "En fe de lo cual firmo la presente carta en la ciudad de Los Teques, Estado "
        "Miranda, a los ____ días del mes de ________________ de ________.",
        size=10, space_after=10)

    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])

# ============================================================
# Build documento
# ============================================================
def build_descripcion_cargo(cfg):
    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=1.8)
    add_membrete(doc, "DESCRIPCIÓN DE CARGO", cfg["cargo_corto"], version="Versión 3.2  ·  RR.HH.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("DESCRIPCIÓN DE CARGO Y FUNCIONES")
    style_run(r, size=14, bold=True, color=TEAL_DARK)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(cfg["cargo"])
    style_run(r2, size=12, bold=True, color=TEAL_DARK)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(10)
    add_hr(sep, color="0F766E", sz="8")

    render_identificacion(doc, cfg)
    render_objetivo(doc, cfg)
    render_funciones(doc, cfg)
    render_competencias(doc, cfg)
    render_condiciones(doc, cfg)
    render_carta_recepcion(doc, cfg)

    add_footer(section, f"Descripción Cargo — {cfg['cargo_corto']} v3.2")

    out = os.path.join(OUT_DIR, cfg["archivo"])
    doc.save(out)
    return out

# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print(f"Actualizando 3 descripciones de cargo con funciones curadas...")
    for key, cfg in CARGOS.items():
        out = build_descripcion_cargo(cfg)
        size_kb = os.path.getsize(out) / 1024
        print(f"  ✓ {os.path.basename(out)}  ({size_kb:.1f} KB)")
    print("\n3 descripciones actualizadas correctamente.")
