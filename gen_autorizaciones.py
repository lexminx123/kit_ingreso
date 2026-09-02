"""
gen_autorizaciones.py — 3 Autorizaciones RR.HH. (Bloque 07)
GRUPO CAVAL 1003, C.A. — ALIKA PETS — Los Teques, Miranda.

Genera en /home/z/my-project/output/07_AUTORIZACIONES/:
  - Autorizacion_Datos_Personales_LOPDP.docx
      LOPDP (Decreto 1.419, Gaceta 6.210 Extraordinaria, 2014)
      + Constitución art. 60 + LOTTT art. 183
  - Autorizacion_Imagen_Redes_Sociales.docx
      Constitución art. 60 + LOPDP art. 5
  - Autorizacion_Vigilancia_Camaras.docx
      Constitución art. 60 + LOPCYMAT arts. 46/74 + Ley Delitos Informáticos

Cada autorización tiene 8-10 capítulos con contenido legal completo.
Al final: CARTA DE RECEPCIÓN + add_signature_block
["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"].

Versión 3.0 — usa _common.py. Firma operativa: Esnatlim Elena Simoza,
C.I. V-17.976.287, Directora Gerente.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _common import (
    setup_a4_portrait, add_membrete, add_doc_title, add_section, add_chapter,
    add_para, add_bullet, add_article, add_signature_block, add_reception_letter,
    add_footer,
    set_cell_bg, set_cell_borders, set_cell_margins, write_cell, style_run, add_hr,
    TEAL_DARK, TEAL_HDR_BG, GRAY_ALT, SLATE_BG, WHITE, BLACK, GRAY_TEXT, GRAY_MUTED,
)

OUT_DIR = "/home/z/my-project/output/07_AUTORIZACIONES"
os.makedirs(OUT_DIR, exist_ok=True)


# ============================================================
# Helper: tabla IDENTIFICACIÓN DEL TRABAJADOR
# ============================================================
def add_worker_identification(doc, intro_text=None):
    """Tabla 4 filas × 2 columnas con labels para Nombre/C.I./Cargo/Fecha."""
    if intro_text:
        add_para(doc, intro_text, size=10, space_after=4)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(5.5)
    tbl.columns[1].width = Cm(11.1)
    rows = [
        ("Nombre completo del trabajador(a)", "_" * 50),
        ("Cédula de identidad",                 "V- ____________________"),
        ("Cargo / Departamento",                "_" * 40),
        ("Fecha de la autorización",            "____ / ____ / ________"),
    ]
    for i, (label, value) in enumerate(rows):
        c_lab = tbl.rows[i].cells[0]
        c_val = tbl.rows[i].cells[1]
        c_lab.width = Cm(5.5)
        c_val.width = Cm(11.1)
        write_cell(c_lab, label, size=9, bold=True, color=TEAL_DARK,
                   align=WD_ALIGN_PARAGRAPH.LEFT, bg=SLATE_BG)
        write_cell(c_val, value, size=9, color=GRAY_MUTED,
                   align=WD_ALIGN_PARAGRAPH.LEFT,
                   bg=WHITE if i % 2 == 0 else GRAY_ALT)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


# ============================================================
# 1. AUTORIZACIÓN PARA EL TRATAMIENTO DE DATOS PERSONALES (LOPDP)
# ============================================================
def gen_datos_personales_lopdp(doc):
    """8 capítulos conforme LOPDP Decreto 1.419 / Gaceta 6.210 / 2014."""
    setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc,
                 doc_label="AUTORIZACIÓN LOPDP",
                 doc_sublabel="Tratamiento de Datos Personales",
                 version="Versión 3.0  ·  RR.HH.")
    add_doc_title(doc, "AUTORIZACIÓN PARA EL TRATAMIENTO DE DATOS PERSONALES")

    add_para(doc,
        "Conforme a la Constitución de la República Bolivariana de Venezuela "
        "(artículo 60), la Ley Orgánica de Protección de Datos Personales — "
        "LOPDP (Decreto N° 1.419, Gaceta Oficial N° 6.210 Extraordinaria, "
        "19 de noviembre de 2014) y el artículo 183 de la Ley Orgánica del "
        "Trabajo, los Trabajadores y las Trabajadoras (LOTTT), el(la) "
        "trabajador(a) que suscribe otorga su consentimiento libre, previo, "
        "expreso e informado para que GRUPO CAVAL 1003, C.A. (ALIKA PETS) "
        "realice el tratamiento de sus datos personales en los términos "
        "siguientes:",
        size=10, space_after=6)

    add_worker_identification(doc)

    # CAP I — Consentimiento expreso
    add_chapter(doc, "I", "DEL CONSENTIMIENTO EXPRESO")
    add_article(doc, "1", "Objeto",
        "El(la) trabajador(a) autoriza de manera libre, previa, expresa e "
        "informada a GRUPO CAVAL 1003, C.A. (en adelante, LA EMPRESA), para "
        "que recolecte, use, almacene, transfiera y, en general, realice "
        "cualquier acto de tratamiento de sus datos personales, conforme a "
        "los artículos 4, 5, 7 y 9 de la LOPDP.")
    add_article(doc, "2", "Carácter del consentimiento",
        "El presente consentimiento se otorga de forma voluntaria, sin "
        "coacción de ninguna naturaleza, en el marco de la relación laboral "
        "vigente, y podrá ser revocado en cualquier momento conforme al "
        "artículo 22 de la LOPDP, sin afectar la licitud del tratamiento "
        "realizado con anterioridad a la revocación.")

    # CAP II — Finalidades
    add_chapter(doc, "II", "DE LAS FINALIDADES DEL TRATAMIENTO")
    add_para(doc, "LA EMPRESA tratará los datos personales del(la) "
                  "trabajador(a) para las siguientes finalidades:", size=10)
    finalidades = [
        ("Gestión laboral", "registro de ingreso, expediente, evaluaciones, "
                            "promociones, traslados y cese."),
        ("Nómina y beneficios", "cálculo y pago de salario, bono de "
                                "alimentación, vacaciones, utilidades, "
                                "prestaciones sociales (LOTTT art. 142)."),
        ("Cumplimiento legal", "inscripción y aportes al IVSS, FAOV-BVV, "
                               "INCES, Registro de Trabajadores, Ministerios "
                               "del Trabajo y de Salud, SENAC, SENCAMER."),
        ("Salud y seguridad laboral", "vigilancia epidemiológica NT-02-2008, "
                                      "exámenes médicos pre-empleo, periódicos "
                                      "y de egreso; registro de incidentes y "
                                      "enfermedades ocupacionales (LOPCYMAT)."),
        ("Comunicaciones internas", "difusión de comunicados, capacitaciones, "
                                    "evaluaciones, encuestas de clima "
                                    "organizacional y notificaciones oficiales."),
        ("Vigilancia y control", "control de acceso físico al centro de "
                                  "trabajo, control de asistencia y uso de "
                                  "cámaras conforme a la autorización "
                                  "específica de vigilancia."),
        ("Evaluación de desempeño", "registro de indicadores, metas, "
                                     "capacitaciones y resultados de gestión."),
        ("Prestaciones y reclamaciones", "trámite de prestaciones sociales, "
                                          "seguros, reclamaciones ante IVSS, "
                                          "FAOV, banca y proveedores de "
                                          "beneficios."),
    ]
    for lead, rest in finalidades:
        add_bullet(doc, rest, bold_lead=lead + ".  ")

    # CAP III — Datos sensibles
    add_chapter(doc, "III", "DE LOS DATOS SENSIBLES")
    add_para(doc,
        "En virtud del artículo 7 de la LOPDP, LA EMPRESA podrá tratar datos "
        "sensibles del(la) trabajador(a) única y exclusivamente para los "
        "fines de salud ocupacional, seguridad y cumplimiento de obligaciones "
        "legales. Se consideran datos sensibles:", size=10)
    add_bullet(doc, "Datos biométricos: huellas dactilares, fotografía para "
                    "carné identificatorio y control de acceso.",
               bold_lead="Biométricos.  ")
    add_bullet(doc, "Datos de salud: antecedentes médicos, resultados de "
                    "exámenes pre-empleo, periódicos y de egreso (NT-02-2008), "
                    "registro de accidentes y enfermedades ocupacionales.",
               bold_lead="Salud.  ")
    add_bullet(doc, "Datos de contacto de emergencia: nombres, teléfonos y "
                    "parentesco de personas a notificar en caso de accidente "
                    "o emergencia.",
               bold_lead="Emergencia.  ")

    # CAP IV — Derechos ARCO+
    add_chapter(doc, "IV", "DE LOS DERECHOS DEL titular (ARCO+)")
    add_para(doc,
        "Conforme a los artículos 14 al 22 de la LOPDP, el(la) trabajador(a) "
        "como titular de los datos personales tiene los siguientes derechos:",
        size=10)
    derechos = [
        ("Acceso", "solicitar y obtener información sobre sus datos "
                   "tratados por LA EMPRESA."),
        ("Rectificación", "solicitar la corrección de datos inexactos, "
                          "incompletos o desactualizados."),
        ("Cancelación", "solicitar la supresión de datos cuando no sean "
                        "necesarios para las finalidades autorizadas o "
                        "haya sido revocado el consentimiento, salvo "
                        "obligación legal de conservación."),
        ("Oposición", "oponerse al tratamiento de sus datos por motivos "
                      "legítimos, sin perjuicio de las obligaciones "
                      "legales pendientes."),
        ("Revocación", "revocar el consentimiento otorgado mediante "
                       "comunicación escrita dirigida a RR.HH."),
        ("Información", "ser informado sobre las finalidades, "
                        "destinatarios y duración del tratamiento de sus "
                        "datos personales."),
    ]
    for lead, rest in derechos:
        add_bullet(doc, rest, bold_lead=lead + ".  ")
    add_para(doc,
        "El ejercicio de estos derechos se realiza mediante solicitud "
        "escrita presentada ante la Dirección de Recursos Humanos de "
        "LA EMPRESA, la cual deberá responder en un plazo máximo de "
        "quince (15) días hábiles conforme al artículo 23 de la LOPDP.",
        size=10, italic=True, color=GRAY_TEXT)

    # CAP V — Cesión a terceros
    add_chapter(doc, "V", "DE LA CESIÓN A TERCEROS")
    add_para(doc,
        "LA EMPRESA podrá ceder los datos personales del(la) trabajador(a) "
        "exclusivamente a los siguientes destinatarios, cuando sea necesario "
        "para las finalidades autorizadas y para el cumplimiento de "
        "obligaciones legales:", size=10)
    cesionarios = [
        ("Instituto Venezolano de los Seguros Sociales (IVSS)",
         "para inscripción, aportes y prestaciones."),
        ("Banco de Vivienda y Hábitat — FAOV-BVV",
         "para aportes y consultas."),
        ("Servicio Nacional de Aprendizaje — INCES",
         "para aportes y formación."),
        ("Ministerio del Poder Popular del Proceso Social del Trabajo",
         "para notificaciones, inspecciones y registros."),
        ("Banca y entidades financieras",
         "para depósito de nómina, prestaciones y beneficios."),
        ("Proveedoras de salud ocupacional",
         "clínicas, laboratorios y PMSSO para exámenes y vigilancia."),
    ]
    for lead, rest in cesionarios:
        add_bullet(doc, rest, bold_lead=lead + ":  ")
    add_para(doc,
        "LA EMPRESA no cederá datos personales a terceros con fines "
        "comerciales o publicitarios sin consentimiento expreso y previo "
        "del(la) trabajador(a).", size=10)

    # CAP VI — Medidas de seguridad
    add_chapter(doc, "VI", "DE LAS MEDIDAS DE SEGURIDAD")
    add_para(doc,
        "LA EMPRESA adopta las medidas técnicas, organizativas y legales "
        "necesarias para garantizar la seguridad, confidencialidad e "
        "integridad de los datos personales, conforme al artículo 24 de la "
        "LOPDP. Estas medidas incluyen:", size=10)
    medidas = [
        "Acceso restringido a datos personales: solo personal autorizado "
        "de RR.HH., Gerencia y Contabilidad, con consentimiento registrado.",
        "Almacenamiento en medios físicos (archivos cerrados) y digitales "
        "(servidores con autenticación, copias de seguridad cifradas).",
        "Bitácora de acceso a datos sensibles, especialmente salud y "
        "biométricos.",
        "Capacitación periódica del personal en LOPDP y manejo seguro de "
        "información.",
        "Procedimiento de bloqueo y cancelación de datos al cese de la "
        "relación laboral.",
        "Protocolo de respuesta a brechas de seguridad: notificación al "
        "titular en un plazo no mayor a setenta y dos (72) horas.",
    ]
    for m in medidas:
        add_bullet(doc, m)

    # CAP VII — Duración
    add_chapter(doc, "VII", "DE LA DURACIÓN DEL TRATAMIENTO")
    add_article(doc, "7", "Vigencia del consentimiento",
        "El presente consentimiento se mantiene durante la vigencia de la "
        "relación laboral y, posteriormente, durante el plazo de "
        "conservación del expediente laboral establecido en el artículo 183 "
        "de la LOTTT (diez — 10 — años contados a partir de la terminación "
        "de la relación laboral).")
    add_article(doc, "8", "Finalidad posterior al cese",
        "Vencido el plazo de conservación legal, los datos serán "
        "cancelados, bloqueados o anonimizados, salvo que exista "
        "obligación legal de conservación adicional o el(la) titular haya "
        "otorgado consentimiento expreso para tratamiento posterior.")

    # CAP VIII — Revocabilidad
    add_chapter(doc, "VIII", "DE LA REVOCABILIDAD DEL CONSENTIMIENTO")
    add_article(doc, "9", "Revocación por escrito",
        "El(la) trabajador(a) podrá revocar el presente consentimiento "
        "mediante comunicación escrita dirigida a la Dirección de Recursos "
        "Humanos de LA EMPRESA, conforme al artículo 22 de la LOPDP.")
    add_article(doc, "10", "Efectos de la revocación",
        "La revocación surtirá efecto en un plazo máximo de quince (15) días "
        "hábiles, sin perjuicio de la obligación legal que LA EMPRESA tenga "
        "de conservar o tratar determinados datos en cumplimiento de "
        "disposiciones del IVSS, FAOV, INCES, Ministerio del Trabajo u otras "
        "autoridades competentes.")

    # Carta de recepción
    add_reception_letter(
        doc,
        "Autorización para el Tratamiento de Datos Personales (LOPDP)",
    )

    # Footer
    add_footer(doc.sections[0], "Autorización LOPDP v3.0")


# ============================================================
# 2. AUTORIZACIÓN DE USO DE IMAGEN Y VOZ
# ============================================================
def gen_imagen_redes_sociales(doc):
    """9 capítulos conforme Constitución art. 60 + LOPDP art. 5."""
    setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc,
                 doc_label="AUTORIZACIÓN DE IMAGEN",
                 doc_sublabel="Uso de Imagen y Voz",
                 version="Versión 3.0  ·  RR.HH.")
    add_doc_title(doc, "AUTORIZACIÓN DE USO DE IMAGEN Y VOZ")

    add_para(doc,
        "Conforme al artículo 60 de la Constitución de la República Bolivariana "
        "de Venezuela y al artículo 5 de la Ley Orgánica de Protección de Datos "
        "Personales (LOPDP, Decreto 1.419, Gaceta 6.210 Extraordinaria, 2014), "
        "el(la) trabajador(a) que suscribe otorga su consentimiento libre, "
        "previo, expreso e informado a GRUPO CAVAL 1003, C.A. (ALIKA PETS) "
        "para el uso de su imagen, voz y representación audiovisual, en los "
        "términos siguientes:",
        size=10, space_after=6)

    add_worker_identification(doc)

    # CAP I — Consentimiento expreso
    add_chapter(doc, "I", "DEL CONSENTIMIENTO EXPRESO")
    add_article(doc, "1", "Objeto",
        "El(la) trabajador(a) autoriza a LA EMPRESA para que capte, "
        "reproduzca, edite, almacene, difunda y, en general, utilice su "
        "imagen, voz y representación audiovisual, con fines institucionales, "
        "comerciales y de comunicación, conforme al artículo 60 de la "
        "Constitución y el artículo 5 de la LOPDP.")
    add_article(doc, "2", "Carácter voluntario",
        "Esta autorización se otorga de forma voluntaria, sin que su "
        "concesión condicione el ingreso, permanencia o promoción del(la) "
        "trabajador(a), ni genere obligación de contraprestación económica "
        "distinta a su salario ordinario.")

    # CAP II — Tipos de captación
    add_chapter(doc, "II", "DE LOS TIPOS DE CAPTACIÓN")
    add_para(doc, "La presente autorización comprende los siguientes tipos "
                  "de captación audiovisual y sonora:", size=10)
    tipos = [
        ("Fotografías", "individuales, grupales o en actividad."),
        ("Videos", "institucionales, de capacitación, de servicios, "
                   "testimoniales."),
        ("Streaming / Live", "transmisiones en vivo por redes sociales o "
                              "sitio web."),
        ("Eventos", "inauguraciones, campañas, jornadas de vacunación, "
                    "ferias, charlas."),
        ("Capacitaciones", "registro de cursos, talleres, demostraciones "
                           "técnicas."),
        ("Con mascotas", "fotos y videos del(la) trabajador(a) prestando "
                          "atención a mascotas de clientes."),
        ("Grabaciones de voz", "para cuñas, podcasts, mensajes "
                                "institucionales."),
    ]
    for lead, rest in tipos:
        add_bullet(doc, rest, bold_lead=lead + ".  ")

    # CAP III — Finalidades
    add_chapter(doc, "III", "DE LAS FINALIDADES")
    add_para(doc, "El uso autorizado de la imagen y voz comprenderá las "
                  "siguientes finalidades:", size=10)
    finalidades = [
        ("Redes sociales corporativas", "Instagram, Facebook, TikTok, "
                                          "YouTube, WhatsApp Business y "
                                          "canales oficiales de ALIKA PETS."),
        ("Sitio web institucional", "www.alikapets.com y dominios "
                                     "asociados."),
        ("Marketing y publicidad", "volantes, banners, vallas, "
                                    "promociones, anuncios digitales."),
        ("Capacitación interna", "material formativo para personal "
                                  "nuevo y existente."),
        ("Comunicaciones institucionales", "gaceta interna, boletines, "
                                            "comunicados oficiales."),
        ("Medios de comunicación", "entrevistas, notas de prensa, "
                                    "reportajes cuando LA EMPRESA lo "
                                    "considere oportuno."),
    ]
    for lead, rest in finalidades:
        add_bullet(doc, rest, bold_lead=lead + ".  ")

    # CAP IV — Duración
    add_chapter(doc, "IV", "DE LA DURACIÓN DE LA AUTORIZACIÓN")
    add_article(doc, "3", "Vigencia",
        "La presente autorización se mantiene durante la vigencia de la "
        "relación laboral y por un plazo de cinco (5) años contados a "
        "partir de la terminación de la misma, conforme a lo dispuesto en "
        "el artículo 183 de la LOTTT sobre conservación del expediente "
        "laboral y en la LOPDP.")
    add_article(doc, "4", "Material ya publicado",
        "Vencido el plazo, el material audiovisual ya publicado podrá "
        "mantenerse en los archivos institucionales con fines documentales "
        "y de archivo histórico, sin nueva difusión comercial sin "
        "consentimiento expreso del titular.")

    # CAP V — Sin contraprestación
    add_chapter(doc, "V", "DE LA AUSENCIA DE CONTRAPRESTACIÓN ECONÓMICA")
    add_para(doc,
        "El(la) trabajador(a) declara que la presente autorización se "
        "concede de manera gratuita, sin que LA EMPRESA esté obligada a "
        "efectuar pago alguno por el uso de su imagen o voz, distinto al "
        "salario y beneficios contractuales ordinarios. Cualquier "
        "contraprestación adicional será objeto de acuerdo escrito "
        "separado.", size=10)

    # CAP VI — Revocación
    add_chapter(doc, "VI", "DE LA REVOCACIÓN")
    add_article(doc, "5", "Plazo de anticipación",
        "El(la) trabajador(a) podrá revocar la presente autorización "
        "mediante comunicación escrita presentada ante la Dirección de "
        "Recursos Humanos, con al menos treinta (30) días continuos de "
        "anticipación a la fecha en que deba hacerse efectiva.")
    add_article(doc, "6", "Efectos",
        "La revocación no afectará el material audiovisual ya publicado "
        "con anterioridad a su efectividad, pero LA EMPRESA se abstendrá "
        "de realizar nuevas publicaciones que incluyan la imagen o voz "
        "del(la) trabajador(a) revocante.")

    # CAP VII — Protección de la dignidad y reputación
    add_chapter(doc, "VII", "DE LA PROTECCIÓN A LA DIGNIDAD Y REPUTACIÓN")
    add_para(doc,
        "LA EMPRESA se compromete a utilizar la imagen y voz del(la) "
        "trabajador(a) con respeto a su dignidad, honra, reputación y "
        "vida privada, conforme al artículo 60 de la Constitución. En "
        "ningún caso se realizarán montajes, manipulaciones o usos que "
        "puedan lesionar la imagen personal o profesional del(la) "
        "trabajador(a) o que constituyan actos de discriminación, "
        "acoso, burla o menoscabo.", size=10)

    # CAP VIII — Protección a menores y mascotas ajenas
    add_chapter(doc, "VIII", "DE LA PROTECCIÓN DE MENORES Y MASCOTAS AJENAS")
    add_article(doc, "7", "Menores de edad",
        "Cuando en la captación aparezcan menores de edad (familiares del "
        "personal o hijos de clientes), LA EMPRESA deberá contar con la "
        "autorización escrita de sus representantes legales para su "
        "publicación, conforme al artículo 7 de la LOPDP y al artículo "
        "60 de la Constitución.")
    add_article(doc, "8", "Mascotas de clientes",
        "La captación de mascotas de clientes requiere autorización "
        "expresa del propietario, conforme al procedimiento interno de "
        "ALIKA PETS. La publicación en redes sociales solo procederá "
        "tras la obtención de dicha autorización.")

    # CAP IX — Disposiciones finales
    add_chapter(doc, "IX", "DE LAS DISPOSICIONES FINALES")
    add_article(doc, "9", "Legislación aplicable",
        "La presente autorización se rige por la Constitución de la "
        "República Bolivariana de Venezuela, la LOPDP, la LOTTT y demás "
        "normas venezolanas aplicables.")
    add_article(doc, "10", "Jurisdicción",
        "Para cualquier controversia derivada de la presente autorización, "
        "las partes se someten a la jurisdicción de los Tribunales con "
        "competencia en materia de Protección del Trabajo de la "
        "Circunscripción Judicial del Estado Miranda.")

    # Carta de recepción
    add_reception_letter(
        doc,
        "Autorización de Uso de Imagen y Voz (Constitución art. 60 y LOPDP)",
    )

    # Footer
    add_footer(doc.sections[0], "Autorización Imagen/Voz v3.0")


# ============================================================
# 3. AUTORIZACIÓN DE VIGILANCIA POR CÁMARAS DE VIDEO
# ============================================================
def gen_vigilancia_camaras(doc):
    """10 capítulos: Const. art. 60 + LOPCYMAT arts. 46/74 + Ley Delitos Informáticos."""
    setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc,
                 doc_label="AUTORIZACIÓN CÁMARAS",
                 doc_sublabel="Vigilancia por Cámaras de Video",
                 version="Versión 3.0  ·  RR.HH.")
    add_doc_title(doc, "AUTORIZACIÓN DE VIGILANCIA POR CÁMARAS DE VIDEO")

    add_para(doc,
        "Conforme al artículo 60 de la Constitución de la República Bolivariana "
        "de Venezuela, los artículos 46 y 74 de la Ley Orgánica de Prevención, "
        "Condiciones y Medio Ambiente de Trabajo (LOPCYMAT), la Ley Orgánica de "
        "Protección de Datos Personales (LOPDP) y la Ley Especial contra los "
        "Delitos Informáticos, el(la) trabajador(a) que suscribe otorga su "
        "consentimiento libre, previo, expreso e informado para que GRUPO "
        "CAVAL 1003, C.A. (ALIKA PETS) instale y opere un sistema de cámaras "
        "de videovigilancia en el centro de trabajo, en los términos "
        "siguientes:",
        size=10, space_after=6)

    add_worker_identification(doc)

    # CAP I — De las zonas vigiladas
    add_chapter(doc, "I", "DE LAS ZONAS VIGILADAS")
    add_para(doc,
        "El sistema de videovigilancia cubre las siguientes zonas del centro "
        "de trabajo, ubicado en Av. Francisco de Miranda, Local N° 1, Los "
        "Teques, Miranda:", size=10)
    zonas = [
        "Área de ventas (tienda de mascotas) — mostrador, góndolas, caja.",
        "Área de clínica veterinaria — recepción, sala de espera, pasillos.",
        "Quirófano (solo entrada y monitoreo periférico, NO procedimientos).",
        "Hospitalización — monitoreo general de pacientes.",
        "Peluquería canina — bañera y mesa de peluquería.",
        "Depósito y bodega — control de inventario.",
        "Recepción general — ingreso y egreso de personal y visitantes.",
        "Estacionamiento externo — control de vehículos y seguridad.",
    ]
    for z in zonas:
        add_bullet(doc, z)
    add_para(doc,
        "Todas las zonas vigiladas contarán con cartel visible "
        "\u00abZona vigilada por cámaras de video\u00bb conforme a la "
        "LOPDP y la LOPCYMAT.", size=10, italic=True, color=GRAY_TEXT)

    # CAP II — De las zonas excluidas
    add_chapter(doc, "II", "DE LAS ZONAS EXCLUIDAS")
    add_para(doc,
        "En virtud del artículo 60 de la Constitución y el artículo 74 de la "
        "LOPCYMAT, queda expresamente PROHIBIDA la instalación de cámaras en "
        "las siguientes zonas íntimas o de descanso del personal:", size=10)
    excluidas = [
        "Baños y servicios sanitarios.",
        "Vestuarios y zonas de cambio de ropa.",
        "Áreas de descanso, comedor y sala de lactancia.",
        "Consultorios médicos durante la atención (salvo consentimiento "
        "expreso del paciente y del médico veterinario, con finalidad "
        "estrictamente clínica).",
    ]
    for e in excluidas:
        add_bullet(doc, e, color=None)
    add_para(doc,
        "La violación de esta prohibición será considerada falta grave "
        "conforme al Reglamento Interno y podrá dar lugar a las sanciones "
        "previstas en el artículo 79 de la LOTTT.",
        size=10, italic=True, color=None)

    # CAP III — De las finalidades
    add_chapter(doc, "III", "DE LAS FINALIDADES DE LA VIGILANCIA")
    add_para(doc, "El sistema de videovigilancia tiene como finalidades "
                  "legítimas las siguientes:", size=10)
    finalidades = [
        ("Seguridad", "proteger la integridad física del personal, "
                       "clientes, visitantes y mascotas."),
        ("Control de inventario", "prevenir y detectar hurtos, "
                                   "mermas o pérdidas en tienda, "
                                   "depósito y peluquería."),
        ("Prevención de incidentes", "anticipar y registrar accidentes "
                                      "de trabajo, incidentes, "
                                      "casi-accidentes y eventos "
                                      "adversos (LOPCYMAT art. 73)."),
        ("Atención de reclamaciones", "verificar hechos denunciados por "
                                       "clientes o trabajadores, "
                                       "garantizando el debido "
                                       "proceso."),
        ("Capacitación", "uso del material, previa anonimización, "
                          "para formación interna en seguridad y "
                          "procedimientos."),
    ]
    for lead, rest in finalidades:
        add_bullet(doc, rest, bold_lead=lead + ".  ")

    # CAP IV — De la retención
    add_chapter(doc, "IV", "DEL PERÍODO DE RETENCIÓN DE LAS GRABACIONES")
    add_article(doc, "1", "Retención general",
        "Las grabaciones se conservarán por un período máximo de noventa "
        "(90) días continuos, tras el cual serán automáticamente "
        "sobreescritas, conforme a lo establecido en la LOPDP y la "
        "buena práctica en materia de minimización de datos.")
    add_article(doc, "2", "Excepción por investigación",
        "Cuando una grabación sea requerida como evidencia en una "
        "investigación interna, denuncia, reclamación, inspección del "
        "trabajo, fiscalía o proceso judicial, será extraída del sistema, "
        "resguardada en medio seguro y conservada por el plazo necesario "
        "hasta la conclusión del asunto.")

    # CAP V — Del acceso restringido
    add_chapter(doc, "V", "DEL ACCESO A LAS GRABACIONES")
    add_para(doc,
        "El acceso a las grabaciones está restringido, conforme al "
        "principio de necesidad y confidencialidad, a las siguientes "
        "personas y roles:", size=10)
    accesos = [
        ("Gerencia", "para fines de supervisión general."),
        ("Dirección de Recursos Humanos",
         "para investigaciones internas disciplinarias."),
        ("Comité de Seguridad y Salud Laboral",
         "para investigaciones de accidentes (LOPCYMAT art. 46)."),
        ("Autoridades competentes",
         "Inspectoría del Trabajo, CICPC, Ministerio Público, "
         "Tribunales, cuando medie requerimiento formal."),
    ]
    for lead, rest in accesos:
        add_bullet(doc, rest, bold_lead=lead + ".  ")
    add_para(doc,
        "El acceso queda registrado en una bitácora con fecha, hora, "
        "responsable y finalidad, conforme al artículo 24 de la LOPDP.",
        size=10, italic=True, color=GRAY_TEXT)

    # CAP VI — De los usos prohibidos
    add_chapter(doc, "VI", "DE LOS USOS PROHIBIDOS")
    add_para(doc,
        "Quedan expresamente prohibidos los siguientes usos del sistema de "
        "videovigilancia:", size=10)
    prohibidos = [
        "La evaluación cotidiana del desempeño laboral del personal.",
        "La vigilancia personal, íntima o de costumbres fuera del "
        "ámbito laboral.",
        "La divulgación, publicación o cesión de las grabaciones a "
        "terceros sin causa legítima y autorización de la Gerencia.",
        "La grabación encubierta en zonas excluidas (CAP II).",
        "El uso de las imágenes para fines comerciales, publicitarios o "
        "de entretenimiento sin consentimiento expreso del(la) "
        "trabajador(a) (ver Autorización de Imagen y Voz).",
    ]
    for p in prohibidos:
        add_bullet(doc, p)

    # CAP VII — De la política de audio
    add_chapter(doc, "VII", "DE LA POLÍTICA DE AUDIO")
    add_article(doc, "3", "Prohibición general de grabación de audio",
        "El sistema de videovigilancia NO grabará audio en zonas "
        "internas, salvo en zonas públicas del centro de trabajo "
        "(recepción, estacionamiento externo) con cartel visible que "
        "informe \u00abZona con grabación de audio\u00bb.")
    add_article(doc, "4", "Excepción legal",
        "Cuando exista sospecha razonable de comisión de un delito, "
        "LA EMPRESA podrá activar la grabación de audio previa "
        "autorización escrita de la Gerencia, conforme a la Ley "
        "Especial contra los Delitos Informáticos.")

    # CAP VIII — Del derecho del trabajador a la revisión
    add_chapter(doc, "VIII", "DEL DERECHO DEL TRABAJADOR A SOLICITAR REVISIÓN")
    add_article(doc, "5", "Causa justificada",
        "El(la) trabajador(a) podrá solicitar por escrito a la "
        "Dirección de Recursos Humanos la revisión de las grabaciones "
        "que lo(la) afecten, cuando exista causa justificada: "
        "incidente de trabajo, acusación disciplinaria, queja de "
        "cliente, denuncia de acoso o cualquier situación que afecte "
        "sus derechos laborales.")
    add_article(doc, "6", "Plazo",
        "LA EMPRESA atenderá la solicitud en un plazo máximo de "
        "quince (15) días hábiles, programando la visualización de "
        "las imágenes en presencia del(la) trabajador(a), de un "
        "representante de RR.HH. y, si lo solicita, de un miembro del "
        "Comité de Seguridad y Salud Laboral.")

    # CAP IX — De la señalización
    add_chapter(doc, "IX", "DE LA SEÑALIZACIÓN OBLIGATORIA")
    add_para(doc,
        "LA EMPRESA colocará y mantendrá carteles visibles en todas las "
        "zonas vigiladas, conforme a la LOPDP y la LOPCYMAT. Los carteles "
        "incluirán:", size=10)
    carteles = [
        "Leyenda \u00abZona vigilada por cámaras de video\u00bb.",
        "Identificación del responsable del tratamiento (GRUPO CAVAL "
        "1003, C.A.) y datos de contacto de RR.HH.",
        "Referencia a la finalidad de la vigilancia y el plazo de "
        "retención (90 días).",
        "Mención de los derechos del titular (LOPDP arts. 14-22).",
    ]
    for c in carteles:
        add_bullet(doc, c)

    # CAP X — De las disposiciones finales
    add_chapter(doc, "X", "DE LAS DISPOSICIONES FINALES")
    add_article(doc, "7", "Marco normativo",
        "La presente autorización se rige por la Constitución de la "
        "República Bolivariana de Venezuela (art. 60), la LOPCYMAT "
        "(arts. 46, 73, 74), la LOPDP (Decreto 1.419), la Ley Especial "
        "contra los Delitos Informáticos, la LOTTT y demás normas "
        "venezolanas aplicables.")
    add_article(doc, "8", "Jurisdicción",
        "Para cualquier controversia derivada de la presente autorización, "
        "las partes se someten a la jurisdicción de los Tribunales con "
        "competencia en materia de Protección del Trabajo de la "
        "Circunscripción Judicial del Estado Miranda.")
    add_article(doc, "9", "Vigencia",
        "La presente autorización se mantiene durante la vigencia de la "
        "relación laboral y constituye requisito para el ingreso y "
        "permanencia en el centro de trabajo.")

    # Carta de recepción
    add_reception_letter(
        doc,
        "Autorización de Vigilancia por Cámaras de Video (LOPCYMAT arts. 46 y 74)",
    )

    # Footer
    add_footer(doc.sections[0], "Autorización Vigilancia/Cámaras v3.0")


# ============================================================
# Main
# ============================================================
def main():
    print("=" * 70)
    print("gen_autorizaciones.py — ALIKA PETS / Grupo Caval 1003, C.A.")
    print("=" * 70)

    docs_a_generar = [
        ("Autorizacion_Datos_Personales_LOPDP.docx",
         gen_datos_personales_lopdp),
        ("Autorizacion_Imagen_Redes_Sociales.docx",
         gen_imagen_redes_sociales),
        ("Autorizacion_Vigilancia_Camaras.docx",
         gen_vigilancia_camaras),
    ]

    generated = []
    for filename, fn in docs_a_generar:
        doc = Document()
        try:
            fn(doc)
            out_path = os.path.join(OUT_DIR, filename)
            doc.save(out_path)
            size = os.path.getsize(out_path) / 1024
            print(f"  ✓ {filename:50s}  {size:6.1f} KB")
            generated.append(out_path)
        except Exception as e:
            print(f"  ✗ Error generando {filename}: {e}")
            import traceback; traceback.print_exc()
    print(f"\nTotal generados: {len(generated)} / {len(docs_a_generar)}")
    print(f"Carpeta: {OUT_DIR}")
    return generated


if __name__ == "__main__":
    main()
