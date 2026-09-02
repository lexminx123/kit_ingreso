"""
gen_checklist_v3.py — Checklist Maestro de Ingreso v3.0
A4 horizontal, 6 páginas, con firma de Directora Gerente.
"""
import os, sys
sys.path.insert(0, "/home/z/my-project/output")
from _common import *
from docx.shared import Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/z/my-project/output/00_Checklist_Maestro_Ingreso_Trabajador.docx"

def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(4)

    # ===== Membrete =====
    hdr = doc.add_table(rows=1, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.autofit = False
    hdr.columns[0].width = Cm(17)
    hdr.columns[1].width = Cm(9)
    for c in hdr.rows[0].cells:
        set_cell_borders(c, color="FFFFFF", sz="0")

    left = hdr.rows[0].cells[0]
    left.text = ""
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ALIKA PETS")
    style_run(r, size=22, bold=True, color=TEAL_DARK)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("Grupo Caval 1003, C.A.  ·  RIF: J501662533")
    style_run(r2, size=10, color=BLACK)
    p3 = left.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    r3 = p3.add_run("Av. Francisco de Miranda, Local N° 1, Los Teques, Miranda  ·  Zona Postal 1201")
    style_run(r3, size=9, color=GRAY_TEXT)

    right = hdr.rows[0].cells[1]
    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    rr = rp.add_run("CHECKLIST MAESTRO")
    style_run(rr, size=14, bold=True, color=TEAL_DARK)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_before = Pt(0)
    rr2 = rp2.add_run("Kit de Ingreso del Trabajador")
    style_run(rr2, size=10, italic=True, color=GRAY_TEXT)
    rp3 = right.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp3.paragraph_format.space_before = Pt(0)
    rr3 = rp3.add_run("Versión 3.0  ·  RR.HH. / Dirección")
    style_run(rr3, size=8, color=GRAY_MUTED)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after = Pt(6)
    add_hr(sep, color="0F766E", sz="14")

    # ===== Datos del Trabajador =====
    p = doc.add_paragraph()
    r = p.add_run("DATOS DEL TRABAJADOR")
    style_run(r, size=11, bold=True, color=TEAL_DARK)
    p.paragraph_format.space_after = Pt(4)

    emp = doc.add_table(rows=3, cols=6)
    emp.alignment = WD_TABLE_ALIGNMENT.CENTER
    emp.autofit = False
    widths = [Cm(3.2), Cm(5.8), Cm(3.2), Cm(5.8), Cm(3.2), Cm(4.8)]
    for i, w in enumerate(widths):
        emp.columns[i].width = w
    for row in emp.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]

    emp_data = [
        [("Nombre y Apellido", ""), ("Cédula de Identidad", ""), ("Fecha de Ingreso", "")],
        [("Cargo", ""), ("Departamento / Área", ""), ("Salario Mensual (Bs.)", "")],
        [("Supervisor Inmediato", ""), ("Tipo de Contrato", ""), ("Fecha de Culminación", "")],
    ]
    for ri, row_data in enumerate(emp_data):
        for ci, (label, val) in enumerate(row_data):
            c_label = emp.rows[ri].cells[ci*2]
            c_val   = emp.rows[ri].cells[ci*2+1]
            write_cell(c_label, label, size=8, bold=True, color=TEAL_DARK, bg=SLATE_BG)
            c_val.text = ""
            pp = c_val.paragraphs[0]
            pp.paragraph_format.space_after = Pt(0)
            rr = pp.add_run(" ________________________")
            style_run(rr, size=9)
            c_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(c_val)
            set_cell_margins(c_val)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ===== Instrucciones =====
    p = doc.add_paragraph()
    r = p.add_run("INSTRUCCIONES DE USO")
    style_run(r, size=11, bold=True, color=TEAL_DARK)
    p.paragraph_format.space_after = Pt(3)

    inst = (
        "1. Este checklist controla la entrega y firma de cada documento del Kit de Ingreso.\n"
        "2. Estado:  P = Pendiente  ·  E = Entregado  ·  F = Firmado  ·  N/A = No aplica al cargo.\n"
        "3. Los documentos marcados como OBLIGATORIO son exigidos por norma vigente. Su omisión genera responsabilidad para la empresa.\n"
        "4. Los documentos marcados como RECOMENDADO refuerzan la protección legal y operativa.\n"
        "5. Siga el ORDEN DE ENTREGA recomendado (sección final). No firme el contrato sin verificar los documentos del Bloque 06.\n"
        "6. Al finalizar, archive el checklist firmado junto con las copias de todos los documentos en el expediente del trabajador."
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(inst.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        style_run(r, size=9, color=BLACK)

    # ===== Tabla Maestra =====
    p = doc.add_paragraph()
    r = p.add_run("TABLA MAESTRA DE DOCUMENTOS DEL KIT")
    style_run(r, size=11, bold=True, color=TEAL_DARK)
    p.paragraph_format.space_after = Pt(4)

    docs = [
        ("00", "CONTROL",                       "Checklist Maestro de Ingreso (este documento)",        "OBLIGATORIO",  "Control interno RR.HH.",                  "Dirección"),
        ("01", "SOLICITUD",                     "Solicitud de Empleo / Ficha de Ingreso Ampliada",      "OBLIGATORIO",  "Art. 57 LOTTT (identificación)",          "Dirección"),
        ("02a","CONTRATOS",                     "Contrato de Trabajo - Gerente",                        "OBLIGATORIO",  "Art. 57-64 LOTTT",                        "Dirección / Rep. Legal"),
        ("02b","CONTRATOS",                     "Contrato de Trabajo - Encargado de Clinica",           "OBLIGATORIO",  "Art. 57-64 LOTTT",                        "Dirección / Rep. Legal"),
        ("02c","CONTRATOS",                     "Contrato de Trabajo - Encargado de Tienda",            "OBLIGATORIO",  "Art. 57-64 LOTTT",                        "Dirección / Rep. Legal"),
        ("02d","CONTRATOS",                     "Contrato de Trabajo - Medico Veterinario",             "OBLIGATORIO",  "Art. 57-64 LOTTT",                        "Dirección / Rep. Legal"),
        ("02e","CONTRATOS",                     "Contrato de Trabajo - Auxiliar Veterinario",           "OBLIGATORIO",  "Art. 57-64 LOTTT",                        "Dirección / Rep. Legal"),
        ("02f","CONTRATOS",                     "Contrato de Trabajo - Dog Groomer",                    "OBLIGATORIO",  "Art. 57-64 LOTTT",                        "Dirección / Rep. Legal"),
        ("02g","CONTRATOS",                     "Contrato de Trabajo - Asistente de Ventas",            "OBLIGATORIO",  "Art. 57-64 LOTTT",                        "Dirección / Rep. Legal"),
        ("03a","DESCRIPCION DE CARGOS",         "Funciones del Cargo (segun rol) + Carta de Recepcion", "OBLIGATORIO",  "Art. 26 LOTTT",                           "Dirección / Jefe directo"),
        ("04a","PRESTACIONES SOCIALES",         "Autorizacion para Deposito de Prestaciones",           "OBLIGATORIO",  "Art. 143 LOTTT",                          "Dirección / Contabilidad"),
        ("04b","PRESTACIONES SOCIALES",         "Designacion de Beneficiarios (prestaciones)",          "OBLIGATORIO",  "Art. 137 LOTTT",                          "Dirección"),
        ("05a","SEGURIDAD LABORAL",             "Notificacion de Riesgos (especifica por rol)",         "OBLIGATORIO",  "Art. 56 LOPCYMAT",                        "Dirección / SST"),
        ("05b","SEGURIDAD LABORAL",             "Hoja de Recorrido Habitual del Trabajador",            "OBLIGATORIO",  "Art. 69 LOPCYMAT",                        "Dirección"),
        ("05c","SEGURIDAD LABORAL",             "Acta de Entrega de Uniformes y EPP",                   "OBLIGATORIO",  "Art. 53 LOPCYMAT",                        "Dirección / SST"),
        ("05d","SEGURIDAD LABORAL",             "Examen Medico Pre-Empleo",                             "OBLIGATORIO",  "Art. 32 LOPCYMAT + NT-02-2008",           "Dirección / SST"),
        ("05e","SEGURIDAD LABORAL",             "Cartilla de Bioseguridad Veterinaria",                 "OBLIGATORIO",  "NT-01-2008 (vet/aux/groomer)",            "Dirección / SST"),
        ("06a","REGISTROS LEGALES",             "Constancia Inscripcion IVSS (Forma 14-100)",           "OBLIGATORIO",  "Ley del Seguro Social",                   "Dirección / Trabajador"),
        ("06b","REGISTROS LEGALES",             "Constancia Inscripcion FAOV-BVV",                      "OBLIGATORIO",  "Ley del Bono Vacacional",                 "Dirección / Trabajador"),
        ("06c","REGISTROS LEGALES",             "Constancia Inscripcion INCES",                         "OBLIGATORIO",  "Ley del INCES",                           "Dirección / Trabajador"),
        ("06d","REGISTROS LEGALES",             "Constancia Seguro Riesgos Laborales (IVSS-PMSSO)",     "OBLIGATORIO",  "Art. 134 LOPCYMAT  (CRITICO)",            "Dirección / Trabajador"),
        ("06e","REGISTROS LEGALES",             "Constancia Vacunacion Antirrabica Pre-Exposicion",     "RECOMENDADO",  "Profilaxis ocupacional (vet/aux/groomer)","Dirección / SST"),
        ("07a","AUTORIZACIONES",                "Autorizacion de Datos Personales (LOPDP)",             "OBLIGATORIO",  "Ley Org. Proteccion Datos Personales",    "Dirección"),
        ("07b","AUTORIZACIONES",                "Autorizacion de Uso de Imagen (redes sociales)",       "RECOMENDADO",  "Marketing / Comunicaciones",              "Dirección / Marketing"),
        ("07c","AUTORIZACIONES",                "Autorizacion de Vigilancia por Camaras",               "RECOMENDADO",  "Ley sobre videos en el trabajo",          "Dirección"),
        ("08a","POLITICAS INTERNAS",            "Reglamento Interno (recepcion firmada)",               "OBLIGATORIO",  "Art. 191 LOTTT",                          "Dirección"),
        ("08b","POLITICAS INTERNAS",            "Codigo de Conducta y Etica",                           "RECOMENDADO",  "Politica empresarial",                    "Dirección"),
        ("08c","POLITICAS INTERNAS",            "Politica de Confidencialidad",                         "OBLIGATORIO",  "Proteccion de informacion",               "Dirección"),
        ("08d","POLITICAS INTERNAS",            "Politica de Uso de Redes Sociales",                    "RECOMENDADO",  "Politica empresarial",                    "Dirección"),
        ("08e","POLITICAS INTERNAS",            "Protocolo de Sustancias Controladas (vet/aux)",        "OBLIGATORIO",  "SENAC (ketamina/opioides)",               "Direccion Medica"),
        ("08f","POLITICAS INTERNAS",            "Protocolo de Mordeduras y Zoonosis",                   "OBLIGATORIO",  "NT-01-2008 (vet/aux/groomer)",            "Dirección / SST"),
        ("08g","POLITICAS INTERNAS",            "Procedimiento de Reporte de Incidentes",               "OBLIGATORIO",  "Art. 73 LOPCYMAT",                        "Dirección / SST"),
        ("09", "CIERRE",                        "Carta de Aceptacion General del Kit",                  "OBLIGATORIO",  "Cierre del expediente",                   "Dirección"),
    ]

    n_rows = len(docs) + 1
    tbl = doc.add_table(rows=n_rows, cols=9)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    col_widths = [Cm(0.9), Cm(2.6), Cm(6.0), Cm(2.2), Cm(4.5), Cm(3.0), Cm(1.4), Cm(1.7), Cm(3.2)]
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = w

    headers = ["#", "BLOQUE", "DOCUMENTO", "OBLIGATORIEDAD", "FUNDAMENTO / NORMA",
               "RESPONSABLE", "ESTADO", "FECHA", "OBSERVACIONES"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        write_cell(c, h, size=8, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
        c.width = col_widths[ci]

    for ri, row in enumerate(docs, start=1):
        num, bloque, docname, oblig, fund, resp = row
        is_alt = (ri % 2 == 0)
        bg = GRAY_ALT if is_alt else WHITE
        if "REGISTROS LEGALES" in bloque:
            bg = AMBER_ALT if is_alt else AMBER_BG

        vals = [num, bloque, docname, oblig, fund, resp, "☐", "__/__/____", ""]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
        for ci, (val, al) in enumerate(zip(vals, aligns)):
            c = tbl.rows[ri].cells[ci]
            c.width = col_widths[ci]
            col = None
            bold = False
            if ci == 3:
                if "OBLIGATORIO" in val:
                    col = RED_CRIT
                    bold = True
                else:
                    col = GREEN_OK
            if ci == 1:
                bold = True
                col = TEAL_DARK
            if ci == 0:
                bold = True
            write_cell(c, val, size=8, bold=bold, color=col, align=al, bg=bg)

    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), 'true')
    trPr.append(th)
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit')
        cs.set(qn('w:val'), 'true')
        trPr.append(cs)

    # ===== Leyenda =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("LEYENDA")
    style_run(r, size=10, bold=True, color=TEAL_DARK)

    ley = doc.add_paragraph()
    ley.paragraph_format.line_spacing = 1.2
    for txt, col, bold in [
        ("OBLIGATORIO ", RED_CRIT, True),
        ("= exigido por norma vigente (su omision genera responsabilidad).     ", BLACK, False),
        ("RECOMENDADO ", GREEN_OK, True),
        ("= refuerza la proteccion legal/operativa.     ", BLACK, False),
        ("ESTADO:  P=Pendiente  ·  E=Entregado  ·  F=Firmado  ·  N/A=No aplica.", BLACK, False),
    ]:
        r = ley.add_run(txt)
        style_run(r, size=9, bold=bold, color=col)

    # ===== Orden de Entrega =====
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("ORDEN RECOMENDADO DE ENTREGA")
    style_run(r, size=12, bold=True, color=TEAL_DARK)
    p.paragraph_format.space_after = Pt(6)

    fases = [
        ("FASE 0 - PRE-INGRESO  (antes del primer dia de trabajo)",
         "Verificar antecedentes y aptitud antes de firmar contrato.",
         ["01  Solicitud de Empleo / Ficha de Ingreso Ampliada",
          "Verificacion de referencias laborales y personales",
          "06a Constancia IVSS (Forma 14-100) - el trabajador debe consignarla",
          "06b Constancia FAOV - el trabajador debe consignarla",
          "06c Constancia INCES - el trabajador debe consignarla",
          "05d Examen Medico Pre-Empleo (programar cita)"]),
        ("FASE 1 - FIRMA DE CONTRATO  (Dia 1)",
         "Formalizacion de la relacion laboral. No iniciar labores sin estos documentos firmados.",
         ["02a-02g Contrato de Trabajo correspondiente al cargo",
          "03a  Funciones del Cargo + Carta de Recepcion firmada",
          "04a  Autorizacion para Deposito de Prestaciones Sociales",
          "04b  Designacion de Beneficiarios",
          "06d  Constancia Seguro Riesgos Laborales (PMSSO) - CRITICO, no omitir"]),
        ("FASE 2 - INDUCCION Y SEGURIDAD  (Dia 1 a Dia 3)",
         "Cumplimiento LOPCYMAT. Todo trabajador debe recibir y firmar antes de iniciar labores operativas.",
         ["05a  Notificacion de Riesgos (especifica por rol)",
          "05b  Hoja de Recorrido Habitual del Trabajador",
          "05c  Acta de Entrega de Uniformes y EPP",
          "08a  Reglamento Interno (recepcion firmada)",
          "08c  Politica de Confidencialidad",
          "08g  Procedimiento de Reporte de Incidentes"]),
        ("FASE 3 - REGISTROS Y AUTORIZACIONES  (Semana 1)",
         "Tramites administrativos y protecciones de datos.",
         ["07a  Autorizacion de Datos Personales (LOPDP)",
          "07b  Autorizacion de Uso de Imagen",
          "07c  Autorizacion de Vigilancia por Camaras",
          "08b  Codigo de Conducta y Etica",
          "08d  Politica de Uso de Redes Sociales"]),
        ("FASE 4 - PROTOCOLOS ESPECIFICOS VETERINARIOS  (Semana 1-2)",
         "Solo para roles con exposicion biologica o manejo de sustancias controladas: Veterinario, Auxiliar Veterinario, Dog Groomer.",
         ["05e  Cartilla de Bioseguridad Veterinaria",
          "06e  Constancia Vacunacion Antirrabica Pre-Exposicion",
          "08e  Protocolo de Sustancias Controladas (vet/aux)",
          "08f  Protocolo de Mordeduras y Zoonosis"]),
        ("FASE 5 - CIERRE Y ARCHIVO  (Semana 2)",
         "Verificacion final del expediente completo.",
         ["Verificar que TODOS los documentos del checklist esten firmados",
          "09  Carta de Aceptacion General del Kit (firma del trabajador)",
          "Archivar expediente fisico + copia digital",
          "Registrar al trabajador en nomina"]),
    ]

    for titulo, desc, items in fases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(titulo)
        style_run(r, size=10, bold=True, color=TEAL_DARK)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(desc)
        style_run(r2, size=9, italic=True, color=GRAY_TEXT)
        for it in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run("☐  ")
            style_run(r, size=10, bold=True, color=TEAL_DARK)
            r2 = p.add_run(it)
            style_run(r2, size=9, color=BLACK)

    # ===== Resumen de Cobertura =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("RESUMEN DE COBERTURA DEL KIT")
    style_run(r, size=11, bold=True, color=TEAL_DARK)

    res = doc.add_table(rows=6, cols=4)
    res.alignment = WD_TABLE_ALIGNMENT.CENTER
    res.autofit = False
    rwidths = [Cm(7), Cm(4), Cm(4), Cm(10)]
    for i, w in enumerate(rwidths):
        res.columns[i].width = w

    res_data = [
        [("Bloque", True), ("Documentos", True), ("Estado actual", True), ("Comentario", True)],
        [("Contratos por rol (7 cargos)", False), ("7", False), ("7/7  (100%)", False), ("Completos. Razon social Grupo Caval 1003.", False)],
        [("Seguridad laboral LOPCYMAT", False), ("9", False), ("9/9  (100%)", False), ("Completa: riesgo por rol, EPP, examen, bioseguridad.", False)],
        [("Registros legales (IVSS/FAOV/INCES/PMSSO)", False), ("5", False), ("5/5  (100%)", False), ("CRITICO: PMSSO destacado en ambar.", False)],
        [("Autorizaciones (LOPDP, imagen, camaras)", False), ("3", False), ("3/3  (100%)", False), ("Completas conforme a LOPDP.", False)],
        [("Politicas internas y protocolos vet.", False), ("7", False), ("7/7  (100%)", False), ("Reglamento, codigo, protocolos vet, etc.", False)],
    ]
    for ri, row in enumerate(res_data):
        for ci, (txt, bold) in enumerate(row):
            c = res.rows[ri].cells[ci]
            c.width = rwidths[ci]
            if ri == 0:
                write_cell(c, txt, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                           align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
            else:
                bg = GRAY_ALT if ri % 2 == 0 else WHITE
                col = RED_CRIT if ("CRITICO" in txt) else BLACK
                write_cell(c, txt, size=9, bold=bold, color=col, bg=bg,
                           align=WD_ALIGN_PARAGRAPH.LEFT if ci in (0,3) else WD_ALIGN_PARAGRAPH.CENTER)

    # ===== Firma de Control =====
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("CONTROL Y CIERRE DEL EXPEDIENTE")
    style_run(r, size=11, bold=True, color=TEAL_DARK)

    add_para(doc,
        "Declaro que he revisado el presente checklist y que todos los documentos marcados como "
        "OBLIGATORIO han sido entregados al trabajador y firmados por este. El expediente se "
        "encuentra completo y archivado.",
        size=9, italic=True, color=GRAY_TEXT, space_after=8)

    # Firma: Trabajador + Directora Gerente (2 columnas, no 3)
    fir = doc.add_table(rows=2, cols=2)
    fir.alignment = WD_TABLE_ALIGNMENT.CENTER
    fir.autofit = False
    fwidths = [Cm(12.75), Cm(12.75)]
    for i, w in enumerate(fwidths):
        fir.columns[i].width = w
    for row in fir.rows:
        for i, c in enumerate(row.cells):
            c.width = fwidths[i]

    fir_labels = ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"]
    # Datos por columna: (nombre, cargo, ci)
    fir_data = [
        (None, None, None),  # Trabajador — campos vacíos
        ("ESNATLIM ELENA SIMOZA", "Directora Gerente", "V-17.976.287"),
    ]
    for ci, lab in enumerate(fir_labels):
        c = fir.rows[0].cells[ci]
        c.width = fwidths[ci]
        write_cell(c, lab, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        c2 = fir.rows[1].cells[ci]
        c2.width = fwidths[ci]
        c2.text = ""
        pp = c2.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(28)
        pp.paragraph_format.space_after = Pt(2)
        r = pp.add_run("______________________________")
        style_run(r, size=9)
        pp2 = c2.add_paragraph()
        pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp2.paragraph_format.space_after = Pt(0)
        name, cargo, ci = fir_data[ci]
        if name:
            r2 = pp2.add_run(name)
            style_run(r2, size=8, bold=True, color=BLACK)
            pp3 = c2.add_paragraph()
            pp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp3.paragraph_format.space_after = Pt(0)
            r3 = pp3.add_run(f"C.I. {ci}")
            style_run(r3, size=8, color=GRAY_TEXT)
            pp4 = c2.add_paragraph()
            pp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp4.paragraph_format.space_after = Pt(0)
            r4 = pp4.add_run(cargo)
            style_run(r4, size=8, italic=True, color=GRAY_TEXT)
        else:
            r2 = pp2.add_run("Firma · C.I.: ____________________")
            style_run(r2, size=8, italic=True, color=GRAY_TEXT)
        pp5 = c2.add_paragraph()
        pp5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp5.paragraph_format.space_before = Pt(2)
        pp5.paragraph_format.space_after = Pt(0)
        r5 = pp5.add_run("Fecha: ____ / ____ / ________")
        style_run(r5, size=8, italic=True, color=GRAY_TEXT)
        set_cell_borders(c)
        set_cell_borders(c2)
        set_cell_margins(c)
        set_cell_margins(c2)

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("ALIKA PETS  ·  Grupo Caval 1003, C.A.  ·  RIF J501662533  ·  Checklist Maestro de Ingreso v3.0  ·  Página ")
    style_run(fr, size=8, color=GRAY_MUTED)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r2 = fp.add_run()
    r2._r.append(fldChar1)
    r2._r.append(instrText)
    r2._r.append(fldChar2)
    style_run(r2, size=8, color=GRAY_MUTED)

    doc.save(OUT)
    return OUT

if __name__ == "__main__":
    out = build()
    size_kb = os.path.getsize(out) / 1024
    print(f"✓ {out}  ({size_kb:.1f} KB)")
