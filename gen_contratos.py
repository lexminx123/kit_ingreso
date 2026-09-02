"""
gen_contratos.py — Genera los 7 contratos individuales de trabajo de ALIKA PETS
(Grupo Caval 1003, C.A. — RIF J501662533).

Versión 3.0
-----------
- Razón social: GRUPO CAVAL 1003, C.A. (J501662533)
- Domicilio: Los Teques, Miranda (alineado con RIF fiscal)
- Salario básico: Bs. 210/mes (Gaceta 7.026 Extraordinaria del 28/04/2025)
- Cestaticket: Bs. 1.500/mes
- Bono de transporte: Bs. 200/mes (no salarial)
- 7 cláusulas comunes + cláusula de historias clínicas (Vet/Aux/Groomer)
- Penalidades simplificadas a 3 niveles (verbal, escrita, rescisión)
- Firma final: LA EMPRESA / EL(LA) TRABAJADOR(A) — el módulo _common.py
  mapea automáticamente "LA EMPRESA" a la DIRECTORA GERENTE
  (Esnatlim Elena Simoza — C.I. V-17.976.287).

Roles:
  02a Gerente                  — indefinido, prueba 30 días
  02b Encargado de Clínica     — 6 meses + prórroga, prueba 30 días
  02c Encargado de Tienda      — 6 meses + prórroga, prueba 30 días
  02d Médico Veterinario       — indefinido + cláusula historias, prueba 30 días
  02e Auxiliar Veterinario     — 3 meses + prórroga + cláusula historias, prueba 15 días
  02f Dog Groomer              — indefinido + cláusula historias, prueba 30 días
  02g Asistente de Ventas      — 3 meses + prórroga, prueba 15 días
"""
import os
import sys

# Asegurar import del _common.py
HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import _common
from _common import (
    setup_a4_portrait, add_membrete, add_doc_title, add_section,
    add_para, add_bullet, add_signature_block, add_footer,
    set_cell_bg, set_cell_borders, set_cell_margins, write_cell, style_run,
    add_hr, add_reception_letter,
    TEAL_DARK, TEAL_HDR_BG, GRAY_ALT, SLATE_BG, WHITE,
    BLACK, GRAY_TEXT, GRAY_MUTED, RED_CRIT, GREEN_OK,
    EMPRESA, RIF_EMP, MARCA, DOMICILIO_EMP,
    REP_LEGAL_NOMBRE, REP_LEGAL_CARGO, REP_LEGAL_CI,
    DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI,
)

# ============================================================
# Datos legales y económicos (actualizados Gaceta 7.026 Ext. 28/04/2025)
# ============================================================
SALARIO_BASICO = "Bs. 210,00"
CESTATICKET    = "Bs. 1.500,00"
BONO_TRANSPORTO = "Bs. 200,00"
DOMICILIO_PROCESAL = "Los Teques, Estado Miranda"

# ============================================================
# Configuración por rol
# ============================================================
ROLES = [
    {
        "code": "02a",
        "file": "02a_Contrato_Gerente.docx",
        "cargo": "GERENTE",
        "titulo": "CONTRATO INDIVIDUAL DE TRABAJO — GERENTE",
        "doc_label": "CONTRATO INDIVIDUAL DE TRABAJO",
        "doc_sublabel": "Gerente General",
        "footer_label": "Contrato Gerente v3.0",
        "duracion": "INDEFINIDA",
        "duracion_detalle": (
            "El presente contrato se celebra por tiempo indeterminado, conforme a lo "
            "previsto en el artículo 65 de la Ley Orgánica del Trabajo, los Trabajadores "
            "y las Trabajadoras (LOTTT). La relación laboral comenzará a regir a partir "
            "de la fecha de inicio de actividades del trabajador en la empresa."
        ),
        "prueba": "30 días continuos",
        "prueba_detalle": (
            "Conforme al artículo 22 de la LOTTT, se establece un período de prueba de "
            "treinta (30) días continuos, durante el cual cualquiera de las partes "
            "podrá dar por terminada la relación laboral sin necesidad de preaviso ni "
            "indemnización alguna."
        ),
        "funciones": [
            "Planificar, dirigir y controlar las operaciones generales de la empresa (clínica veterinaria, tienda de mascotas y peluquería canina).",
            "Supervisar al personal de las tres áreas (clínica, tienda y peluquería) y velar por el cumplimiento de los protocolos internos.",
            "Elaborar y ejecutar el plan estratégico comercial, presupuesto anual y metas de venta.",
            "Gestionar relaciones con proveedores, distribuidores y laboratorios veterinarios.",
            "Garantizar el cumplimiento de las normas legales: LOTTT, LOPCYMAT, LOPDP, NT-01-2008, NT-02-2008 y demás normativa venezolana aplicable.",
            "Autorizar gastos, pagos a proveedores y movimientos bancarios conforme a los límites aprobados por la Junta Directiva.",
            "Conducir reuniones de coordinación semanal con encargados de área.",
            "Atender y resolver escalamientos de clientes cuando corresponda.",
            "Rendir informe mensual de gestión ante la Vicepresidencia.",
        ],
        "incluye_historias": False,
    },
    {
        "code": "02b",
        "file": "02b_Contrato_Encargado_Clinica.docx",
        "cargo": "ENCARGADO DE CLÍNICA",
        "titulo": "CONTRATO INDIVIDUAL DE TRABAJO — ENCARGADO DE CLÍNICA",
        "doc_label": "CONTRATO INDIVIDUAL DE TRABAJO",
        "doc_sublabel": "Encargado(a) de Clínica Veterinaria",
        "footer_label": "Contrato Encargado de Clínica v3.0",
        "duracion": "6 MESES + PRÓRROGA",
        "duracion_detalle": (
            "Conforme a los artículos 64 y 65 de la LOTTT, el presente contrato se celebra "
            "por tiempo determinado de seis (6) meses, en razón de la naturaleza del cargo "
            "de encargado(a) de clínica, cuya actividad admite esta modalidad. El contrato "
            "podrá prorrogarse por igual período hasta por el límite legal de un (1) año; "
            "vencido este plazo, la relación se considerará por tiempo indeterminado."
        ),
        "prueba": "30 días continuos",
        "prueba_detalle": (
            "Conforme al artículo 22 de la LOTTT, se establece un período de prueba de "
            "treinta (30) días continuos, durante el cual cualquiera de las partes "
            "podrá dar por terminada la relación laboral sin necesidad de preaviso ni "
            "indemnización."
        ),
        "funciones": [
            "Coordinar la atención clínica y quirúrgica de pacientes bajo supervisión del médico veterinario tratante.",
            "Gestionar la agenda de citas, hospitalización y cirugías.",
            "Supervisar el inventario de medicamentos, insumos médicos y material quirúrgico de la clínica.",
            "Velar por el cumplimiento de las normas de bioseguridad (NT-01-2008) y vigilancia médica (NT-02-2008).",
            "Atender y resolver reclamaciones de clientes de la clínica.",
            "Llevar el control de historias clínicas y su archivo conforme al artículo 23 de la Ley de Ejercicio de la Medicina Veterinaria.",
            "Coordinar el turno del personal auxiliar y de enfermería veterinaria.",
            "Reportar a Gerencia los indicadores de atención clínica y egresos hospitalarios.",
        ],
        "incluye_historias": False,
    },
    {
        "code": "02c",
        "file": "02c_Contrato_Encargado_Tienda.docx",
        "cargo": "ENCARGADO DE TIENDA",
        "titulo": "CONTRATO INDIVIDUAL DE TRABAJO — ENCARGADO DE TIENDA",
        "doc_label": "CONTRATO INDIVIDUAL DE TRABAJO",
        "doc_sublabel": "Encargado(a) de Tienda de Mascotas",
        "footer_label": "Contrato Encargado de Tienda v3.0",
        "duracion": "6 MESES + PRÓRROGA",
        "duracion_detalle": (
            "Conforme a los artículos 64 y 65 de la LOTTT, el presente contrato se celebra "
            "por tiempo determinado de seis (6) meses, en razón de la naturaleza del cargo "
            "de encargado(a) de tienda. El contrato podrá prorrogarse por igual período "
            "hasta por el límite legal de un (1) año; vencido este plazo, la relación se "
            "considerará por tiempo indeterminado."
        ),
        "prueba": "30 días continuos",
        "prueba_detalle": (
            "Conforme al artículo 22 de la LOTTT, se establece un período de prueba de "
            "treinta (30) días continuos, durante el cual cualquiera de las partes "
            "podrá dar por terminada la relación laboral sin necesidad de preaviso ni "
            "indemnización."
        ),
        "funciones": [
            "Coordinar la operación diaria de la tienda de mascotas: atención al cliente, ventas, facturación y arqueo de caja.",
            "Supervisar el inventario de productos (alimentos, accesorios, higiene, juguetes, farmacia OTC).",
            "Realizar y controlar pedidos a proveedores y recepción de mercancía.",
            "Velar por la exhibición, rotación de productos y control de fechas de vencimiento.",
            "Garantizar la atención conforme a las políticas de servicio al cliente y manejo de quejas.",
            "Coordinar el turno del personal de ventas de la tienda.",
            "Reportar a Gerencia los indicadores comerciales diarios y mensuales (ventas, devoluciones, mermas).",
            "Cumplir y hacer cumplir las normas LOPCYMAT aplicables al área de tienda (ergonomía, manipulación manual de cargas).",
        ],
        "incluye_historias": False,
    },
    {
        "code": "02d",
        "file": "02d_Contrato_Medico_Veterinario.docx",
        "cargo": "MÉDICO VETERINARIO",
        "titulo": "CONTRATO INDIVIDUAL DE TRABAJO — MÉDICO VETERINARIO",
        "doc_label": "CONTRATO INDIVIDUAL DE TRABAJO",
        "doc_sublabel": "Médico(a) Veterinario(a) Colegiado(a)",
        "footer_label": "Contrato Médico Veterinario v3.0",
        "duracion": "INDEFINIDA",
        "duracion_detalle": (
            "El presente contrato se celebra por tiempo indeterminado, conforme a lo "
            "previsto en el artículo 65 de la LOTTT, en razón de la naturaleza permanente "
            "de las funciones del médico veterinario en la clínica. La relación laboral "
            "comenzará a regir a partir de la fecha de inicio de actividades del "
            "trabajador en la empresa."
        ),
        "prueba": "30 días continuos",
        "prueba_detalle": (
            "Conforme al artículo 22 de la LOTTT, se establece un período de prueba de "
            "treinta (30) días continuos, durante el cual cualquiera de las partes "
            "podrá dar por terminada la relación laboral sin necesidad de preaviso ni "
            "indemnización."
        ),
        "funciones": [
            "Realizar consultas clínicas, diagnósticas y tratamientos de pacientes (caninos, felinos y otras especies menores).",
            "Ejecutar cirugías y procedimientos médicos-quirúrgicos conforme a los protocolos internos.",
            "Llevar y mantener actualizadas las historias clínicas de los pacientes atendidos.",
            "Prescribir y administrar medicamentos, incluidas sustancias controladas, conforme a la normativa del SENAC y la Ley Orgánica de Drogas.",
            "Atender emergencias y guardias conforme al cronograma establecido.",
            "Aplicar las normas de bioseguridad NT-01-2008 y vigilancia médica NT-02-2008.",
            "Atender y resolver consultas de los dueños o responsables de los pacientes.",
            "Mantener actualizado el registro de vacunas y desparasitaciones.",
            "Reportar enfermedades de notificación obligatoria ante las autoridades sanitarias competentes.",
            "Cumplir el Código de Ética del Colegio de Médicos Veterinarios de Venezuela.",
        ],
        "incluye_historias": True,
    },
    {
        "code": "02e",
        "file": "02e_Contrato_Auxiliar_Veterinario.docx",
        "cargo": "AUXILIAR VETERINARIO",
        "titulo": "CONTRATO INDIVIDUAL DE TRABAJO — AUXILIAR VETERINARIO",
        "doc_label": "CONTRATO INDIVIDUAL DE TRABAJO",
        "doc_sublabel": "Auxiliar Veterinario",
        "footer_label": "Contrato Auxiliar Veterinario v3.0",
        "duracion": "3 MESES + PRÓRROGA",
        "duracion_detalle": (
            "Conforme a los artículos 64 y 65 de la LOTTT, el presente contrato se celebra "
            "por tiempo determinado de tres (3) meses, en razón de la naturaleza del cargo "
            "de auxiliar veterinario. El contrato podrá prorrogarse por períodos iguales "
            "hasta por el límite legal; vencido este plazo, la relación se considerará "
            "por tiempo indeterminado."
        ),
        "prueba": "15 días continuos",
        "prueba_detalle": (
            "Conforme al artículo 22 de la LOTTT, se establece un período de prueba de "
            "quince (15) días continuos, durante el cual cualquiera de las partes "
            "podrá dar por terminada la relación laboral sin necesidad de preaviso ni "
            "indemnización."
        ),
        "funciones": [
            "Asistir al médico veterinario en consultas, cirugías y procedimientos médicos.",
            "Restringir, contener y manejar a los pacientes animales con técnicas de bienestar animal.",
            "Preparar material quirúrgico, instrumental y equipos médicos para procedimientos.",
            "Administrar medicamentos y vacunas bajo indicación del médico veterinario tratante.",
            "Apoyar en la toma de muestras y procesamiento de exámenes de laboratorio básicos.",
            "Alimentar, hidratar y limpiar jaulas de hospitalización conforme al protocolo de bioseguridad.",
            "Registrar y archivar historias clínicas bajo supervisión del médico veterinario.",
            "Aplicar las normas de bioseguridad NT-01-2008, uso de EPP y manejo de residuos COVENIN 2747-93.",
            "Atender y orientar a los dueños o responsables de los pacientes en ausencia del veterinario.",
        ],
        "incluye_historias": True,
    },
    {
        "code": "02f",
        "file": "02f_Contrato_Dog_Groomer.docx",
        "cargo": "DOG GROOMER",
        "titulo": "CONTRATO INDIVIDUAL DE TRABAJO — DOG GROOMER",
        "doc_label": "CONTRATO INDIVIDUAL DE TRABAJO",
        "doc_sublabel": "Peluquero(a) Canino(a)",
        "footer_label": "Contrato Dog Groomer v3.0",
        "duracion": "INDEFINIDA",
        "duracion_detalle": (
            "El presente contrato se celebra por tiempo indeterminado, conforme a lo "
            "previsto en el artículo 65 de la LOTTT, en razón de la naturaleza permanente "
            "de las funciones del dog groomer en la peluquería canina. La relación "
            "laboral comenzará a regir a partir de la fecha de inicio de actividades del "
            "trabajador en la empresa."
        ),
        "prueba": "30 días continuos",
        "prueba_detalle": (
            "Conforme al artículo 22 de la LOTTT, se establece un período de prueba de "
            "treinta (30) días continuos, durante el cual cualquiera de las partes "
            "podrá dar por terminada la relación laboral sin necesidad de preaviso ni "
            "indemnización."
        ),
        "funciones": [
            "Realizar baños, cortes, peinados y arreglo estético de caninos y felinos conforme a las indicaciones del dueño y estándares de raza.",
            "Aplicar técnicas de manejo y bienestar animal para minimizar el estrés del paciente.",
            "Identificar y reportar al médico veterinario anomalías cutáneas, parasitarias o de salud general detectadas durante el servicio.",
            "Mantener limpio y desinfectado el área de peluquería, bañeras, mesas y herramientas conforme a NT-01-2008.",
            "Llevar registro de servicios prestados y ficha de cada paciente.",
            "Velar por el buen estado de máquinas, tijeras, secadoras y demás herramientas de trabajo.",
            "Aplicar las normas de bioseguridad y uso de EPP (guantes, bata, calzado antideslizante).",
            "Atender y orientar al cliente sobre el cuidado estético y sanitario de su mascota.",
        ],
        "incluye_historias": True,
    },
    {
        "code": "02g",
        "file": "02g_Contrato_Asistente_Ventas.docx",
        "cargo": "ASISTENTE DE VENTAS",
        "titulo": "CONTRATO INDIVIDUAL DE TRABAJO — ASISTENTE DE VENTAS",
        "doc_label": "CONTRATO INDIVIDUAL DE TRABAJO",
        "doc_sublabel": "Asistente de Ventas (Tienda)",
        "footer_label": "Contrato Asistente de Ventas v3.0",
        "duracion": "3 MESES + PRÓRROGA",
        "duracion_detalle": (
            "Conforme a los artículos 64 y 65 de la LOTTT, el presente contrato se celebra "
            "por tiempo determinado de tres (3) meses, en razón de la naturaleza del cargo "
            "de asistente de ventas. El contrato podrá prorrogarse por períodos iguales "
            "hasta por el límite legal; vencido este plazo, la relación se considerará "
            "por tiempo indeterminado."
        ),
        "prueba": "15 días continuos",
        "prueba_detalle": (
            "Conforme al artículo 22 de la LOTTT, se establece un período de prueba de "
            "quince (15) días continuos, durante el cual cualquiera de las partes "
            "podrá dar por terminada la relación laboral sin necesidad de preaviso ni "
            "indemnización."
        ),
        "funciones": [
            "Atender al cliente en la tienda de mascotas, asesorando sobre productos y servicios.",
            "Operar la caja registradora, facturar y realizar arqueos diarios.",
            "Apoyar en la recepción, almacenamiento y exhibición de mercancía.",
            "Verificar fechas de vencimiento y rotación de productos en percha.",
            "Mantener el orden, limpieza y exhibición de la tienda.",
            "Informar al encargado de tienda sobre faltantes, mermas o devoluciones.",
            "Promocionar productos y servicios de la clínica, tienda y peluquería.",
            "Cumplir las normas LOPCYMAT y LOPDP aplicables a su puesto.",
        ],
        "incluye_historias": False,
    },
]


# ============================================================
# Helpers específicos de contratos
# ============================================================
def add_party_intro(doc):
    """Bloque introductorio de partes del contrato."""
    add_para(doc,
        "Entre los suscritos: GRUPO CAVAL 1003, C.A., sociedad mercantil inscrita ante el "
        "Registro Mercantil correspondiente, bajo el N° de RIF J501662533, en adelante "
        "denominada \u00abLA EMPRESA\u00bb, representada en este acto por su Directora Gerente, "
        "ciudadana ESNATLIM ELENA SIMOZA, titular de la cédula de identidad N° "
        "V-17.976.287, por una parte; y por la otra, el ciudadano(a) "
        "_____________________________________, titular de la cédula de identidad N° "
        "V-___________________, en adelante \u00abEL(LA) TRABAJADOR(A)\u00bb, quienes declaran "
        "ser mayores de edad y hábiles para contratar, han convenido en celebrar el "
        "presente Contrato Individual de Trabajo, el cual se regirá por las siguientes "
        "cláusulas y por la Ley Orgánica del Trabajo, los Trabajadores y las Trabajadoras "
        "(LOTTT) y demás normas aplicables:",
        size=10, space_after=6)


def add_clause_header(doc, number, title):
    """Encabezado de cláusula: 'CLÁUSULA N° X — TÍTULO'."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run(f"CLÁUSULA {number}.  ")
    style_run(r1, size=10, bold=True, color=TEAL_DARK)
    r2 = p.add_run(title)
    style_run(r2, size=10, bold=True, color=BLACK)
    add_hr(p, color="CBD5E1", sz="4")


def add_jornada_table(doc):
    """Tabla de jornada y horario de trabajo."""
    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(10.6)
    rows = [
        ("Jornada diurna", "Lunes a sábado, de 8:00 a.m. a 4:00 p.m. (con 1 hora de descanso)"),
        ("Jornada mixta", "Lunes a sábado, de 12:00 m. a 8:00 p.m. (con 1 hora de descanso)"),
        ("Horario asignado", "El horario específico será fijado por LA EMPRESA según las necesidades del servicio, conforme al artículo 171 LOTTT."),
    ]
    for i, (k, v) in enumerate(rows):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(6); c1.width = Cm(10.6)
        write_cell(c0, k, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(c1, v, size=9, color=BLACK, bg=(GRAY_ALT if i % 2 else WHITE))
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def add_salario_table(doc):
    """Tabla resumen del salario y conceptos salariales/no salariales."""
    from docx.shared import RGBColor as _RC
    WHITE = _RC(0xFF, 0xFF, 0xFF)
    headers = ["Concepto", "Monto mensual", "Naturaleza"]
    rows = [
        ("Salario básico",          SALARIO_BASICO,   "Salarial (base prestaciones)"),
        ("Bono de alimentación (Cestaticket)", CESTATICKET, "No salarial (Ley Programa Alimentación Trabajador)"),
        ("Bono de transporte",      BONO_TRANSPORTO,  "No salarial"),
    ]
    tbl = doc.add_table(rows=1+len(rows), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [6.0, 4.6, 6.0]
    for i, w in enumerate(widths):
        tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.width = Cm(widths[ci])
        write_cell(c, h, size=9, bold=True, color=WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE_BG
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            c.width = Cm(widths[ci])
            bold = (ci == 0)
            write_cell(c, val, size=9, bold=bold, color=BLACK, bg=bg,
                       align=WD_ALIGN_PARAGRAPH.LEFT if ci != 1 else WD_ALIGN_PARAGRAPH.CENTER)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


# Color blanco reutilizable para encabezados de tablas
def _white():
    from docx.shared import RGBColor as _RC
    return _RC(0xFF, 0xFF, 0xFF)


# Fondo blanco para filas alternas en tablas (alias local de WHITE string)
WHITE_BG = "FFFFFF"


def add_penalidades_table(doc):
    """Tabla simplificada de penalidades — 3 niveles."""
    WHITE_C = _white()
    headers = ["Nivel", "Tipo de falta", "Medida disciplinaria"]
    rows = [
        ("1", "LEVE — Incumplimientos menores de procedimiento, impuntualidad ocasional, omisiones leves.",
         "Amonestación VERBAL por el jefe inmediato. Se deja constancia en el expediente del trabajador."),
        ("2", "GRAVE — Reincidencia de faltas leves, incumplimiento de protocolos de bioseguridad o atención al cliente, descuido de inventario.",
         "Amonestación ESCRITA firmada por el trabajador y la Directora Gerente. Se archiva en expediente."),
        ("3", "MUY GRAVE — Faltas del artículo 79 LOTTT: robo, abandono del trabajo, indisciplina, acoso, revelación de secretos, daño intencional a bienes de la empresa o de pacientes.",
         "RESCISIÓN del contrato por causa justificada conforme al artículo 79 LOTTT, sin derecho a prestaciones adicionales a las legalmente causadas."),
    ]
    tbl = doc.add_table(rows=1+len(rows), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [1.4, 7.6, 7.6]
    for i, w in enumerate(widths):
        tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.width = Cm(widths[ci])
        write_cell(c, h, size=9, bold=True, color=WHITE_C,
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE_BG
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            c.width = Cm(widths[ci])
            color_txt = BLACK
            bold = False
            if ci == 0:
                if "1" in val:
                    color_txt = GREEN_OK; bold = True
                elif "2" in val:
                    color_txt = TEAL_DARK; bold = True
                elif "3" in val:
                    color_txt = RED_CRIT; bold = True
            elif ci == 1 and "MUY GRAVE" in val:
                color_txt = RED_CRIT; bold = True
            elif ci == 2 and "RESCISIÓN" in val:
                color_txt = RED_CRIT; bold = True
            write_cell(c, val, size=9, bold=bold, color=color_txt, bg=bg,
                       align=WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    # Nota legal bajo la tabla
    add_para(doc,
        "Nota: ningún descuento salarial se aplicará sin la autorización previa del Inspector del "
        "Trabajo, conforme al artículo 59 de la LOTTT. La vía disciplinaria no sustituye las "
        "acciones civiles o penales que correspondan.",
        size=8, italic=True, color=GRAY_TEXT, space_after=6)


def add_clause_historias(doc):
    """Cláusula adicional — Propiedad de historias clínicas y no competencia."""
    add_clause_header(doc, "8", "PROPIEDAD DE HISTORIAS CLÍNICAS, CONFIDENCIALIDAD Y NO COMPETENCIA")
    add_para(doc,
        "EL(LA) TRABAJADOR(A) reconoce y acepta que las historias clínicas, fichas de "
        "pacientes, registros médicos, protocolos clínicos y quirúrgicos, recetas y "
        "cualesquiera otros documentos generados en el ejercicio de sus funciones, son "
        "propiedad exclusiva de LA EMPRESA, conforme al artículo 23 de la Ley de Ejercicio "
        "de la Medicina Veterinaria y al Reglamento Interno. En consecuencia:",
        size=10, space_after=4)
    add_bullet(doc,
        "Las historias clínicas y fichas de pacientes no podrán ser sustraídas, copiadas, "
        "fotografiadas, duplicadas o entregadas a terceros sin autorización escrita de LA EMPRESA.",
        bold_lead="8.1  ")
    add_bullet(doc,
        "La información clínica de los pacientes y los datos de los clientes son confidenciales "
        "y están protegidos por la LOPDP; su divulgación injustificada se considerará falta "
        "grave conforme a la cláusula de penalidades.",
        bold_lead="8.2  ")
    add_bullet(doc,
        "EL(LA) TRABAJADOR(A) se obliga a no prestar servicios profesionales, remunerados o "
        "no, a la clientela de LA EMPRESA durante los doce (12) meses siguientes a la "
        "terminación de la relación laboral dentro del ámbito territorial del Estado Miranda, "
        "cuando dichos servicios compitan con la actividad de LA EMPRESA. Esta cláusula se "
        "limita estrictamente a la clientela atendida durante el último año de la relación.",
        bold_lead="8.3  ")
    add_bullet(doc,
        "El incumplimiento de cualquiera de las obligaciones de esta cláusula dará derecho a "
        "LA EMPRESA a exigir la indemnización de daños y perjuicios correspondiente, sin "
        "perjuicio de las acciones civiles y penales a que hubiere lugar.",
        bold_lead="8.4  ")


# ============================================================
# Generador de contrato individual
# ============================================================
def gen_contrato(role, out_dir):
    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)

    add_membrete(doc,
                 doc_label=role["doc_label"],
                 doc_sublabel=role["doc_sublabel"],
                 version="Versión 3.0  ·  RR.HH.")

    add_doc_title(doc, role["titulo"])

    # Preámbulo de partes
    add_party_intro(doc)

    # ===== CLÁUSULA 1 — PARTES =====
    add_clause_header(doc, "1", "PARTES")
    add_para(doc,
        f"LA EMPRESA: {EMPRESA}, sociedad mercantil de domicilio en {DOMICILIO_EMP}, "
        f"RIF N° {RIF_EMP}, marca comercial {MARCA}, dedicada a la actividad de clínica "
        f"veterinaria, tienda de mascotas y peluquería canina. Representada en este acto "
        f"por su Directora Gerente, ciudadana {DIRECTORA_NOMBRE}, titular de la cédula de "
        f"identidad N° {DIRECTORA_CI}.",
        size=10, space_after=4)
    add_para(doc,
        "EL(LA) TRABAJADOR(A): _______________________________________, venezolano(a), "
        "mayor de edad, de estado civil __________, titular de la cédula de identidad N° "
        "V-___________________, domiciliado(a) en _________________________________, "
        "Estado ____________.",
        size=10, space_after=6)

    # ===== CLÁUSULA 2 — OBJETO Y CARGO =====
    add_clause_header(doc, "2", "OBJETO DEL CONTRATO Y CARGO")
    add_para(doc,
        f"LA EMPRESA contrata los servicios personales de EL(LA) TRABAJADOR(A) para "
        f"desempeñar el cargo de {role['cargo']}, quien acepta y se obliga a prestar sus "
        f"servicios de manera personal, bajo subordinación y dependencia, conforme a las "
        f"instrucciones y directrices que imparta LA EMPRESA o sus representantes. Las "
        f"funciones específicas del cargo son las siguientes:",
        size=10, space_after=4)
    for fn in role["funciones"]:
        add_bullet(doc, fn, size=10)
    add_para(doc,
        "LA EMPRESA se reserva el derecho de reasignar funciones complementarias acordes "
        "con el cargo, así como de trasladar al trabajador entre las áreas de clínica, "
        "tienda y peluquería cuando lo requieran las necesidades del servicio, sin que "
        "ello implique modificación del salario acordado.",
        size=10, space_after=6)

    # ===== CLÁUSULA 3 — DURACIÓN =====
    add_clause_header(doc, "3", f"DURACIÓN: {role['duracion']}")
    add_para(doc, role["duracion_detalle"], size=10, space_after=4)
    add_para(doc, role["prueba_detalle"], size=10, space_after=6)

    # ===== CLÁUSULA 4 — JORNADA Y HORARIO =====
    add_clause_header(doc, "4", "JORNADA Y HORARIO DE TRABAJO")
    add_para(doc,
        "La jornada ordinaria de trabajo será la prevista en el artículo 173 de la LOTTT "
        "(jornada diurna, mixta o nocturna, según corresponda), con un (1) día de descanso "
        "semanal, conforme al artículo 184 de la LOTTT:",
        size=10, space_after=4)
    add_jornada_table(doc)
    add_para(doc,
        "Las horas extraordinarias se pagarán con el recargo del 50 % (diurna) o del 95 % "
        "(nocturna) sobre el salario normal, conforme a los artículos 198 y 199 de la LOTTT, "
        "y solo procederán cuando sean expresamente autorizadas por LA EMPRESA.",
        size=10, space_after=6)

    # ===== CLÁUSULA 5 — SALARIO Y PRESTACIONES =====
    add_clause_header(doc, "5", "SALARIO, BONOS Y PRESTACIONES SOCIALES")
    add_para(doc,
        f"EL(LA) TRABAJADOR(A) devengará un salario básico mensual de {SALARIO_BASICO} "
        f"(Salario Mínimo Nacional vigente conforme a la Gaceta Oficial N° 7.026 "
        f"Extraordinaria de fecha 28 de abril de 2025). Adicionalmente, recibirá los "
        f"siguientes conceptos:",
        size=10, space_after=4)
    add_salario_table(doc)
    add_para(doc,
        "Las prestaciones sociales se calcularán conforme a los artículos 142 y 143 de la "
        "LOTTT, sobre la base del salario normal devengado. LA EMPRESA depositará las "
        "prestaciones sociales mensualmente en la cuenta individual del trabajador, dentro "
        "de los primeros cinco (5) días hábiles del mes siguiente, conforme al artículo 143 "
        "de la LOTTT. Igualmente, LA EMPRESA cotizará al IVSS, FAOV, INCES y demás "
        "organismos de seguridad social conforme a la ley.",
        size=10, space_after=6)

    # ===== CLÁUSULA 6 — OBLIGACIONES Y PROHIBICIONES =====
    add_clause_header(doc, "6", "OBLIGACIONES Y PROHIBICIONES")
    add_para(doc, "Son obligaciones de EL(LA) TRABAJADOR(A):", size=10, bold=True, space_after=2)
    obl = [
        "Cumplir con las órdenes e instrucciones de LA EMPRESA en todo lo concerniente al trabajo.",
        "Concurrir al trabajo en el horario establecido y permanecer en él durante la jornada.",
        "Observar las normas de bioseguridad NT-01-2008 y vigilancia médica NT-02-2008, así como el uso obligatorio del EPP asignado.",
        "Mantener en buen estado los equipos, herramientas, mobiliario e instalaciones de la empresa.",
        "Guardar secreto sobre la información técnica, comercial y administrativa de LA EMPRESA.",
        "Participar en los programas de capacitación y formación que LA EMPRESA establezca.",
        "Cumplir el Reglamento Interno y el Código de Conducta de la empresa.",
    ]
    for o in obl:
        add_bullet(doc, o, size=10)
    add_para(doc, "Son prohibiciones para EL(LA) TRABAJADOR(A):", size=10, bold=True,
             space_before=4, space_after=2)
    proh = [
        "Trabajar en estado de embriaguez o bajo influencia de sustancias estupefacientes o psicotrópicas.",
        "Sustraer, ocultar o dañar herramientas, mercancía, equipo, pacientes o documentos de LA EMPRESA.",
        "Abandonar el trabajo durante la jornada sin causa justificada y sin autorización del superior.",
        "Ejecutar labores por cuenta propia o de terceros dentro del horario laboral.",
        "Divulgar información confidencial, datos de clientes o historias clínicas de pacientes.",
    ]
    for p in proh:
        add_bullet(doc, p, size=10, color=RED_CRIT)
    add_para(doc, "", size=4, space_after=2)

    # ===== CLÁUSULA 7 — PENALIDADES =====
    add_clause_header(doc, "7", "PENALIDADES Y RÉGIMEN DISCIPLINARIO")
    add_para(doc,
        "El incumplimiento de las obligaciones aquí previstas se sancionará conforme al "
        "siguiente régimen disciplinario simplificado de tres (3) niveles, sin perjuicio de "
        "las previsiones del Reglamento Interno y de las causales de despido justificado "
        "del artículo 79 de la LOTTT:",
        size=10, space_after=4)
    add_penalidades_table(doc)

    # ===== CLÁUSULA 8 (solo Vet/Aux/Groomer) — HISTORIAS CLÍNICAS =====
    if role["incluye_historias"]:
        add_clause_historias(doc)
        next_clause = 9
    else:
        next_clause = 8

    # ===== CLÁUSULA FINAL — LOPDP =====
    add_clause_header(doc, str(next_clause), "PROTECCIÓN DE DATOS PERSONALES (LOPDP)")
    add_para(doc,
        "Conforme a la Ley Orgánica de Protección de Datos Personales (Decreto 1.419, "
        "Gaceta 6.210 Extraordinaria, 2014), EL(LA) TRABAJADOR(A) autoriza a LA EMPRESA "
        "para el tratamiento de sus datos personales (nombre, cédula, dirección, teléfono, "
        "datos biométricos, datos de salud) con finalidades de gestión laboral, nómina, "
        "cumplimiento de obligaciones ante IVSS, FAOV, INCES, PMSSO y demás organismos, así "
        "como para el archivo laboral previsto en el artículo 183 de la LOTTT. Esta "
        "autorización se mantendrá vigente durante la relación laboral y hasta por diez (10) "
        "años después de su terminación. EL(LA) TRABAJADOR(A) podrá ejercer sus derechos "
        "ARCO+ (Acceso, Rectificación, Cancelación, Oposición, Revocación, Información) "
        "mediante solicitud escrita dirigida a la Dirección de Gestión Humana.",
        size=10, space_after=6)

    # ===== CLÁUSULA FINAL+1 — TERMINACIÓN Y DOMICILIO =====
    final_clause = next_clause + 1
    add_clause_header(doc, str(final_clause), "TERMINACIÓN, LITIGIOS Y DOMICILIO")
    add_para(doc,
        "El presente contrato terminará por las causales previstas en los artículos 71, 72 "
        "y 79 de la LOTTT. Las controversias que se susciten con motivo de la interpretación "
        "o ejecución del presente contrato serán sustanciadas ante la Inspectoría del "
        "Trabajo con competencia en el Estado Miranda, o ante la jurisdicción laboral "
        "ordinaria. Para todos los efectos legales, las partes eligen como domicilio "
        f"especial y procesal la ciudad de {DOMICILIO_PROCESAL}, a cuyos tribunales se "
        "someten expresamente.",
        size=10, space_after=6)

    # ===== CIERRE — Lugar, fecha y firmas =====
    add_para(doc,
        f"Se hacen dos (2) ejemplares de un mismo tenor y a un solo efecto, en la ciudad "
        f"de {DOMICILIO_PROCESAL}, a los ____ días del mes de ________________ de ________.",
        size=10, space_before=6, space_after=10)

    add_section(doc, "FIRMAS DE LAS PARTES")
    add_signature_block(doc, ["LA EMPRESA", "EL(LA) TRABAJADOR(A)"])

    # Footer
    add_footer(section, role["footer_label"])

    # Guardar
    out_path = os.path.join(out_dir, role["file"])
    doc.save(out_path)
    return out_path


# ============================================================
# Main
# ============================================================
def main():
    out_dir = os.path.join(HERE, "02_CONTRATOS")
    os.makedirs(out_dir, exist_ok=True)
    print(f"[gen_contratos] Generando 7 contratos en: {out_dir}\n")
    for role in ROLES:
        try:
            path = gen_contrato(role, out_dir)
            size_kb = os.path.getsize(path) / 1024
            print(f"  OK  {role['code']}  {role['file']:42s}  {size_kb:5.1f} KB")
        except Exception as e:
            print(f"  ERROR {role['code']}  {role['file']}  →  {e}")
            raise
    print("\n[gen_contratos] Finalizado.")


if __name__ == "__main__":
    main()
