"""
gen_discriminacion_remuneracion.py — Documento de política de remuneraciones
que discrimina un sueldo total de $250/mes en:
  - $50  Salario base (salario — incidencia en prestaciones)
  - $80  Cestaticket / Bono de Alimentación (salario en Venezuela — pagos quincenales $40)
  - $40  Bono de Buen Vivir (no salarial — bono extraordinario)
  - $40  Bono de Transporte (no salarial — reembolso de gastos)
  - $40  Otros beneficios no salariales (no salarial)

Incluye justificación legal de cada concepto conforme a LOTTT, Ley de Alimentación,
jurisprudencia TSJ y política interna.
"""
import os, sys
sys.path.insert(0, "/home/z/my-project/output")
from _common import *
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

OUT = "/home/z/my-project/output/Discriminacion_Remuneraciones.docx"

def build():
    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "POLÍTICA DE REMUNERACIONES", "Discriminación de sueldo",
                 version="Versión 1.0  ·  Dirección")
    add_doc_title(doc, "POLÍTICA DE REMUNERACIONES Y DISCRIMINACIÓN SALARIAL")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Conforme a LOTTT arts. 104, 142-146 · Ley de Alimentación · Jurisprudencia TSJ")
    style_run(r, size=10, italic=True, color=GRAY_TEXT)

    # ============================================================
    # 1. DATOS DEL TRABAJADOR Y RESUMEN EJECUTIVO
    # ============================================================
    add_section(doc, "1. DATOS DEL TRABAJADOR(A) Y RESUMEN")
    tbl = doc.add_table(rows=3, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(3.5), Cm(5.0), Cm(3.5), Cm(5.0)]):
        tbl.columns[i].width = w
    datos = [
        ("Nombre del trabajador:", "_____________________", "Cédula de Identidad:", "V-____________"),
        ("Cargo:", "_____________________", "Fecha de ingreso:", "____/____/______"),
        ("Total remuneración mensual:", "USD 250,00", "Salario base (para prestaciones):", "USD 50,00"),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(datos):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        c = tbl.rows[ri].cells[0]; c.width = Cm(3.5)
        write_cell(c, l1, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[1]; c.width = Cm(5.0)
        write_cell(c, v1, size=9, bold=(ri==2), color=(RED_CRIT if ri==2 else BLACK), bg=bg)
        c = tbl.rows[ri].cells[2]; c.width = Cm(3.5)
        write_cell(c, l2, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[3]; c.width = Cm(5.0)
        write_cell(c, v2, size=9, bold=(ri==2), color=(GREEN_OK if ri==2 else BLACK), bg=bg)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

    # ============================================================
    # 2. RESUMEN EJECUTIVO
    # ============================================================
    add_section(doc, "2. RESUMEN EJECUTIVO")
    add_para(doc,
        "La presente política discrimina la remuneración mensual total del trabajador en "
        "cinco (5) conceptos, con el objetivo de clarificar cuáles constituyen salario a "
        "efectos del cálculo de prestaciones sociales y demás beneficios laborales, y cuáles "
        "son beneficios no salariales conforme a la Ley Orgánica del Trabajo, los Trabajadores "
        "y las Trabajadoras (LOTTT), la Ley de Alimentación para los Trabajadores y la "
        "jurisprudencia vinculante del Tribunal Supremo de Justicia (TSJ).")

    add_para(doc,
        "Total remuneración mensual: USD 250,00  ·  Salario base para prestaciones: USD 50,00",
        size=11, bold=True, color=TEAL_DARK, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=4, space_after=4)
    add_para(doc,
        "Diferencia no salarial: USD 200,00 (80% del total) — justificada conforme al artículo 104 de la LOTTT.",
        size=10, italic=True, color=GRAY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # ============================================================
    # 3. TABLA DE DISCRIMINACIÓN DETALLADA
    # ============================================================
    add_section(doc, "3. DISCRIMINACIÓN DETALLADA DE LA REMUNERACIÓN MENSUAL")

    tbl = doc.add_table(rows=7, cols=7)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    col_widths = [Cm(0.6), Cm(4.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(2.5), Cm(4.4)]
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = w

    headers = ["#", "CONCEPTO", "MONTO USD", "MENSUAL", "QUINCENAL", "CARÁCTER", "FUNDAMENTO LEGAL"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = col_widths[ci]
        write_cell(c, h, size=8, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)

    conceptos = [
        ("1", "Salario Base",
         "50,00", "50,00", "25,00",
         "SALARIAL",
         "LOTTT Art. 104 — Salario base ordinario. Incide en prestaciones, utilidades, vacaciones, bono vacacional."),
        ("2", "Bono de Alimentación (Cestaticket)",
         "80,00", "80,00", "40,00",
         "SALARIAL",
         "Ley de Alimentación + TSJ sent. 2014 — Pagos quincenales de USD 40,00. En Venezuela se considera salario a todos los efectos (TSJ vinculante)."),
        ("3", "Bono de Buen Vivir",
         "40,00", "40,00", "20,00",
         "NO SALARIAL",
         "LOTTT Art. 104 (parágrafo) — Bono extraordinario ocasional como incentivo o apoyo. No forma parte del salario regular ni incide en prestaciones."),
        ("4", "Bono de Transporte",
         "40,00", "40,00", "20,00",
         "NO SALARIAL",
         "LOTTT Art. 104 — Reembolso de gastos de traslado al trabajo. No es contraprestación por servicios, sino compensación de gastos."),
        ("5", "Otros beneficios no salariales",
         "40,00", "40,00", "20,00",
         "NO SALARIAL",
         "LOTTT Art. 104 — Beneficios sociales asistenciales (ayudas, bonos especiales, asignaciones eventuales). No son salario."),
    ]

    for ri, (n, con, monto, mens, quinc, caract, fund) in enumerate(conceptos, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        # Número
        c = tbl.rows[ri].cells[0]; c.width = col_widths[0]
        write_cell(c, n, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        # Concepto
        c = tbl.rows[ri].cells[1]; c.width = col_widths[1]
        write_cell(c, con, size=8, bold=True, color=TEAL_DARK, bg=bg)
        # Monto USD
        c = tbl.rows[ri].cells[2]; c.width = col_widths[2]
        write_cell(c, "$ " + monto, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        # Mensual
        c = tbl.rows[ri].cells[3]; c.width = col_widths[3]
        write_cell(c, "$ " + mens, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        # Quincenal
        c = tbl.rows[ri].cells[4]; c.width = col_widths[4]
        write_cell(c, "$ " + quinc, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        # Carácter (salarial/no salarial)
        c = tbl.rows[ri].cells[5]; c.width = col_widths[5]
        if "SALARIAL" in caract and "NO" not in caract:
            write_cell(c, "SALARIAL", size=8, bold=True, color=RED_CRIT,
                       align=WD_ALIGN_PARAGRAPH.CENTER, bg="FEE2E2")
        else:
            write_cell(c, "NO SALARIAL", size=8, bold=True, color=GREEN_OK,
                       align=WD_ALIGN_PARAGRAPH.CENTER, bg="DCFCE7")
        # Fundamento
        c = tbl.rows[ri].cells[6]; c.width = col_widths[6]
        write_cell(c, fund, size=7, italic=True, color=GRAY_TEXT, bg=bg)

    # Fila TOTAL
    c = tbl.rows[6].cells[0]; c.width = col_widths[0]
    write_cell(c, "", size=8, bg=TEAL_HDR_BG)
    set_cell_bg(c, TEAL_HDR_BG)
    c = tbl.rows[6].cells[1]; c.width = col_widths[1]
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("TOTAL REMUNERACIÓN MENSUAL")
    style_run(r, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[2]; c.width = col_widths[2]
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 250,00")
    style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[3]; c.width = col_widths[3]
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 250,00")
    style_run(r, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[4]; c.width = col_widths[4]
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 125,00")
    style_run(r, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[5]; c.width = col_widths[5]
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 50 SAL + $ 200 NS")
    style_run(r, size=7, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[6]; c.width = col_widths[6]
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Salario = USD 50,00  ·  No salarial = USD 200,00")
    style_run(r, size=7, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_para(doc,
        "USD 50,00 SALARIAL (incide en prestaciones) + USD 200,00 NO SALARIAL (NO incide en prestaciones).",
        size=10, bold=True, color=TEAL_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # ============================================================
    # 4. DETALLE Y JUSTIFICACIÓN DE CADA CONCEPTO
    # ============================================================
    add_section(doc, "4. DETALLE Y JUSTIFICACIÓN LEGAL DE CADA CONCEPTO")

    # 4.1 Salario Base
    add_para(doc, "4.1  Salario Base — USD 50,00 mensuales (USD 25,00 quincenales)",
             size=11, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_para(doc,
        "Concepto: Remuneración básica ordinaria pagada al trabajador por la prestación de "
        "sus servicios, conforme al artículo 104 de la LOTTT. Se paga quincenalmente (50% "
        "cada quincena = USD 25,00 cada 15 días), conforme al artículo 169 de la LOTTT.",
        size=10, space_after=3)
    add_para(doc,
        "Carácter: SALARIAL. Forma parte del salario base a todos los efectos legales: "
        "cálculo de prestaciones sociales (Art. 142-146 LOTTT), utilidades (Art. 174), "
        "vacaciones y bono vacacional (Art. 192 y 196), indemnizaciones por despido "
        "injustificado (Art. 125), y demás beneficios laborales.",
        size=10, space_after=3)
    add_para(doc,
        "Justificación: El salario base se fija en USD 50,00 conforme al contrato individual "
        "de trabajo. Aunque el Ejecutivo Nacional fijó el salario mínimo en Bs. 210 mensuales "
        "(Gaceta 7.026 Extraordinaria del 28/04/2025), las partes convienen una remuneración "
        "superior en divisas, equivalente a USD 50,00 mensuales al tipo de cambio BCV vigente "
        "al momento del pago. Este monto cumple y excede el mínimo legal.",
        size=10, space_after=6)

    # 4.2 Cestaticket
    add_para(doc, "4.2  Bono de Alimentación (Cestaticket) — USD 80,00 mensuales (USD 40,00 quincenales)",
             size=11, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_para(doc,
        "Concepto: Beneficio de alimentación previsto en la Ley de Alimentación para los "
        "Trabajadores y las Trabajadoras. Se paga mediante tickets, tarjetas electrónicas o "
        "en efectivo, entregados quincenalmente (USD 40,00 cada 15 días).",
        size=10, space_after=3)
    add_para(doc,
        "Carácter: SALARIAL. Conforme a la sentencia vinculante N° 0414 del 22/05/2014 de la "
        "Sala Social del Tribunal Supremo de Justicia, el cestaticket PAGADO EN EFECTIVO o "
        "MEDIANTE INSTRUMENTO NEGOCIABLE constituye salario a todos los efectos legales. "
        "Por tanto, INCIDE en el cálculo de prestaciones sociales, utilidades, vacaciones, "
        "bono vacacional e indemnizaciones.",
        size=10, space_after=3)
    add_para(doc,
        "Justificación: El monto de USD 80,00 mensuales se fija de común acuerdo entre las "
        "partes y se paga quincenalmente (USD 40,00 cada quincena) por comodidad operativa y "
        "para acompañar el pago del salario base. La empresa opta por pago en efectivo/transferencia "
        "para facilitar al trabajador la disposición del beneficio.",
        size=10, space_after=6)

    # 4.3 Bono de Buen Vivir
    add_para(doc, "4.3  Bono de Buen Vivir — USD 40,00 mensuales (USD 20,00 quincenales)",
             size=11, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_para(doc,
        "Concepto: Bono extraordinario de carácter asistencial, otorgado por la empresa como "
        "incentivo y apoyo al trabajador para contribuir a su bienestar integral. Se paga "
        "quincenalmente (USD 20,00 cada 15 días).",
        size=10, space_after=3)
    add_para(doc,
        "Carácter: NO SALARIAL. Conforme al artículo 104 (parágrafo) de la LOTTT, los bonos "
        "extraordinarios de carácter ocasional, otorgados por la empresa como incentivo o "
        "apoyo en circunstancias específicas, NO se consideran parte del salario regular y, "
        "por ende, NO afectan el cálculo de las prestaciones sociales ni demás beneficios.",
        size=10, space_after=3)
    add_para(doc,
        "Justificación: Este bono NO constituye contraprestación por los servicios prestados, "
        "sino un beneficio social-asistencial otorgado discrecionalmente por la empresa para "
        "contribuir al bienestar del trabajador y su familia. Su pago está condicionado a la "
        "permanencia del trabajador en la empresa y a la disposición presupuestaria. La empresa "
        "se reserva el derecho de modificar o suspender este beneficio, lo cual refuerza su "
        "carácter no salarial.",
        size=10, space_after=6)

    # 4.4 Bono de Transporte
    add_para(doc, "4.4  Bono de Transporte — USD 40,00 mensuales (USD 20,00 quincenales)",
             size=11, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_para(doc,
        "Concepto: Asistencia económica proporcionada por la empresa al trabajador como "
        "reembolso de los gastos de traslado al lugar de trabajo. Se paga quincenalmente "
        "(USD 20,00 cada 15 días).",
        size=10, space_after=3)
    add_para(doc,
        "Carácter: NO SALARIAL. Conforme al artículo 104 de la LOTTT, este bono se clasifica "
        "como reembolso de gastos y NO como salario por servicios prestados. No incide en el "
        "cálculo de prestaciones sociales ni demás beneficios laborales.",
        size=10, space_after=3)
    add_para(doc,
        "Justificación: El bono de transporte tiene por objeto cubrir el costo del traslado "
        "del trabajador desde su domicilio hasta el centro de trabajo y viceversa. Al no "
        "constituir contraprestación por servicios, sino compensación de un gasto necesario "
        "para prestar el servicio, no tiene incidencia salarial conforme a la doctrina "
        "y jurisprudencia venezolana.",
        size=10, space_after=6)

    # 4.5 Otros beneficios no salariales
    add_para(doc, "4.5  Otros beneficios no salariales — USD 40,00 mensuales (USD 20,00 quincenales)",
             size=11, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_para(doc,
        "Concepto: Conjunto de beneficios sociales-asistenciales otorgados por la empresa al "
        "trabajador, tales como: ayudas para alimentos, dotación de uniformes, beneficios "
        "sociales eventuales, incentivos por productividad, bonos especiales por fechas "
        "significativas, asistencia económica para situaciones específicas, entre otros. Se "
        "paga quincenalmente (USD 20,00 cada 15 días).",
        size=10, space_after=3)
    add_para(doc,
        "Carácter: NO SALARIAL. Conforme al artículo 104 de la LOTTT, los beneficios sociales "
        "asistenciales NO se consideran salario. Estos beneficios tienen naturaleza "
        "compensatoria o asistencial y no remuneran el trabajo prestado.",
        size=10, space_after=3)
    add_para(doc,
        "Justificación: Este concepto agrupa diversos beneficios otorgados por la empresa "
        "que, por su naturaleza asistencial, social o eventual, no constituyen contraprestación "
        "por servicios. Su pago está condicionado a la disposición de la empresa y al "
        "cumplimiento de ciertas condiciones por parte del trabajador, lo cual refuerza su "
        "carácter no salarial. La empresa se reserva el derecho de redistribuir este monto "
        "entre diferentes beneficios según las necesidades del trabajador.",
        size=10, space_after=6)

    # ============================================================
    # 5. CÁLCULO DE PRESTACIONES SOCIALES
    # ============================================================
    add_section(doc, "5. CÁLCULO DE PRESTACIONES SOCIALES")

    add_para(doc,
        "Conforme al artículo 142 de la LOTTT, las prestaciones sociales se calculan sobre "
        "el SALARIO base del trabajador. En este caso, solo los conceptos SALARIALES "
        "(Salario Base USD 50,00 + Cestaticket USD 80,00) forman parte del salario a efectos "
        "del cálculo de prestaciones. Los conceptos NO SALARIALES (USD 200,00) NO inciden.",
        size=10, space_after=6)

    # Tabla de cálculo
    tbl = doc.add_table(rows=5, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(5.0), Cm(3.5), Cm(4.0), Cm(4.5)]):
        tbl.columns[i].width = w
    headers = ["CONCEPTO", "BASE DE CÁLCULO", "FÓRMULA", "MONTO POR AÑO"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = [Cm(5.0), Cm(3.5), Cm(4.0), Cm(4.5)][ci]
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    rows = [
        ("Prestaciones de antigüedad (Art. 142-143)",
         "Salario diario = USD 130/30 = USD 4,33",
         "5 días/mes × 12 meses = 60 días/año",
         "USD 260,00 / año"),
        ("Utilidades (Art. 174)",
         "Salario diario = USD 4,33",
         "45 días/año (mínimo legal)",
         "USD 195,00 / año"),
        ("Vacaciones (Art. 192)",
         "Salario diario = USD 4,33",
         "15 días hábiles/año",
         "USD 65,00 / año"),
        ("Bono vacacional (Art. 196)",
         "Salario diario = USD 4,33",
         "7 días × (15/6) = 17,5 días/año",
         "USD 75,75 / año"),
    ]
    for ri, row in enumerate(rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]; c.width = [Cm(5.0), Cm(3.5), Cm(4.0), Cm(4.5)][ci]
            write_cell(c, val, size=9, bold=(ci==0), color=(TEAL_DARK if ci==0 else BLACK), bg=bg)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc,
        "Nota: La base de cálculo (USD 130,00 = salario base USD 50,00 + cestaticket USD 80,00) "
        "es el SALARIO a todos los efectos legales. La diferencia (USD 200,00 mensuales) NO "
        "se toma en cuenta para prestaciones por ser beneficios no salariales.",
        size=9, italic=True, color=GRAY_TEXT, space_after=8)

    # ============================================================
    # 6. FORMA DE PAGO
    # ============================================================
    add_section(doc, "6. FORMA Y OPORTUNIDAD DE PAGO")
    add_para(doc,
        "El pago de la remuneración se realizará QUINCENALMENTE, los días 15 y último de "
        "cada mes, mediante transferencia bancaria a la cuenta del trabajador, conforme al "
        "artículo 169 de la LOTTT. El detalle del pago quincenal es:",
        size=10, space_after=4)

    tbl = doc.add_table(rows=7, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(6.0), Cm(3.0), Cm(3.5), Cm(4.5)]):
        tbl.columns[i].width = w
    headers = ["CONCEPTO", "MENSUAL", "QUINCENAL", "FORMA DE PAGO"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = [Cm(6.0), Cm(3.0), Cm(3.5), Cm(4.5)][ci]
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    pagos = [
        ("Salario Base", "$ 50,00", "$ 25,00", "Transferencia bancaria"),
        ("Bono de Alimentación (Cestaticket)", "$ 80,00", "$ 40,00", "Transferencia bancaria"),
        ("Bono de Buen Vivir", "$ 40,00", "$ 20,00", "Transferencia bancaria"),
        ("Bono de Transporte", "$ 40,00", "$ 20,00", "Transferencia bancaria"),
        ("Otros beneficios no salariales", "$ 40,00", "$ 20,00", "Transferencia bancaria"),
    ]
    for ri, (con, mens, quinc, forma) in enumerate(pagos, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        c = tbl.rows[ri].cells[0]; c.width = Cm(6.0)
        write_cell(c, con, size=9, bg=bg)
        c = tbl.rows[ri].cells[1]; c.width = Cm(3.0)
        write_cell(c, mens, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[2]; c.width = Cm(3.5)
        write_cell(c, quinc, size=9, bold=True, color=TEAL_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[3]; c.width = Cm(4.5)
        write_cell(c, forma, size=8, italic=True, color=GRAY_TEXT, bg=bg)
    # Total
    c = tbl.rows[6].cells[0]; c.width = Cm(6.0)
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run("TOTAL")
    style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[1]; c.width = Cm(3.0)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 250,00")
    style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[2]; c.width = Cm(3.5)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 125,00")
    style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[3]; c.width = Cm(4.5)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Transferencia bancaria")
    style_run(r, size=9, italic=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)

    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # ============================================================
    # 7. MODIFICACIONES
    # ============================================================
    add_section(doc, "7. MODIFICACIONES Y ACTUALIZACIONES")
    add_para(doc,
        "La empresa se reserva el derecho de modificar los montos de los conceptos NO "
        "SALARIALES (Bono de Buen Vivir, Bono de Transporte y Otros beneficios no salariales) "
        "según sus posibilidades presupuestarias, sin que ello constituya alteración del "
        "salario base del trabajador. Cualquier modificación al SALARIO BASE (USD 50,00) o "
        "al CESTATICKET (USD 80,00) se realizará de común acuerdo entre las partes, conforme "
        "a lo establecido en el contrato individual de trabajo.",
        size=10, space_after=8)

    # ============================================================
    # 8. DECLARACIÓN Y FIRMAS
    # ============================================================
    add_section(doc, "8. DECLARACIÓN Y FIRMAS")
    add_para(doc,
        "El trabajador declara haber leído y comprendido la presente política de remuneraciones, "
        "entendiendo que su remuneración mensual total es de USD 250,00, discriminados en USD "
        "50,00 de salario base (salarial) y USD 200,00 de beneficios no salariales. Acepta que "
        "esta discriminación se realizará conforme a la normativa legal vigente y que las "
        "prestaciones sociales se calcularán sobre la base salarial (USD 130,00 mensuales = "
        "salario base + cestaticket).",
        size=10, space_after=6)
    add_para(doc,
        "En la ciudad de Los Teques, Estado Miranda, a los ____ días del mes de "
        "________________ de ________.", size=10, space_after=10)

    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])
    add_footer(section, "Política de Remuneraciones v1.0")

    doc.save(OUT)
    return OUT

if __name__ == "__main__":
    out = build()
    size_kb = os.path.getsize(out) / 1024
    print(f"✓ {out}")
    print(f"  Tamaño: {size_kb:.1f} KB")
