"""
gen_documentos_finales_v3.py — Genera los 7 documentos restantes del kit v3.0
con la firma correcta de DIRECTORA GERENTE (Esnatlim Elena Simoza).

Documentos:
  1. 00_Checklist_Maestro_Ingreso_Trabajador.docx
  2. 01_Solicitud_de_Empleo/Solicitud_de_Empleo_Ficha_Ingreso.docx (ya generada v3, reutilizar)
  3. 05_SEGURIDAD_LABORAL/Hoja_Recorrido_Habitual.docx
  4. 05_SEGURIDAD_LABORAL/Acta_Entrega_EPP.docx
  5. 05_SEGURIDAD_LABORAL/Examen_Medico_Pre_Empleo.docx
  6. 06_REGISTROS_LEGALES/Checklist_Constancias_IVSS_FAOV_INCES_PMSSO.docx
  7. 09_CIERRE/Carta_Aceptacion_General.docx
"""
import os, sys
sys.path.insert(0, "/home/z/my-project/output")
from _common import *
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 1. HOJA DE RECORRIDO HABITUAL DEL TRABAJADOR
# ============================================================
def gen_hoja_recorrido():
    out = "/home/z/my-project/output/05_SEGURIDAD_LABORAL/Hoja_Recorrido_Habitual.docx"
    os.makedirs(os.path.dirname(out), exist_ok=True)

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "SEGURIDAD LABORAL", "Recorrido habitual", version="Versión 3.0  ·  SST")
    add_doc_title(doc, "HOJA DE RECORRIDO HABITUAL DEL TRABAJADOR")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Conforme al Art. 69 LOPCYMAT (accidentes in itinere)")
    style_run(r, size=10, italic=True, color=GRAY_TEXT)

    add_section(doc, "OBJETIVO")
    add_para(doc,
        "Establecer y documentar el trayecto habitual del trabajador hacia y desde el "
        "centro de trabajo, cumpliendo con el artículo 69 de la LOPCYMAT, a fin de "
        "determinar el recorrido seguro que deberá ser utilizado para efectos de "
        "reconocimiento de accidentes de trabajo in itinere.")

    add_section(doc, "DATOS DEL TRABAJADOR(A)")
    tbl = doc.add_table(rows=3, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(3.5), Cm(5.0), Cm(3.5), Cm(5.0)]):
        tbl.columns[i].width = w
    datos = [
        ("Nombre completo:", "Cédula de Identidad:"),
        ("Cargo:", "Teléfono de contacto:"),
        ("Dirección de habitación:", "Fecha de actualización:"),
    ]
    for ri, (l1, l2) in enumerate(datos):
        c = tbl.rows[ri].cells[0]; c.width = Cm(3.5)
        write_cell(c, l1, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[1]; c.width = Cm(5.0)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*25)
        set_cell_borders(c); set_cell_margins(c)
        c = tbl.rows[ri].cells[2]; c.width = Cm(3.5)
        write_cell(c, l2, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[3]; c.width = Cm(5.0)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*25)
        set_cell_borders(c); set_cell_margins(c)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_section(doc, "DATOS DEL CENTRO DE TRABAJO")
    add_bullet(doc, "GRUPO CAVAL 1003, C.A. (ALIKA PETS).", bold_lead="Nombre de la empresa:  ")
    add_bullet(doc, "Av. Francisco de Miranda, Local N° 1, Sector Francisco de Miranda, Los Teques, Estado Miranda, Zona Postal 1201.", bold_lead="Dirección del centro de trabajo:  ")
    add_bullet(doc, "_____ minutos aproximadamente.", bold_lead="Tiempo aproximado de recorrido hacia el centro de trabajo:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "DETALLES DEL RECORRIDO HABITUAL")

    def add_text_block(label, n_lines=3):
        add_para(doc, label, size=10, bold=True, color=BLACK, space_after=2)
        c = doc.add_table(rows=1, cols=1)
        c.alignment = WD_TABLE_ALIGNMENT.CENTER
        c.columns[0].width = Cm(17)
        cc = c.rows[0].cells[0]
        cc.text = ""
        pp = cc.paragraphs[0]
        for _ in range(n_lines):
            pp2 = cc.add_paragraph()
            pp2.paragraph_format.space_after = Pt(14)
        set_cell_borders(cc)
        set_cell_margins(cc, top=200, bottom=200)
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_text_block("Ruta habitual hacia el centro de trabajo (describa el trayecto exacto desde su dirección de habitación hasta el centro de trabajo):")
    add_text_block("Rutas alternas hacia el centro de trabajo (en caso de imprevisto):")
    add_text_block("Medio de transporte habitual (vehículo propio, transporte público, bicicleta, caminata, etc.):", n_lines=1)

    add_section(doc, "DETALLES DEL RETORNO")
    add_text_block("Ruta habitual para el retorno al domicilio (describa el trayecto desde el centro de trabajo hasta su dirección de habitación):")
    add_para(doc, "Tiempo aproximado de recorrido para el retorno:", size=10, bold=True, space_after=2)
    add_para(doc, "_____ minutos aproximadamente.", size=10, space_after=4)
    add_text_block("Rutas alternas para el retorno al domicilio (en caso de imprevisto):")

    add_section(doc, "OBSERVACIONES IMPORTANTES")
    add_bullet(doc, "En caso de utilizar una ruta alterna por motivos imprevistos, el trabajador deberá notificar inmediatamente al gerente o supervisor de la empresa, indicando la razón para el cambio de recorrido.", bold_lead="1.  ")
    add_bullet(doc, "El trabajador estará obligado a justificar la necesidad de tomar una ruta distinta, en concordancia con lo establecido en el Artículo 69 de la LOPCYMAT. La no justificación del cambio puede conllevar a que el accidente no sea reconocido como accidente de trabajo.", bold_lead="2.  ")
    add_bullet(doc, "En caso de modificación de su dirección de habitación, el trabajador deberá notificarlo formalmente a la empresa dentro de los tres días hábiles siguientes para actualizar su recorrido habitual en el registro.", bold_lead="3.  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "FIRMAS")
    add_para(doc, "En la ciudad de Los Teques, Estado Miranda, a los ____ días del mes de ________________ de ________.",
             size=10, space_after=8)
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])
    add_footer(section, "Hoja de Recorrido Habitual v3.0")

    doc.save(out)
    return out


# ============================================================
# 2. ACTA DE ENTREGA DE UNIFORMES Y EPP
# ============================================================
def gen_acta_epp():
    out = "/home/z/my-project/output/05_SEGURIDAD_LABORAL/Acta_Entrega_EPP.docx"
    os.makedirs(os.path.dirname(out), exist_ok=True)

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "SEGURIDAD LABORAL", "Entrega de EPP", version="Versión 3.0  ·  SST")
    add_doc_title(doc, "ACTA DE ENTREGA DE UNIFORMES Y EQUIPOS DE PROTECCIÓN PERSONAL")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Conforme al Art. 53 LOPCYMAT (Numeral 4)")
    style_run(r, size=10, italic=True, color=GRAY_TEXT)

    add_para(doc,
        "Yo, _____________________________________________, titular de la cédula de identidad "
        "N° V-___________________, a los ____ días del mes de ________________ de ________, "
        "se procede a realizar la entrega de uniformes de trabajo y Equipos de Protección "
        "Personal (EPP) al(la) empleado(a) de GRUPO CAVAL 1003, C.A. (ALIKA PETS), en "
        "cumplimiento de los derechos establecidos en la LOPCYMAT, específicamente en el "
        "Título IV, Capítulo I, Artículo 53, que regula los derechos y deberes de los "
        "trabajadores y trabajadoras, y establece el derecho de los mismos a un ambiente "
        "laboral seguro y propicio, incluyendo la provisión de implementos de protección "
        "personal necesarios y adecuados a sus labores.")

    add_section(doc, "DATOS DE IDENTIFICACIÓN")
    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(3.5), Cm(5.0), Cm(3.5), Cm(5.0)]):
        tbl.columns[i].width = w
    datos = [
        ("Nombre del trabajador:", "Cédula de Identidad:"),
        ("Cargo:", "Fecha de entrega:"),
    ]
    for ri, (l1, l2) in enumerate(datos):
        c = tbl.rows[ri].cells[0]; c.width = Cm(3.5)
        write_cell(c, l1, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[1]; c.width = Cm(5.0)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*25)
        set_cell_borders(c); set_cell_margins(c)
        c = tbl.rows[ri].cells[2]; c.width = Cm(3.5)
        write_cell(c, l2, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[3]; c.width = Cm(5.0)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*25)
        set_cell_borders(c); set_cell_margins(c)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_section(doc, "INVENTARIO DE PRENDAS Y EQUIPOS ENTREGADOS")
    tbl2 = doc.add_table(rows=13, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.autofit = False
    for i, w in enumerate([Cm(0.8), Cm(8.5), Cm(2.5), Cm(5.2)]):
        tbl2.columns[i].width = w
    headers = ["#", "PRENDA / EQUIPO", "CANTIDAD", "ESTADO / OBSERVACIONES"]
    for ci, h in enumerate(headers):
        c = tbl2.rows[0].cells[ci]; c.width = [Cm(0.8), Cm(8.5), Cm(2.5), Cm(5.2)][ci]
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)

    items = [
        "Camisa corporativa (uniforme)",
        "Pantalón corporativo (uniforme)",
        "Gorra con logo (si aplica)",
        "Chaqueta cortaviento (si aplica)",
        "Bata/gabacha impermeable (clínica/groomer)",
        "Guantes de nitrilo (caja)",
        "Mascarilla quirúrgica/desechable",
        "Gafas protectoras",
        "Botas cerradas antideslizantes",
        "Gorro quirúrgico (clínica)",
        "Delantal plomado + collar tiroideo (vet, si aplica)",
        "Otros: ____________________",
    ]
    for ri, item in enumerate(items, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        c = tbl2.rows[ri].cells[0]; c.width = Cm(0.8)
        write_cell(c, str(ri), size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl2.rows[ri].cells[1]; c.width = Cm(8.5)
        write_cell(c, item, size=9, bg=bg)
        c = tbl2.rows[ri].cells[2]; c.width = Cm(2.5)
        write_cell(c, "_____", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl2.rows[ri].cells[3]; c.width = Cm(5.2)
        write_cell(c, "_____________________", size=9, bg=bg)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_section(doc, "OBSERVACIONES")
    add_bullet(doc, "La entrega se realiza para dar cumplimiento al numeral 4 del Artículo 53 de la LOPCYMAT, en el cual se establece que el trabajador tiene derecho a ser provisto de los implementos y equipos de protección personal adecuados según las condiciones presentes en su lugar de trabajo.", bold_lead="1.  ")
    add_bullet(doc, "El trabajador ha sido informado sobre el uso adecuado de los uniformes e implementos de protección entregados y el cuidado necesario para su preservación. Será considerado una falta grave el hecho de no cuidar, deteriorar o maltratar intencionalmente dichos uniformes o implementos.", bold_lead="2.  ")
    add_bullet(doc, "En caso de que el daño se deba a mal uso o negligencia, el trabajador podrá ser responsable de cubrir los gastos de reposición o reparación de los mismos, dependiendo de la gravedad del daño ocasionado. Los descuentos por pérdidas o daños se tramitarán conforme al artículo 59 de la LOTTT (con autorización del Inspector del Trabajo).", bold_lead="3.  ")
    add_bullet(doc, "Con esta entrega, la empresa garantiza el cumplimiento de las normas de seguridad e higiene que protegen la integridad física y mental del trabajador, de acuerdo con el derecho a un ambiente de trabajo seguro, conforme a la LOPCYMAT.", bold_lead="4.  ")
    add_bullet(doc, "El trabajador se compromete a usar los EPP durante toda su jornada laboral y a reportar inmediatamente cualquier deterioro, pérdida o necesidad de reposición.", bold_lead="5.  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "CONFORMIDAD")
    add_para(doc,
        "En señal de conformidad, firman la presente acta el representante de la empresa y "
        "el(la) trabajador(a) en la ciudad de Los Teques, Estado Miranda, en la fecha "
        "indicada arriba.",
        size=10, italic=True, space_after=8)

    # Firma: SST es función de la Directora Gerente
    add_signature_block(doc, ["LA EMPRESA (SST)", "EL(LA) TRABAJADOR(A)"])
    add_footer(section, "Acta de Entrega EPP v3.0")

    doc.save(out)
    return out


# ============================================================
# 3. EXAMEN MÉDICO PRE-EMPLEO
# ============================================================
def gen_examen_medico():
    out = "/home/z/my-project/output/05_SEGURIDAD_LABORAL/Examen_Medico_Pre_Empleo.docx"
    os.makedirs(os.path.dirname(out), exist_ok=True)

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=1.8)
    add_membrete(doc, "SEGURIDAD LABORAL", "Vigilancia médica", version="Versión 3.0  ·  SST")
    add_doc_title(doc, "EXAMEN MÉDICO PRE-EMPLEO")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Conforme al artículo 32 de la LOPCYMAT y Norma Técnica NT-02-2008")
    style_run(r, size=10, italic=True, color=GRAY_TEXT)

    add_para(doc,
        "El presente examen médico tiene por objeto evaluar el estado de salud del "
        "trabajador antes del inicio de la relación laboral, a fin de: (i) constatar su "
        "aptitud para el cargo; (ii) establecer el estado de salud basal para futuras "
        "comparaciones; (iii) detectar condiciones que requieran restricciones laborales; "
        "y (iv) cumplir con lo establecido en la LOPCYMAT y la Norma Técnica NT-02-2008 "
        "sobre Vigilancia de la Salud de los Trabajadores.")

    add_section(doc, "1. DATOS DEL TRABAJADOR(A)")
    tbl = doc.add_table(rows=4, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(3.5), Cm(5.0), Cm(3.5), Cm(5.0)]):
        tbl.columns[i].width = w
    datos = [
        ("Nombre completo:", "Cédula de Identidad:"),
        ("Fecha de nacimiento:", "Edad:"),
        ("Cargo a desempeñar:", "Departamento / Área:"),
        ("Fecha del examen:", "Lugar del examen:"),
    ]
    for ri, (l1, l2) in enumerate(datos):
        c = tbl.rows[ri].cells[0]; c.width = Cm(3.5)
        write_cell(c, l1, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[1]; c.width = Cm(5.0)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*25)
        set_cell_borders(c); set_cell_margins(c)
        c = tbl.rows[ri].cells[2]; c.width = Cm(3.5)
        write_cell(c, l2, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[3]; c.width = Cm(5.0)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*25)
        set_cell_borders(c); set_cell_margins(c)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_section(doc, "2. ANTECEDENTES PERSONALES Y FAMILIARES")
    add_para(doc, "Marque con X si aplica:", size=10, italic=True, color=GRAY_TEXT, space_after=3)
    antecedentes = [
        "Diabetes Mellitus", "Hipertensión arterial", "Asma / alergias",
        "Enfermedades cardíacas", "Enfermedades renales", "Enfermedades hepáticas",
        "Trastornos psiquiátricos", "Cirugías previas", "Hospitalizaciones previas",
        "Fracturas / traumatismos", "Convulsiones / epilepsia", "Problemas de visión",
        "Problemas auditivos", "Problemas de columna", "Enfermedades de la piel",
        "Tuberculosis", "Hepatitis", "VIH / SIDA",
        "Zoonosis previas (rabia, leptospirosis, toxoplasmosis, brucelosis)",
        "Vacuna antirrábica pre-exposición (aplica a vet/aux/groomer)",
        "Vacuna antitetánica vigente",
        "Otro (especificar): ______________________________",
    ]
    for ant in antecedentes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.2
        rb = p.add_run("☐  ")
        style_run(rb, size=10, bold=True, color=TEAL_DARK)
        r = p.add_run(ant)
        style_run(r, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "3. EXAMEN FÍSICO")
    tbl = doc.add_table(rows=8, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(4.5), Cm(4.5), Cm(4.5), Cm(4.5)]):
        tbl.columns[i].width = w
    headers = ["Parámetro", "Valor", "Parámetro", "Valor"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = Cm(4.5)
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    examenes = [
        ("Peso (kg):", "Talla (m):"),
        ("IMC:", "Presión arterial:"),
        ("Frecuencia cardíaca:", "Frecuencia respiratoria:"),
        ("Temperatura (°C):", "Perímetro abdominal:"),
        ("Agudeza visual sin lentes:", "Agudeza visual con lentes:"),
        ("Agudeza auditiva:", "Reflejos osteotendinosos:"),
        ("Examen de piel:", "Examen de columna:"),
    ]
    for ri, (l1, l2) in enumerate(examenes, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        c = tbl.rows[ri].cells[0]; c.width = Cm(4.5)
        write_cell(c, l1, size=9, bold=True, color=TEAL_DARK, bg=bg)
        c = tbl.rows[ri].cells[1]; c.width = Cm(4.5)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*20); set_cell_borders(c); set_cell_margins(c); set_cell_bg(c, bg)
        c = tbl.rows[ri].cells[2]; c.width = Cm(4.5)
        write_cell(c, l2, size=9, bold=True, color=TEAL_DARK, bg=bg)
        c = tbl.rows[ri].cells[3]; c.width = Cm(4.5)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*20); set_cell_borders(c); set_cell_margins(c); set_cell_bg(c, bg)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_section(doc, "4. EXÁMENES DE LABORATORIO SOLICITADOS")
    labs = [
        "Hemograma completo", "Glicemia en ayunas", "Examen completo de orina",
        "Examen de heces (parasitológico)", "Prueba de embarazo (si aplica)",
        "Serología (VDRL/RPR)", "Pruebas de función hepática (TGO, TGP)",
        "Radiografía de tórax (PA y lateral)", "Electrocardiograma (ECG)",
        "Prueba de Mantoux (tuberculosis)", "Marcadores virales (HBsAg, Anti-HCV) — vet/aux",
        "Título de anticuerpos antirrábicos — vet/aux/groomer",
    ]
    for lab in labs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.2
        rb = p.add_run("☐  ")
        style_run(rb, size=10, bold=True, color=TEAL_DARK)
        r = p.add_run(lab)
        style_run(r, size=10)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "5. EVALUACIÓN DE APTITUD PARA EL CARGO")
    add_para(doc, "Resultado del examen:", size=10, bold=True, space_after=3)
    opciones = [
        ("APTO sin restricciones", "El trabajador se encuentra en condiciones de salud adecuadas para desempeñar el cargo sin limitaciones."),
        ("APTO con restricciones", "El trabajador es apto pero requiere restricciones específicas que se detallan abajo."),
        ("APTO condicionado a reevaluación", "El trabajador requiere reevaluación médica en un plazo determinado."),
        ("NO APTO", "El trabajador no reúne las condiciones de salud para el cargo. No se recomienda su ingreso."),
    ]
    for op, desc in opciones:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run("☐  ")
        style_run(rb, size=11, bold=True, color=TEAL_DARK)
        rl = p.add_run(op + ":  ")
        style_run(rl, size=10, bold=True, color=BLACK)
        r = p.add_run(desc)
        style_run(r, size=10, color=GRAY_TEXT)

    add_para(doc, "Restricciones recomendadas (si aplica):", size=10, bold=True, space_before=4, space_after=3)
    c = doc.add_table(rows=1, cols=1)
    c.alignment = WD_TABLE_ALIGNMENT.CENTER
    c.columns[0].width = Cm(17)
    cc = c.rows[0].cells[0]
    cc.text = ""
    pp = cc.paragraphs[0]
    for _ in range(3):
        pp2 = cc.add_paragraph()
        pp2.paragraph_format.space_after = Pt(12)
    set_cell_borders(cc)
    set_cell_margins(cc, top=200, bottom=200)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "6. OBSERVACIONES DEL MÉDICO OCUPACIONAL")
    c = doc.add_table(rows=1, cols=1)
    c.alignment = WD_TABLE_ALIGNMENT.CENTER
    c.columns[0].width = Cm(17)
    cc = c.rows[0].cells[0]
    cc.text = ""
    pp = cc.paragraphs[0]
    for _ in range(5):
        pp2 = cc.add_paragraph()
        pp2.paragraph_format.space_after = Pt(12)
    set_cell_borders(cc)
    set_cell_margins(cc, top=200, bottom=200)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "7. DATOS DEL MÉDICO Y FIRMAS")
    t_med = doc.add_table(rows=1, cols=2)
    t_med.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_med.autofit = False
    for i in range(2):
        t_med.columns[i].width = Cm(8.5)
    left_data = [
        ("Médico ocupacional:", ""),
        ("C.I. del médico:", ""),
        ("N° colegiado CMV-CVM:", ""),
        ("Especialidad:", ""),
    ]
    right_data = [
        ("N° de registro MPPS:", ""),
        ("Fecha del examen:", ""),
        ("Próxima evaluación:", ""),
        ("Firma y sello:", ""),
    ]
    cl = t_med.rows[0].cells[0]
    cl.text = ""
    for i, (l, v) in enumerate(left_data):
        pp = cl.paragraphs[0] if i == 0 else cl.add_paragraph()
        pp.paragraph_format.space_after = Pt(6)
        rl = pp.add_run(l + " ")
        style_run(rl, size=9, bold=True, color=TEAL_DARK)
        rv = pp.add_run("_"*25)
        style_run(rv, size=9)
    set_cell_borders(cl); set_cell_margins(cl); set_cell_bg(cl, GRAY_LIGHT)
    cr = t_med.rows[0].cells[1]
    cr.text = ""
    for i, (l, v) in enumerate(right_data):
        pp = cr.paragraphs[0] if i == 0 else cr.add_paragraph()
        pp.paragraph_format.space_after = Pt(6)
        rl = pp.add_run(l + " ")
        style_run(rl, size=9, bold=True, color=TEAL_DARK)
        rv = pp.add_run("_"*25)
        style_run(rv, size=9)
    set_cell_borders(cr); set_cell_margins(cr); set_cell_bg(cr, GRAY_LIGHT)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc,
        "El resultado del presente examen es confidencial conforme al artículo 23 de la "
        "LOPCYMAT y será archivado en el expediente médico del trabajador, accesible "
        "únicamente al Comité de Seguridad y Salud Laboral y a las autoridades competentes.",
        size=9, italic=True, color=GRAY_TEXT, space_after=8)

    # Firma del examen: Médico + Trabajador + Directora Gerente (SST)
    add_signature_block(doc, ["MÉDICO OCUPACIONAL", "EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])
    add_footer(section, "Examen Médico Pre-Empleo v3.0")

    doc.save(out)
    return out


# ============================================================
# 4. CHECKLIST DE REGISTROS LEGALES
# ============================================================
def gen_registros_legales():
    out = "/home/z/my-project/output/06_REGISTROS_LEGALES/Checklist_Constancias_IVSS_FAOV_INCES_PMSSO.docx"
    os.makedirs(os.path.dirname(out), exist_ok=True)

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "REGISTROS LEGALES", "Constancias obligatorias", version="Versión 3.0  ·  RR.HH.")
    add_doc_title(doc, "CHECKLIST DE REGISTROS LEGALES")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("IVSS · FAOV-BVV · INCES · PMSSO")
    style_run(r, size=11, italic=True, color=GRAY_TEXT)

    add_para(doc,
        "El presente checklist documenta la verificación de los registros del trabajador "
        "ante los organismos del Estado venezolano. Su omisión genera responsabilidad "
        "directa para GRUPO CAVAL 1003, C.A. y puede impedir el pago de prestaciones, "
        "seguros sociales y demás beneficios. Debe completarse ANTES del primer día de "
        "trabajo (Fase 0 del proceso de ingreso).")

    add_section(doc, "DATOS DEL TRABAJADOR(A)")
    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(3.5), Cm(5.0), Cm(3.5), Cm(5.0)]):
        tbl.columns[i].width = w
    datos = [
        ("Nombre completo:", "Cédula de Identidad:"),
        ("Cargo:", "Fecha de ingreso:"),
    ]
    for ri, (l1, l2) in enumerate(datos):
        c = tbl.rows[ri].cells[0]; c.width = Cm(3.5)
        write_cell(c, l1, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[1]; c.width = Cm(5.0)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*25)
        set_cell_borders(c); set_cell_margins(c)
        c = tbl.rows[ri].cells[2]; c.width = Cm(3.5)
        write_cell(c, l2, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG)
        c = tbl.rows[ri].cells[3]; c.width = Cm(5.0)
        c.text = ""; c.paragraphs[0].add_run(" " + "_"*25)
        set_cell_borders(c); set_cell_margins(c)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_section(doc, "REGISTROS A VERIFICAR")
    tbl = doc.add_table(rows=6, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(0.8), Cm(5.5), Cm(4.5), Cm(1.8), Cm(2.0), Cm(3.0)]):
        tbl.columns[i].width = w
    headers = ["#", "Registro", "Marco legal", "Estado", "Fecha", "Observaciones"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = [Cm(0.8), Cm(5.5), Cm(4.5), Cm(1.8), Cm(2.0), Cm(3.0)][ci]
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)

    registros = [
        ("1", "Inscripción IVSS (Forma 14-100)", "Ley del Seguro Social Obligatorio",
         "Constancia emitida por el IVSS que acredita la afiliación del trabajador al régimen de seguridad social."),
        ("2", "Inscripción FAOV-BVV", "Ley del Banco de Vivienda y Hábitat (RBVV)",
         "Constancia que acredita inscripción en el Fondo de Ahorro Obligatorio para la Vivienda."),
        ("3", "Inscripción INCES", "Ley del INCES",
         "Constancia que acredita inscripción del trabajador en el Instituto Nacional de Capacitación y Educación Socialista."),
        ("4", "Seguro de Riesgos Laborales (PMSSO)", "Art. 134 LOPCYMAT — CRÍTICO",
         "Sin este seguro, la empresa responde con su patrimonio por cualquier accidente laboral o enfermedad ocupacional."),
        ("5", "Constancia Vacunación Antirrábica Pre-Exposición", "Recomendación OMS / Norma Vet.",
         "Aplica a Veterinario, Auxiliar Vet. y Dog Groomer. Esquema: 3 dosis + refuerzo anual."),
    ]

    for ri, (n, reg, marc, obs) in enumerate(registros, start=1):
        bg = AMBER_BG if "CRÍTICO" in marc else (GRAY_ALT if ri % 2 == 0 else WHITE)
        c = tbl.rows[ri].cells[0]; c.width = Cm(0.8)
        write_cell(c, n, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        col = RED_CRIT if "CRÍTICO" in marc else BLACK
        c = tbl.rows[ri].cells[1]; c.width = Cm(5.5)
        write_cell(c, reg, size=9, bold=("CRÍTICO" in marc), color=col, bg=bg)
        c = tbl.rows[ri].cells[2]; c.width = Cm(4.5)
        write_cell(c, marc, size=9, italic=True, color=col, bg=bg)
        c = tbl.rows[ri].cells[3]; c.width = Cm(1.8)
        write_cell(c, "☐ P\n☐ E\n☐ F", size=8, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[4]; c.width = Cm(2.0)
        write_cell(c, "__/__/__", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[5]; c.width = Cm(3.0)
        write_cell(c, obs, size=8, bg=bg)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc,
        "ESTADO:  P = Pendiente  ·  E = Entregado por el trabajador  ·  F = Verificado y firmado.",
        size=9, italic=True, color=GRAY_TEXT, space_after=6)

    add_section(doc, "PROCEDIMIENTO DE INSCRIPCIÓN")
    procs = [
        ("IVSS (Forma 14-100):",
         "El trabajador debe acudir a la oficina del IVSS más cercana con cédula y RIF. Si "
         "es primera vez, solicitar inscripción. Si ya está inscrito, solicitar actualización "
         "de datos laborales. La empresa debe reportar la incorporación del trabajador en "
         "los primeros 5 días del mes siguiente al ingreso."),
        ("FAOV-BVV:",
         "El trabajador debe inscribirse en el banco operador de su preferencia (Banco del "
         "Pueblo, BFC, etc.) con cédula, RIF y constancia de trabajo. La empresa debe "
         "retener el 1% del salario y el 2% adicional (3% total) y depositarlo dentro de "
         "los primeros 5 días del mes siguiente."),
        ("INCES:",
         "La empresa debe inscribir al trabajador y aportar el 2% de la nómina (1% "
         "aportación patronal + 0.5% aporte trabajador + 0.5% adicional patronal). El "
         "trabajador puede solicitar constancia individual."),
        ("PMSSO (Seguro Riesgos Laborales) — CRÍTICO:",
         "La empresa debe inscribirse en el IVSS-PMSSO como empresa con riesgos laborales "
         "y reportar a cada trabajador. Sin esta inscripción, cualquier accidente de "
         "trabajo o enfermedad ocupacional es responsabilidad directa de la empresa con su "
         "patrimonio. La inscripción es OBLIGATORIA antes del primer día de labores."),
    ]
    for label, txt in procs:
        col = RED_CRIT if "CRÍTICO" in label else BLACK
        add_bullet(doc, txt, bold_lead=label + "  ", color=col)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "CONSECUENCIAS DE NO CUMPLIR")
    sanciones = [
        ("IVSS:", "Multa administrativa + responsabilidad solidaria de la empresa por prestaciones médicas del trabajador no asegurado."),
        ("FAOV-BVV:", "Multa + obligación de enterar los aportes retenidos con intereses moratorios."),
        ("INCES:", "Multa + obligación de enterar los aportes retenidos."),
        ("PMSSO:", "RESPONSABILIDAD ILIMITADA de la empresa por accidentes laborales y enfermedades ocupacionales. Sin límite de monto."),
    ]
    for label, txt in sanciones:
        col = RED_CRIT if "ILIMITADA" in txt else BLACK
        add_bullet(doc, txt, bold_lead=label + "  ", color=col)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_section(doc, "VERIFICACIÓN Y CIERRE")
    add_para(doc,
        "Declaro que he verificado la totalidad de los registros arriba indicados y que "
        "todos los aplicables al trabajador se encuentran en regla. Las constancias "
        "originales están archivadas en el expediente del trabajador.",
        size=10, italic=True, space_after=6)
    add_para(doc,
        "En la ciudad de Los Teques, Estado Miranda, a los ____ días del mes de "
        "________________ de ________.", size=10, space_after=8)
    add_signature_block(doc, ["DIRECTORA GERENTE", "EL(LA) TRABAJADOR(A)"])
    add_footer(section, "Checklist Registros Legales v3.0")

    doc.save(out)
    return out


# ============================================================
# 5. CARTA DE ACEPTACIÓN GENERAL
# ============================================================
def gen_carta_aceptacion():
    out = "/home/z/my-project/output/09_CIERRE/Carta_Aceptacion_General.docx"
    os.makedirs(os.path.dirname(out), exist_ok=True)

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.2)
    add_membrete(doc, "CIERRE DE EXPEDIENTE", "Aceptación general del Kit", version="Versión 3.0  ·  RR.HH.")
    add_doc_title(doc, "CARTA DE ACEPTACIÓN GENERAL DEL KIT DE INGRESO")

    add_para(doc,
        "Ciudad y fecha:  Los Teques, ____ de ________________ de ________",
        size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

    add_para(doc,
        "Ciudadana\nESNATLIM ELENA SIMOZA\nDirectora Gerente\nGRUPO CAVAL 1003, C.A. (ALIKA PETS)\nSu despacho.-",
        size=10, space_after=8)

    add_para(doc, "Ref: Aceptación general del Kit de Ingreso.", size=10, bold=True, space_after=8)

    add_para(doc,
        "Yo, _____________________________________________, titular de la cédula de identidad "
        "N° V-___________________, en mi condición de trabajador(a) de GRUPO CAVAL 1003, "
        "C.A. (ALIKA PETS), ingresado(a) en el cargo de _________________________ con fecha "
        "____ / ____ / ________, por medio de la presente declaro:")

    declaraciones = [
        ("PRIMERO:  ",
         "Que he recibido copia íntegra y legible de todos los documentos que conforman el "
         "Kit de Ingreso del Trabajador, conforme al Checklist Maestro firmado por separado, "
         "incluyendo, sin carácter limitativo: contrato individual de trabajo, descripción "
         "de cargo, reglamento interno, código de conducta, política de confidencialidad, "
         "política de uso de redes sociales, autorización para depósito de prestaciones "
         "sociales, designación de beneficiarios, notificación de riesgos, hoja de recorrido "
         "habitual, acta de entrega de EPP, cartilla de bioseguridad veterinaria (cuando "
         "aplique), protocolo de sustancias controladas (cuando aplique), protocolo de "
         "mordeduras y zoonosis (cuando aplique), procedimiento de reporte de incidentes, "
         "autorización de tratamiento de datos personales (LOPDP), autorización de uso de "
         "imagen, y autorización de vigilancia por cámaras."),
        ("SEGUNDO:  ",
         "Que he leído en su totalidad cada uno de los documentos recibidos, he recibido "
         "explicaciones verbales sobre su contenido por parte de la Dirección, y he tenido "
         "la oportunidad de formular preguntas, las cuales fueron respondidas "
         "satisfactoriamente."),
        ("TERCERO:  ",
         "Que comprendo mis obligaciones, prohibiciones y deberes, y me comprometo a "
         "cumplirlos fielmente durante toda la vigencia de mi relación laboral con la "
         "empresa, así como las normas de seguridad y salud laboral establecidas en la "
         "LOPCYMAT y en los protocolos internos."),
        ("CUARTO:  ",
         "Que entiendo que el incumplimiento de las disposiciones contenidas en los "
         "documentos del Kit podrá dar lugar a las sanciones disciplinarias previstas en el "
         "Reglamento Interno y, según la gravedad, a la terminación de la relación laboral "
         "por causa justificada conforme al artículo 79 de la LOTTT."),
        ("QUINTO:  ",
         "Que autorizo expresamente a GRUPO CAVAL 1003, C.A. para el tratamiento de mis "
         "datos personales conforme a la LOPDP, así como el uso de mi imagen conforme a la "
         "autorización firmada por separado."),
        ("SEXTO:  ",
         "Que acepto que la presente firma constituye prueba fehaciente de la entrega y "
         "recepción de todos los documentos del Kit, renunciando a alegar desconocimiento "
         "de su contenido en el futuro."),
        ("SÉPTIMO:  ",
         "Que reconozco que GRUPO CAVAL 1003, C.A. ha cumplido con sus obligaciones de "
         "información, capacitación y entrega de equipos conforme a la LOPCYMAT, y que "
         "estoy en condiciones de iniciar mis labores."),
    ]
    for label, txt in declaraciones:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        rl = p.add_run(label)
        style_run(rl, size=10, bold=True, color=TEAL_DARK)
        r = p.add_run(txt)
        style_run(r, size=10)

    add_para(doc, "", space_after=4)
    add_para(doc,
        "En fe de lo cual firmo la presente carta en la ciudad de Los Teques, Estado "
        "Miranda, en la fecha indicada arriba.",
        size=10, space_after=12)

    # 2 firmas: Trabajador + Directora Gerente
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])
    add_footer(section, "Carta de Aceptación General v3.0")

    doc.save(out)
    return out


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print("Generando 5 documentos finales v3.0...")
    outs = [
        gen_hoja_recorrido(),
        gen_acta_epp(),
        gen_examen_medico(),
        gen_registros_legales(),
        gen_carta_aceptacion(),
    ]
    for o in outs:
        size_kb = os.path.getsize(o) / 1024
        print(f"  ✓ {o}  ({size_kb:.1f} KB)")
    print("\nTodos los documentos generados correctamente.")
