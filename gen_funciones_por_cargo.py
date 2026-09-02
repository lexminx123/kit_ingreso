"""
gen_funciones_por_cargo.py — 7 Descripciones de Cargo (Bloque 03)
GRUPO CAVAL 1003, C.A. — ALIKA PETS — Los Teques, Miranda.

Genera en /home/z/my-project/output/03_DESCRIPCION_DE_CARGOS/:
  03a_Funciones_Gerente.docx
  03b_Funciones_Encargado_Clinica.docx
  03c_Funciones_Encargado_Tienda.docx
  03d_Funciones_Medico_Veterinario.docx
  03e_Funciones_Auxiliar_Veterinario.docx
  03f_Funciones_Dog_Groomer.docx
  03g_Funciones_Asistente_Ventas.docx

Estructura por cargo:
  - Membrete "DESCRIPCIÓN DE CARGO" + subtítulo del cargo
  - Sección 1: Identificación del Cargo (tabla 6 filas)
  - Sección 2: Objetivo General
  - Sección 3: Funciones Específicas (organizadas por área, bullets)
  - Sección 4: Competencias y Requisitos
  - Sección 5: Condiciones de Trabajo y Riesgos
  - Carta de Recepción y Aceptación (firma TRABAJADOR + DIRECTORA GERENTE)
  - Footer "Descripción Cargo — [CARGO] v3.0"

Versión 3.0 — usa _common.py (módulo reutilizable).
Firma operativa: Esnatlim Elena Simoza, C.I. V-17.976.287, Directora Gerente.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# Importar utilidades compartidas
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _common import (
    setup_a4_portrait, add_membrete, add_doc_title, add_section, add_chapter,
    add_para, add_bullet, add_article, add_signature_block, add_reception_letter,
    add_footer, add_sanciones_table,
    set_cell_bg, set_cell_borders, set_cell_margins, write_cell, style_run, add_hr,
    TEAL_DARK, TEAL_HDR_BG, GRAY_ALT, SLATE_BG, WHITE, BLACK, GRAY_TEXT, GRAY_MUTED,
)

OUT_DIR = "/home/z/my-project/output/03_DESCRIPCION_DE_CARGOS"
os.makedirs(OUT_DIR, exist_ok=True)


# ============================================================
# Helper: tabla de Identificación del Cargo
# ============================================================
def add_identificacion_table(doc, titulo, departamento, reporta_a, supervisa_a,
                             nivel_jerarquico, tipo_contrato):
    """Tabla 6 filas x 2 columnas: label teal + valor."""
    tbl = doc.add_table(rows=6, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(5.5)
    tbl.columns[1].width = Cm(11.1)

    rows = [
        ("Título del cargo",        titulo),
        ("Departamento",            departamento),
        ("Reporta a",               reporta_a),
        ("Supervisa a",             supervisa_a),
        ("Nivel jerárquico",        nivel_jerarquico),
        ("Tipo de contrato",        tipo_contrato),
    ]
    for i, (label, value) in enumerate(rows):
        c_lab = tbl.rows[i].cells[0]
        c_val = tbl.rows[i].cells[1]
        c_lab.width = Cm(5.5)
        c_val.width = Cm(11.1)
        write_cell(c_lab, label, size=9, bold=True, color=TEAL_DARK,
                   align=WD_ALIGN_PARAGRAPH.LEFT, bg=SLATE_BG)
        write_cell(c_val, value, size=9, color=BLACK,
                   align=WD_ALIGN_PARAGRAPH.LEFT, bg=WHITE if i % 2 == 0 else GRAY_ALT)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


# ============================================================
# Helper: bloque de funciones por área
# ============================================================
def add_funciones_area(doc, area_title, bullets):
    """Sub-encabezado teal + lista de bullets con descripción corta al inicio."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(area_title)
    style_run(r, size=10, bold=True, color=TEAL_DARK)
    for item in bullets:
        if isinstance(item, tuple):
            lead, rest = item
            add_bullet(doc, rest, bold_lead=lead + ".  ")
        else:
            add_bullet(doc, item)


# ============================================================
# Helper: competencias y requisitos
# ============================================================
def add_competencias_block(doc, tecnicas, blandas, academicas_legales):
    """Tres sub-bloques de competencias."""
    add_section(doc, "Competencias técnicas")
    for t in tecnicas:
        add_bullet(doc, t)
    add_section(doc, "Competencias blandas (actitudinales)")
    for b in blandas:
        add_bullet(doc, b)
    add_section(doc, "Requisitos académicos y legales")
    for a in academicas_legales:
        add_bullet(doc, a)


# ============================================================
# Helper: condiciones de trabajo y riesgos
# ============================================================
def add_condiciones_block(doc, lugar, jornada, epp, riesgos):
    """Bloque de condiciones laborales + EPP + riesgos."""
    add_para(doc, f"Lugar de trabajo: {lugar}", size=10, bold=False,
             color=BLACK, space_before=2, space_after=2)
    add_para(doc, f"Jornada: {jornada}", size=10, color=BLACK, space_after=2)
    add_para(doc, "Equipos de protección personal (EPP) obligatorios:",
             size=10, bold=True, color=TEAL_DARK, space_before=4, space_after=1)
    for e in epp:
        add_bullet(doc, e)
    add_para(doc, "Riesgos principales del cargo:",
             size=10, bold=True, color=TEAL_DARK, space_before=4, space_after=1)
    for r in riesgos:
        add_bullet(doc, r)


# ============================================================
# Datos de cada cargo
# ============================================================
CARGOS = {

    # ---------- 03a GERENTE ----------
    "gerente": {
        "file": "03a_Funciones_Gerente.docx",
        "sublabel": "Gerente General",
        "titulo": "GERENTE GENERAL",
        "departamento": "Dirección General",
        "reporta_a": "Directora Gerente",
        "supervisa_a": "Encargado de Clínica, Encargado de Tienda, Médico(s) Veterinario(s), Auxiliares Veterinarios, Dog Groomer(s), Asistente(s) de Ventas, personal administrativo",
        "nivel_jerarquico": "Ejecutivo superior — Nivel 1",
        "tipo_contrato": "Tiempo indeterminado (LOTTT art. 65) — período de prueba 30 días",
        "objetivo": (
            "Planificar, dirigir, coordinar y controlar la operación integral de la clínica "
            "veterinaria, tienda de mascotas y peluquería canina de GRUPO CAVAL 1003, C.A. "
            "(ALIKA PETS), asegurando la calidad de los servicios, el cumplimiento de las "
            "obligaciones legales venezolanas (LOTTT, LOPCYMAT, LOPDP, Normas NT-01/NT-02, "
            "Ley de Ejercicio de la Medicina Veterinaria), la rentabilidad económica del "
            "negocio y la satisfacción del cliente interno y externo."
        ),
        "funciones": [
            ("Dirección y supervisión general", [
                "Establecer los objetivos estratégicos, operativos y comerciales de la empresa a corto, mediano y largo plazo.",
                "Dirigir y supervisar el cumplimiento de las políticas internas, reglamento, código de conducta y protocolos veterinarios.",
                "Coordinar la integración operativa entre las tres unidades de negocio: clínica, tienda y peluquería canina.",
                "Representar a la empresa ante clientes, proveedores, autoridades sanitarias (SENCAMER, SENAC), IVSS, INCES, FAOV, Inspectoría del Trabajo y comunidades locales.",
            ]),
            ("Gestión del personal", [
                "Reclutar, seleccionar, contratar, inducir, capacitar y evaluar al personal de las tres áreas.",
                "Asignar funciones, horarios, rotaciones, vacaciones y permisos conforme a la LOTTT.",
                "Aplicar el régimen disciplinario del Reglamento Interno y, según la gravedad, iniciar procedimientos de rescisión contractual por causa justificada (LOTTT art. 79).",
                "Velar por el clima organizacional, la prevención del acoso laboral y la igualdad de oportunidades.",
            ]),
            ("Control de inventarios y compras", [
                "Autorizar las órdenes de compra de medicamentos, insumos médicos, alimentos, accesorios, productos de peluquería y consumibles.",
                "Supervisar los inventarios físicos mensuales y la conciliación con el sistema de gestión.",
                "Controlar el libro foliado de sustancias controladas (opioides, anestésicos, benzodiacepinas) conforme al SENAC.",
                "Verificar el cumplimiento de la cadena de frío y las condiciones de almacenamiento.",
            ]),
            ("Control financiero y contable", [
                "Supervisar el flujo de caja diario, depósitos bancarios, conciliaciones y arqueos de caja de las tres unidades.",
                "Revisar y firmar la emisión de facturas, notas de crédito, retenciones de IVA e ISLR.",
                "Aprobar pagos a proveedores, nómina, prestaciones sociales y beneficios contractuales.",
                "Controlar la morosidad, créditos a clientes y cuentas por cobrar.",
            ]),
            ("Indicadores y mejora continua (KPIs)", [
                "Definir y monitorear KPIs: facturación por área, ticket promedio, margen bruto, ocupación de quirófano, tasa de hospitalización, rotación de inventario, satisfacción del cliente, ausentismo, rotación de personal.",
                "Realizar reuniones mensuales de gestión con cada jefe de área para revisión de indicadores y planes de mejora.",
                "Implementar acciones correctivas, preventivas y de mejora derivadas de auditorías internas, quejas de clientes o reportes de incidentes.",
            ]),
            ("Cumplimiento normativo", [
                "Garantizar el cumplimiento de la LOTTT, LOPCYMAT, LOPDP, Ley del Seguro Social, Ley del INCES, Ley del BVV (FAOV), NT-01-2008 (Bioseguridad), NT-02-2008 (Vigilancia Médica), Ley de Ejercicio de la Medicina Veterinaria, Ley Orgánica de Drogas y demás normativa aplicable.",
                "Mantener actualizado el Comité de Seguridad y Salud Laboral (art. 46 LOPCYMAT) y el Programa de Seguridad y Salud en el Trabajo.",
                "Supervisar la declaración y pago puntual de impuestos, retenciones y contribuciones parafiscales.",
                "Atender inspecciones del trabajo, fiscales, sanitarias y ambientales.",
            ]),
        ],
        "tecnicas": [
            "Manejo avanzado de herramientas de gestión: Excel, sistemas POS, software veterinario, sistema de nómina.",
            "Conocimientos de contabilidad básica, costos, presupuestos y análisis financiero.",
            "Conocimiento del marco legal venezolano aplicable al sector veterinario y comercial.",
            "Gestión de inventarios, compras y cadena de suministro.",
            "Interpretación de indicadores de gestión (KPIs) y cuadros de mando.",
        ],
        "blandas": [
            "Liderazgo, capacidad de decisión y delegación efectiva.",
            "Comunicación asertiva con clientes, proveedores y equipo.",
            "Resolución de conflictos y manejo de situaciones de presión.",
            "Pensamiento estratégico y orientación a resultados.",
            "Integridad, confidencialidad y ética profesional.",
        ],
        "academicas_legales": [
            "Título universitario en Administración, Contaduría Pública, Ingeniería de Producción, Medicina Veterinaria o afines (preferencial).",
            "Experiencia comprobada mínima de cinco (5) años en cargos de gerencia general o de operaciones.",
            "Inscripción vigente en el IVSS, FAOV, INCES y registro mercantil actualizado.",
            "Curso vigente de Técnico Superior en Seguridad y Salud Laboral o delegado del Comité SST (LOPCYMAT art. 46).",
        ],
        "condiciones": {
            "lugar": (
                "Oficina administrativa de la sede en Av. Francisco de Miranda, Local N° 1, "
                "Los Teques, Miranda; con circulación por las áreas de clínica, tienda y peluquería."
            ),
            "jornada": (
                "Jornada diurna mixta de lunes a sábado, 40 horas semanales distribuidas conforme a "
                "LOTTT art. 173. Rotación de guardias los fines de semana para supervisión remota."
            ),
            "epp": [
                "Uniforme corporativo y carné identificatorio.",
                "Calzado cerrado antideslizante.",
                "Uso ocasional de mascarilla y guantes al ingresar a áreas clínicas o de depósito.",
            ],
            "riesgos": [
                "Ergonómico por trabajo prolongado en computadora y posturas sedentes.",
                "Estrés laboral por toma de decisiones y atención simultánea de múltiples frentes.",
                "Biológico eventual al circular por áreas de hospitalización y quirófano (NT-01-2008).",
                "Riesgo psicosocial por carga de responsabilidad y conflictos interpersonales.",
            ],
        },
    },

    # ---------- 03b ENCARGADO CLÍNICA ----------
    "encargado_clinica": {
        "file": "03b_Funciones_Encargado_Clinica.docx",
        "sublabel": "Encargado(a) de Clínica Veterinaria",
        "titulo": "ENCARGADO(A) DE CLÍNICA VETERINARIA",
        "departamento": "Clínica Veterinaria",
        "reporta_a": "Gerente General",
        "supervisa_a": "Médico(s) Veterinario(s), Auxiliar(es) Veterinario(s), recepcionistas del área clínica",
        "nivel_jerarquico": "Jefatura de área — Nivel 2",
        "tipo_contrato": "Determinado por 6 meses (LOTTT art. 65) — período de prueba 30 días",
        "objetivo": (
            "Coordinar el funcionamiento diario de la unidad de clínica veterinaria, garantizando "
            "la calidad técnica de los servicios médicos, la atención oportuna y humana al cliente, "
            "el cumplimiento de los protocolos veterinarios (NT-01-2008, NT-02-2008), la "
            "disponibilidad de insumos y medicamentos, y el adecuado manejo de historias clínicas, "
            "presupuestos y registros de sustancias controladas."
        ),
        "funciones": [
            ("Dirección del área clínica", [
                "Abrir y cerrar el área clínica según el horario asignado, verificando condiciones de bioseguridad y equipos.",
                "Distribuir consultas, cirugías, hospitalización y urgencias entre el médico veterinario y auxiliares.",
                "Velar por el cumplimiento del Protocolo de Sustancias Controladas, Bioseguridad y Mordeduras/Zoonosis.",
            ]),
            ("Atención al cliente", [
                "Recibir al cliente y a la mascota, registrar el motivo de consulta y antecedentes en la historia clínica.",
                "Brindar información clara sobre servicios, horarios, costos y tiempos estimados.",
                "Atender y escalar quejas o reclamaciones conforme al procedimiento interno.",
            ]),
            ("Presupuestos y cobros", [
                "Elaborar presupuestos detallados de consultas, procedimientos, cirugías y hospitalización.",
                "Registrar cobros, emitir facturas, manejar caja chica del área clínica y entregar comprobantes.",
                "Conciliar al cierre del día los ingresos del área con el sistema y el depósito bancario.",
            ]),
            ("Inventario médico", [
                "Controlar existencias de medicamentos, insumos, reactivos de laboratorio y material quirúrgico.",
                "Realizar pedidos oportunos a proveedores y verificar la cadena de frío al recibir.",
                "Llevar el libro foliado de sustancias controladas con doble cerradura (SENAC).",
                "Reportar mensualmente al Gerente las mermas, vencimientos y reposiciones necesarias.",
            ]),
            ("Higiene y bioseguridad", [
                "Supervisar la limpieza y desinfección de consultorios, quirófano, hospitalización y depósito conforme a NT-01-2008.",
                "Velar por la segregación y disposición de residuos COVENIN 2747-93 (rojo, amarillo, negro, verde).",
                "Garantizar el uso obligatorio de EPP por el personal y la señalización de zonas.",
            ]),
            ("Políticas y registros", [
                "Garantizar que toda consulta, procedimiento y medicación quede registrado en la historia clínica.",
                "Cumplir y hacer cumplir la política de confidencialidad y propiedad de las historias clínicas (Ley Ejercicio Medicina Veterinaria art. 23).",
                "Mantener actualizados los consentimientos informados y autorizaciones de hospitalización/cirugía.",
                "Reportar al Gerente cualquier incidente, accidente o evento adverso en el formato establecido.",
            ]),
        ],
        "tecnicas": [
            "Manejo de software veterinario (historia clínica electrónica, agenda, facturación).",
            "Conocimiento de farmacología veterinaria básica, antagonistas y dosis de emergencia.",
            "Manejo de protocolos de bioseguridad NT-01-2008 y residuos COVENIN 2747-93.",
            "Operación básica de equipos: autoclave, centrifuga, microscopio, monitor de signos vitales.",
            "Conocimiento de sustancias controladas y manejo del libro foliado SENAC.",
        ],
        "blandas": [
            "Liderazgo de equipo, empatía y manejo de clientes en situaciones de estrés.",
            "Comunicación clara y asertiva con clientes y personal médico.",
            "Capacidad de organización y priorización bajo presión.",
            "Orientación al detalle y al cumplimiento de protocolos.",
            "Trabajo en equipo y disposición para aprender.",
        ],
        "academicas_legales": [
            "Técnico Superior Universitario (TSU) en Medicina Veterinaria, Veterinaria, Enfermería Veterinaria o afines (preferencial).",
            "Experiencia mínima de tres (3) años en clínica veterinaria, con al menos uno en supervisión.",
            "Curso de bioseguridad NT-01-2008 y manejo de sustancias controladas.",
            "Inscripción vigente en IVSS, FAOV, INCES y PMSSO.",
        ],
        "condiciones": {
            "lugar": (
                "Área de clínica veterinaria en la sede de Av. Francisco de Miranda, Local N° 1, "
                "Los Teques, Miranda. Incluye recepción, consultorios, quirófano, hospitalización "
                "y depósito médico."
            ),
            "jornada": (
                "Jornada mixta de lunes a sábado, 44 horas semanales con un día libre rotativo "
                "(LOTTT art. 173). Disponibilidad para cubrir guardias de fin de semana y "
                "hospitalizados."
            ),
            "epp": [
                "Uniforme clínico (scrub) institucional, carné identificatorio.",
                "Guantes de nitrilo, mascarilla quirúrgica, gafas de protección.",
                "Calzado cerrado antideslizante y gorro quirúrgico en sala de operaciones.",
                "Delantal plomado para radiografías (cuando aplique).",
            ],
            "riesgos": [
                "Biológico por exposición a fluidos, secreciones y zoonosis (rabia, leptospirosis, tiña).",
                "Ergonómico por levantamiento de animales, posturas prolongadas y traslados.",
                "Por mordedura, arañazo o patada de animales (Protocolo de Mordeduras/Zoonosis).",
                "Químico por anestésicos, desinfectantes y fármacos.",
                "Psicosocial por trabajo bajo presión y atención de urgencias.",
            ],
        },
    },

    # ---------- 03c ENCARGADO TIENDA ----------
    "encargado_tienda": {
        "file": "03c_Funciones_Encargado_Tienda.docx",
        "sublabel": "Encargado(a) de Tienda de Mascotas",
        "titulo": "ENCARGADO(A) DE TIENDA DE MASCOTAS",
        "departamento": "Tienda de Mascotas",
        "reporta_a": "Gerente General",
        "supervisa_a": "Asistente(s) de Ventas y cajero(s) del área de tienda",
        "nivel_jerarquico": "Jefatura de área — Nivel 2",
        "tipo_contrato": "Determinado por 6 meses (LOTTT art. 65) — período de prueba 30 días",
        "objetivo": (
            "Administrar la operación diaria de la tienda de mascotas — venta de alimentos, "
            "accesorios, productos de higiene y antiparasitarios — asegurando la disponibilidad "
            "de inventario, la correcta atención al cliente, el manejo transparente de la caja, "
            "la conservación de la cadena de frío y la exhibición atractiva de los productos "
            "conforme a las políticas internas de ALIKA PETS."
        ),
        "funciones": [
            ("Manejo de caja", [
                "Abrir y cerrar caja diariamente, registrando monto inicial, ingresos, egresos y arqueo final.",
                "Cobrar ventas con efectivo, punto de venta, transferencia o tarjeta, emitiendo factura y comprobante.",
                "Realizar depósitos bancarios diarios y conciliar con el sistema de ventas.",
                "Reportar cualquier diferencia, faltante o exceso al Gerente.",
            ]),
            ("Control de inventario", [
                "Mantener actualizado el kárdex de productos, incluyendo entradas, salidas y mermas.",
                "Realizar conteos físicos parciales semanales y un inventario general mensual.",
                "Verificar fechas de vencimiento y rotación FIFO de alimentos y medicamentos OTC.",
                "Coordinar recepción de mercancía con proveedores, verificando cantidades y calidad.",
                "Reponer góndolas, exhibidores y neveras cuidando la cadena de frío.",
            ]),
            ("Atención al cliente", [
                "Saludar y asesorar al cliente en la selección de productos adecuados para su mascota.",
                "Informar sobre promociones, programas de fidelización, devoluciones y garantías.",
                "Atender y escalar quejas o reclamos conforme al procedimiento interno.",
                "Mantener una actitud cordial, presentable y proactiva.",
            ]),
            ("Mantenimiento del área", [
                "Velar por la limpieza y orden de góndolas, exhibidores, neveras, mostrador y bodega.",
                "Revisar funcionamiento de neveras, aire acondicionado, iluminación y equipos POS.",
                "Reportar fallas técnicas, daños o necesidades de mantenimiento al Gerente.",
            ]),
            ("Políticas y registros", [
                "Cumplir y hacer cumplir las políticas internas: reglamento, código de conducta, confidencialidad.",
                "Registrar todas las transacciones en el sistema POS, sin excepciones.",
                "Aplicar el procedimiento de reporte de incidentes en caso de robo, hurto, asalto o accidente.",
                "Velar por la correcta exhibición de precios, rotulación y señalización.",
            ]),
        ],
        "tecnicas": [
            "Manejo de sistemas POS y software de inventario.",
            "Conocimiento básico de nutrición, antiparasitarios y accesorios para mascotas.",
            "Operación de equipos: lector de código de barras, terminal bancario, nevera de cadena de frío.",
            "Manejo de arqueos de caja, conciliaciones y depósitos bancarios.",
            "Conocimientos básicos de exhibición y merchandising.",
        ],
        "blandas": [
            "Atención cordial y orientación al cliente.",
            "Honestidad, responsabilidad y manejo ético del dinero.",
            "Capacidad de organización y trabajo bajo presión.",
            "Trabajo en equipo y buena comunicación con el área de clínica y peluquería.",
            "Disposición para aprender sobre nuevos productos y marcas.",
        ],
        "academicas_legales": [
            "Bachiller o TSU en Administración, Comercio, Mercadeo o afines (preferencial).",
            "Experiencia mínima de dos (2) años en tienda, retail o caja.",
            "Inscripción vigente en IVSS, FAOV, INCES y PMSSO.",
            "Curso básico de seguridad y salud laboral (LOPCYMAT).",
        ],
        "condiciones": {
            "lugar": (
                "Área de tienda de mascotas en la sede de Av. Francisco de Miranda, Local N° 1, "
                "Los Teques, Miranda. Incluye mostrador, góndolas, bodega y caja."
            ),
            "jornada": (
                "Jornada mixta de lunes a sábado, 44 horas semanales con un día libre rotativo "
                "(LOTTT art. 173). Posibilidad de jornada extendida en temporadas altas."
            ),
            "epp": [
                "Uniforme corporativo y carné identificatorio.",
                "Calzado cerrado antideslizante.",
                "Guantes para manipulación de productos de limpieza o bodega.",
            ],
            "riesgos": [
                "Ergonómico por levantamiento de cajas, posturas prolongadas y traslados.",
                "Riesgo de asalto o hurto al manejar caja (Protocolo de Reporte de Incidentes).",
                "Caídas al mismo nivel por derrames o piso mojado.",
                "Cortes o pinchazos con cajas, esquineros o cintas.",
                "Químico por manipulación de productos de limpieza y antiparasitarios.",
            ],
        },
    },

    # ---------- 03d MÉDICO VETERINARIO ----------
    "medico_veterinario": {
        "file": "03d_Funciones_Medico_Veterinario.docx",
        "sublabel": "Médico(a) Veterinario(a)",
        "titulo": "MÉDICO(A) VETERINARIO(A)",
        "departamento": "Clínica Veterinaria",
        "reporta_a": "Encargado(a) de Clínica / Gerente General",
        "supervisa_a": "Auxiliar(es) Veterinario(s) en procedimientos médicos y quirúrgicos",
        "nivel_jerarquico": "Profesional técnico — Nivel 3",
        "tipo_contrato": "Tiempo indeterminado (LOTTT art. 65) — período de prueba 30 días",
        "objetivo": (
            "Prestar atención médica veterinaria integral a los pacientes de ALIKA PETS — "
            "consultas, diagnósticos, cirugías, hospitalización, urgencias y medicina "
            "preventiva — conforme a la Ley de Ejercicio de la Medicina Veterinaria, las "
            "Normas Técnicas NT-01-2008 (Bioseguridad) y NT-02-2008 (Vigilancia Médica), "
            "el Código de Ética del Colegio de Médicos Veterinarios de Venezuela y los "
            "protocolos internos de la clínica."
        ),
        "funciones": [
            ("Consultas y diagnóstico", [
                "Realizar anamnesis, examen físico completo y diagnóstico diferencial de los pacientes.",
                "Solicitar e interpretar exámenes de laboratorio (sangre, orina, heces, citología) y estudios de imagen (radiografía, ecografía).",
                "Indicar y explicar al cliente el plan diagnóstico, terapéutico y pronóstico.",
                "Registrar la consulta en la historia clínica electrónica con detalle profesional.",
            ]),
            ("Cirugía", [
                "Programar y realizar cirugías programadas y de emergencia (ovariohisterectomía, orquiectomía, cesárea, laparotomía exploratoria, hernioplastias, entre otras).",
                "Aplicar el protocolo de asepsia, antisepsia, anestesia y monitoreo transoperatorio.",
                "Verificar lista de chequeo quirúrgico, consentimiento informado y ayuno del paciente.",
                "Supervisar la recuperación postoperatoria y entregar indicaciones al cliente y al auxiliar.",
            ]),
            ("Urgencias y hospitalización", [
                "Atender urgencias durante el turno y como guardia (paro cardiorrespiratorio, traumatismos, intoxicaciones, distocias).",
                "Realizar rondas diarias a pacientes hospitalizados, registrando evolución y ajustando tratamientos.",
                "Decidir altas, derivaciones a especialistas o eutanasia conforme al Código de Ética Veterinaria.",
            ]),
            ("Medicina preventiva", [
                "Aplicar esquemas de vacunación y desparasitación para especies domésticas (perros, gatos, conejos, aves).",
                "Asesorar al cliente en tenencia responsable, nutrición, comportamiento y bienestar animal.",
                "Realizar campañas de esterilización y educación a la comunidad.",
            ]),
            ("Coordinación y comunicación", [
                "Coordinar con el Encargado de Clínica la agenda de cirugías y la disponibilidad de quirófano.",
                "Mantener comunicación fluida con el cliente durante hospitalizaciones y postoperatorios.",
                "Participar en reuniones clínicas, ateneos y discusión de casos complejos.",
            ]),
            ("Sustancias controladas y registros", [
                "Prescribir y administrar opioides, anestésicos y benzodiacepinas conforme a la Ley Orgánica de Drogas y resoluciones del SENAC.",
                "Registras en el libro foliado cada dispensación, administración y descarte con doble firma.",
                "Cumplir con la cadena de custodia y el reporte trimestral al SENAC.",
                "Mantener actualizado el carnet del Colegio de Médicos Veterinarios y el CMVV.",
            ]),
        ],
        "tecnicas": [
            "Diagnóstico clínico, quirúrgico e imagenológico veterinario.",
            "Anestesiología veterinaria y monitoreo transoperatorio.",
            "Interpretación de laboratorio clínico veterinario.",
            "Manejo de equipos: autoclave, rayos X, ecógrafo, monitor multiparámétrico, ventilador mecánico.",
            "Manejo de protocolos de bioseguridad NT-01-2008 y residuos COVENIN 2747-93.",
            "Software de historia clínica electrónica veterinaria.",
        ],
        "blandas": [
            "Empatía y comunicación con clientes en situaciones emocionalmente complejas.",
            "Trabajo en equipo con auxiliares y otros médicos.",
            "Capacidad de decisión bajo presión en urgencias.",
            "Vocación de servicio y ética profesional.",
            "Actualización continua en medicina veterinaria.",
        ],
        "academicas_legales": [
            "Título de Médico Veterinario emitido por universidad reconocida.",
            "Colegiación vigente en el Colegio de Médicos Veterinarios de Venezuela y registro en el CMVV.",
            "Curso de bioseguridad NT-01-2008 vigente.",
            "Curso de manejo de sustancias controladas (SENAC).",
            "Inscripción vigente en IVSS, FAOV, INCES y PMSSO.",
            "Curso de RCP veterinaria y manejo de emergencias (preferencial).",
        ],
        "condiciones": {
            "lugar": (
                "Consultorios, quirófano, hospitalización y laboratorio de la clínica veterinaria "
                "en Av. Francisco de Miranda, Local N° 1, Los Teques, Miranda."
            ),
            "jornada": (
                "Jornada mixta de lunes a sábado, 44 horas semanales con un día libre rotativo "
                "(LOTTT art. 173). Guardias nocturnas y de fin de semana rotativas."
            ),
            "epp": [
                "Uniforme clínico (scrub) institucional, carné identificatorio.",
                "Guantes de nitrilo, mascarilla quirúrgica, gafas de protección y gorro quirúrgico.",
                "Calzado cerrado antideslizante.",
                "Delantal plomado para radiografías.",
                "Bata estéril para cirugía.",
            ],
            "riesgos": [
                "Biológico por exposición a fluidos, secreciones y zoonosis (rabia, leptospirosis, toxoplasmosis, tiña, brucelosis).",
                "Por mordedura, arañazo o patada de animales (Protocolo de Mordeduras/Zoonosis).",
                "Químico por anestésicos inhalatorios (isoflurano), desinfectantes y fármacos citostáticos.",
                "Radiológico por exposición a rayos X (vigilancia dosimétrica NT-02-2008).",
                "Ergonómico por levantamiento de animales y posturas quirúrgicas prolongadas.",
                "Psicosocial por trabajo bajo presión, guardias y manejo de duelos con clientes.",
            ],
        },
    },

    # ---------- 03e AUXILIAR VETERINARIO ----------
    "auxiliar_veterinario": {
        "file": "03e_Funciones_Auxiliar_Veterinario.docx",
        "sublabel": "Auxiliar Veterinario",
        "titulo": "AUXILIAR VETERINARIO",
        "departamento": "Clínica Veterinaria",
        "reporta_a": "Médico(a) Veterinario(a) / Encargado(a) de Clínica",
        "supervisa_a": "No aplica — personal operativo de apoyo",
        "nivel_jerarquico": "Técnico operativo — Nivel 4",
        "tipo_contrato": "Determinado por 3 meses (LOTTT art. 65) — período de prueba 15 días",
        "objetivo": (
            "Apoyar al médico veterinario en consultas, procedimientos, cirugías, "
            "hospitalización y laboratorio, garantizando el manejo adecuado de los pacientes, "
            "el cumplimiento de los protocolos de bioseguridad (NT-01-2008) y residuos "
            "(COVENIN 2747-93), la atención cordial al cliente y el registro completo de "
            "los procedimientos en la historia clínica."
        ),
        "funciones": [
            ("Consultas", [
                "Recibir y pesar al paciente, tomar temperatura, frecuencia cardíaca y respiratoria.",
                "Preparar el consultorio con materiales, limpiar y desinfectar entre pacientes.",
                "Asistir al médico veterinario en la contención y manejo del animal durante el examen.",
                "Registrar signos vitales y antecedentes en la historia clínica electrónica.",
            ]),
            ("Tratamientos", [
                "Administrar medicamentos por vía oral, tópica, subcutánea, intramuscular o intravenosa según indicación médica.",
                "Aplicar vacunas, antiparasitarios y fluidoterapia con doble chequeo de dosis.",
                "Colocar y retirar catéteres intravenosos, vendajes y drenajes.",
                "Apoyar en la colocación de sondas urinarias y esofágicas.",
            ]),
            ("Laboratorio", [
                "Tomar muestras de sangre, orina, heces y secreciones para análisis.",
                "Operar centrifuga, microscopio y analyzadores de laboratorio.",
                "Realizar frotis, tinciones y observación microscópica bajo supervisión.",
                "Registrar resultados en la historia clínica.",
            ]),
            ("Cirugía", [
                "Preparar al paciente: ayuno, bañado, rasurado, voie venosa.",
                "Preparar el quirófano: material estéril, instrumental, anestésicos y monitores.",
                "Asistir al cirujano como instrumentista y circulante.",
                "Monitorear signos vitales transoperatorios bajo supervisión del médico.",
                "Cuidar la recuperación postoperatoria y registrar la evolución.",
            ]),
            ("Hospitalización", [
                "Atender a los pacientes hospitalizados: alimentación, hidratación, paseo, higiene y medicación.",
                "Realizar rondas de control cada turno y registrar evolución en la historia clínica.",
                "Limpiar y desinfectar jaulas, mantas y comederos entre pacientes.",
                "Mantener comunicación con el médico sobre cualquier cambio relevante.",
            ]),
            ("Atención al cliente", [
                "Recibir al cliente y a la mascota en la recepción del área clínica.",
                "Explicar indicaciones postoperatorias, postvacunales y de medicación.",
                "Cobrar servicios del área clínica cuando el encargado no esté disponible.",
                "Atender y escalar quejas o reclamos conforme al procedimiento interno.",
            ]),
            ("Higiene y bioseguridad", [
                "Cumplir y hacer cumplir los protocolos NT-01-2008 (Bioseguridad) y COVENIN 2747-93 (residuos).",
                "Segregar residuos en bolsas rojas, amarillas, negras y verdes según corresponda.",
                "Mantener limpio y desinfectado el quirófano, hospitalización y depósito.",
                "Lavar y esterilizar instrumental (autoclave) y verificar indicadores biológicos.",
            ]),
            ("Políticas y registros", [
                "Cumplir las políticas internas: reglamento, código de conducta, confidencialidad.",
                "Registrar todos los procedimientos y medicaciones en la historia clínica.",
                "Aplicar el procedimiento de reporte de incidentes ante cualquier accidente o evento adverso.",
                "Mantener confidencialidad sobre historias clínicas y datos de clientes (LOPDP).",
            ]),
        ],
        "tecnicas": [
            "Manejo y contención de animales domésticos (perros, gatos, conejos, aves).",
            "Técnicas de venopunción, cateterismo y administración de medicamentos.",
            "Operación de autoclave, microscopio y centrifuga.",
            "Conocimiento de bioseguridad NT-01-2008 y manejo de residuos.",
            "Software básico de historia clínica veterinaria.",
            "Técnicas de asepsia y antisepsia.",
        ],
        "blandas": [
            "Empatía y trato cordial con animales y sus dueños.",
            "Trabajo en equipo y disposición para aprender.",
            "Capacidad de organización y atención al detalle.",
            "Calma bajo presión, especialmente en urgencias.",
            "Responsabilidad y puntualidad.",
        ],
        "academicas_legales": [
            "TSU en Medicina Veterinaria, Veterinaria (trunco), Enfermería Veterinaria o afines (preferencial).",
            "Experiencia mínima de seis (6) meses en clínica veterinaria (aceptable pasantía).",
            "Curso de bioseguridad NT-01-2008 (preferencial).",
            "Inscripción vigente en IVSS, FAOV, INCES y PMSSO.",
        ],
        "condiciones": {
            "lugar": (
                "Consultorios, quirófano, hospitalización y laboratorio de la clínica veterinaria "
                "en Av. Francisco de Miranda, Local N° 1, Los Teques, Miranda."
            ),
            "jornada": (
                "Jornada mixta de lunes a sábado, 44 horas semanales con un día libre rotativo "
                "(LOTTT art. 173). Posibilidad de cubrir guardias de fin de semana para hospitalizados."
            ),
            "epp": [
                "Uniforme clínico (scrub) institucional, carné identificatorio.",
                "Guantes de nitrilo, mascarilla quirúrgica, gafas de protección.",
                "Calzado cerrado antideslizante.",
                "Gorro quirúrgico y bata estéril en sala de operaciones.",
            ],
            "riesgos": [
                "Biológico por exposición a fluidos, secreciones y zoonosis (rabia, leptospirosis, tiña, sarna).",
                "Por mordedura, arañazo o patada de animales (Protocolo de Mordeduras/Zoonosis).",
                "Ergonómico por levantamiento de animales y posturas prolongadas.",
                "Químico por desinfectantes, anestésicos y fármacos.",
                "Cortopunzante por agujas, bisturíes yrotura de material de vidrio.",
                "Psicosocial por trabajo bajo presión y carga emocional.",
            ],
        },
    },

    # ---------- 03f DOG GROOMER ----------
    "dog_groomer": {
        "file": "03f_Funciones_Dog_Groomer.docx",
        "sublabel": "Dog Groomer (Peluquero Canino)",
        "titulo": "DOG GROOMER (PELUQUERO CANINO)",
        "departamento": "Peluquería Canina",
        "reporta_a": "Encargado(a) de Clínica / Gerente General",
        "supervisa_a": "No aplica — personal operativo especializado",
        "nivel_jerarquico": "Técnico especializado — Nivel 3",
        "tipo_contrato": "Tiempo indeterminado (LOTTT art. 65) — período de prueba 30 días",
        "objetivo": (
            "Brindar servicios de peluquería, baño, corte y arreglo estético de mascotas — "
            "principalmente caninos y felinos — asegurando el bienestar animal, la "
            "seguridad del paciente y del personal, la calidad del servicio y el cumplimiento "
            "de los protocolos de bioseguridad NT-01-2008 y de manejo de residuos "
            "COVENIN 2747-93."
        ),
        "funciones": [
            ("Cortes y peluquería", [
                "Realizar cortes de raza (poodle, schnauzer, yorkshire, shih tzu, terriers, etc.) y cortes higiénicos.",
                "Aplicar técnicas de tijera, máquina y rastrillo según el tipo de pelo y la raza.",
                "Asesorar al cliente sobre el corte más adecuado según el estilo de vida del animal.",
            ]),
            ("Baño y secado", [
                "Bañar con champús y acondicionadores adecuados al tipo de piel y pelaje.",
                "Aplicar tratamientos antipulgas y antiparasitarios tópicos según indicación.",
                "Secar con toalla y secador profesional, cepillando y desenredando el pelaje.",
                "Limpiar oídos y cortar uñas con técnica segura.",
            ]),
            ("Higiene complementaria", [
                "Limpieza de glándulas anales (bajo indicación del médico veterinario).",
                "Limpieza dental básica con ultrasonido o cepillado.",
                "Aplicación de perfumes, laqueado y terminaciones estéticas.",
            ]),
            ("Evaluación y manejo animal", [
                "Evaluar el estado general del animal al ingreso: piel, orejas, ojos, aparato locomotor, comportamiento.",
                "Reportar al médico veterinario cualquier hallazgo anómalo (masas, dermatitis, otitis, cojeras).",
                "Aplicar técnicas de contención y manejo ético (refuerzo positivo, bozal si es necesario).",
                "Decidir si un animal no está en condiciones de ser peluqueado (agresividad, estrés, salud).",
            ]),
            ("Limpieza y mantenimiento de equipos", [
                "Desinfectar bañera, mesa de peluquería, jaulas, tijeras, máquinas y cepillos entre pacientes (NT-01-2008).",
                "Segregar pelos, residuos y material biológico en bolsas conforme COVENIN 2747-93.",
                "Mantener secadores, máquinas y cuchillas en buen estado, con afilado periódico.",
                "Reportar fallas técnicas al Encargado de Clínica.",
            ]),
        ],
        "tecnicas": [
            "Técnicas de corte por raza y corte higiénico (máquina, tijera, rastrillo).",
            "Manejo y contención ética de caninos y felinos.",
            "Conocimiento de tipos de piel, pelaje y dermatología básica veterinaria.",
            "Operación de secadores profesionales, máquinas y bañeras.",
            "Protocolos de bioseguridad NT-01-2008 y manejo de residuos.",
        ],
        "blandas": [
            "Paciencia y empatía con los animales.",
            "Trato cordial con los clientes.",
            "Capacidad de observación y atención al detalle estético.",
            "Calma y manejo de animales estresados o agresivos.",
            "Trabajo en equipo con el área clínica.",
        ],
        "academicas_legales": [
            "Curso certificado de Dog Groomer o peluquería canina (preferencial).",
            "Experiencia mínima de un (1) año como peluquero canino.",
            "Curso básico de bioseguridad NT-01-2008 (preferencial).",
            "Inscripción vigente en IVSS, FAOV, INCES y PMSSO.",
        ],
        "condiciones": {
            "lugar": (
                "Área de peluquería canina en la sede de Av. Francisco de Miranda, Local N° 1, "
                "Los Teques, Miranda. Incluye bañera, mesa de peluquería, jaulas de espera y "
                "depósito de insumos."
            ),
            "jornada": (
                "Jornada mixta de lunes a sábado, 44 horas semanales con un día libre rotativo "
                "(LOTTT art. 173). Posibilidad de jornada extendida en temporadas altas."
            ),
            "epp": [
                "Uniforme impermeable (delantal) y carné identificatorio.",
                "Guantes de nitrilo, mascarilla quirúrgica y gafas de protección.",
                "Calzado cerrado antideslizante y resistente al agua.",
                "Delantal térmico y guantes térmicos para secado.",
                "Tapones auditivos por exposición a secadores.",
            ],
            "riesgos": [
                "Biológico por exposición a fluidos, pelos y zoonosis (tiña, sarna, leptospirosis).",
                "Por mordedura, arañazo o patada de animales (Protocolo de Mordeduras/Zoonosis).",
                "Ergonómico por levantamiento de animales, posturas prolongadas y movimientos repetitivos.",
                "Eléctrico por uso de secadores y máquinas en ambiente húmedo.",
                "Acústico por exposición a secadores y ladridos.",
                "Químico por champús, antiparasitarios y desinfectantes.",
            ],
        },
    },

    # ---------- 03g ASISTENTE DE VENTAS ----------
    "asistente_ventas": {
        "file": "03g_Funciones_Asistente_Ventas.docx",
        "sublabel": "Asistente de Ventas",
        "titulo": "ASISTENTE DE VENTAS",
        "departamento": "Tienda de Mascotas",
        "reporta_a": "Encargado(a) de Tienda / Gerente General",
        "supervisa_a": "No aplica — personal operativo de apoyo",
        "nivel_jerarquico": "Operativo — Nivel 4",
        "tipo_contrato": "Determinado por 3 meses (LOTTT art. 65) — período de prueba 15 días",
        "objetivo": (
            "Atender a los clientes de la tienda de mascotas, asesorarlos en la selección de "
            "productos, realizar transacciones de venta, mantener el orden y la exhibición de "
            "la mercancía, apoyar la carga y descarga de productos, y cumplir las políticas "
            "internas de ALIKA PETS, garantizando una experiencia de compra satisfactoria."
        ),
        "funciones": [
            ("Atención al cliente", [
                "Saludar y recibir al cliente al ingresar a la tienda.",
                "Asesorar en la selección de alimentos, accesorios, productos de higiene y antiparasitarios según la especie, edad y condición de la mascota.",
                "Informar sobre promociones, programas de fidelización y disponibilidad de productos.",
                "Atender y escalar quejas o reclamos al Encargado de Tienda.",
            ]),
            ("Transacciones", [
                "Cobrar las ventas con efectivo, punto de venta, transferencia o tarjeta.",
                "Emitir facturas y comprobantes, y entregar el vuelto correcto.",
                "Apoyar en el arqueo de caja al cierre, cuando el Encargado lo solicite.",
                "Registrar todas las operaciones en el sistema POS.",
            ]),
            ("Carga y descarga", [
                "Recibir la mercancía de proveedores, verificar cantidades y condiciones.",
                "Trasladar cajas a la bodega y a las góndolas, cuidando la cadena de frío.",
                "Estibar y organizar los productos conforme al planograma.",
                "Rotular precios y exhibir productos.",
            ]),
            ("Mantenimiento del área", [
                "Mantener limpio y ordenado el mostrador, las góndolas, las neveras y la bodega.",
                "Revisar fechas de vencimiento y reportar productos por rotar.",
                "Reportar fallas técnicas, daños o necesidades de mantenimiento al Encargado.",
            ]),
            ("Registro y control", [
                "Registrar entradas y salidas de mercancía en el sistema.",
                "Comunicar al Encargado los niveles bajos de inventario.",
                "Etiquetar y separar productos dañados o devueltos.",
                "Participar en los inventarios físicos mensuales.",
            ]),
            ("Cumplimiento de normas", [
                "Cumplir las políticas internas: reglamento, código de conducta, confidencialidad.",
                "Aplicar el procedimiento de reporte de incidentes ante robo, asalto o accidente.",
                "Velar por el uso correcto del uniforme y la presentación personal.",
                "Mantener la confidencialidad de la información de clientes y la empresa (LOPDP).",
            ]),
        ],
        "tecnicas": [
            "Operación de sistemas POS, lector de código de barras y terminal bancario.",
            "Conocimiento básico de productos para mascotas: alimentación, accesorios, antiparasitarios.",
            "Manejo básico de inventarios y merchandising.",
            "Operación de neveras y manejo de cadena de frío.",
            "Técnicas de atención al cliente y ventas.",
        ],
        "blandas": [
            "Cordialidad, empatía y orientación al cliente.",
            "Honestidad y manejo ético del dinero.",
            "Capacidad de trabajo en equipo.",
            "Organización y atención al detalle.",
            "Disposición para aprender sobre productos y mascotas.",
        ],
        "academicas_legales": [
            "Bachiller (preferencial TSU en Comercio, Mercadeo o afines).",
            "Experiencia mínima de seis (6) meses en atención al cliente o retail (aceptable pasantía).",
            "Inscripción vigente en IVSS, FAOV, INCES y PMSSO.",
            "Curso básico de seguridad y salud laboral (LOPCYMAT).",
        ],
        "condiciones": {
            "lugar": (
                "Tienda de mascotas en la sede de Av. Francisco de Miranda, Local N° 1, "
                "Los Teques, Miranda. Incluye mostrador, góndolas, bodega y caja."
            ),
            "jornada": (
                "Jornada mixta de lunes a sábado, 44 horas semanales con un día libre rotativo "
                "(LOTTT art. 173). Posibilidad de jornada extendida en temporadas altas."
            ),
            "epp": [
                "Uniforme corporativo y carné identificatorio.",
                "Calzado cerrado antideslizante.",
                "Guantes para manipulación de bodega y productos de limpieza.",
            ],
            "riesgos": [
                "Ergonómico por levantamiento de cajas, posturas prolongadas y traslados.",
                "Riesgo de asalto o hurto al manejar caja (Protocolo de Reporte de Incidentes).",
                "Caídas al mismo nivel por derrames o piso mojado.",
                "Cortes o pinchazos con cajas, esquineros o cintas.",
                "Químico por manipulación de productos de limpieza y antiparasitarios.",
            ],
        },
    },
}


# ============================================================
# Generador principal de un cargo
# ============================================================
def generar_cargo(cargo_key, datos):
    """Genera el .docx para un cargo."""
    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)

    # Membrete
    add_membrete(doc,
                 doc_label="DESCRIPCIÓN DE CARGO",
                 doc_sublabel=datos["sublabel"],
                 version="Versión 3.0  ·  RR.HH.")

    # Título del documento
    add_doc_title(doc, f"DESCRIPCIÓN DE CARGO\n{datos['titulo']}")

    # ===== Sección 1: Identificación del Cargo =====
    add_section(doc, "1.  IDENTIFICACIÓN DEL CARGO")
    add_identificacion_table(
        doc,
        titulo=datos["titulo"],
        departamento=datos["departamento"],
        reporta_a=datos["reporta_a"],
        supervisa_a=datos["supervisa_a"],
        nivel_jerarquico=datos["nivel_jerarquico"],
        tipo_contrato=datos["tipo_contrato"],
    )

    # ===== Sección 2: Objetivo General =====
    add_section(doc, "2.  OBJETIVO GENERAL DEL CARGO")
    add_para(doc, datos["objetivo"], size=10, space_after=6)

    # ===== Sección 3: Funciones Específicas =====
    add_section(doc, "3.  FUNCIONES ESPECÍFICAS")
    add_para(doc,
             "Las funciones del cargo se organizan por áreas de desempeño. Cada área "
             "agrupa las responsabilidades específicas que el(la) trabajador(a) deberá "
             "ejecutar durante su jornada laboral:",
             size=10, italic=True, color=GRAY_TEXT, space_after=4)
    for area_title, bullets in datos["funciones"]:
        add_funciones_area(doc, area_title, bullets)

    # ===== Sección 4: Competencias y Requisitos =====
    add_section(doc, "4.  COMPETENCIAS Y REQUISITOS")
    add_competencias_block(doc,
                           datos["tecnicas"],
                           datos["blandas"],
                           datos["academicas_legales"])

    # ===== Sección 5: Condiciones de Trabajo y Riesgos =====
    add_section(doc, "5.  CONDICIONES DE TRABAJO Y RIESGOS")
    cond = datos["condiciones"]
    add_condiciones_block(doc,
                          lugar=cond["lugar"],
                          jornada=cond["jornada"],
                          epp=cond["epp"],
                          riesgos=cond["riesgos"])

    # ===== Carta de Recepción y Aceptación =====
    doc_name = f"Descripción de Cargo — {datos['titulo']}"
    # Override the reception letter text to include "designado(a) para el cargo de [CARGO]"
    _add_reception_letter_custom(doc, doc_name, datos["titulo"])

    # ===== Footer =====
    footer_label = f"Descripción Cargo — {datos['sublabel']} v3.0"
    add_footer(section, footer_label)

    # Guardar
    out_path = os.path.join(OUT_DIR, datos["file"])
    doc.save(out_path)
    return out_path


# ============================================================
# Carta de recepción con texto personalizado
# (igual que add_reception_letter pero con la mención al cargo)
# ============================================================
def _add_reception_letter_custom(doc, doc_name_full, cargo_titulo):
    """Carta de recepción con mención 'designado(a) para el cargo de [CARGO]'."""
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(8)
    sep.paragraph_format.space_after = Pt(6)
    add_hr(sep, color="0F766E", sz="8")

    add_section(doc, "CARTA DE RECEPCIÓN Y ACEPTACIÓN")

    add_para(doc,
        f"Yo, _____________________________________________, titular de la cédula de "
        f"identidad N° V-___________________, en mi condición de trabajador(a) de "
        f"GRUPO CAVAL 1003, C.A. (ALIKA PETS), designado(a) para el cargo de "
        f"{cargo_titulo}, declaro por medio de la presente que:",
        size=10, space_after=4)

    add_bullet(doc,
        f"He recibido copia íntegra y legible del documento titulado "
        f"\u00ab{doc_name_full}\u00bb, así como las explicaciones verbales necesarias "
        f"para su correcta comprensión.",
        bold_lead="PRIMERO:  ")

    add_bullet(doc,
        f"He leído en su totalidad el contenido del mencionado documento, comprendo mis "
        f"funciones, responsabilidades, competencias y condiciones de trabajo, y me "
        f"comprometo a cumplirlas fielmente durante toda la vigencia de mi relación "
        f"laboral con la empresa.",
        bold_lead="SEGUNDO:  ")

    add_bullet(doc,
        f"Entiendo que el incumplimiento de las disposiciones aquí contenidas podrá "
        f"dar lugar a las sanciones disciplinarias previstas en el Reglamento Interno "
        f"de Trabajo y, según la gravedad, a la terminación de la relación laboral por "
        f"causa justificada conforme al artículo 79 de la Ley Orgánica del Trabajo, los "
        f"Trabajadores y las Trabajadoras (LOTTT).",
        bold_lead="TERCERO:  ")

    add_bullet(doc,
        f"Acepto que la presente firma constituye prueba fehaciente de la entrega y "
        f"recepción del documento, renunciando a alegar desconocimiento de su contenido.",
        bold_lead="CUARTO:  ")

    add_para(doc, "", size=6, space_after=2)
    add_para(doc,
        "En fe de lo cual firmo la presente carta en la ciudad de Los Teques, Estado "
        "Miranda, a los ____ días del mes de ________________ de ________.",
        size=10, space_after=10)

    # Firma: Trabajador + Directora Gerente (rol operativo)
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])


# ============================================================
# Main
# ============================================================
def main():
    print("=" * 70)
    print("gen_funciones_por_cargo.py — ALIKA PETS / Grupo Caval 1003, C.A.")
    print("=" * 70)
    generated = []
    for key, datos in CARGOS.items():
        try:
            path = generar_cargo(key, datos)
            size = os.path.getsize(path) / 1024
            print(f"  ✓ {datos['file']:50s}  {size:6.1f} KB")
            generated.append(path)
        except Exception as e:
            print(f"  ✗ Error generando {datos['file']}: {e}")
            import traceback; traceback.print_exc()
    print(f"\nTotal generados: {len(generated)} / {len(CARGOS)}")
    print(f"Carpeta: {OUT_DIR}")
    return generated


if __name__ == "__main__":
    main()
