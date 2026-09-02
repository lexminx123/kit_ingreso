"""
gen_honorarios_profesionales.py — Genera 3 documentos:
  1. Concepto_Honorarios_Profesionales.docx (justificación jurídica técnica)
  2. Contrato_Honorarios_Medico_Veterinario.docx
  3. Contrato_Honorarios_Dog_Groomer.docx

Los honorarios profesionales se discriminan en USD 250/mes:
  - USD 50: Honorarios base (servicios profesionales)
  - USD 80: Bono de alimentación (no constituye salario)
  - USD 40: Bono de Buen Vivir
  - USD 40: Bono de Transporte
  - USD 40: Otros beneficios no salariales

Marco jurídico venezolano:
  - LOTTT Art. 22, 26, 40 (definición de relación de dependencia vs. trabajo autónomo)
  - Ley de Ejercicio de la Medicina Veterinaria (profesión liberal)
  - Código de Comercio (servicios profesionales independientes)
  - Ley de Impuesto sobre la Renta (retención del 3% honorarios profesionales)
  - Jurisprudencia TSJ sobre honorarios profesionales
"""
import os, sys
sys.path.insert(0, "/home/z/my-project/output")
from _common import *
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# ============================================================
# 1. CONCEPTO TÉCNICO-JURÍDICO
# ============================================================
def gen_concepto_honorarios():
    out = "/home/z/my-project/output/Concepto_Honorarios_Profesionales.docx"

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.2)
    add_membrete(doc, "CONCEPTO JURÍDICO", "Honorarios profesionales",
                 version="Versión 1.0  ·  Dirección")
    add_doc_title(doc, "CONCEPTO TÉCNICO-JURÍDICO SOBRE HONORARIOS PROFESIONALES")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Médico(a) Veterinario(a) y Dog Groomer — ALIKA PETS / Grupo Caval 1003, C.A.")
    style_run(r, size=10, italic=True, color=GRAY_TEXT)

    # Encabezado formal de concepto
    add_para(doc,
        "Ciudad y fecha:  Los Teques, ____ de ________________ de ________",
        size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

    add_para(doc,
        "Ciudadana\nESNATLIM ELENA SIMOZA\nDirectora Gerente\nGRUPO CAVAL 1003, C.A. (ALIKA PETS)\nSu despacho.-",
        size=10, space_after=8)

    add_para(doc, "Ref: Concepto técnico-jurídico sobre la figura de honorarios profesionales.",
             size=10, bold=True, space_after=10)

    add_para(doc,
        "Recibida su solicitud, procedemos a emitir el presente concepto técnico-jurídico con el "
        "objeto de establecer la viabilidad, alcance y características de la figura contractual de "
        "HONORARIOS PROFESIONALES para los servicios del Médico(a) Veterinario(a) y del Dog "
        "Groomer en ALIKA PETS.", size=10, space_after=8)

    # ============================================================
    # I. ANTECEDENTES
    # ============================================================
    add_section(doc, "I. ANTECEDENTES")
    add_para(doc,
        "GRUPO CAVAL 1003, C.A., identificada con RIF N° J501662533, marca comercial ALIKA PETS, "
        "dedicada a la actividad de clínica veterinaria, tienda de mascotas y peluquería canina, "
        "con domicilio en Av. Francisco de Miranda, Local N° 1, Los Teques, Estado Miranda, "
        "requiere formalizar la vinculación con dos (2) perfiles profesionales con "
        "características particulares: (i) Médico(a) Veterinario(a), profesión liberal regulada "
        "por la Ley de Ejercicio de la Medicina Veterinaria; y (ii) Dog Groomer (peluquero(a) "
        "canino), oficio técnico especializado con autonomía en la ejecución.")

    add_para(doc,
        "Ambos perfiles desarrollan labores que requieren un alto grado de autonomía técnica, "
        "juicio profesional independiente y disposición de tiempo flexible, lo que hace "
        "conveniente evaluar la figura de honorarios profesionales como alternativa a la "
        "relación de trabajo dependiente prevista en la LOTTT.", space_after=8)

    # ============================================================
    # II. MARCO JURÍDICO APLICABLE
    # ============================================================
    add_section(doc, "II. MARCO JURÍDICO APLICABLE")
    add_para(doc, "La figura de honorarios profesionales encuentra fundamento en:", space_after=3)
    add_bullet(doc,
        "Reconoce la figura del trabajador no dependiente (Art. 40) y excluye del régimen "
        "laboral a quienes presten servicios sin subordinación ni dependencia.",
        bold_lead="Ley Orgánica del Trabajo, los Trabajadores y las Trabajadoras (LOTTT) Art. 22 y 40:  ")
    add_bullet(doc,
        "Las personas que ejerzan profesiones liberales o técnicas mediante el cobro de "
        "honorarios no están sujetas al régimen de la LOTTT, salvo que demuestren la existencia "
        "de la relación de dependencia (subordinación, horario, exclusividad, salario).",
        bold_lead="Ley de Ejercicio de la Medicina Veterinaria:  ")
    add_bullet(doc,
        "Regula la prestación de servicios profesionales independientes y la responsabilidad "
        "civil del profesional.",
        bold_lead="Código de Comercio y Código Civil:  ")
    add_bullet(doc,
        "Los honorarios profesionales se gravan con retención del 3% en cabeza del contratista, "
        "y deben declararse en el Impuesto sobre la Renta (ISLR) del profesional.",
        bold_lead="Ley de Impuesto sobre la Renta (ISLR) Art. 27:  ")
    add_bullet(doc,
        "El profesional independiente puede estar inscrito en el Régimen de Incapacitados o "
        "como contribuyente formal del IVA, según el caso.",
        bold_lead="Providencias SENIAT y normativa tributaria:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # ============================================================
    # III. DISTINCIÓN ENTRE RELACIÓN DE TRABAJO Y HONORARIOS
    # ============================================================
    add_section(doc, "III. DISTINCIÓN ENTRE RELACIÓN DE TRABAJO Y HONORARIOS PROFESIONALES")
    add_para(doc,
        "La jurisprudencia venezolana (TSJ Sala Social, sent. N° 0224 del 26/04/2011, entre "
        "otras) ha establecido que la diferencia esencial entre una relación de trabajo "
        "dependiente y un contrato de honorarios profesionales radica en la AUSENCIA o "
        "PRESENCIA de los siguientes elementos:", space_after=4)

    tbl = doc.add_table(rows=8, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(4.5), Cm(6.5), Cm(6.0)]):
        tbl.columns[i].width = w
    headers = ["ELEMENTO", "RELACIÓN DE TRABAJO (LOTTT)", "HONORARIOS PROFESIONALES"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = [Cm(4.5), Cm(6.5), Cm(6.0)][ci]
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)

    distinciones = [
        ("Subordinación",
         "El trabajador recibe órdenes directas del empleador.",
         "El profesional ejecuta con autonomía técnica, sin subordinación."),
        ("Horario",
         "Jornada fija (Art. 168-180 LOTTT).",
         "Horario flexible, definido por el profesional para cumplir los servicios."),
        ("Exclusividad",
         "El trabajador no puede atender otros clientes.",
         "El profesional puede atender otros clientes (salvo cláusula de no competencia parcial)."),
        ("Remuneración",
         "Salario mensual/quincenal, pagadero conforme a LOTTT.",
         "Honorarios por servicio, mensualidad, o porcentaje. No es salario."),
        ("Riesgo",
         "El empleador asume el riesgo del negocio.",
         "El profesional asume su propio riesgo técnico y civil."),
        ("Prestaciones",
         "Prestaciones sociales, utilidades, vacaciones, bono vacacional.",
         "No generan prestaciones sociales. Solo el pago acordado."),
        ("Terminación",
         "Causales del Art. 79 LOTTT, indemnizaciones.",
         "Resolución por incumplimiento contractual, sin indemnizaciones laborales."),
    ]
    for ri, (el, rt, hp) in enumerate(distinciones, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        c = tbl.rows[ri].cells[0]; c.width = Cm(4.5)
        write_cell(c, el, size=9, bold=True, color=TEAL_DARK, bg=bg)
        c = tbl.rows[ri].cells[1]; c.width = Cm(6.5)
        write_cell(c, rt, size=9, bg=bg)
        c = tbl.rows[ri].cells[2]; c.width = Cm(6.0)
        write_cell(c, hp, size=9, bg=bg)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_para(doc,
        "CRITERIO RECTOR (TSJ): La simple denominación del contrato no determina la naturaleza "
        "de la relación. Lo que define si existe relación de trabajo o contrato de honorarios es "
        "la REALIDAD DE LOS HECHOS, particularmente la existencia o no de los elementos de "
        "subordinación, horario fijo y exclusividad. Un contrato mal denominado como de "
        "honorarios, que en la práctica configure subordinación, será declarado RELACIÓN DE "
        "TRABAJO por los tribunales con todas las consecuencias laborales (prestaciones, "
        "antigüedad, indemnizaciones).",
        size=10, italic=True, color=GRAY_TEXT, space_after=8)

    # ============================================================
    # IV. RECOMENDACIÓN TÉCNICA
    # ============================================================
    add_section(doc, "IV. RECOMENDACIÓN TÉCNICA")
    add_para(doc,
        "Con base en las particularidades de cada perfil, se emiten las siguientes "
        "recomendaciones:", space_after=4)

    add_para(doc, "IV.1. Médico(a) Veterinario(a)", size=11, bold=True, color=TEAL_DARK, space_after=3)
    add_para(doc,
        "RECOMENDABLE la figura de HONORARIOS PROFESIONALES. La profesión veterinaria es una "
        "profesión liberal reconocida por la Ley de Ejercicio de la Medicina Veterinaria, "
        "con autonomía técnica plena. El MV decide diagnóstico, tratamiento y procedimientos; "
        "la clínica solo provee el espacio, equipos y pacientes. La responsabilidad civil "
        "profesional es del MV (con su seguro de responsabilidad civil profesional). El MV "
        "puede atender otros consultorios, clínicas o particulares fuera del horario pactado.",
        size=10, space_after=4)

    add_para(doc,
        "REQUISITOS para que la figura sea válida:", size=10, bold=True, space_after=2)
    add_bullet(doc, "El MV debe estar inscrito en el Colegio de Médicos Veterinarios de Venezuela (CMV).")
    add_bullet(doc, "El MV debe tener RIF propio y emitir facturas a nombre de ALIKA PETS.")
    add_bullet(doc, "El MV declara sus ingresos en ISLR como profesional independiente.")
    add_bullet(doc, "ALIKA PETS retiene el 3% de ISLR sobre los honorarios pagados (Art. 27 ISLR).")
    add_bullet(doc, "El horario debe ser flexible (turnos acordados, no jornada fija de 8h).")
    add_bullet(doc, "El MV puede atender otros consultorios (cláusula de NO EXCLUSIVIDAD absoluta).")
    add_bullet(doc, "El MV asume su propio seguro de responsabilidad civil profesional.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc, "IV.2. Dog Groomer (Peluquero Canino)", size=11, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_para(doc,
        "VIABLE la figura de HONORARIOS PROFESIONALES, con ciertas reservas. El oficio de "
        "peluquero canino no es una profesión liberal regulada por una ley de ejercicio "
        "específica, pero el groomer ejecuta sus servicios con autonomía técnica (cortes, "
        "baños, manejo del animal) y puede atender a otros clientes. Se recomienda documentar "
        "cuidadosamente la ausencia de los elementos de subordinación para evitar que un juez "
        "re-califique la relación como laboral dependiente.",
        size=10, space_after=4)

    add_para(doc,
        "REQUISITOS para que la figura sea válida:", size=10, bold=True, space_after=2)
    add_bullet(doc, "El groomer debe tener RIF propio y emitir recibos/facturas a ALIKA PETS.")
    add_bullet(doc, "El groomer declara sus ingresos en ISLR (puede acogerse al régimen simplificado).")
    add_bullet(doc, "ALIKA PETS retiene el 1% (no profesional técnico con título) o el 3% si es profesional.")
    add_bullet(doc, "El horario debe ser flexible (citas acordadas, no jornada fija).")
    add_bullet(doc, "El groomer puede atender otros clientes en su tiempo no pactado con ALIKA PETS.")
    add_bullet(doc, "El groomer debe llevar sus propias herramientas (tijeras, cuchillas) salvo acuerdo.")
    add_bullet(doc, "El groomer asume el riesgo de su actividad (lesiones por mordeduras, cortes).")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # ============================================================
    # V. DISCRIMINACIÓN DE HONORARIOS Y BENEFICIOS
    # ============================================================
    add_section(doc, "V. DISCRIMINACIÓN DE HONORARIOS Y BENEFICIOS")
    add_para(doc,
        "El monto mensual de USD 250,00 que ALIKA PETS pagará al profesional se discrimina así:",
        size=10, space_after=4)

    tbl = doc.add_table(rows=7, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(0.6), Cm(6.5), Cm(2.0), Cm(2.0), Cm(5.5)]):
        tbl.columns[i].width = w
    headers = ["#", "CONCEPTO", "MENSUAL", "QUINCENAL", "NATURALEZA"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = [Cm(0.6), Cm(6.5), Cm(2.0), Cm(2.0), Cm(5.5)][ci]
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    conceptos = [
        ("1", "Honorarios profesionales base", "$ 50,00", "$ 25,00",
         "HONORARIOS — Contraprestación por servicios profesionales."),
        ("2", "Bono de Alimentación (Cestaticket)", "$ 80,00", "$ 40,00",
         "BENEFICIO — Asistencia alimentaria. No es salario en honorarios."),
        ("3", "Bono de Buen Vivir", "$ 40,00", "$ 20,00",
         "BENEFICIO NO REMUNERATIVO — Bono extraordinario asistencial."),
        ("4", "Bono de Transporte", "$ 40,00", "$ 20,00",
         "BENEFICIO NO REMUNERATIVO — Reembolso de gastos de traslado."),
        ("5", "Otros beneficios", "$ 40,00", "$ 20,00",
         "BENEFICIO NO REMUNERATIVO — Ayudas, incentivos, dotación."),
    ]
    for ri, (n, con, mens, quinc, nat) in enumerate(conceptos, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        c = tbl.rows[ri].cells[0]; c.width = Cm(0.6)
        write_cell(c, n, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[1]; c.width = Cm(6.5)
        write_cell(c, con, size=9, bold=True, color=TEAL_DARK, bg=bg)
        c = tbl.rows[ri].cells[2]; c.width = Cm(2.0)
        write_cell(c, mens, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[3]; c.width = Cm(2.0)
        write_cell(c, quinc, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[4]; c.width = Cm(5.5)
        write_cell(c, nat, size=8, italic=True, color=GRAY_TEXT, bg=bg)
    # Fila total
    c = tbl.rows[6].cells[0]; c.width = Cm(0.6)
    c.text = ""
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[1]; c.width = Cm(6.5)
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run("TOTAL MENSUAL")
    style_run(r, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[2]; c.width = Cm(2.0)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 250,00")
    style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[3]; c.width = Cm(2.0)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 125,00")
    style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[4]; c.width = Cm(5.5)
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run("Pago quincenal")
    style_run(r, size=8, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    add_para(doc,
        "IMPORTANTE: Por tratarse de HONORARIOS PROFESIONALES y NO de relación de trabajo, "
        "estos pagos NO generan prestaciones sociales, antigüedad, utilidades, vacaciones, "
        "bono vacacional ni indemnizaciones laborales. El profesional solo tiene derecho al "
        "pago de los honorarios y beneficios acordados mientras dure el contrato.",
        size=10, bold=True, color=RED_CRIT, space_after=8)

    # ============================================================
    # VI. OBLIGACIONES TRIBUTARIAS
    # ============================================================
    add_section(doc, "VI. OBLIGACIONES TRIBUTARIAS")
    add_para(doc, "ALIKA PETS (el contratante) debe:", size=10, bold=True, color=TEAL_DARK, space_after=3)
    add_bullet(doc, "Retener el 3% del ISLR sobre los honorarios pagados al Médico Veterinario (profesional universitario).")
    add_bullet(doc, "Retener el 1% del ISLR sobre los honorarios pagados al Dog Groomer (no profesional universitario, encomendado).")
    add_bullet(doc, "Emitir comprobante de retención mensual y entregarlo al profesional.")
    add_bullet(doc, "Declarar y enterar las retenciones al SENIAT dentro de los primeros 15 días del mes siguiente.")
    add_bullet(doc, "Llevar registro de pagos en libro de compras/ventas.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc, "El profesional contratado debe:", size=10, bold=True, color=TEAL_DARK, space_after=3)
    add_bullet(doc, "Tener RIF propio vigente.")
    add_bullet(doc, "Emitir factura o recibo por cada pago recibido (consecutivos, con N° de control, RIF).")
    add_bullet(doc, "Declarar y pagar el ISLR anualmente (declaración definitiva antes del 31/03).")
    add_bullet(doc, "Si es profesional universitario (MV), inscribirse en el CMV y mantener colegiatura vigente.")
    add_bullet(doc, "El MV debe tener seguro de responsabilidad civil profesional.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

    # ============================================================
    # VII. RIESGOS Y CAUTELAS
    # ============================================================
    add_section(doc, "VII. RIESGOS Y CAUTELAS")
    add_para(doc,
        "Se identifican los siguientes riesgos jurídicos que deben ser mitigados con la "
        "correcta ejecución del contrato:", space_after=4)
    add_bullet(doc,
        "Si en la práctica se configuran los elementos de subordinación (órdenes directas "
        "continuadas, horario fijo, prohibición de atender otros clientes), el juez puede "
        "re-calificar la relación como laboral dependiente, con todas las consecuencias "
        "(prestaciones, indemnizaciones, multas).",
        bold_lead="Riesgo de re-calificación laboral:  ")
    add_bullet(doc,
        "Si ALIKA PETS no retiene el ISLR correspondiente, será responsable solidaria por "
        "el impuesto no retenido más intereses y multas.",
        bold_lead="Riesgo tributario:  ")
    add_bullet(doc,
        "El MV responde civilmente por su actuar profesional (mala praxis). El groomer "
        "responde por lesiones a las mascotas. ALIKA PETS responde subsidiariamente solo si "
        "se demuestra culpa de la empresa (ej. equipos defectuosos).",
        bold_lead="Riesgo de responsabilidad civil:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)

    add_para(doc,
        "Para mitigar estos riesgos, los contratos deben: (i) dejar constancia expresa de la "
        "ausencia de subordinación; (ii) fijar horarios por turnos acordados, no jornada fija; "
        "(iii) reconocer la libertad de atender otros clientes; (iv) exigir facturación mensual "
        "del profesional; (v) retener el ISLR correspondiente; y (vi) documentar la autonomía "
        "técnica del profesional en la ejecución de los servicios.",
        size=10, italic=True, color=GRAY_TEXT, space_after=8)

    # ============================================================
    # VIII. CONCLUSIÓN
    # ============================================================
    add_section(doc, "VIII. CONCLUSIÓN")
    add_para(doc,
        "Es JURÍDICAMENTE VIABLE la figura de HONORARIOS PROFESIONALES para vincular a ALIKA "
        "PETS con el Médico(a) Veterinario(a) y el Dog Groomer, siempre que se cumplan "
        "estrictamente los siguientes requisitos:")
    add_bullet(doc, "Suscripción de contrato de honorarios profesionales por escrito, con constancia expresa de autonomía.")
    add_bullet(doc, "Ausencia real de subordinación, horario fijo y exclusividad absoluta.")
    add_bullet(doc, "Cumplimiento de las obligaciones tributarias (retenciones ISLR, facturación).")
    add_bullet(doc, "Inscripción del profesional en el CMV (para MV) y mantenimiento de RIF propio (para ambos).")
    add_bullet(doc, "Suscripción de seguro de responsabilidad civil profesional (MV) y declaración anual ISLR.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc,
        "Se anexan a este concepto técnico los modelos de contrato de honorarios profesionales "
        "para Médico(a) Veterinario(a) y para Dog Groomer, debidamente alineados con los "
        "criterios aquí expuestos.", size=10, italic=True, color=GRAY_TEXT, space_after=8)

    add_para(doc,
        "En la ciudad de Los Teques, Estado Miranda, a la fecha indicada arriba.",
        size=10, space_after=10)

    add_signature_block(doc, ["DIRECTORA GERENTE", "ASESORÍA LEGAL"])

    add_footer(section, "Concepto Honorarios Profesionales v1.0")

    doc.save(out)
    return out


# ============================================================
# 2. CONTRATO DE HONORARIOS — MÉDICO VETERINARIO
# ============================================================
def gen_contrato_honorarios_vet():
    out = "/home/z/my-project/output/Contrato_Honorarios_Medico_Veterinario.docx"
    os.makedirs(os.path.dirname(out), exist_ok=True)

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "CONTRATO DE HONORARIOS", "Profesional Veterinario(a)",
                 version="Versión 1.0  ·  Dirección")
    add_doc_title(doc, "CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Médico(a) Veterinario(a) — Honorarios Profesionales")
    style_run(r, size=11, italic=True, color=GRAY_TEXT)

    add_para(doc,
        "Entre los suscritos: GRUPO CAVAL 1003, C.A., sociedad mercantil de domicilio en "
        "Av. Francisco de Miranda, Local N° 1, Sector Francisco de Miranda, Los Teques, "
        "Estado Miranda, Zona Postal 1201, RIF N° J501662533, marca comercial ALIKA PETS, "
        "dedicada a la actividad de clínica veterinaria, tienda de mascotas y peluquería "
        "canina, en adelante «LA EMPRESA», representada en este acto por su Directora "
        "Gerente, ciudadana ESNATLIM ELENA SIMOZA, titular de la cédula de identidad N° "
        "V-17.976.287, por una parte; y por la otra, el(la) ciudadano(a) "
        "___________________________________________, venezolano(a), mayor de edad, titular "
        "de la cédula de identidad N° V-___________________, inscrito(a) en el Colegio de "
        "Médicos Veterinarios de Venezuela (CMV) bajo el N° __________________, con RIF "
        "personal N° __________________, en adelante «EL PROFESIONAL», quienes declaran ser "
        "mayores de edad y hábiles para contratar, han convenido en celebrar el presente "
        "CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES, el cual se regirá por las "
        "siguientes cláusulas:", size=10, space_after=8)

    # CLÁUSULA 1: OBJETO
    add_section(doc, "CLÁUSULA PRIMERA: OBJETO")
    add_para(doc,
        "EL PROFESIONAL se obliga a prestar a LA EMPRESA servicios profesionales de "
        "MEDICINA VETERINARIA, en su carácter de profesional liberal independiente, "
        "comprendiendo entre otras: consultas clínicas generales y especializadas; "
        "procedimientos quirúrgicos; atención de urgencias y emergencias; realización e "
        "interpretación de exámenes complementarios (radiografías, ecografías, "
        "laboratorios); supervisión de pacientes hospitalizados; aplicación de vacunas y "
        "pautas de desparasitación; redacción y firma de recetas, certificados e "
        "informes médicos; y demás servicios propios del ejercicio profesional veterinario.")
    add_para(doc,
        "Los servicios se prestarán CON PLENA AUTONOMÍA TÉCNICA. EL PROFESIONAL decidirá "
        "libremente el diagnóstico, tratamiento, procedimiento y medicación de cada "
        "paciente, conforme a su juicio clínico y a las normas del ejercicio profesional. "
        "LA EMPRESA reconoce esta autonomía técnica y NO impartirá instrucciones que "
        "afecten el criterio profesional del PROFESIONAL.", space_after=6)

    # CLÁUSULA 2: NATURALEZA DEL CONTRATO
    add_section(doc, "CLÁUSULA SEGUNDA: NATURALEZA DEL CONTRATO")
    add_para(doc,
        "Las partes declaran expresamente que el presente contrato es de PRESTACIÓN DE "
        "SERVICIOS PROFESIONALES BAJO LA MODALIDAD DE HONORARIOS, conforme al artículo 40 "
        "de la Ley Orgánica del Trabajo, los Trabajadores y las Trabajadoras (LOTTT), a "
        "la Ley de Ejercicio de la Medicina Veterinaria y a la jurisprudencia del Tribunal "
        "Supremo de Justicia. En consecuencia:")
    add_bullet(doc,
        "EL PROFESIONAL actúa como profesional liberal independiente, sin relación de "
        "subordinación ni dependencia respecto a LA EMPRESA.",
        bold_lead="1. No subordinación:  ")
    add_bullet(doc,
        "EL PROFESIONAL organiza su tiempo y métodos de trabajo de manera autónoma, "
        "sujeto únicamente a los horarios de turnos acordados y a la entrega de resultados.",
        bold_lead="2. Autonomía técnica:  ")
    add_bullet(doc,
        "EL PROFESIONAL puede prestar servicios a otros consultorios, clínicas y "
        "particulares, en horarios distintos a los pactados con LA EMPRESA.",
        bold_lead="3. No exclusividad:  ")
    add_bullet(doc,
        "La remuneración se pagará en calidad de HONORARIOS PROFESIONALES, sin generar "
        "prestaciones sociales, utilidades, vacaciones, bono vacacional ni indemnizaciones "
        "laborales de ninguna naturaleza.",
        bold_lead="4. No salarial:  ")
    add_bullet(doc,
        "EL PROFESIONAL asume el riesgo técnico y civil de su actuación profesional, "
        "debiendo mantener vigente póliza de Seguro de Responsabilidad Civil Profesional.",
        bold_lead="5. Asumen riesgo:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # CLÁUSULA 3: HORARIOS Y TURNOS
    add_section(doc, "CLÁUSULA TERCERA: HORARIOS Y TURNOS")
    add_para(doc,
        "Las partes acuerdan que EL PROFESIONAL prestará sus servicios en turnos "
        "flexibles, conforme al siguiente esquema:")
    add_bullet(doc, "Días: de lunes a sábado (con un día de descanso acordado semanal).")
    add_bullet(doc, "Turno: _____ horas a _____ horas (con receso de 1 hora para almuerzo).")
    add_bullet(doc, "Guardias de emergencia: ____ veces al mes (según calendario rotativo).")
    add_para(doc,
        "EL PROFESIONAL podrá modificar la distribución semanal de sus turnos previa "
        "coordinación con LA EMPRESA, siempre que asegure la cobertura mínima pactada. "
        "Esta flexibilidad horaria es esencial para preservar la naturaleza no laboral "
        "del contrato.", space_after=6)

    # CLÁUSULA 4: HONORARIOS Y BENEFICIOS
    add_section(doc, "CLÁUSULA CUARTA: HONORARIOS Y BENEFICIOS")
    add_para(doc,
        "Por los servicios profesionales prestados, LA EMPRESA pagará a EL PROFESIONAL "
        "una remuneración mensual total de USD 250,00 (DÓLARES DE LOS ESTADOS UNIDOS DE "
        "AMÉRICA DOSCIENTOS CINCUENTA CON 00/100), discriminada así:",
        space_after=4)

    tbl = doc.add_table(rows=7, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(0.6), Cm(7.5), Cm(2.5), Cm(2.5)]):
        tbl.columns[i].width = w
    headers = ["#", "CONCEPTO", "MENSUAL", "QUINCENAL"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = [Cm(0.6), Cm(7.5), Cm(2.5), Cm(2.5)][ci]
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    conceptos = [
        ("1", "Honorarios profesionales base (servicios veterinarios)", "$ 50,00", "$ 25,00"),
        ("2", "Bono de Alimentación (Cestaticket)", "$ 80,00", "$ 40,00"),
        ("3", "Bono de Buen Vivir (asistencial)", "$ 40,00", "$ 20,00"),
        ("4", "Bono de Transporte (reembolso gastos)", "$ 40,00", "$ 20,00"),
        ("5", "Otros beneficios (uniformes, dotación, incentivos)", "$ 40,00", "$ 20,00"),
    ]
    for ri, (n, con, mens, quinc) in enumerate(conceptos, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        c = tbl.rows[ri].cells[0]; c.width = Cm(0.6)
        write_cell(c, n, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[1]; c.width = Cm(7.5)
        write_cell(c, con, size=9, bg=bg)
        c = tbl.rows[ri].cells[2]; c.width = Cm(2.5)
        write_cell(c, mens, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[3]; c.width = Cm(2.5)
        write_cell(c, quinc, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
    c = tbl.rows[6].cells[0]; c.width = Cm(0.6)
    c.text = ""; set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[1]; c.width = Cm(7.5)
    c.text = ""
    p = c.paragraphs[0]; r = p.add_run("TOTAL MENSUAL")
    style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[2]; c.width = Cm(2.5)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 250,00"); style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[3]; c.width = Cm(2.5)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 125,00"); style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc,
        "FORMA DE PAGO: Quincenal, los días 15 y último de cada mes, mediante "
        "transferencia bancaria a la cuenta del PROFESIONAL (Banco: _______________, "
        "Cuenta N°: _______________). El pago de honorarios estará sujeto a la "
        "retención del 3% por concepto de Impuesto sobre la Renta (ISLR) conforme al "
        "artículo 27 de la Ley de ISLR, retención que LA EMPRESA enterará al SENIAT "
        "dentro de los primeros 15 días del mes siguiente y entregará comprobante al "
        "PROFESIONAL.", size=10, space_after=4)

    add_para(doc,
        "FACTURACIÓN: EL PROFESIONAL se obliga a emitir factura o recibo por cada pago "
        "recibido, con sus datos de RIF, número de control, número de factura y "
        "descripción del servicio. Sin factura no procederá el pago.",
        size=10, italic=True, color=GRAY_TEXT, space_after=6)

    # CLÁUSULA 5: OBLIGACIONES DEL PROFESIONAL
    add_section(doc, "CLÁUSULA QUINTA: OBLIGACIONES DEL PROFESIONAL")
    add_bullet(doc, "Prestar los servicios profesionales con diligencia, calidad y oportunidad.")
    add_bullet(doc, "Mantener vigente la inscripción en el CMV y la colegiatura profesional.")
    add_bullet(doc, "Mantener póliza de Seguro de Responsabilidad Civil Profesional vigente y entregar copia a LA EMPRESA.")
    add_bullet(doc, "Cumplir los protocolos clínicos, de bioseguridad (NT-01-2008) y manejo de sustancias controladas (SENAC).")
    add_bullet(doc, "Llevar y mantener actualizadas las historias clínicas de los pacientes atendidos, las cuales son propiedad de LA EMPRESA.")
    add_bullet(doc, "Redactar y firmar recetas, certificados, consentimientos informados e informes médicos.")
    add_bullet(doc, "Atender urgencias y emergencias según turnos asignados.")
    add_bullet(doc, "Mantener confidencialidad de la información de LA EMPRESA, clientes y pacientes (durante el contrato y por 5 años después).")
    add_bullet(doc, "Emitir facturas por cada pago y declarar el ISLR anualmente.")
    add_bullet(doc, "No contactar clientes de LA EMPRESA para ofrecer servicios externos durante la vigencia del contrato y por 12 meses después de su terminación (cláusula de no competencia parcial post-contractual).")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # CLÁUSULA 6: OBLIGACIONES DE LA EMPRESA
    add_section(doc, "CLÁUSULA SEXTA: OBLIGACIONES DE LA EMPRESA")
    add_bullet(doc, "Pagar los honorarios y beneficios en la forma y oportunidad pactadas.")
    add_bullet(doc, "Proveer el espacio físico, equipos, instrumental, medicamentos e insumos necesarios para la prestación de los servicios.")
    add_bullet(doc, "Garantizar el mantenimiento de las instalaciones y equipos.")
    add_bullet(doc, "Retener y enterar el 3% de ISLR al SENIAT, y entregar comprobante al PROFESIONAL.")
    add_bullet(doc, "Respetar la autonomía técnica del PROFESIONAL y no impartir instrucciones que afecten su criterio clínico.")
    add_bullet(doc, "Proporcionar acceso a las historias clínicas y registros de los pacientes.")
    add_bullet(doc, "Mantener el Sistema de Vigilancia Médica conforme a la NT-02-2008.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # CLÁUSULA 7: PROPIEDAD INTELECTUAL Y HISTORIAS CLÍNICAS
    add_section(doc, "CLÁUSULA SÉPTIMA: PROPIEDAD DE HISTORIAS CLÍNICAS")
    add_para(doc,
        "Las historias clínicas veterinarias, registros médicos, radiografías, resultados "
        "de laboratorio, fotografías clínicas de pacientes y demás documentación médica "
        "generada durante la prestación del servicio son PROPIEDAD EXCLUSIVA DE LA EMPRESA. "
        "EL PROFESIONAL no podrá sustraerlas, copiarlas, fotografiarlas, divulgarlas ni "
        "utilizarlas para fines personales o de terceros, ni durante ni después de la "
        "vigencia del presente contrato.")
    add_para(doc,
        "EL PROFESIONAL entregará a LA EMPRESA todas las historias clínicas y registros "
        "al término del contrato. La violación de esta cláusula generará responsabilidad "
        "civil por daños y perjuicios.", space_after=6)

    # CLÁUSULA 8: DURACIÓN
    add_section(doc, "CLÁUSULA OCTAVA: DURACIÓN")
    add_para(doc,
        "El presente contrato tendrá una duración de DOCE (12) MESES, contados a partir "
        "del ____ de ________________ de ______, hasta el ____ de ________________ de ______. "
        "Podrá prorrogarse por mutuo acuerdo mediante addendum suscrito por las partes con "
        "al menos 30 días de anticipación a su vencimiento.", space_after=6)

    # CLÁUSULA 9: TERMINACIÓN
    add_section(doc, "CLÁUSULA NOVENA: TERMINACIÓN")
    add_para(doc, "El contrato podrá terminar por:")
    add_bullet(doc, "Vencimiento del plazo pactado, sin necesidad de notificación.")
    add_bullet(doc, "Resolución por mutuo acuerdo, mediante acta suscrita por las partes.")
    add_bullet(doc, "Resolución unilateral por incumplimiento de cualquiera de las partes, previa notificación escrita con 15 días de anticipación.")
    add_bullet(doc, "Resolución inmediata por causa grave (mala praxis profesional, violación de confidencialidad, abandono de servicios).")
    add_para(doc,
        "La terminación del contrato NO genera derecho a prestaciones sociales, "
        "indemnizaciones laborales, ni cualquier otro concepto de naturaleza salarial. "
        "Solo procederá el pago de los honorarios pendientes por servicios efectivamente "
        "prestados a la fecha de terminación.", space_after=6)

    # CLÁUSULA 10: RESPONSABILIDAD CIVIL
    add_section(doc, "CLÁUSULA DÉCIMA: RESPONSABILIDAD CIVIL PROFESIONAL")
    add_para(doc,
        "EL PROFESIONAL responde civil y profesionalmente por los actos, omisiones y "
        "decisiones clínicas que adopte en el ejercicio de su profesión. LA EMPRESA no "
        "asume responsabilidad solidaria por la actividad profesional del PROFESIONAL, "
        "salvo que se demuestre culpa directa de la empresa (equipos defectuosos, "
        "instalaciones inseguras).")
    add_para(doc,
        "EL PROFESIONAL declara tener vigente póliza de Seguro de Responsabilidad Civil "
        "Profesional con cobertura no menor a USD 50.000,00, y entrega copia de la póliza "
        "a LA EMPRESA al momento de la suscripción del presente contrato.",
        space_after=6)

    # CLÁUSULA 11: CONFIDENCIALIDAD
    add_section(doc, "CLÁUSULA DÉCIMA PRIMERA: CONFIDENCIALIDAD")
    add_para(doc,
        "EL PROFESIONAL se obliga a mantener en reserva toda la información de LA EMPRESA, "
        "sus clientes, pacientes, proveedores y estrategias comerciales, durante la "
        "vigencia del contrato y por CINCO (5) AÑOS después de su terminación. La "
        "violación de esta cláusula generará responsabilidad civil por daños y perjuicios.",
        space_after=6)

    # CLÁUSULA 12: LOPDP
    add_section(doc, "CLÁUSULA DÉCIMA SEGUNDA: PROTECCIÓN DE DATOS PERSONALES")
    add_para(doc,
        "Las partes se comprometen a cumplir la Ley Orgánica de Protección de Datos "
        "Personales (LOPDP) en el tratamiento de los datos personales de clientes, "
        "pacientes y personal. El profesional autoriza a LA EMPRESA el tratamiento de sus "
        "datos personales con fines administrativos y tributarios, conforme a la "
        "Autorización firmada por separado.", space_after=6)

    # CLÁUSULA 13: DOMICILIO Y JURISDICCIÓN
    add_section(doc, "CLÁUSULA DÉCIMA TERCERA: DOMICILIO Y JURISDICCIÓN")
    add_para(doc,
        "Para todos los efectos derivados del presente contrato, las partes eligen como "
        "domicilio procesal especial, con exclusión de cualquier otro, la ciudad de Los "
        "Teques, Estado Miranda, a cuyos tribunales declaran someterse. Las controversias "
        "derivadas del presente contrato se sustanciarán por la vía ordinaria civil, "
        "correspondiendo a los tribunales civiles competentes el conocimiento de las mismas. "
        "Se excluye expresamente la jurisdicción laboral por tratarse de un contrato de "
        "naturaleza civil y mercantil.", space_after=8)

    add_para(doc,
        "Se hacen dos (02) ejemplares de un mismo tenor y un mismo efecto, en la ciudad de "
        "Los Teques, a los ____ días del mes de ________________ de ________.",
        size=10, space_after=10)

    add_signature_block(doc, ["LA EMPRESA", "EL PROFESIONAL"])
    add_footer(section, "Contrato Honorarios Médico Veterinario v1.0")

    doc.save(out)
    return out


# ============================================================
# 3. CONTRATO DE HONORARIOS — DOG GROOMER
# ============================================================
def gen_contrato_honorarios_groomer():
    out = "/home/z/my-project/output/Contrato_Honorarios_Dog_Groomer.docx"
    os.makedirs(os.path.dirname(out), exist_ok=True)

    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "CONTRATO DE HONORARIOS", "Peluquero Canino",
                 version="Versión 1.0  ·  Dirección")
    add_doc_title(doc, "CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Dog Groomer (Peluquero Canino) — Honorarios Profesionales")
    style_run(r, size=11, italic=True, color=GRAY_TEXT)

    add_para(doc,
        "Entre los suscritos: GRUPO CAVAL 1003, C.A., sociedad mercantil de domicilio en "
        "Av. Francisco de Miranda, Local N° 1, Sector Francisco de Miranda, Los Teques, "
        "Estado Miranda, Zona Postal 1201, RIF N° J501662533, marca comercial ALIKA PETS, "
        "dedicada a la actividad de clínica veterinaria, tienda de mascotas y peluquería "
        "canina, en adelante «LA EMPRESA», representada en este acto por su Directora "
        "Gerente, ciudadana ESNATLIM ELENA SIMOZA, titular de la cédula de identidad N° "
        "V-17.976.287, por una parte; y por la otra, el(la) ciudadano(a) "
        "___________________________________________, venezolano(a), mayor de edad, titular "
        "de la cédula de identidad N° V-___________________, con RIF personal N° "
        "__________________, con experiencia técnica certificada en peluquería canina y "
        "felina (anexar certificaciones o constancias), en adelante «EL PROFESIONAL», "
        "quienes declaran ser mayores de edad y hábiles para contratar, han convenido en "
        "celebrar el presente CONTRATO DE PRESTACIÓN DE SERVICIOS PROFESIONALES, el cual "
        "se regirá por las siguientes cláusulas:", size=10, space_after=8)

    # CLÁUSULA 1: OBJETO
    add_section(doc, "CLÁUSULA PRIMERA: OBJETO")
    add_para(doc,
        "EL PROFESIONAL se obliga a prestar a LA EMPRESA servicios de PELUQUERÍA CANINA "
        "Y FELINA, en su carácter de técnico especializado independiente, comprendiendo: "
        "cortes de pelo según raza, tipo de pelaje o solicitud del cliente; baño, secado "
        "y cepillado profesional; corte y limado de uñas; limpieza del canal auditivo "
        "externo; evaluación inicial del estado del animal; manejo y contención amigable "
        "del animal; limpieza y desinfección del área de trabajo entre pacientes; y "
        "mantenimiento de las herramientas de peluquería.")
    add_para(doc,
        "Los servicios se prestarán CON PLENA AUTONOMÍA TÉCNICA. EL PROFESIONAL decide "
        "libremente las técnicas de corte, sujeción y manejo del animal, conforme a las "
        "buenas prácticas de peluquería canina y al bienestar animal. LA EMPRESA reconoce "
        "esta autonomía técnica.", space_after=6)

    # CLÁUSULA 2: NATURALEZA
    add_section(doc, "CLÁUSULA SEGUNDA: NATURALEZA DEL CONTRATO")
    add_para(doc,
        "Las partes declaran expresamente que el presente contrato es de PRESTACIÓN DE "
        "SERVICIOS BAJO LA MODALIDAD DE HONORARIOS PROFESIONALES, conforme al artículo 40 "
        "de la LOTTT y la jurisprudencia del Tribunal Supremo de Justicia. En consecuencia:")
    add_bullet(doc, "EL PROFESIONAL actúa como técnico independiente, sin relación de subordinación.",
               bold_lead="1. No subordinación:  ")
    add_bullet(doc, "EL PROFESIONAL organiza su trabajo de manera autónoma, sujeto a las citas acordadas.",
               bold_lead="2. Autonomía técnica:  ")
    add_bullet(doc, "EL PROFESIONAL puede atender a otros clientes y prestar servicios externos en horarios distintos a los pactados.",
               bold_lead="3. No exclusividad:  ")
    add_bullet(doc, "La remuneración se pagará en calidad de HONORARIOS, sin generar prestaciones sociales.",
               bold_lead="4. No salarial:  ")
    add_bullet(doc, "EL PROFESIONAL asume el riesgo de su actividad (mordeduras, cortes, lesiones a mascotas).",
               bold_lead="5. Asume riesgo:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    # CLÁUSULA 3: HORARIOS
    add_section(doc, "CLÁUSULA TERCERA: HORARIOS Y TURNOS")
    add_para(doc,
        "Las partes acuerdan que EL PROFESIONAL prestará sus servicios con la siguiente "
        "flexibilidad horaria:")
    add_bullet(doc, "Días: de lunes a sábado (con un día de descanso acordado semanal).")
    add_bullet(doc, "Turno: _____ horas a _____ horas (con receso de 1 hora).")
    add_bullet(doc, "Citas asignadas por LA EMPRESA, previa coordinación con EL PROFESIONAL.")
    add_para(doc,
        "EL PROFESIONAL podrá aceptar o reprogramar citas, siempre que asegure la "
        "cobertura mínima pactada de ___ servicios/día. Esta flexibilidad es esencial "
        "para preservar la naturaleza no laboral del contrato.", space_after=6)

    # CLÁUSULA 4: HONORARIOS
    add_section(doc, "CLÁUSULA CUARTA: HONORARIOS Y BENEFICIOS")
    add_para(doc,
        "Por los servicios prestados, LA EMPRESA pagará a EL PROFESIONAL una remuneración "
        "mensual total de USD 250,00 (DÓLARES DE LOS ESTADOS UNIDOS DE AMÉRICA "
        "DOSCIENTOS CINCUENTA CON 00/100), discriminada así:", space_after=4)

    tbl = doc.add_table(rows=7, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate([Cm(0.6), Cm(7.5), Cm(2.5), Cm(2.5)]):
        tbl.columns[i].width = w
    headers = ["#", "CONCEPTO", "MENSUAL", "QUINCENAL"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.width = [Cm(0.6), Cm(7.5), Cm(2.5), Cm(2.5)][ci]
        write_cell(c, h, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    conceptos = [
        ("1", "Honorarios profesionales base (peluquería canina)", "$ 50,00", "$ 25,00"),
        ("2", "Bono de Alimentación (Cestaticket)", "$ 80,00", "$ 40,00"),
        ("3", "Bono de Buen Vivir (asistencial)", "$ 40,00", "$ 20,00"),
        ("4", "Bono de Transporte (reembolso gastos)", "$ 40,00", "$ 20,00"),
        ("5", "Otros beneficios (uniformes, dotación, incentivos)", "$ 40,00", "$ 20,00"),
    ]
    for ri, (n, con, mens, quinc) in enumerate(conceptos, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        c = tbl.rows[ri].cells[0]; c.width = Cm(0.6)
        write_cell(c, n, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[1]; c.width = Cm(7.5)
        write_cell(c, con, size=9, bg=bg)
        c = tbl.rows[ri].cells[2]; c.width = Cm(2.5)
        write_cell(c, mens, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
        c = tbl.rows[ri].cells[3]; c.width = Cm(2.5)
        write_cell(c, quinc, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg)
    c = tbl.rows[6].cells[0]; c.width = Cm(0.6)
    c.text = ""; set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[1]; c.width = Cm(7.5)
    c.text = ""
    p = c.paragraphs[0]; r = p.add_run("TOTAL MENSUAL")
    style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[2]; c.width = Cm(2.5)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 250,00"); style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    c = tbl.rows[6].cells[3]; c.width = Cm(2.5)
    c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("$ 125,00"); style_run(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    set_cell_bg(c, TEAL_HDR_BG); set_cell_borders(c); set_cell_margins(c)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc,
        "FORMA DE PAGO: Quincenal, los días 15 y último de cada mes, mediante "
        "transferencia bancaria a la cuenta del PROFESIONAL (Banco: _______________, "
        "Cuenta N°: _______________). El pago estará sujeto a la retención del 1% por "
        "concepto de ISLR (encomendado no profesional universitario) conforme al artículo "
        "27 de la Ley de ISLR, retención que LA EMPRESA enterará al SENIAT y entregará "
        "comprobante al PROFESIONAL.", size=10, space_after=4)

    add_para(doc,
        "FACTURACIÓN: EL PROFESIONAL se obliga a emitir recibo o factura por cada pago "
        "recibido, con sus datos de RIF, número de control y descripción del servicio. "
        "Sin factura no procederá el pago.", size=10, italic=True, color=GRAY_TEXT, space_after=6)

    # CLÁUSULA 5: OBLIGACIONES DEL PROFESIONAL
    add_section(doc, "CLÁUSULA QUINTA: OBLIGACIONES DEL PROFESIONAL")
    add_bullet(doc, "Prestar los servicios con diligencia, calidad, oportunidad y respeto al bienestar animal.")
    add_bullet(doc, "Realizar cortes según raza, tipo de pelaje o solicitud del cliente.")
    add_bullet(doc, "Evaluar el estado del animal antes del servicio y reportar cualquier hallazgo al Médico Veterinario.")
    add_bullet(doc, "Suspender el servicio e informar al MV si el animal presenta signos de asfixia, fatiga extrema, síncope o agresividad inmanejable.")
    add_bullet(doc, "Aplicar técnicas de manejo y contención amigable, libres de crueldad.")
    add_bullet(doc, "Limpiar, desinfectar y ordenar la mesa de peluquería, bañera y jaulas entre cada paciente.")
    add_bullet(doc, "Mantener sus propias herramientas de corte (tijeras, cuchillas, máquinas) en buen estado.")
    add_bullet(doc, "Mantener confidencialidad de la información de LA EMPRESA y sus clientes (durante el contrato y por 3 años después).")
    add_bullet(doc, "Emitir facturas por cada pago y declarar ISLR anualmente.")
    add_bullet(doc, "No contactar clientes de LA EMPRESA para ofrecer servicios externos durante el contrato y por 12 meses después.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # CLÁUSULA 6: OBLIGACIONES DE LA EMPRESA
    add_section(doc, "CLÁUSULA SEXTA: OBLIGACIONES DE LA EMPRESA")
    add_bullet(doc, "Pagar los honorarios y beneficios en la forma y oportunidad pactadas.")
    add_bullet(doc, "Proveer el espacio físico de peluquería, bañera, mesa, jaulas, secadora, productos (champús, acondicionadores) y servicios básicos.")
    add_bullet(doc, "Coordinar las citas y asignarlas previamente.")
    add_bullet(doc, "Mantener el área de peluquería limpia, segura y en condiciones adecuadas.")
    add_bullet(doc, "Retener y enterar el 1% de ISLR al SENIAT, entregar comprobante al PROFESIONAL.")
    add_bullet(doc, "Respetar la autonomía técnica del PROFESIONAL en los servicios de peluquería.")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

    # CLÁUSULA 7: HERRAMIENTAS
    add_section(doc, "CLÁUSULA SÉPTIMA: HERRAMIENTAS Y EQUIPOS")
    add_para(doc,
        "EL PROFESIONAL utilizará sus propias herramientas de corte (tijeras, cuchillas, "
        "máquinas portátiles, peines, cepillos) que son de su propiedad. LA EMPRESA "
        "proveerá la bañera, mesa de peluquería, secadora, jaulas, champús, "
        "acondicionadores, toallas y demás insumos. Las herramientas de EL PROFESIONAL "
        "se identificarán y se entregarán al término del contrato.", space_after=6)

    # CLÁUSULA 8: RESPONSABILIDAD
    add_section(doc, "CLÁUSULA OCTAVA: RESPONSABILIDAD CIVIL")
    add_para(doc,
        "EL PROFESIONAL responde civilmente por los daños causados a las mascotas durante "
        "el servicio, salvo que se demuestre que el daño fue causado por equipos "
        "defectuosos o instalaciones inseguras proporcionadas por LA EMPRESA. EL "
        "PROFESIONAL debe mantener póliza de Responsabilidad Civil (opcional pero "
        "recomendada) y reportar inmediatamente cualquier incidente.")
    add_para(doc,
        "En caso de mordedura o arañazo al PROFESIONAL durante el servicio, este será "
        "responsable de su atención médica, dado que se trata de un profesional "
        "independiente que asume el riesgo de su actividad. Se recomienda mantener "
        "vacuna antirrábica pre-exposición vigente.", space_after=6)

    # CLÁUSULA 9: DURACIÓN
    add_section(doc, "CLÁUSULA NOVENA: DURACIÓN")
    add_para(doc,
        "El presente contrato tendrá una duración de DOCE (12) MESES, contados a partir "
        "del ____ de ________________ de ______, hasta el ____ de ________________ de ______. "
        "Podrá prorrogarse por mutuo acuerdo mediante addendum suscrito por las partes con "
        "al menos 30 días de anticipación a su vencimiento.", space_after=6)

    # CLÁUSULA 10: TERMINACIÓN
    add_section(doc, "CLÁUSULA DÉCIMA: TERMINACIÓN")
    add_para(doc, "El contrato podrá terminar por:")
    add_bullet(doc, "Vencimiento del plazo pactado.")
    add_bullet(doc, "Resolución por mutuo acuerdo, mediante acta suscrita por las partes.")
    add_bullet(doc, "Resolución unilateral por incumplimiento, previa notificación escrita con 15 días de anticipación.")
    add_bullet(doc, "Resolución inmediata por causa grave (maltrato animal, robo, violación de confidencialidad).")
    add_para(doc,
        "La terminación NO genera derecho a prestaciones sociales ni indemnizaciones "
        "laborales. Solo procederá el pago de honorarios pendientes por servicios "
        "efectivamente prestados.", space_after=6)

    # CLÁUSULA 11: CONFIDENCIALIDAD
    add_section(doc, "CLÁUSULA DÉCIMA PRIMERA: CONFIDENCIALIDAD")
    add_para(doc,
        "EL PROFESIONAL se obliga a mantener en reserva toda la información de LA EMPRESA, "
        "sus clientes y pacientes, durante la vigencia del contrato y por TRES (3) AÑOS "
        "después de su terminación.", space_after=6)

    # CLÁUSULA 12: LOPDP
    add_section(doc, "CLÁUSULA DÉCIMA SEGUNDA: PROTECCIÓN DE DATOS PERSONALES")
    add_para(doc,
        "Las partes cumplirán la LOPDP. EL PROFESIONAL autoriza a LA EMPRESA el "
        "tratamiento de sus datos personales con fines administrativos y tributarios.",
        space_after=6)

    # CLÁUSULA 13: DOMICILIO Y JURISDICCIÓN
    add_section(doc, "CLÁUSULA DÉCIMA TERCERA: DOMICILIO Y JURISDICCIÓN")
    add_para(doc,
        "Para todos los efectos, las partes eligen como domicilio procesal especial la "
        "ciudad de Los Teques, Estado Miranda, a cuyos tribunales civiles declaran "
        "someterse. Se excluye expresamente la jurisdicción laboral.", space_after=8)

    add_para(doc,
        "Se hacen dos (02) ejemplares de un mismo tenor y un mismo efecto, en la ciudad de "
        "Los Teques, a los ____ días del mes de ________________ de ________.",
        size=10, space_after=10)

    add_signature_block(doc, ["LA EMPRESA", "EL PROFESIONAL"])
    add_footer(section, "Contrato Honorarios Dog Groomer v1.0")

    doc.save(out)
    return out


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print("Generando 3 documentos de honorarios profesionales...")
    outs = [
        gen_concepto_honorarios(),
        gen_contrato_honorarios_vet(),
        gen_contrato_honorarios_groomer(),
    ]
    for o in outs:
        size_kb = os.path.getsize(o) / 1024
        print(f"  ✓ {o}  ({size_kb:.1f} KB)")
    print("\nTodos los documentos generados correctamente.")
