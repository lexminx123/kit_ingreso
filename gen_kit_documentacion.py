"""
gen_kit_documentacion.py — Genera el documento Word "Kit de Ingresos.docx"
que contiene TODOS los códigos Python que generan los formatos del kit.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "/home/z/my-project/output"
OUT_FILE = "/home/z/my-project/output/Kit_de_Ingresos.docx"

# Paleta teal
TEAL_DARK   = RGBColor(0x0F, 0x76, 0x6E)
TEAL_HDR_BG = "0F766E"
GRAY_ALT    = "F3F4F6"
GRAY_LIGHT  = "F9FAFB"
SLATE_BG    = "E2E8F0"
WHITE       = "FFFFFF"
BLACK       = RGBColor(0x11, 0x18, 0x27)
GRAY_TEXT   = RGBColor(0x55, 0x65, 0x75)
GRAY_MUTED  = RGBColor(0x90, 0x98, 0xA8)
CODE_BG    = "F8FAFC"
CODE_BORDER= "E2E8F0"

# ============================================================
# Utilidades básicas
# ============================================================
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="CBD5E1", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, v in (('top',top),('bottom',bottom),('left',left),('right',right)):
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tc_pr.append(tcMar)

def style_run(run, size=10, bold=False, color=None, font="Calibri", italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)

def add_hr(paragraph, color="0F766E", sz="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    p_pr.append(pBdr)

def add_para(doc, text, size=10, bold=False, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
             italic=False, font="Calibri", line_spacing=1.3, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if indent is not None:
        p.paragraph_format.left_indent = indent
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color, font=font, italic=italic)
    return p

def add_bullet(doc, text, size=10, bold_lead=None, indent=Cm(0.6)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    rb = p.add_run("•  ")
    style_run(rb, size=size, bold=True, color=TEAL_DARK)
    if bold_lead:
        rl = p.add_run(bold_lead)
        style_run(rl, size=size, bold=True, color=BLACK)
    r = p.add_run(text)
    style_run(r, size=size, color=BLACK)

# ============================================================
# Renderizado de código Python en Word
# ============================================================
def add_code_block(doc, code_text, filename=""):
    """Añade un bloque de código Python al documento, con fondo claro, fuente monoespaciada y números de línea."""
    lines = code_text.split("\n")
    # Tabla 1x1 para contener el código con fondo y borde
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(17)
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(17)
    cell.text = ""
    # Configuración de celda (fondo gris claro + bordes)
    set_cell_bg(cell, CODE_BG)
    set_cell_borders(cell, color=CODE_BORDER, sz="6")
    set_cell_margins(cell, top=120, bottom=120, left=200, right=200)

    # Primera línea: nombre del archivo
    if filename:
        p_hdr = cell.paragraphs[0]
        p_hdr.paragraph_format.space_after = Pt(2)
        p_hdr.paragraph_format.line_spacing = 1.0
        r_hdr = p_hdr.add_run("# " + filename)
        style_run(r_hdr, size=8, bold=True, color=TEAL_DARK, font="Consolas")
        # línea separadora
        p_sep = cell.add_paragraph()
        p_sep.paragraph_format.space_after = Pt(2)
        p_sep.paragraph_format.line_spacing = 1.0
        r_sep = p_sep.add_run("# " + "="*70)
        style_run(r_sep, size=8, color=GRAY_MUTED, font="Consolas")
    else:
        # primer párrafo vacío
        cell.paragraphs[0].text = ""

    # Cada línea de código en su párrafo (Consolas 8pt, line-height compacto)
    for i, line in enumerate(lines):
        if i == 0 and filename:
            # ya añadimos header, esta es la primera línea de código
            p = cell.add_paragraph()
        elif i == 0 and not filename:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)

        # Número de línea a la izquierda (en gris claro)
        r_num = p.add_run(f"{i+1:4d}  ")
        style_run(r_num, size=7, color=GRAY_MUTED, font="Consolas")
        # Línea de código
        # Reemplazar tabs por 4 espacios
        r_code = p.add_run(line.replace("\t", "    ") if line else " ")
        style_run(r_code, size=8, color=BLACK, font="Consolas")

    # Espacio después del bloque
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

# ============================================================
# Construcción del documento
# ============================================================
def build_document():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(4)

    # ============================================================
    # PORTADA
    # ============================================================
    # Tabla 2 columnas: izquierda marca, derecha info
    hdr = doc.add_table(rows=1, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.autofit = False
    hdr.columns[0].width = Cm(12)
    hdr.columns[1].width = Cm(5)
    for c in hdr.rows[0].cells:
        set_cell_borders(c, color="FFFFFF", sz="0")

    left = hdr.rows[0].cells[0]
    left.text = ""
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ALIKA PETS")
    style_run(r, size=28, bold=True, color=TEAL_DARK)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("Grupo Caval 1003, C.A.  ·  RIF: J501662533")
    style_run(r2, size=11, color=BLACK)
    p3 = left.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    r3 = p3.add_run("Av. Francisco de Miranda, Local N° 1, Los Teques, Miranda  ·  Zona Postal 1201")
    style_run(r3, size=9, color=GRAY_TEXT)

    right = hdr.rows[0].cells[1]
    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    rr = rp.add_run("DOCUMENTACIÓN")
    style_run(rr, size=11, bold=True, color=TEAL_DARK)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_before = Pt(0)
    rr2 = rp2.add_run("Códigos fuente Python")
    style_run(rr2, size=9, italic=True, color=GRAY_TEXT)
    rp3 = right.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp3.paragraph_format.space_before = Pt(0)
    rr3 = rp3.add_run("Versión 3.1")
    style_run(rr3, size=8, color=GRAY_MUTED)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after = Pt(20)
    add_hr(sep, color="0F766E", sz="14")

    # Título grande centrado
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("KIT DE INGRESOS")
    style_run(r, size=36, bold=True, color=TEAL_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("del Trabajador")
    style_run(r, size=24, bold=True, color=TEAL_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Códigos Python completos para regenerar los 39 formatos del kit")
    style_run(r, size=12, italic=True, color=GRAY_TEXT)

    # Stats en tabla 4 columnas
    for _ in range(2):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    stats = doc.add_table(rows=1, cols=4)
    stats.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats.autofit = False
    for i in range(4):
        stats.columns[i].width = Cm(4)
    stats_data = [
        ("11", "scripts Python"),
        ("~7,950", "líneas de código"),
        ("39", "formatos generados"),
        ("100%", "kit completo"),
    ]
    for ci, (num, lbl) in enumerate(stats_data):
        c = stats.rows[0].cells[ci]
        c.width = Cm(4)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(num)
        style_run(r, size=22, bold=True, color=TEAL_DARK)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(lbl)
        style_run(r2, size=9, color=GRAY_TEXT)
        set_cell_bg(c, GRAY_LIGHT)
        set_cell_borders(c, color="E2E8F0")
        set_cell_margins(c)

    # Salto de página después de portada
    doc.add_page_break()

    # ============================================================
    # ÍNDICE
    # ============================================================
    p = doc.add_paragraph()
    r = p.add_run("ÍNDICE DE SCRIPTS")
    style_run(r, size=18, bold=True, color=TEAL_DARK)
    p.paragraph_format.space_after = Pt(4)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(10)
    add_hr(sep, color="0F766E", sz="10")

    scripts_info = [
        ("1", "_common.py", "Módulo de utilidades compartidas (membrete, firmas, paleta, helpers)",
         "547 líneas", "Reutilizado por todos los demás scripts"),
        ("2", "gen_solicitud_v3.py", "Solicitud de Empleo / Ficha de Ingreso (2 páginas, compacta)",
         "526 líneas", "01_Solicitud_de_Empleo/Solicitud_de_Empleo_Ficha_Ingreso.docx"),
        ("3", "gen_contratos.py", "7 Contratos individuales de trabajo (uno por cargo)",
         "700 líneas", "02_CONTRATOS/02a-02g_*.docx (7 archivos)"),
        ("4", "gen_funciones_por_cargo.py", "7 Descripciones de Cargo con carta de recepción",
         "994 líneas", "03_DESCRIPCION_DE_CARGOS/03a-03g_*.docx (7 archivos)"),
        ("5", "gen_prestaciones.py", "2 documentos de prestaciones (autorización + beneficiarios)",
         "400 líneas", "04_PRESTACIONES/ (2 archivos)"),
        ("6", "gen_riesgos_por_rol.py", "5 Notificaciones de Riesgos específicas por cargo",
         "739 líneas", "05_SEGURIDAD_LABORAL/Notificacion_Riesgos_*.docx (5 archivos)"),
        ("7", "gen_protocolos_vet.py", "4 protocolos veterinarios (bioseguridad + sustancias + mordeduras + incidentes)",
         "1,281 líneas", "05_SEGURIDAD_LABORAL/Cartilla_*.docx + 08_POLITICAS_INTERNAS/Protocolo_*.docx"),
        ("8", "gen_documentos_finales_v3.py", "5 documentos finales (recorrido, EPP, examen médico, registros legales, carta aceptación)",
         "720 líneas", "05_SEGURIDAD_LABORAL + 06_REGISTROS_LEGALES + 09_CIERRE"),
        ("9", "gen_autorizaciones.py", "3 autorizaciones (LOPDP, imagen, cámaras)",
         "718 líneas", "07_AUTORIZACIONES/ (3 archivos)"),
        ("10", "gen_politicas.py", "4 políticas internas (reglamento, código, confidencialidad, redes sociales)",
         "958 líneas", "08_POLITICAS_INTERNAS/ (4 archivos)"),
        ("11", "gen_checklist_v3.py", "Checklist Maestro de Ingreso (A4 horizontal)",
         "465 líneas", "00_Checklist_Maestro_Ingreso_Trabajador.docx"),
    ]

    # Tabla índice
    tbl = doc.add_table(rows=1+len(scripts_info), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    col_widths = [Cm(0.8), Cm(5.5), Cm(4.5), Cm(2.5), Cm(4.5)]
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = w

    headers = ["#", "SCRIPT", "DESCRIPCIÓN", "LÍNEAS", "ARCHIVOS GENERADOS"]
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.width = col_widths[ci]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, size=9, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        set_cell_bg(c, TEAL_HDR_BG)
        set_cell_borders(c)
        set_cell_margins(c)

    for ri, (num, name, desc, lines, archs) in enumerate(scripts_info, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        vals = [(num, WD_ALIGN_PARAGRAPH.CENTER, True),
                (name, WD_ALIGN_PARAGRAPH.LEFT, True),
                (desc, WD_ALIGN_PARAGRAPH.LEFT, False),
                (lines, WD_ALIGN_PARAGRAPH.CENTER, False),
                (archs, WD_ALIGN_PARAGRAPH.LEFT, False)]
        for ci, (val, al, bold) in enumerate(vals):
            c = tbl.rows[ri].cells[ci]
            c.width = col_widths[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = al
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            col = TEAL_DARK if ci == 1 else (BLACK if ci != 3 else GRAY_TEXT)
            r = p.add_run(val)
            style_run(r, size=9, bold=bold, color=col, font="Consolas" if ci == 1 else "Calibri")
            set_cell_bg(c, bg)
            set_cell_borders(c)
            set_cell_margins(c)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)

    # ============================================================
    # INSTRUCCIONES DE USO
    # ============================================================
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("INSTRUCCIONES DE USO")
    style_run(r, size=18, bold=True, color=TEAL_DARK)
    p.paragraph_format.space_after = Pt(4)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(10)
    add_hr(sep, color="0F766E", sz="10")

    add_para(doc, "Requisitos previos", size=12, bold=True, color=TEAL_DARK, space_before=4, space_after=3)
    add_bullet(doc, "Versión 3.0+ para usar f-strings.", bold_lead="Python 3.6+:  ")
    add_bullet(doc, "Librería para generar documentos Word. Instalar con: pip install python-docx", bold_lead="python-docx:  ")
    add_bullet(doc, "Para convertir a PDF (verificación visual). Instalar con: apt install libreoffice", bold_lead="LibreOffice:  ")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    add_para(doc, "Cómo regenerar todos los formatos", size=12, bold=True, color=TEAL_DARK, space_before=4, space_after=3)
    add_para(doc,
        "1. Copia los 11 scripts .py a una carpeta de trabajo (por ejemplo /home/z/my-project/output/).\n"
        "2. Ejecuta cada script en orden:\n"
        "   python3 _common.py                  # (no genera nada, es módulo)\n"
        "   python3 gen_solicitud_v3.py          # genera 1 documento\n"
        "   python3 gen_contratos.py            # genera 7 contratos\n"
        "   python3 gen_funciones_por_cargo.py   # genera 7 descripciones de cargo\n"
        "   python3 gen_prestaciones.py         # genera 2 documentos\n"
        "   python3 gen_riesgos_por_rol.py       # genera 5 notificaciones de riesgos\n"
        "   python3 gen_protocolos_vet.py        # genera 4 protocolos veterinarios\n"
        "   python3 gen_documentos_finales_v3.py # genera 5 documentos finales\n"
        "   python3 gen_autorizaciones.py       # genera 3 autorizaciones\n"
        "   python3 gen_politicas.py             # genera 4 políticas internas\n"
        "   python3 gen_checklist_v3.py          # genera 1 checklist maestro\n"
        "3. Total: 39 documentos .docx generados en las carpetas correspondientes.",
        size=10, space_after=4)

    add_para(doc, "Personalización", size=12, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_bullet(doc, "Edita las constantes EMPRESA, RIF_EMP, MARCA, DOMICILIO_EMP al inicio del archivo _common.py", bold_lead="Datos de la empresa:  ")
    add_bullet(doc, "Edita DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI al inicio del archivo _common.py", bold_lead="Directora Gerente:  ")
    add_bullet(doc, "Edita SALARIO_BS y CESTATICKET_BS en gen_contratos.py (Bs. 210 + Bs. 1.500 vigentes abril 2025)", bold_lead="Salarios:  ")
    add_bullet(doc, "Los textos de cada script están hardcodeados en español venezolano. Edítalos según tus necesidades.", bold_lead="Textos legales:  ")

    add_para(doc, "Estructura de carpetas generada", size=12, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_para(doc,
        "output/\n"
        "├── 00_Checklist_Maestro_Ingreso_Trabajador.docx\n"
        "├── 01_Solicitud_de_Empleo/\n"
        "│   └── Solicitud_de_Empleo_Ficha_Ingreso.docx\n"
        "├── 02_CONTRATOS/                  (7 contratos)\n"
        "├── 03_DESCRIPCION_DE_CARGOS/       (7 descripciones)\n"
        "├── 04_PRESTACIONES/                (2 documentos)\n"
        "├── 05_SEGURIDAD_LABORAL/           (9 documentos)\n"
        "├── 06_REGISTROS_LEGALES/           (1 checklist)\n"
        "├── 07_AUTORIZACIONES/              (3 autorizaciones)\n"
        "├── 08_POLITICAS_INTERNAS/          (7 políticas)\n"
        "└── 09_CIERRE/                      (1 carta de aceptación)",
        size=9, font="Consolas", space_after=4)

    add_para(doc, "Notas importantes", size=12, bold=True, color=TEAL_DARK, space_before=6, space_after=3)
    add_bullet(doc, "Todos los scripts importan el módulo _common.py, que debe estar en la misma carpeta.", bold_lead="Dependencia:  ")
    add_bullet(doc, "Las firmas se mapean automáticamente: 'RR.HH.' → Directora Gerente (Esnatlim Elena Simoza, C.I. V-17.976.287).", bold_lead="Mapeo de firmas:  ")
    add_bullet(doc, "Todos los documentos usan la paleta teal #0F766E y membrete ALIKA PETS / Grupo Caval 1003.", bold_lead="Estilo:  ")
    add_bullet(doc, "Los formatos están conformes a LOTTT, LOPCYMAT, LOPDP, NT-01-2008, NT-02-2008 (Venezuela).", bold_lead="Normativa:  ")
    add_bullet(doc, "Recomendado que un abogado laboralista revise los 7 contratos antes del primer uso formal.", bold_lead="Revisión legal:  ")

    # ============================================================
    # CÓDIGOS COMPLETOS
    # ============================================================
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("CÓDIGOS PYTHON COMPLETOS")
    style_run(r, size=18, bold=True, color=TEAL_DARK)
    p.paragraph_format.space_after = Pt(4)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(8)
    add_hr(sep, color="0F766E", sz="10")

    add_para(doc,
        "A continuación se incluyen los 11 scripts Python completos en orden de dependencia. "
        "El primer script (_common.py) es módulo de utilidades compartidas y debe copiarse junto "
        "con los demás en la misma carpeta para que funcionen las importaciones.",
        size=10, italic=True, color=GRAY_TEXT, space_after=10)

    # Orden de scripts (dependencias primero)
    scripts_to_include = [
        "_common.py",
        "gen_solicitud_v3.py",
        "gen_contratos.py",
        "gen_funciones_por_cargo.py",
        "gen_prestaciones.py",
        "gen_riesgos_por_rol.py",
        "gen_protocolos_vet.py",
        "gen_documentos_finales_v3.py",
        "gen_autorizaciones.py",
        "gen_politicas.py",
        "gen_checklist_v3.py",
    ]

    section_num = 0
    for script_name in scripts_to_include:
        script_path = os.path.join(OUT_DIR, script_name)
        section_num += 1

        # Header de sección con número, nombre, descripción
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r1 = p.add_run(f"SCRIPT {section_num} de {len(scripts_to_include)}  ·  ")
        style_run(r1, size=10, bold=True, color=GRAY_TEXT)
        r2 = p.add_run(script_name)
        style_run(r2, size=12, bold=True, color=TEAL_DARK)

        # Línea separadora
        sep_p = doc.add_paragraph()
        sep_p.paragraph_format.space_before = Pt(0)
        sep_p.paragraph_format.space_after = Pt(6)
        sep_p.paragraph_format.keep_with_next = True
        add_hr(sep_p, color="CBD5E1", sz="6")

        # Leer contenido del script
        try:
            with open(script_path, "r", encoding="utf-8") as f:
                code_text = f.read()
        except Exception as e:
            code_text = f"# Error leyendo {script_name}: {e}"

        # Añadir bloque de código
        add_code_block(doc, code_text, filename=script_name)

    # ============================================================
    # CIERRE
    # ============================================================
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("FIN DEL DOCUMENTO")
    style_run(r, size=14, bold=True, color=TEAL_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("Kit de Ingresos del Trabajador — Documentación técnica completa v3.1")
    style_run(r, size=10, italic=True, color=GRAY_TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("ALIKA PETS  ·  Grupo Caval 1003, C.A.  ·  RIF J501662533")
    style_run(r, size=10, color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("Directora Gerente: Esnatlim Elena Simoza  ·  C.I. V-17.976.287")
    style_run(r, size=9, color=GRAY_TEXT)

    # ============================================================
    # FOOTER con número de página
    # ============================================================
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("ALIKA PETS  ·  Grupo Caval 1003, C.A.  ·  RIF J501662533  ·  Kit de Ingresos — Documentación v3.1  ·  Página ")
    style_run(fr, size=8, color=GRAY_MUTED)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r2 = fp.add_run()
    r2._r.append(fldChar1)
    r2._r.append(instrText)
    r2._r.append(fldChar2)
    style_run(r2, size=8, color=GRAY_MUTED)

    doc.save(OUT_FILE)
    return OUT_FILE

if __name__ == "__main__":
    print("Generando Kit_de_Ingresos.docx con todos los códigos Python...")
    out = build_document()
    size_kb = os.path.getsize(out) / 1024
    print(f"✓ {out}")
    print(f"  Tamaño: {size_kb:.1f} KB")
