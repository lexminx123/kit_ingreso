"""
gen_politicas.py — Genera 4 políticas internas ALIKA PETS / Grupo Caval 1003, C.A.

Documentos generados en /home/z/my-project/output/08_POLITICAS_INTERNAS/:
  1. Reglamento_Interno.docx           (Art. 191 LOTTT, 11 capítulos, tabla sanciones)
  2. Codigo_Conducta.docx              (carta Directora Gerente + principios + anti-soborno + recepción)
  3. Politica_Confidencialidad.docx    (PROPIEDAD HISTORIAS CLÍNICAS + no competencia 12m)
  4. Politica_Uso_Redes_Sociales.docx  (uso personal/corporativo, contenido prohibido)

Versión 3.0 — _common.py mapea automáticamente "RR.HH." → "DIRECTORA GERENTE" (Esnatlim Simoza).
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# importar utilidades compartidas
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _common import (
    setup_a4_portrait, add_membrete, add_doc_title, add_chapter, add_section,
    add_article, add_para, add_bullet, add_sanciones_table, add_signature_block,
    add_signature_block_rep_legal, add_reception_letter, add_footer, add_hr,
    set_cell_bg, set_cell_borders, set_cell_margins, write_cell,
    TEAL_DARK, TEAL_HDR_BG, GRAY_ALT, AMBER_BG, SLATE_BG, WHITE, BLACK,
    GRAY_TEXT, GRAY_MUTED, RED_CRIT, GREEN_OK,
    EMPRESA, RIF_EMP, MARCA, DOMICILIO_EMP,
    REP_LEGAL_NOMBRE, REP_LEGAL_CARGO, REP_LEGAL_CI,
    DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI,
)

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "08_POLITICAS_INTERNAS")
os.makedirs(OUT_DIR, exist_ok=True)

FOOTER_BASE = "Políticas Internas  ·  v3.0"

# ============================================================
# Helper: tabla simple con encabezado teal y filas alternas
# ============================================================
def add_data_table(doc, headers, rows, col_widths_cm=None, header_size=9, body_size=9):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        if col_widths_cm:
            c.width = Cm(col_widths_cm[ci])
        write_cell(c, h, size=header_size, bold=True,
                   color=RGBColor(0xFF, 0xFF, 0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            if col_widths_cm:
                c.width = Cm(col_widths_cm[ci])
            write_cell(c, str(val), size=body_size, color=BLACK, bg=bg,
                       align=WD_ALIGN_PARAGRAPH.LEFT if ci > 0 else WD_ALIGN_PARAGRAPH.CENTER)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl

# ============================================================
# Helper: cuadro destacado (aviso)
# ============================================================
def add_alert_box(doc, lines, bg=AMBER_BG, border="F59E0B"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(16.6)
    c = tbl.rows[0].cells[0]
    c.width = Cm(16.6)
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    for i, ln in enumerate(lines):
        if i > 0:
            br = p.add_run()
            br.add_break()
        if ln.startswith("**"):
            r = p.add_run(ln.strip("*"))
            style_run_local(r, size=9, bold=True, color=RED_CRIT)
        else:
            r = p.add_run(ln)
            style_run_local(r, size=9, color=BLACK)
    set_cell_bg(c, bg)
    set_cell_borders(c, color=border, sz="8")
    set_cell_margins(c, top=120, bottom=120, left=160, right=160)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

def style_run_local(run, size=10, bold=False, color=None, font="Calibri", italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

# ============================================================
# DOCUMENTO 1: REGLAMENTO INTERNO
# ============================================================
def gen_reglamento_interno():
    doc = Document()
    setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "REGLAMENTO INTERNO", "Régimen Disciplinario y Normativo", "v3.0  ·  RR.HH.")
    add_doc_title(doc, "REGLAMENTO INTERNO DE TRABAJO")
    add_para(doc,
        f"En cumplimiento de lo establecido en el artículo 191 de la Ley Orgánica del Trabajo, "
        f"los Trabajadores y las Trabajadoras (LOTTT), la empresa {EMPRESA} (marca comercial "
        f"{MARCA}), en adelante LA EMPRESA, dedicada a la prestación de servicios de clínica "
        f"veterinaria, tienda de mascotas y peluquería canina, con domicilio en {DOMICILIO_EMP}, "
        f"establece el presente Reglamento Interno de Trabajo, de cumplimiento obligatorio para "
        f"todo el personal bajo cualquier modalidad de contratación.",
        size=10, space_after=6)
    add_para(doc,
        "El presente Reglamento entra en vigencia a partir de su publicación mediante fijación "
        "en lugar visible del centro de trabajo y entrega de copia a cada trabajador(a), con "
        "constancia escrita de recepción conforme al artículo 192 de la LOTTT.",
        size=10, space_after=8)

    # ===== CAPÍTULO I — DISPOSICIONES GENERALES =====
    add_chapter(doc, "I", "DISPOSICIONES GENERALES")
    add_article(doc, 1, "Ámbito de aplicación",
        "Las disposiciones del presente Reglamento Interno son de carácter obligatorio para "
        "todo el personal de LA EMPRESA, sin distinción de jerarquía, cargo, modalidad de "
        "contratación, jornada o turno, incluyendo personal fijo, contratado, eventual, de "
        "prueba, pasantes y personal en formación.")
    add_article(doc, 2, "Finalidad",
        "El presente Reglamento tiene por finalidad establecer las normas de orden técnico, "
        "disciplinario y de conducta que regulan la prestación de servicios en LA EMPRESA, "
        "así como garantizar un ambiente de trabajo seguro, armónico, productivo y respetuoso "
        "de la dignidad de las personas, conforme a la Constitución de la República Bolivariana "
        "de Venezuela, la LOTTT, la LOPCYMAT y demás leyes aplicables.")
    add_article(doc, 3, "Principios",
        "Son principios rectores del presente Reglamento: la legalidad, la igualdad y no "
        "discriminación, el respeto a la dignidad humana, la buena fe, la primacía de la "
        "realidad, la protección al trabajo, la interdicción de la arbitrariedad, la "
        "adaptación del trabajo a la persona, la debida diligencia y la transparencia.")
    add_article(doc, 4, "Modificación",
        "LA EMPRESA se reserva el derecho de modificar el presente Reglamento cuando lo "
        "estime conveniente, conforme a las necesidades del servicio y a la legislación "
        "vigente. Toda modificación será notificada por escrito al personal con al menos "
        "quince (15) días hábiles de anticipación a su entrada en vigencia y se entregará "
        "copia actualizada con constancia de recepción.")

    # ===== CAPÍTULO II — ADMISIÓN, INGRESO Y CONTRATACIÓN =====
    add_chapter(doc, "II", "ADMISIÓN, INGRESO Y CONTRATACIÓN")
    add_article(doc, 5, "Requisitos de ingreso",
        "Toda persona que aspire a laborar en LA EMPRESA deberá consignar: cédula de identidad "
        "vigente, currículum vitae actualizado, constancia de trabajo anterior, referencias "
        "personales y laborales, constancias de IVSS, FAOV, INCES y PMSSO, título académico o "
        "certificado de competencia, inscripción en el Colegio Profesional correspondiente "
        "(médicos veterinarios), RIF personal y datos de cuenta bancaria activa.")
    add_article(doc, 6, "Período de prueba",
        "Conforme al artículo 22 de la LOTTT, el período de prueba será de treinta (30) días "
        "para contratos por tiempo indeterminado y de catorce (14) días para contratos a plazo "
        "determinado, durante los cuales cualquiera de las partes podrá dar por terminada la "
        "relación laboral sin causa justificada y sin derecho a prestaciones sociales, salvo "
        "indemnización de antigüedad proporcional conforme al artículo 125 de la LOTTT.")
    add_article(doc, 7, "Contrato de trabajo",
        "Toda relación laboral se formalizará por escrito en un contrato individual de trabajo "
        "que deberá contener las menciones obligatorias del artículo 71 de la LOTTT. El "
        "contrato se celebrará por duplicado, entregándose un ejemplar al trabajador(a) "
        "contra firma de recepción.")

    # ===== CAPÍTULO III — JORNADA, HORARIO Y DESCANSOS =====
    add_chapter(doc, "III", "JORNADA, HORARIO Y DESCANSOS")
    add_article(doc, 8, "Jornada de trabajo",
        "La jornada de trabajo será de cinco (5) días a la semana, conforme al artículo 173 "
        "de la LOTTT. La jornada diurna será de ocho (8) horas, la mixta de siete y media "
        "(7,5) horas y la nocturna de siete (7) horas, sin exceder cuarenta (40) horas "
        "semanales en cualquiera de sus modalidades.")
    add_article(doc, 9, "Horario",
        "El horario de atención al público es de lunes a sábado, en turnos rotativos "
        "establecidos por la Gerencia. Cada trabajador(a) recibirá su programación de turnos "
        "con al menos siete (7) días de anticipación. Los cambios de turno excepcionales "
        "serán comunicados con cuarenta y ocho (48) horas de anticipación, salvo casos "
        "fortuitos o de fuerza mayor.")
    add_article(doc, 10, "Descanso intrajornada",
        "Todo trabajador(a) tiene derecho a un descanso intrajornada de treinta (30) a "
        "sesenta (60) minutos, según jornada, conforme al artículo 184 de la LOTTT. Este "
        "descanso no se computará como tiempo efectivo de trabajo.")
    add_article(doc, 11, "Días feriados y de descanso",
        "Los días domingo son de descanso obligatorio remunerado. Los días feriados nacionales "
        "se regirán por la Ley de Feriados Nacionales. Cuando por necesidades del servicio "
        "(atención de emergencias veterinarias) deba laborarse un día de descanso u horas "
        "extraordinarias, se aplicará la compensación prevista en los artículos 171 y 171-A "
        "de la LOTTT (recargo del 50% en día feriado, 95% si fuere nocturno).")

    # ===== CAPÍTULO IV — DERECHOS Y DEBERES =====
    add_chapter(doc, "IV", "DERECHOS DEL TRABAJADOR(A)")
    add_article(doc, 12, "Derechos individuales",
        "Son derechos de cada trabajador(a): percibir la remuneración pactada en moneda de "
        "curso legal, disfrutar de jornada y horario establecidos, descanso semanal y "
        "feriados, vacaciones anuales remuneradas, bono de alimentación (Ley de Alimentación "
        "para los Trabajadores y las Trabajadoras), prestaciones sociales conforme a la "
        "LOTTT, utilidades, y un ambiente de trabajo seguro y saludable conforme a la "
        "LOPCYMAT.")
    add_article(doc, 13, "Derechos colectivos",
        "Los trabajadores y trabajadoras tienen derecho a la libertad sindical, a la "
        "negociación colectiva, a la huelga conforme a la ley, a reunirse en asamblea fuera "
        "del horario de trabajo y a elegir delegados de prevención ante el Comité de "
        "Seguridad y Salud Laboral conforme al artículo 46 de la LOPCYMAT.")
    add_article(doc, 14, "Capacitación",
        "Todo trabajador(a) tiene derecho a recibir capacitación y formación profesional "
        "permanente, conforme al artículo 56 de la LOTTT y a la Ley del INCES. LA EMPRESA "
        "facilitará al menos una (1) jornada de capacitación trimestral al personal.")

    add_chapter(doc, "V", "DEBERES Y OBLIGACIONES DEL TRABAJADOR(A)")
    add_article(doc, 15, "Deberes generales",
        "Son deberes de todo el personal: cumplir las órdenes e instrucciones legítimas de "
        "los superiores, observar buena conducta, conservar y usar correctamente las "
        "herramientas, equipos, suministros e instalaciones, mantener el aseo personal y "
        "del puesto de trabajo, asistir puntualmente, observar las normas de seguridad y "
        "salud laboral, y guardar lealtad y reserva sobre la información de LA EMPRESA.")
    add_bullet(doc, "Cumplir con el horario establecido, registrando la entrada y salida en el sistema habilitado.")
    add_bullet(doc, "Permanecer en el puesto durante la jornada, salvo autorización del supervisor.")
    add_bullet(doc, "Atender con cortesía, respeto y eficiencia a los clientes, pacientes y compañeros.")
    add_bullet(doc, "Portar el uniforme y los elementos de protección personal (EPP) establecidos.")
    add_bullet(doc, "Manipular con cuidado a los animales bajo atención y seguir los protocolos veterinarios.")
    add_bullet(doc, "Registrar con veracidad y oportunidad las historias clínicas, ventas y procedimientos.")
    add_bullet(doc, "Cuidar y mantener en buen estado el equipo, instrumental, medicamentos e instalaciones.")
    add_bullet(doc, "Reportar inmediatamente cualquier accidente de trabajo, incidente o condición insegura.")
    add_article(doc, 16, "Deberes especiales del personal veterinario y auxiliar",
        "El personal médico veterinario y auxiliar veterinario deberá cumplir adicionalmente "
        "con: el Código de Ética del Colegio Médico Veterinario de Venezuela (CMVV), la "
        "Norma Técnica NT-01-2008 de Bioseguridad, los protocolos internos de manejo de "
        "sustancias controladas, profilaxis antirrábica y residuos biomédicos, así como "
        "mantener vigente la inscripción en el CMVV y la certificación de vacunación "
        "antirrábica pre-exposición.")

    # ===== CAPÍTULO VI — PROHIBICIONES =====
    add_chapter(doc, "VI", "PROHIBICIONES AL PERSONAL")
    add_article(doc, 17, "Prohibiciones generales",
        "Conforme al artículo 79 de la LOTTT, queda terminantemente prohibido al personal:")
    add_bullet(doc, "Faltar al trabajo sin causa justificada o sin aviso previo al supervisor.")
    add_bullet(doc, "Presentarse al trabajo en estado de ebriedad o bajo efectos de sustancias psicotrópicas.")
    add_bullet(doc, "Consumir, distribuir o traficar sustancias estupefacientes o psicotrópicas en el centro de trabajo.")
    add_bullet(doc, "Portar armas de fuego, blancas o contundentes dentro de las instalaciones sin autorización.")
    add_bullet(doc, "Fumar o encender fuego en áreas no autorizadas (quirófano, depósito, farmacia, hospitalización).")
    add_bullet(doc, "Sustraer, ocultar o darnos bienes de LA EMPRESA, de los clientes o de las mascotas.")
    add_bullet(doc, "Maltratar física, verbal o psicológicamente a cualquier animal, compañero, cliente o proveedor.")
    add_bullet(doc, "Abandonar el puesto de trabajo sin causa justificada o sin dejar reemplazo.")
    add_bullet(doc, "Dormir durante la jornada de trabajo, en especial en turnos de hospitalización o guardia.")
    add_bullet(doc, "Divulgar información confidencial, técnica, comercial o de historias clínicas de pacientes.")
    add_bullet(doc, "Realizar trabajos particulares durante la jornada o usar equipos de LA EMPRESA para fines ajenos.")
    add_bullet(doc, "Aceptar dádivas, regalos o comisiones de proveedores o clientes (conflictos de interés).")
    add_bullet(doc, "Discriminar por raza, sexo, orientación sexual, edad, religión, ideología política o discapacidad.")
    add_bullet(doc, "Hostigar sexualmente o por cualquier causa a compañeros, subalternos, clientes o proveedores.")
    add_bullet(doc, "Publicar en redes sociales contenidos que afecten la imagen de LA EMPRESA o revelen información interna.")
    add_article(doc, 18, "Prohibiciones especiales",
        "Queda prohibido al personal veterinario y auxiliar: practicar abortos inducidos sin "
        "causa terapéutica documentada, eutanasias no autorizadas por el propietario del "
        "animal o sin protocolo veterinario, recetar sustancias controladas sin fundamento "
        "clínico, usar anestésicos u opioides sin control de inventario y sin doble chequeo, "
        "y abandonar un procedimiento quirúrgico sin causa justificada.")

    # ===== CAPÍTULO VII — ESCALA DE FALTAS Y SANCIONES =====
    add_chapter(doc, "VII", "ESCALA DE FALTAS")
    add_article(doc, 19, "Clasificación",
        "Las faltas cometidas por el personal se clasifican en leves, graves y muy graves, "
        "conforme a la naturaleza del hecho, las circunstancias y los perjuicios causados. "
        "La calificación corresponde a la Gerencia, previa opinión del Comité de "
        "Convivencia cuando aplique, conforme al artículo 191 de la LOTTT.")
    add_article(doc, 20, "Faltas leves",
        "Son faltas leves aquellas que no afecten significativamente la marcha del servicio "
        "ni causen perjuicio grave, tales como: impuntualidad de menos de quince (15) "
        "minutos, descuido menor en el aseo del puesto, no portar el uniforme correctamente, "
        "olvidar el carné identificativo, conversar en exceso durante la atención al "
        "público, descuidos menores de procedimiento.")
    add_article(doc, 21, "Faltas graves",
        "Son faltas graves: la impuntualidad reiterada (más de tres veces en un mes), "
        "inasistencia injustificada al trabajo, negligencia en el manejo de animales que "
        "cause lesiones, descuido de instrumental o equipo, incumplimiento de protocolos "
        "de bioseguridad, abandono temporal del puesto, incumplimiento del régimen de "
        "sustancias controladas, riñas o alteraciones del orden, uso indebido de redes "
        "sociales con afectación a la imagen de LA EMPRESA, falsedad en registros clínicos "
        "o de ventas.")
    add_article(doc, 22, "Faltas muy graves",
        "Son faltas muy graves: el robo o hurto, fraude, maltrato animal deliberado, "
        "agredir físicamente a compañeros o clientes, acoso sexual o laboral, presentarse "
        "bajo efectos de sustancias, tráfico de sustancias estupefacientes o psicotrópicas, "
        "divulgación de información confidencial o de historias clínicas, apropiación de "
        "fondos o especies, sabotaje, abandono del puesto en emergencia, así como la "
        "reincidencia en faltas graves.")

    add_chapter(doc, "VIII", "ESCALA DE SANCIONES DISCIPLINARIAS")
    add_article(doc, 23, "Tipos de sanciones",
        "Conforme al artículo 191 de la LOTTT y al principio de proporcionalidad, las "
        "sanciones aplicables son tres: amonestación verbal, amonestación escrita y "
        "despido justificado. Queda expresamente prohibido el descuento salarial "
        "automático como sanción, conforme al artículo 59 de la LOTTT; cualquier "
        "descuento por daños requiere autorización previa del Inspector del Trabajo.")
    add_article(doc, 24, "Tabla de sanciones",
        "La siguiente tabla resume el régimen sancionatorio aplicable:")
    sanciones_headers = ["Falta", "Tipo", "Sanción", "Constancia"]
    sanciones_rows = [
        ["Impuntualidad leve (<15 min)", "LEVE", "Amonestación verbal", "Anotación en expediente"],
        ["Descuido menor en aseo/uniforme", "LEVE", "Amonestación verbal", "Expediente del supervisor"],
        ["Impuntualidad reiterada (>3/mes)", "GRAVE", "Amonestación escrita", "Acta firmada por trabajador"],
        ["Inasistencia injustificada", "GRAVE", "Amonestación escrita", "Acta + descuento día no laborado"],
        ["Negligencia con animal (lesión)", "GRAVE", "Amonestación escrita", "Acta + investigación causa raíz"],
        ["Incumplir protocolo bioseguridad", "GRAVE", "Amonestación escrita", "Acta + capacitación reforzada"],
        ["Incumplir control sustancias", "GRAVE", "Amonestación escrita", "Acta + auditoría inventario"],
        ["Reincidencia en falta grave", "MUY GRAVE", "DESPIDO JUSTIFICADO", "Art. 79 LOTTT"],
        ["Robo, hurto, fraude", "MUY GRAVE", "DESPIDO JUSTIFICADO", "Art. 79.3 LOTTT"],
        ["Maltrato animal deliberado", "MUY GRAVE", "DESPIDO JUSTIFICADO", "Art. 79.4 + Ley Protección Animal"],
        ["Acoso sexual o laboral", "MUY GRAVE", "DESPIDO JUSTIFICADO", "Art. 79.6 LOTTT + Ley Especial Dignificación Mujer"],
        ["Presentarse bajo efectos", "MUY GRAVE", "DESPIDO JUSTIFICADO", "Art. 79.7 LOTTT"],
        ["Divulgar información confidencial", "MUY GRAVE", "DESPIDO JUSTIFICADO", "Art. 79.9 LOTTT + LOPDP"],
    ]
    add_sanciones_table(doc, sanciones_headers, sanciones_rows, col_widths_cm=[5.5, 2.5, 4.5, 4.1])
    add_para(doc, "Nota: Toda sanción escrita se notificará al trabajador(a) dentro de los "
             "tres (3) días hábiles siguientes, indicando los hechos, su calificación y la "
             "sanción impuesta. El trabajador(a) tiene derecho a firmar conformidad o dejar "
             "constancia de su disconformidad, sin que ello afecte la aplicación de la "
             "sanción. El expediente disciplinario se conservará por cinco (5) años conforme "
             "al artículo 183 de la LOTTT.",
        size=9, space_before=4, italic=True)

    # ===== CAPÍTULO IX — PROCEDIMIENTO DISCIPLINARIO =====
    add_chapter(doc, "IX", "PROCEDIMIENTO DISCIPLINARIO")
    add_article(doc, 25, "Investigación previa",
        "Antes de aplicar cualquier sanción escrita o despido, LA EMPRESA deberá practicar "
        "una investigación interna que permita esclarecer los hechos, identificar a los "
        "responsables y graduar la falta. La investigación incluirá la toma de declaraciones "
        "del trabajador(a) presuntamente infractor, testigos y supervisores, así como el "
        "análisis de pruebas documentales, videográficas o periciales.")
    add_article(doc, 26, "Derecho a la defensa",
        "El trabajador(a) tiene derecho a ser oído, a aportar pruebas, a estar acompañado "
        "por un representante sindical o compañero de trabajo, y a contestar los cargos "
        "imputados por escrito dentro de los tres (3) días hábiles siguientes a la "
        "notificación. La defensa se ejerce sin interrumpir la prestación del servicio, "
        "salvo suspensión preventiva justificada conforme al artículo 73 de la LOTTT.")
    add_article(doc, 27, "Notificación",
        "La sanción se notificará por escrito al trabajador(a), con copia al expediente y, "
        "cuando exista, al sindicato. La notificación contendrá: fecha, hechos imputados, "
        "calificación de la falta, sanción impuesta, fundamento legal y recurso que proceda.")
    add_article(doc, 28, "Recursos",
        "Contra la sanción impuesta el trabajador(a) podrá ejercer los recursos "
        "administrativos previstos en la ley y acudir ante la Inspectoría del Trabajo o la "
        "jurisdicción laboralista, conforme a los artículos 474 y siguientes de la LOTTT.")
    add_article(doc, "28b", "Prescripción y rehabilitación",
        "Las faltas leves prescriben a los treinta (30) días, las graves a los sesenta (60) "
        "días y las muy graves a los seis (6) meses, conforme al artículo 191 de la LOTTT. "
        "Las amonestaciones escritas se cancelan del expediente del trabajador(a) a los "
        "doce (12) meses sin nuevas faltas. La rehabilitación del trabajador(a) sancionado "
        "es automática al cumplir el plazo sin reincidencia.")
    add_article(doc, "28c", "Comité de Convivencia",
        "Cuando el número de trabajadores alcance el mínimo legal previsto en la LOTTT, "
        "se constituirá un Comité de Convivencia o de Disciplina paritario, que actuará "
        "como órgano asesor en la calificación de faltas y propuesta de sanciones. Sus "
        "integrantes serán electos en asamblea y durarán dos (2) años en sus funciones.")

    # ===== CAPÍTULO X — SEGURIDAD Y SALUD LABORAL =====
    add_chapter(doc, "X", "SEGURIDAD Y SALUD LABORAL")
    add_article(doc, 29, "Política de SST",
        "LA EMPRESA declara como prioridad la seguridad y salud de su personal, conforme a "
        "la LOPCYMAT y la Norma Técnica NT-01-2008. Se mantendrá un Programa de Seguridad "
        "y Salud en el Trabajo bajo la responsabilidad de la Dirección y del Comité de "
        "Seguridad y Salud Laboral, integrado paritariamente conforme al artículo 46 de la "
        "LOPCYMAT.")
    add_article(doc, 30, "Obligaciones del trabajador en SST",
        "Todo trabajador(a) está obligado a: usar correctamente el EPP entregado, cumplir "
        "los protocolos de bioseguridad veterinaria, reportar inmediatamente cualquier "
        "accidente de trabajo o incidente, participar en las capacitaciones, someterse a "
        "los exámenes médicos preventivos (NT-02-2008), y colaborar con el Comité SST.")
    add_article(doc, 31, "Notificación de riesgos",
        "Conforme al artículo 56 de la LOPCYMAT, LA EMPRESA entregará a cada trabajador(a) "
        "la Notificación de Riesgos específica de su cargo, así como la información sobre "
        "las medidas de prevención, control y atención en caso de accidente o enfermedad "
        "ocupacional.")
    add_article(doc, 32, "Prohibición de represalias",
        "Queda prohibido cualquier acto de represalia contra el trabajador(a) que reporte "
        "condiciones inseguras, se niegue a ejecutar tareas que pongan en peligro su vida "
        "o salud, o participe en las actividades del Comité SST, conforme al artículo 26 "
        "de la LOPCYMAT.")

    # ===== CAPÍTULO XI — DISPOSICIONES FINALES =====
    add_chapter(doc, "XI", "DISPOSICIONES FINALES")
    add_article(doc, 33, "Vigencia",
        "El presente Reglamento Interno entra en vigencia desde su fijación en lugar visible "
        "del centro de trabajo y la entrega de copia individual al personal con constancia "
        "de recepción, conforme al artículo 192 de la LOTTT.")
    add_article(doc, 34, "Revisión",
        "El presente Reglamento será revisado anualmente, o cuando modifiquen las leyes "
        "aplicables o las condiciones operativas de LA EMPRESA. Las modificaciones serán "
        "notificadas conforme al artículo 4 del presente Reglamento.")
    add_article(doc, 35, "Prevalencia legal",
        "Las cláusulas del presente Reglamento se interpretarán en armonía con la "
        "Constitución, la LOTTT, la LOPCYMAT, la LOPDP y demás leyes aplicables. En caso "
        "de conflicto, prevalecerá la norma más favorable al trabajador(a), conforme al "
        "principio in dubio pro operario.")
    add_article(doc, 36, "Documentos complementarios",
        "Forman parte integrante del presente Reglamento: el Código de Conducta, la "
        "Política de Confidencialidad, la Política de Uso de Redes Sociales, los Protocolos "
        "Veterinarios (Bioseguridad, Sustancias Controladas, Mordeduras/Zoonosis y Reporte "
        "de Incidentes), la Cartilla de Bioseguridad y las Notificaciones de Riesgos por "
        "cargo, todos entregados al trabajador(a) contra firma de recepción.")
    add_article(doc, 37, "Publicación y difusión",
        f"El presente Reglamento será publicado en lugar visible del centro de trabajo, en "
        f"la cartelera laboral y en el sistema informático de {EMPRESA} accesible al "
        f"personal. Se entregará copia individual al momento del ingreso y cada vez que "
        f"sean modificadas sus disposiciones. La Constancia de Recepción firmada por el "
        f"trabajador(a) se archivará en su expediente personal por cinco (5) años conforme "
        f"al artículo 183 de la LOTTT.")
    add_article(doc, 38, "Capacitación sobre el Reglamento",
        "Al momento del ingreso, el nuevo personal recibirá una sesión de inducción de al "
        "menos una (1) hora sobre el contenido del presente Reglamento, a cargo del área "
        "de Recursos Humanos o de la Dirección. La capacitación se repetirá anualmente "
        "con énfasis en las modificaciones recientes.")
    add_article(doc, 39, "Consulta y aclaratorias",
        "Cualquier duda sobre la interpretación del presente Reglamento será resuelta por "
        "la Dirección de Recursos Humanos. Las aclaratorias formales se emitirán por "
        "escrito y serán de aplicación general, incorporándose a la próxima revisión.")
    add_article(doc, 40, "Aprobación",
        f"Aprobado por la Junta Directiva de {EMPRESA} en sesión ordinaria, y firmado por "
        f"la Directora Gerente en representación operativa de LA EMPRESA.")

    # ===== Carta de recepción (firma DIRECTORA GERENTE automáticamente) =====
    add_reception_letter(doc, "Reglamento Interno de Trabajo")

    # Pie de página
    add_footer(doc.sections[0], FOOTER_BASE)
    out_path = os.path.join(OUT_DIR, "Reglamento_Interno.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# DOCUMENTO 2: CÓDIGO DE CONDUCTA
# ============================================================
def gen_codigo_conducta():
    doc = Document()
    setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "CÓDIGO DE CONDUCTA", "Ética e Integridad", "v3.0  ·  RR.HH.")
    add_doc_title(doc, "CÓDIGO DE CONDUCTA Y ÉTICA EMPRESARIAL")

    # ---- Carta introductoria de la Dirección (firmada por Esnatlim Elena Simoza, Directora Gerente) ----
    add_section(doc, "CARTA DE LA DIRECCIÓN")
    add_para(doc,
        f"Estimado(a) colaborador(a):",
        size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc,
        f"Es para mí un honor darles la bienvenida a {EMPRESA}, empresa que opera bajo la "
        f"marca comercial {MARCA}, dedicada a la atención clínica veterinaria, la "
        f"comercialización de productos para mascotas y los servicios de peluquería canina. "
        f"Desde nuestra fundación hemos asumido el compromiso de prestar un servicio de "
        f"excelencia, basado en el respeto por la vida animal, el trato digno a nuestros "
        f"clientes y el profesionalismo de nuestro equipo humano.",
        size=10, space_after=4)
    add_para(doc,
        "El presente Código de Conducta establece los principios éticos y las reglas de "
        "comportamiento que deben guiar la actuación de todas las personas que laboran en "
        "nuestra empresa. Su cumplimiento no es opcional: constituye la base sobre la cual "
        "se construye la confianza de nuestros clientes, proveedores y de la sociedad en "
        "general. Lo que aquí se establece refleja nuestros valores: respeto, integridad, "
        "compasión animal, transparencia y compromiso.",
        size=10, space_after=4)
    add_para(doc,
        "Solicito a cada uno(a) de ustedes leer detenidamente este documento, comprenderlo "
        "y hacerlo suyo. Cualquier duda será aclarada por la Dirección de Recursos Humanos. "
        "La firma de recepción al final del documento tiene carácter de aceptación formal.",
        size=10, space_after=4)
    add_para(doc, "Atentamente,", size=10, space_after=10)
    # firma Directora Gerente (Esnatlim)
    add_signature_block(doc, ["DIRECTORA GERENTE"])

    # ===== CAPÍTULO I — PRINCIPIOS Y VALORES =====
    add_chapter(doc, "I", "PRINCIPIOS Y VALORES")
    add_article(doc, 1, "Objeto",
        "El presente Código de Conducta tiene por objeto establecer los principios éticos "
        "y las reglas de comportamiento aplicables a todo el personal de LA EMPRESA, "
        "independientemente del cargo, jerarquía o modalidad de contratación.")
    add_article(doc, 2, "Valores institucionales",
        "Son valores fundamentales de LA EMPRESA: el respeto a la vida animal, la "
        "integridad personal, la honestidad, la transparencia, la responsabilidad, la "
        "calidad en el servicio, el trabajo en equipo, la confidencialidad, la no "
        "discriminación y la mejora continua.")

    # ===== CAPÍTULO II — CONDUCTA CON CLIENTES Y MASCOTAS =====
    add_chapter(doc, "II", "CONDUCTA CON LOS CLIENTES Y SUS MASCOTAS")
    add_article(doc, 3, "Atención al cliente",
        "Todo cliente será atendido con cortesía, respeto y prontitud, sin discriminación "
        "alguna. Se brindará información clara, veraz y oportuna sobre los servicios, "
        "costos, alternativas terapéuticas y pronósticos. La negativa injustificada de "
        "atención constituye falta grave.")
    add_article(doc, 4, "Bienestar animal",
        "Todo animal bajo atención será tratado con compasión y respeto. Se aplicarán los "
        "Cinco Libertades del bienestar animal: libre de hambre y sed; libre de "
        "incomodidades; libre de dolor, lesiones o enfermedad; libre de expresar conductas "
        "naturales; y libre de miedo y estrés. El maltrato animal deliberado constituye "
        "falta muy grave y será causal de despido conforme al artículo 79 de la LOTTT y la "
        "Ley para la Protección de la Fauna Doméstica Libre y en Cautiverio.")
    add_article(doc, 5, "Transparencia en honorarios",
        "Los honorarios y precios deberán informarse al cliente antes de iniciar cualquier "
        "servicio, conforme al artículo 50 de la Ley para la Defensa de las Personas en el "
        "Acceso a Bienes y Servicios (INDEPABIS). Queda prohibido cobrar por servicios no "
        "prestados o inflar facturación.")

    # ===== CAPÍTULO III — CONDUCTA ENTRE COMPAÑEROS =====
    add_chapter(doc, "III", "CONDUCTA ENTRE COMPAÑEROS DE TRABAJO")
    add_article(doc, 6, "Trato respetuoso",
        "Las relaciones entre compañeros se basarán en el respeto, la cortesía y la "
        "colaboración. Se prohíbe toda forma de violencia física, verbal, psicológica, "
        "exclusiones discriminatorias, burlas y rumores que afecten el clima laboral.")
    add_article(doc, 7, "Acoso sexual y laboral",
        "Queda terminantemente prohibido el acoso sexual y el acoso laboral en cualquiera "
        "de sus formas (Ley Especial para la Dignificación de la Mujer Trabajadora y "
        "LOPCYMAT art. 23). Todo reporte será atendido con confidencialidad, sin "
        "represalias, y dará lugar a investigación inmediata por la Gerencia y el Comité "
        "SST.")

    # ===== CAPÍTULO IV — CONDUCTA CON PROVEEDORES Y COMPETENCIA =====
    add_chapter(doc, "IV", "RELACIÓN CON PROVEEDORES Y COMPETENCIA")
    add_article(doc, 8, "Selección de proveedores",
        "La selección de proveedores se realizará con base en criterios objetivos de "
        "calidad, precio, oportunidad y trayectoria. Queda prohibido recibir regalos, "
        "comisiones o ventajas de proveedores o clientes que comprometan la imparcialidad "
        "de las decisiones comerciales.")
    add_article(doc, 9, "Competencia leal",
        "LA EMPRESA y su personal observarán las normas de competencia leal. Queda prohibido "
        "desprestigiar a competidores mediante afirmaciones falsas, así como obtener "
        "información confidencial de la competencia por medios indebidos.")

    # ===== CAPÍTULO V — CONFLICTOS DE INTERÉS =====
    add_chapter(doc, "V", "CONFLICTOS DE INTERÉS")
    add_article(doc, 10, "Definición",
        "Existe conflicto de interés cuando los intereses personales del trabajador(a) — "
        "directos o indirectos — se oponen o pudieran oponerse a los intereses de LA "
        "EMPRESA, comprometiendo la imparcialidad de sus decisiones.")
    add_article(doc, 11, "Deber de declaración",
        "Todo trabajador(a) deberá declarar por escrito ante la Gerencia cualquier situación "
        "que pudiera constituir conflicto de interés: vínculos familiares con proveedores o "
        "competidores, participación accionaria en empresas relacionadas, prestación de "
        "servicios a clientes fuera del horario laboral, recibimiento de regalos o "
        "atenciones de terceros.")
    add_article(doc, 12, "Actividades externas",
        "El personal médico veterinario y auxiliar veterinario deberá abstenerse de atender "
        "privadamente a clientes de LA EMPRESA fuera del horario laboral, así como de "
        "derivar pacientes a consultorios externos sin autorización. El incumplimiento "
        "constituye falta grave.")

    # ===== CAPÍTULO VI — ANTI-SOBORNO Y ANTI-CORRUPCIÓN =====
    add_chapter(doc, "VI", "POLÍTICA ANTI-SOBORNO Y ANTI-CORRUPCIÓN")
    add_article(doc, 13, "Prohibición absoluta",
        "Queda terminantemente prohibido a todo el personal ofrecer, prometer, dar, aceptar "
        "o solicitar, directa o indirectamente, dinero, regalos, favores o cualquier ventaja "
        "a funcionarios públicos, clientes, proveedores o terceros para obtener o retainner "
        "negocios o ventajas indebidas. Esta prohibición se extiende a familiares y personas "
        "interpuestas.")
    add_article(doc, 14, "Regalos y atenciones",
        "Solo se aceptarán regalos de cortesía de valor simbólico (menos de cincuenta "
        "dólares estadounidenses o su equivalente) y que no comprometan la imparcialidad "
        "del trabajador(a). Cualquier regalo que exceda este monto deberá ser rechazado y "
        "comunicado a la Gerencia, quien decidirá su destino.")
    add_article(doc, 15, "Sanciones",
        "El incumplimiento de la política anti-soborno constituye falta muy grave y será "
        "causal de despido justificado conforme al artículo 79 de la LOTTT, sin perjuicio "
        "de las acciones penales y civiles a que hubiere lugar, conforme a la Convención de "
        "las Naciones Unidas contra la Corrupción y la legislación venezolana.")
    add_alert_box(doc, [
        "**PROTOCOLO ANTI-SOBORNO — RECORDATORIO OBLIGATORIO**",
        "• Todo regalo, invitación o atención que exceda los USD 50 debe ser DECLARADO por escrito a la Gerencia.",
        "• Está PROHÍBIDO ofrecer dinero o ventajas a funcionarios del SENAC, INDEPABIS, Inspectoría del Trabajo, IVSS, INCES, FAOV, Colegio Médico Veterinario o cualquier autoridad.",
        "• Todo acto de corrupción detectado debe ser reportado en un máximo de 24 horas a la Directora Gerente.",
        "• El reporte de buena fe no generará represalias (LOPCYMAT art. 26).",
    ])

    # ===== CAPÍTULO VII — CONFIDENCIALIDAD Y PROTECCIÓN DE DATOS =====
    add_chapter(doc, "VII", "CONFIDENCIALIDAD Y PROTECCIÓN DE DATOS")
    add_article(doc, 16, "Deber de reserva",
        "Todo trabajador(a) está obligado a guardar reserva sobre la información técnica, "
        "comercial, financiera y operativa de LA EMPRESA, así como sobre los datos de "
        "clientes, proveedores, compañeros y, especialmente, sobre las historias clínicas "
        "de los pacientes veterinarios, conforme a la LOPDP y la Ley de Ejercicio de la "
        "Medicina Veterinaria (art. 23).")
    add_article(doc, 17, "Duración",
        "El deber de confidencialidad se mantiene durante la relación laboral y se extiende "
        "por cinco (5) años después de finalizada, conforme a la Política de "
        "Confidencialidad de LA EMPRESA.")

    # ===== CAPÍTULO VIII — USO DE REDES SOCIALES =====
    add_chapter(doc, "VIII", "USO DE REDES SOCIALES")
    add_article(doc, 18, "Principios",
        "El uso de redes sociales por parte del personal se regirá por la Política de Uso "
        "de Redes Sociales de LA EMPRESA. En resumen: no se publicará contenido que afecte "
        "la imagen de la empresa, no se revelará información confidencial, no se publicarán "
        "fotos de mascotas sin consentimiento del propietario, y no se publicarán opiniones "
        "sobre incidentes antes del comunicado oficial.")

    # ===== CAPÍTULO IX — PROTECCIÓN A PERSONAS VULNERABLES =====
    add_chapter(doc, "IX", "PROTECCIÓN A PERSONAS VULNERABLES")
    add_article(doc, 19, "Principio de protección",
        "LA EMPRESA velará por un trato preferente a mujeres embarazadas, personas con "
        "discapacidad, adultos mayores, adolescentes trabajadores (cuando aplique conforme "
        "a la Ley del Trabajo de los y las Adolescentes) y cualquier persona en condición "
        "de vulnerabilidad, conforme a la Constitución, la LOPCYMAT y la Ley Especial para "
        "la Dignificación de la Mujer Trabajadora.")
    add_article(doc, 20, "Mujeres embarazadas y en lactancia",
        "Las trabajadoras embarazadas o en período de lactancia no desempeñarán labores que "
        "impliquen exposición a radiaciones ionizantes, anestésicos volátiles, quimioterapia "
        "o zoonosis de elevado riesgo, conforme al artículo 78 de la LOPCYMAT y al "
        "protocolo de Bioseguridad Veterinaria. Gozarán del descanso pre y postnatal "
        "previsto en la Ley para la Protección de la Familia, la Maternidad y la Paternidad.")

    # ===== CAPÍTULO X — DENUNCIA Y CANAL ÉTICO =====
    add_chapter(doc, "X", "CANAL DE DENUNCIA ÉTICA")
    add_article(doc, 21, "Canal de denuncia",
        "Cualquier trabajador(a), cliente o proveedor podrá denunciar conductas contrarias "
        "al presente Código a través de los siguientes canales: (i) reporte verbal o "
        "escrito a la Directora Gerente; (ii) buzón físico ubicado en la recepción; "
        "(iii) correo electrónico habilitado por LA EMPRESA para denuncias; (iv) Comité de "
        "Seguridad y Salud Laboral.")
    add_article(doc, 22, "Confidencialidad y anti-represalias",
        "Toda denuncia será tratada con estricta confidencialidad. Queda prohibida cualquier "
        "forma de represalia contra quien denuncie de buena fe una conducta indebida, "
        "conforme al artículo 26 de la LOPCYMAT. La denuncia maliciosa o falsa constituye "
        "falta grave.")

    # ===== CAPÍTULO XI — SANCIONES =====
    add_chapter(doc, "XI", "SANCIONES POR INCUMPLIMIENTO")
    add_article(doc, 23, "Régimen aplicable",
        "El incumplimiento del presente Código será sancionado conforme al Reglamento "
        "Interno de Trabajo (Capítulos VII y VIII) y al artículo 79 de la LOTTT. Las "
        "infracciones graves o muy graves podrán ser causal de despido justificado sin "
        "derecho a indemnización sustitutiva de preaviso ni de antigüedad adicional.")

    # ===== Carta de recepción (firma DIRECTORA GERENTE automáticamente) =====
    add_reception_letter(doc, "Código de Conducta y Ética Empresarial")

    add_footer(doc.sections[0], FOOTER_BASE)
    out_path = os.path.join(OUT_DIR, "Codigo_Conducta.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# DOCUMENTO 3: POLÍTICA DE CONFIDENCIALIDAD
# ============================================================
def gen_politica_confidencialidad():
    doc = Document()
    setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "POLÍTICA DE CONFIDENCIALIDAD", "Protección de Información", "v3.0  ·  RR.HH.")
    add_doc_title(doc, "POLÍTICA DE CONFIDENCIALIDAD Y PROTECCIÓN DE INFORMACIÓN")

    # ===== CAPÍTULO I — DISPOSICIONES GENERALES =====
    add_chapter(doc, "I", "DISPOSICIONES GENERALES")
    add_article(doc, 1, "Objeto",
        "La presente Política tiene por objeto establecer las reglas para la protección, "
        "manejo, conservación y divulgación de la información confidencial de LA EMPRESA, "
        "sus clientes, proveedores, personal y pacientes veterinarios, conforme a la "
        "Constitución (art. 60), la LOPDP (Decreto 1.419), la Ley de Ejercicio de la "
        "Medicina Veterinaria y demás normas aplicables.")
    add_article(doc, 2, "Ámbito de aplicación",
        "Esta Política aplica a todo el personal de LA EMPRESA sin distinción, incluyendo "
        "directivos, contratados, pasantes, personal en prueba, así como a proveedores y "
        "terceros que tengan acceso a información confidencial.")
    add_article(doc, 3, "Aprobación y vigencia",
        f"Aprobada por la Dirección de {EMPRESA}. Entra en vigencia desde su entrega con "
        f"constancia de recepción al personal.")

    # ===== CAPÍTULO II — DEFINICIONES =====
    add_chapter(doc, "II", "DEFINICIONES")
    add_bullet(doc, "Toda información técnica, comercial, financiera, operativa o estratégica de LA EMPRESA.",
        bold_lead="Información Confidencial:  ")
    add_bullet(doc, "Toda información identificable relativa a clientes, proveedores o compañeros (nombre, C.I., teléfono, dirección, etc.).",
        bold_lead="Datos Personales:  ")
    add_bullet(doc, "Datos médicos, clínicos y quirúrgicos de animales atendidos, así como los datos personales de sus propietarios vinculados a tales servicios.",
        bold_lead="Historias Clínicas Veterinarias:  ")
    add_bullet(doc, "Secretos empresariales, listados de clientes y proveedores, costos, márgenes, estrategias comerciales, contratos y procesos técnicos.",
        bold_lead="Secretos Comerciales:  ")

    # ===== CAPÍTULO III — INFORMACIÓN CONFIDENCIAL =====
    add_chapter(doc, "III", "INFORMACIÓN SUJETA A CONFIDENCIALIDAD")
    add_article(doc, 4, "Enumeración enunciativa",
        "Se considera información confidencial, entre otras:")
    add_bullet(doc, "Historias clínicas veterinarias, registros de vacunación, protocolos quirúrgicos, imágenes diagnósticas y cualquier dato clínico de los pacientes.")
    add_bullet(doc, "Datos personales de clientes y de sus mascotas: nombre, C.I., teléfono, dirección, historial de pagos.")
    add_bullet(doc, "Datos personales del personal: documentos de identidad, salarios, historial médico ocupacional, datos bancarios.")
    add_bullet(doc, "Estrategias comerciales, listas de precios, márgenes, descuentos especiales, listado de proveedores y condiciones contractuales.")
    add_bullet(doc, "Protocolos técnicos, procedimientos internos, manuales, software, bases de datos, contraseñas y accesos.")
    add_bullet(doc, "Información sobre proyectos, lanzamientos de productos, alianzas, expansiones, antes de su publicación oficial.")
    add_bullet(doc, "Cualquier información señalada expresamente como confidencial por LA EMPRESA o que por su naturaleza deba ser tratada como tal.")

    # ===== CAPÍTULO IV — PROPIEDAD DE HISTORIAS CLÍNICAS =====
    add_chapter(doc, "IV", "PROPIEDAD DE HISTORIAS CLÍNICAS VETERINARIAS")
    add_alert_box(doc, [
        "**CLÁUSULA ESENCIAL — PROPIEDAD DE LAS HISTORIAS CLÍNICAS**",
        f"Conforme al artículo 23 de la Ley de Ejercicio de la Medicina Veterinaria y "
        f"protocolos internos de {EMPRESA}, todas las historias clínicas veterinarias, "
        f"registros de vacunación, protocolos anestésicos, imágenes diagnósticas y demás "
        f"registros clínicos elaborados durante la relación laboral son PROPIEDAD EXCLUSIVA "
        f"de {EMPRESA} y no del profesional que los elabore.",
        "El personal médico veterinario y auxiliar veterinario NO podrán sustraer, copiar, "
        "fotografiar, transferir o retener historias clínicas u otros registros clínicos "
        "al término de su relación laboral. El incumplimiento constituye falta muy grave y "
        "será causal de despido justificado conforme al artículo 79 de la LOTTT, sin "
        "perjuicio de las acciones civiles y penales correspondientes.",
    ], bg=AMBER_BG)
    add_article(doc, 5, "Titularidad",
        f"Las historias clínicas y demás registros clínicos son propiedad exclusiva de "
        f"{EMPRESA}. Los profesionales que las elaboren actúan en ejercicio de sus "
        f"funciones laborales y no adquieren derecho de propiedad intelectual sobre las "
        f"mismas, conforme al artículo 23 de la Ley de Ejercicio de la Medicina Veterinaria.")
    add_article(doc, 6, "Custodia y archivo",
        "Las historias clínicas se archivarán en el sistema informático de LA EMPRESA o, "
        "en su defecto, en el archivo físico bajo llave, con acceso restringido al "
        "personal autorizado. La conservación se realizará por al menos diez (10) años "
        "conforme al artículo 183 de la LOTTT y la normativa del CMVV.")
    add_article(doc, 7, "Acceso y consulta",
        "El cliente propietario de la mascota tiene derecho a obtener copia de la historia "
        "clínica de su animal, previa solicitud escrita y pago de la tarifa administrativa "
        "establecida. La entrega se hará contra firma de recepción. El personal no podrá "
        "divulgar datos clínicos a terceros sin autorización escrita del propietario.")

    # ===== CAPÍTULO V — OBLIGACIONES DE CONFIDENCIALIDAD =====
    add_chapter(doc, "V", "OBLIGACIONES DEL PERSONAL")
    add_article(doc, 8, "Obligaciones positivas",
        "El personal deberá:")
    add_bullet(doc, "Manejar la información confidencial con la máxima diligencia y exclusivamente para fines laborales.")
    add_bullet(doc, "Guardar y proteger los dispositivos, contraseñas y accesos a sistemas informáticos.")
    add_bullet(doc, "Cerrar sesión al abandonar el puesto de trabajo.")
    add_bullet(doc, "Reportar inmediatamente cualquier sospecha de acceso indebido, pérdida o filtración de información.")
    add_article(doc, 9, "Obligaciones negativas",
        "El personal NO deberá:")
    add_bullet(doc, "Divulgar, compartir, publicar o transmitir información confidencial a terceros, dentro o fuera de LA EMPRESA.")
    add_bullet(doc, "Utilizar la información confidencial para beneficio propio o de terceros.")
    add_bullet(doc, "Sustraer copias de historias clínicas, listados de clientes, protocolos o cualquier documento de LA EMPRESA.")
    add_bullet(doc, "Acceder a sistemas o áreas no autorizadas para su cargo.")
    add_bullet(doc, "Discutir asuntos confidenciales en presencia de clientes, proveedores o terceros no autorizados.")

    # ===== CAPÍTULO VI — DURACIÓN Y POST-CONTRACTUALIDAD =====
    add_chapter(doc, "VI", "DURACIÓN Y OBLIGACIONES POST-CONTRACTUALES")
    add_article(doc, 10, "Duración del deber de confidencialidad",
        "El deber de confidencialidad se mantiene durante toda la relación laboral y se "
        "extiende por cinco (5) años contados a partir de la finalización de la relación, "
        "sea cual fuere la causa de terminación.")
    add_article(doc, 11, "Cláusula de no competencia post-contractual",
        "Tratándose de personal con acceso a información estratégica, cartera de clientes "
        "o historias clínicas, se acuerda una cláusula de no competencia parcial post-"
        "contractual por un plazo máximo de doce (12) meses contados a partir de la "
        "terminación de la relación laboral.")
    add_para(doc, "Alcance y limitaciones de la cláusula:",
        size=10, bold=True, space_after=2)
    add_bullet(doc, "Aplica exclusivamente al Médico Veterinario, Auxiliar Veterinario y Dog Groomer.", bold_lead="Cargos afectados:  ")
    add_bullet(doc, "Dentro de un radio de TRES (3) kilómetros del centro de trabajo de LA EMPRESA en Los Teques, Estado Miranda.", bold_lead="Ámbito geográfico:  ")
    add_bullet(doc, "Clínicas veterinarias, consultorios, tiendas de mascotas y peluquerías caninas que presten los mismos servicios que LA EMPRESA.", bold_lead="Actividades restringidas:  ")
    add_bullet(doc, "No se prohíbe al trabajador ejercer su profesión en establecimientos no competidores directamente, en otras zonas geográficas, ni en el sector público.", bold_lead="Exclusiones:  ")
    add_bullet(doc, "Conforme a los artículos 27 y 28 de la LOTTT, esta cláusula no excederá el plazo razonable para tutelar el legítimo interés de LA EMPRESA, y no impedirá al trabajador ejercer su derecho al trabajo.", bold_lead="Proporcionalidad:  ")
    add_bullet(doc, "El incumplimiento de esta cláusula generará responsabilidad civil por daños y perjuicios, sin perjuicio de las acciones legales correspondientes.", bold_lead="Consecuencias:  ")

    # ===== CAPÍTULO VII — PROTECCIÓN DE DATOS PERSONALES =====
    add_chapter(doc, "VII", "PROTECCIÓN DE DATOS PERSONALES (LOPDP)")
    add_article(doc, 12, "Tratamiento de datos",
        "El tratamiento de datos personales de clientes, trabajadores y terceros se regirá "
        "por la LOPDP (Decreto 1.419). LA EMPRESA adoptará las medidas técnicas, "
        "organizativas y legales necesarias para garantizar la seguridad, confidencialidad "
        "e integridad de los datos.")
    add_article(doc, 13, "Derechos de los titulares",
        "Los titulares de datos personales tienen derecho a acceder, rectificar, cancelar y "
        "opponerse al tratamiento de sus datos, así como a revocar el consentimiento otorgado "
        "y a ser informados sobre las finalidades del tratamiento (derechos ARCO+).")
    add_article(doc, 14, "Medidas de seguridad",
        "LA EMPRESA mantendrá controles de acceso, cifrado, copias de respaldo, "
        "auditorías de uso y bitácoras de accesos. El personal está obligado a cumplir "
        "estrictamente las políticas de seguridad informática y tratamiento de datos.")

    # ===== CAPÍTULO VIII — SANCIONES Y RESPONSABILIDADES =====
    add_chapter(doc, "VIII", "SANCIONES Y RESPONSABILIDADES")
    add_article(doc, 15, "Sanciones laborales",
        "El incumplimiento de la presente Política constituye falta grave o muy grave "
        "según el caso, conforme al Reglamento Interno de Trabajo y al artículo 79 de la "
        "LOTTT. La divulgación deliberada de información confidencial o de historias "
        "clínicas será causal de despido justificado.")
    add_article(doc, 16, "Responsabilidad civil y penal",
        "Sin perjuicio de las sanciones laborales, el responsable responderá civilmente por "
        "los daños y perjuicios causados a LA EMPRESA, a clientes o a terceros, y "
        "penalmente conforme a la Ley Especial contra los Delitos Informáticos, la LOPDP "
        "y demás leyes aplicables.")

    # ===== Carta de recepción (firma DIRECTORA GERENTE automáticamente) =====
    add_reception_letter(doc, "Política de Confidencialidad y Protección de Información")

    add_footer(doc.sections[0], FOOTER_BASE)
    out_path = os.path.join(OUT_DIR, "Politica_Confidencialidad.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# DOCUMENTO 4: POLÍTICA DE USO DE REDES SOCIALES
# ============================================================
def gen_politica_redes_sociales():
    doc = Document()
    setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "POLÍTICA DE REDES SOCIALES", "Uso Personal y Corporativo", "v3.0  ·  RR.HH.")
    add_doc_title(doc, "POLÍTICA DE USO DE REDES SOCIALES")

    # ===== CAPÍTULO I — DISPOSICIONES GENERALES =====
    add_chapter(doc, "I", "DISPOSICIONES GENERALES")
    add_article(doc, 1, "Objeto",
        "La presente Política establece las reglas para el uso personal y corporativo de "
        "redes sociales por parte del personal de LA EMPRESA, con el fin de proteger la "
        "imagen institucional, la confidencialidad de la información y los derechos de "
        "terceros, conforme a la Constitución (arts. 57 y 60), la LOPDP, la LOTTT y la Ley "
        "Especial contra los Delitos Informáticos.")
    add_article(doc, 2, "Ámbito de aplicación",
        "Aplica a todo el personal, dentro y fuera del horario laboral, cuando el contenido "
        "publicado se refiera directa o indirectamente a LA EMPRESA, a sus clientes, "
        "mascotas, compañeros, proveedores o a la actividad veterinaria, comercial o de "
        "peluquería canina que en ella se desarrolla.")
    add_article(doc, 3, "Plataformas incluidas",
        "Se incluyen todas las plataformas sociales: Instagram, Facebook, TikTok, X (Twitter), "
        "WhatsApp, YouTube, LinkedIn, Snapchat, Threads, Telegram y cualquier otra red "
        "social, blog, foro o plataforma de mensajería, presente o futura.")

    # ===== CAPÍTULO II — USO PERSONAL DE REDES SOCIALES =====
    add_chapter(doc, "II", "USO PERSONAL DE REDES SOCIALES")
    add_article(doc, 4, "Derecho al uso personal",
        "LA EMPRESA respeta el derecho del personal al uso personal de redes sociales en "
        "su tiempo libre. Sin embargo, dicho uso no podrá afectar la imagen, reputación, "
        "confidencialidad o intereses legítimos de LA EMPRESA.")
    add_article(doc, 5, "Uso durante el horario laboral",
        "Durante la jornada de trabajo, el uso personal de redes sociales se limita a los "
        "descansos autorizados. Queda prohibido usar el teléfono móvil u otros dispositivos "
        "para fines personales durante la atención a clientes, procedimientos veterinarios "
        "o manejo de animales, salvo emergencia comunicada al supervisor.")
    add_article(doc, 6, "Identificación como empleado",
        "Cuando un trabajador(a) publique contenido relacionado con LA EMPRESA, deberá "
        "hacerlo de manera que no se interprete como comunicado oficial. Se prohíbe "
        "presentarse como portavoz de LA EMPRESA sin autorización escrita de la Gerencia.")

    # ===== CAPÍTULO III — CONTENIDO PROHIBIDO =====
    add_chapter(doc, "III", "CONTENIDO PROHIBIDO")
    add_article(doc, 7, "Publicaciones prohibidas",
        "Queda terminantemente prohibido publicar:")
    add_bullet(doc, "Fotos, videos o información de pacientes veterinarios o de sus propietarios sin consentimiento escrito.")
    add_bullet(doc, "Fotos o videos del interior de las instalaciones (quirófano, hospitalización, depósito, farmacia) sin autorización.")
    add_bullet(doc, "Información sobre historias clínicas, diagnósticos, tratamientos o pronósticos de animales atendidos.")
    add_bullet(doc, "Datos personales de clientes, compañeros o proveedores (nombre, C.I., teléfono, dirección).")
    add_bullet(doc, "Información sobre salarios, contratos, estrategias comerciales, precios, descuentos o listado de proveedores.")
    add_bullet(doc, "Opiniones o comentarios sobre incidentes, accidentes, demandas o investigaciones internas antes del comunicado oficial de LA EMPRESA.")
    add_bullet(doc, "Contenido que afecte la dignidad, reputación o imagen de LA EMPRESA, sus directivos, trabajadores, clientes o proveedores.")
    add_bullet(doc, "Contenido difamatorio, discriminatorio, racista, sexista, violento, obsceno o que infrinja la ley.")
    add_bullet(doc, "Contenido que infrinja derechos de autor, marcas o propiedad intelectual de terceros.")
    add_bullet(doc, "Publicidad o promoción de clínicas competidoras o de productos no autorizados por LA EMPRESA.")

    # ===== CAPÍTULO IV — USO CORPORATIVO DE REDES SOCIALES =====
    add_chapter(doc, "IV", "USO CORPORATIVO DE REDES SOCIALES")
    add_article(doc, 8, "Cuentas corporativas",
        "Las cuentas oficiales de LA EMPRESA en redes sociales (Instagram, Facebook, "
        "TikTok, WhatsApp Business, página web) son administradas exclusivamente por el "
        "personal autorizado por la Gerencia. El acceso a las credenciales es personal e "
        "intransferible.")
    add_article(doc, 9, "Contenido corporativo",
        "El contenido publicado en cuentas corporativas deberá ser aprobado previamente por "
        "la Gerencia. Se promoverá la divulgación de servicios, promociones, contenido "
        "educativo y campañas de bienestar animal, respetando siempre los derechos de "
        "terceros.")
    add_article(doc, 10, "Hashtags y mención",
        "El personal autorizado a publicar deberá usar los hashtags oficiales de LA EMPRESA "
        "y mencionar las cuentas corporativas. Se prohíbe crear cuentas paralelas no "
        "autorizadas que usen el nombre, marca o logotipo de LA EMPRESA.")

    # ===== CAPÍTULO V — IDENTIFICACIÓN Y DESCARGO DE RESPONSABILIDAD =====
    add_chapter(doc, "V", "IDENTIFICACIÓN Y DESCARGO DE RESPONSABILIDAD")
    add_article(doc, 11, "Identificación del trabajador",
        "Cuando el personal publique contenido relacionado con el ejercicio profesional "
        "(por ejemplo, casos veterinarios, técnicas de peluquería) en cuentas personales, "
        "deberá incluir un descargo explícito de que las opiniones son personales y no "
        "representan la posición de LA EMPRESA. Ejemplo sugerido: 'Las opiniones "
        "expresadas son personales y no representan a mi empleador.'")
    add_article(doc, 12, "Protección de mascotas ajenas",
        "No se publicarán fotos o videos de mascotas de clientes sin consentimiento "
        "expreso del propietario. La autorización se documentará por escrito y se "
        "archivará en el expediente del cliente. Las mascotas del personal podrán "
        "publicarse libremente.")
    add_article(doc, 13, "Protección de menores",
        "No se publicarán fotos o videos de menores de edad sin autorización expresa de "
        "sus representantes legales, conforme al artículo 60 de la Constitución y la LOPDP.")

    # ===== CAPÍTULO VI — GESTIÓN DE INCIDENTES =====
    add_chapter(doc, "VI", "GESTIÓN DE INCIDENTES EN REDES SOCIALES")
    add_article(doc, 14, "Reporte obligatorio",
        "El personal que detecte publicaciones que afecten a LA EMPRESA, a clientes, "
        "mascotas o compañeros, deberá reportarlo inmediatamente a la Gerencia por el "
        "canal más rápido disponible (WhatsApp corporativo o correo electrónico). No se "
        "responderá ni se replicará el contenido sin autorización.")
    add_article(doc, 15, "Comunicado oficial",
        "En caso de incidentes relevantes (accidentes, denuncias, hechos de interés "
        "público), LA EMPRESA emitirá un comunicado oficial. El personal se abstendrá de "
        "publicar opiniones o comentarios antes de dicho comunicado.")
    add_article(doc, 16, "Respuesta a comentarios negativos",
        "La respuesta a comentarios negativos o quejas de clientes en redes sociales "
        "corresponde exclusivamente a la Gerencia. El personal no deberá responder "
        "directamente ni eliminar comentarios sin autorización.")

    # ===== CAPÍTULO VII — SANCIONES =====
    add_chapter(doc, "VII", "SANCIONES")
    add_article(doc, 17, "Régimen sancionatorio",
        "El incumplimiento de la presente Política se sancionará conforme al Reglamento "
        "Interno de Trabajo y al artículo 79 de la LOTTT. La publicación deliberada de "
        "información confidencial, historias clínicas o contenido que afecte la imagen de "
        "LA EMPRESA constituye falta muy grave y será causal de despido justificado.")

    # ===== CAPÍTULO VIII — DISPOSICIONES FINALES =====
    add_chapter(doc, "VIII", "DISPOSICIONES FINALES")
    add_article(doc, 18, "Aceptación",
        "La firma de la carta de recepción al final del presente documento tiene carácter "
        "de aceptación formal y explícita de la Política. El personal declara comprender "
        "las obligaciones aquí establecidas y se compromete a cumplirlas.")
    add_article(doc, 19, "Vigencia y revisión",
        "Esta Política entra en vigencia desde su entrega y se revisará anualmente o "
        "cuando las plataformas o la legislación lo requieran.")

    # ===== Carta de recepción (firma DIRECTORA GERENTE automáticamente) =====
    add_reception_letter(doc, "Política de Uso de Redes Sociales")

    add_footer(doc.sections[0], FOOTER_BASE)
    out_path = os.path.join(OUT_DIR, "Politica_Uso_Redes_Sociales.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# Main
# ============================================================
def main():
    print("=" * 64)
    print("gen_politicas.py — Generando 4 políticas internas ALIKA PETS")
    print("=" * 64)
    paths = []
    for fn in (gen_reglamento_interno, gen_codigo_conducta,
               gen_politica_confidencialidad, gen_politica_redes_sociales):
        try:
            p = fn()
            paths.append(p)
            print(f"  ✓ {os.path.basename(p)}  →  {os.path.getsize(p)//1024} KB")
        except Exception as e:
            print(f"  ✗ Error en {fn.__name__}: {e}")
            raise
    print("=" * 64)
    print(f"Total: {len(paths)} documentos generados en {OUT_DIR}")
    return paths


if __name__ == "__main__":
    main()
