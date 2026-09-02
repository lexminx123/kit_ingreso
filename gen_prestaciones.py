"""
gen_prestaciones.py — Genera 2 documentos del Bloque 04 PRESTACIONES:
  1. Autorizacion_Deposito_Prestaciones.docx  (Art. 143 LOTTT)
  2. Designacion_Beneficiarios.docx           (Art. 137 LOTTT)

Versión 3.0
-----------
- Razón social: GRUPO CAVAL 1003, C.A. (J501662533)
- Salario básico actualizado: Bs. 210/mes (Gaceta 7.026 Extraordinaria 28/04/2025)
- Cestaticket actualizado: Bs. 1.500/mes
- Firma: DIRECTORA GERENTE (Esnatlim Elena Simoza — C.I. V-17.976.287)
  · Autorización Prestaciones → Trabajador + DIRECTORA GERENTE
  · Designación Beneficiarios → Trabajador + Testigo + DIRECTORA GERENTE

El módulo _common.py mapea automáticamente "DIRECTORA GERENTE" a Esnatlim,
y deja "EL(LA) TRABAJADOR(A)" y "TESTIGO" como campos vacíos para llenar.
"""
import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import _common
from _common import (
    setup_a4_portrait, add_membrete, add_doc_title, add_section,
    add_para, add_bullet, add_signature_block, add_footer,
    set_cell_bg, set_cell_borders, set_cell_margins, write_cell, style_run,
    add_hr, TEAL_DARK, TEAL_HDR_BG, GRAY_ALT, SLATE_BG, WHITE,
    BLACK, GRAY_TEXT, GRAY_MUTED, RED_CRIT, GREEN_OK,
    EMPRESA, RIF_EMP, MARCA, DOMICILIO_EMP,
    DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI,
)

# ============================================================
# Datos económicos actualizados (Gaceta 7.026 Ext. 28/04/2025)
# ============================================================
SALARIO_BASICO = "Bs. 210,00"
CESTATICKET    = "Bs. 1.500,00"
BONO_TRANSPORTO = "Bs. 200,00"


# ============================================================
# Helpers locales
# ============================================================
def _white():
    from docx.shared import RGBColor as _RC
    return _RC(0xFF, 0xFF, 0xFF)


def add_worker_identification(doc):
    """Tabla con identificación del trabajador (campos vacíos para llenar)."""
    rows = [
        ("Nombre y Apellido del Trabajador(a):", ""),
        ("Cédula de Identidad N°:", ""),
        ("Cargo a desempeñar:", ""),
        ("Fecha de Ingreso:", ""),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(6.5)
    tbl.columns[1].width = Cm(10.1)
    for i, (k, v) in enumerate(rows):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(6.5); c1.width = Cm(10.1)
        write_cell(c0, k, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(c1, v if v else "_" * 30, size=9, color=BLACK,
                   bg=(GRAY_ALT if i % 2 else WHITE))
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


# ============================================================
# 1. AUTORIZACIÓN DEPÓSITO DE PRESTACIONES (Art. 143 LOTTT)
# ============================================================
def gen_autorizacion_prestaciones(out_dir):
    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)

    add_membrete(doc,
                 doc_label="AUTORIZACIÓN DEPÓSITO PRESTACIONES",
                 doc_sublabel="Artículo 143 LOTTT",
                 version="Versión 3.0  ·  RR.HH.")

    add_doc_title(doc, "AUTORIZACIÓN PARA EL DEPÓSITO DE PRESTACIONES SOCIALES")

    # Intro
    add_para(doc,
        f"Yo, _____________________________________________, titular de la cédula de "
        f"identidad N° V-___________________, en mi condición de trabajador(a) de "
        f"{EMPRESA} ({MARCA}), RIF N° {RIF_EMP}, con domicilio en {DOMICILIO_EMP}, "
        f"actuando en pleno ejercicio de mis derechos y sin coacción alguna, declaro "
        f"conocer el contenido del artículo 143 de la Ley Orgánica del Trabajo, los "
        f"Trabajadores y las Trabajadoras (LOTTT) y, en consecuencia, autorizo a la "
        f"empresa para depositar mensualmente mis prestaciones sociales en la modalidad "
        f"que a continuación especifico:",
        size=10, space_after=8)

    # Identificación trabajador
    add_section(doc, "IDENTIFICACIÓN DEL TRABAJADOR(A)")
    add_worker_identification(doc)

    # Conceptos salariales actualizados
    add_section(doc, "CONCEPTOS SALARIALES ACTUALIZADOS (Gaceta 7.026 Ext. 28/04/2025)")
    headers = ["Concepto", "Monto mensual", "Naturaleza"]
    rows = [
        ("Salario básico mensual",                 SALARIO_BASICO,   "Salarial (base prestaciones)"),
        ("Bono de alimentación (Cestaticket)",     CESTATICKET,      "No salarial (Ley Programa Alimentación)"),
        ("Bono de transporte",                     BONO_TRANSPORTO,  "No salarial"),
    ]
    tbl = doc.add_table(rows=1+len(rows), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [6.0, 4.6, 6.0]
    for i, w in enumerate(widths):
        tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.width = Cm(widths[ci])
        write_cell(c, h, size=9, bold=True, color=_white(),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            c.width = Cm(widths[ci])
            bold = (ci == 0)
            write_cell(c, val, size=9, bold=bold, color=BLACK, bg=bg,
                       align=WD_ALIGN_PARAGRAPH.LEFT if ci != 1 else WD_ALIGN_PARAGRAPH.CENTER)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    # Modalidades de depósito (Art. 143)
    add_section(doc, "MODALIDAD DE DEPÓSITO (Artículo 143 LOTTT)")
    add_para(doc,
        "Conforme al artículo 143 de la LOTTT, el trabajador podrá elegir la modalidad "
        "de depósito de sus prestaciones sociales. Marque con una X la opción elegida:",
        size=10, space_after=4)
    opt_rows = [
        ("☐", "OPCIÓN A — Fondo de Ahorro Habitacional (FAOV-BVV)",
         "Las prestaciones sociales se depositarán mensualmente en la cuenta individual de ahorro habitacional del trabajador ante el Banco de la Vivienda y Hábitat (BVV)."),
        ("☐", "OPCIÓN B — Fondo de Prestaciones de la Empresa",
         "Las prestaciones sociales se mantendrán en un fondo contable manejado directamente por la empresa, a nombre del trabajador, generando la tasa de interés pasiva fijada por el BCV."),
    ]
    tbl = doc.add_table(rows=1+len(opt_rows), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [1.0, 6.6, 9.0]
    for i, w in enumerate(widths):
        tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(["☐", "Modalidad", "Descripción"]):
        c = tbl.rows[0].cells[ci]
        c.width = Cm(widths[ci])
        write_cell(c, h, size=9, bold=True, color=_white(),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(opt_rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            c.width = Cm(widths[ci])
            bold = (ci == 1)
            color = TEAL_DARK if ci == 1 else BLACK
            write_cell(c, val, size=9, bold=bold, color=color, bg=bg,
                       align=WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    # Datos de la cuenta FAOV (si aplica)
    add_section(doc, "DATOS DE LA CUENTA FAOV-BVV (completar solo si eligió Opción A)")
    faov_rows = [
        ("Banco receptor:", ""),
        ("N° de cuenta de ahorro habitacional:", ""),
        ("Tipo de cuenta:", "☐  Ahorro Habitacional  ☐  Corriente"),
    ]
    tbl = doc.add_table(rows=len(faov_rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(7.5)
    tbl.columns[1].width = Cm(9.1)
    for i, (k, v) in enumerate(faov_rows):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(7.5); c1.width = Cm(9.1)
        write_cell(c0, k, size=9, bold=True, color=TEAL_DARK, bg=SLATE_BG,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(c1, v if v else "_" * 25, size=9, color=BLACK,
                   bg=(GRAY_ALT if i % 2 else WHITE))
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    # Declaraciones
    add_section(doc, "DECLARACIONES DEL TRABAJADOR(A)")
    decl = [
        "Declaro conocer que las prestaciones sociales me pertenecen desde el momento en que se causan y que el depósito mensual no me priva del derecho a reclamarlas al término de la relación laboral, conforme al artículo 142 de la LOTTT.",
        "Declaro conocer que el monto depositado mensualmente se calculará conforme al artículo 143 de la LOTTT (quince (15) días de salario por cada mes de servicio o alícuota correspondiente).",
        "Autorizo a LA EMPRESA a efectuar los retiros, depósitos y transferencias bancarias necesarias para cumplir con la modalidad de depósito aquí indicada.",
        "Me comprometo a notificar por escrito a LA EMPRESA cualquier modificación de la cuenta bancaria o de la modalidad de depósito elegida, con al menos quince (15) días de anticipación.",
        "Declaro conocer que puedo revocar la presente autorización en cualquier momento, mediante comunicación escrita dirigida a la Dirección de Gestión Humana, surtiendo efectos a partir del mes siguiente a su recepción.",
    ]
    for d in decl:
        add_bullet(doc, d, size=10)

    add_para(doc, "", size=4, space_after=4)

    # Cierre
    add_para(doc,
        "En fe de lo cual firmo la presente autorización en la ciudad de Los Teques, "
        "Estado Miranda, a los ____ días del mes de ________________ de ________.",
        size=10, space_before=6, space_after=10)

    # Firmas: Trabajador + Directora Gerente
    add_section(doc, "FIRMAS")
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])

    # Footer
    add_footer(section, "Autorización Depósito Prestaciones v3.0")

    out_path = os.path.join(out_dir, "Autorizacion_Deposito_Prestaciones.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# 2. DESIGNACIÓN DE BENEFICIARIOS (Art. 137 LOTTT)
# ============================================================
def gen_designacion_beneficiarios(out_dir):
    doc = Document()
    section = setup_a4_portrait(doc, margins_cm=2.0)

    add_membrete(doc,
                 doc_label="DESIGNACIÓN BENEFICIARIOS",
                 doc_sublabel="Artículo 137 LOTTT",
                 version="Versión 3.0  ·  RR.HH.")

    add_doc_title(doc, "DESIGNACIÓN DE BENEFICIARIOS PARA COBRO DE PRESTACIONES SOCIALES")

    # Intro legal
    add_para(doc,
        f"Conforme al artículo 137 de la Ley Orgánica del Trabajo, los Trabajadores y las "
        f"Trabajadoras (LOTTT), yo, _____________________________________________, "
        f"titular de la cédula de identidad N° V-___________________, en mi condición de "
        f"trabajador(a) activo(a) de {EMPRESA} ({MARCA}), RIF N° {RIF_EMP}, con domicilio "
        f"en {DOMICILIO_EMP}, designo a las personas que abajo se indican como mis "
        f"beneficiarios para que, en caso de mi fallecimiento, perciban las prestaciones "
        f"sociales y demás beneficios laborales que me correspondan. Esta designación se "
        f"sujeta al orden de prelación y porcentajes legales previstos en el referido "
        f"artículo 137 de la LOTTT.",
        size=10, space_after=8)

    # Identificación trabajador
    add_section(doc, "IDENTIFICACIÓN DEL TRABAJADOR(A)")
    add_worker_identification(doc)

    # Orden de prelación legal (Art. 137)
    add_section(doc, "ORDEN DE PRELACIÓN LEGAL (Artículo 137 LOTTT)")
    add_para(doc,
        "El artículo 137 de la LOTTT establece el siguiente orden de prelación para el "
        "cobro de las prestaciones sociales en caso de fallecimiento del trabajador:",
        size=10, space_after=4)
    pre_rows = [
        ("1°", "El cónyuge sobreviviente y los hijos menores de dieciocho (18) años, o con discapacidad, sin otras limitaciones."),
        ("2°", "Los hijos mayores de dieciocho (18) años que no tengan derecho a prestación propia."),
        ("3°", "El padre y la madre del trabajador fallecido."),
        ("4°", "Los hermanos menores de dieciocho (18) años, o con discapacidad, del trabajador fallecido."),
        ("5°", "Los hermanos mayores de dieciocho (18) años del trabajador fallecido."),
    ]
    tbl = doc.add_table(rows=1+len(pre_rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(1.6)
    tbl.columns[1].width = Cm(15.0)
    for ci, h in enumerate(["#", "Orden de prelación"]):
        c = tbl.rows[0].cells[ci]
        c.width = Cm(1.6 if ci == 0 else 15.0)
        write_cell(c, h, size=9, bold=True, color=_white(),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(pre_rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            c.width = Cm(1.6 if ci == 0 else 15.0)
            write_cell(c, val, size=9, bold=(ci == 0), color=BLACK, bg=bg,
                       align=WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    # Tabla de beneficiarios designados (5)
    add_section(doc, "DESIGNACIÓN DE BENEFICIARIOS (máximo cinco (5))")
    add_para(doc,
        "Indique a continuación, hasta un máximo de cinco (5) beneficiarios, con sus datos "
        "completos, parentesco con el trabajador y porcentaje de distribución. La suma de "
        "los porcentajes debe ser igual al 100 %.",
        size=10, space_after=4)
    ben_headers = ["N°", "Nombre y Apellido del Beneficiario", "C.I. N°",
                   "Parentesco", "Fecha Nac.", "%"]
    ben_rows = [
        ("1", "", "", "", "", ""),
        ("2", "", "", "", "", ""),
        ("3", "", "", "", "", ""),
        ("4", "", "", "", "", ""),
        ("5", "", "", "", "", ""),
    ]
    tbl = doc.add_table(rows=1+len(ben_rows), cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [0.8, 5.8, 3.0, 2.8, 2.2, 2.0]
    for i, w in enumerate(widths):
        tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(ben_headers):
        c = tbl.rows[0].cells[ci]
        c.width = Cm(widths[ci])
        write_cell(c, h, size=9, bold=True, color=_white(),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(ben_rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            c.width = Cm(widths[ci])
            if ci == 0:
                write_cell(c, val, size=9, bold=True, color=TEAL_DARK, bg=bg,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                write_cell(c, val, size=9, color=BLACK, bg=bg,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

    add_para(doc,
        "Distribución total: _________ %   (debe sumar 100 %)",
        size=10, bold=True, color=TEAL_DARK, space_after=6)

    # Declaraciones
    add_section(doc, "DECLARACIONES DEL TRABAJADOR(A)")
    decl = [
        "Declaro que los beneficiarios arriba designados son las personas que deseo reconozcan como herederos laborales para el cobro de mis prestaciones sociales y demás beneficios laborales, en caso de mi fallecimiento.",
        "Declaro conocer que la presente designación puede ser modificada por mí en cualquier momento, mediante comunicación escrita presentada ante la Dirección de Gestión Humana de LA EMPRESA.",
        "Declaro conocer que, en caso de controversia entre los beneficiarios designados sobre el derecho a ser considerados como tales, LA EMPRESA depositará las prestaciones sociales ante un Juez de Primera Instancia del Trabajo, dentro del octavo (8°) día hábil siguiente a la terminación de la relación laboral, conforme al artículo 137 de la LOTTT.",
        "Declaro conocer que la presente designación no afecta los derechos que por ley correspondan a los herederos legítimos conforme al Código Civil venezolano.",
        "Acepto que la presente designación es válida solo para los efectos del artículo 137 de la LOTTT y no constituye testamento ni disposición de bienes distintos a los laborales.",
    ]
    for d in decl:
        add_bullet(doc, d, size=10)

    add_para(doc, "", size=4, space_after=4)

    # Cierre
    add_para(doc,
        "En fe de lo cual firmo la presente designación en la ciudad de Los Teques, "
        "Estado Miranda, a los ____ días del mes de ________________ de ________.",
        size=10, space_before=6, space_after=10)

    # Firmas: Trabajador + Testigo + Directora Gerente
    add_section(doc, "FIRMAS")
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "TESTIGO", "DIRECTORA GERENTE"])

    # Footer
    add_footer(section, "Designación Beneficiarios v3.0")

    out_path = os.path.join(out_dir, "Designacion_Beneficiarios.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# Main
# ============================================================
def main():
    out_dir = os.path.join(HERE, "04_PRESTACIONES")
    os.makedirs(out_dir, exist_ok=True)
    print(f"[gen_prestaciones] Generando 2 documentos en: {out_dir}\n")

    try:
        p1 = gen_autorizacion_prestaciones(out_dir)
        print(f"  OK  Autorizacion_Deposito_Prestaciones.docx  {os.path.getsize(p1)/1024:5.1f} KB")
    except Exception as e:
        print(f"  ERROR Autorización → {e}")
        raise

    try:
        p2 = gen_designacion_beneficiarios(out_dir)
        print(f"  OK  Designacion_Beneficiarios.docx            {os.path.getsize(p2)/1024:5.1f} KB")
    except Exception as e:
        print(f"  ERROR Designación → {e}")
        raise

    print("\n[gen_prestaciones] Finalizado.")


if __name__ == "__main__":
    main()
