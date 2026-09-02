"""
gen_protocolos_vet.py — Genera 4 protocolos veterinarios ALIKA PETS / Grupo Caval 1003, C.A.

Documentos generados:
  1. /05_SEGURIDAD_LABORAL/Cartilla_Bioseguridad_Veterinaria.docx       (9 pp, NT-01-2008, 12 cap.)
  2. /08_POLITICAS_INTERNAS/Protocolo_Sustancias_Controladas.docx       (6 pp, SENAC)
  3. /08_POLITICAS_INTERNAS/Protocolo_Mordeduras_Zoonosis.docx          (6 pp, OMS Tipos 1-4, PEP Essen)
  4. /08_POLITICAS_INTERNAS/Procedimiento_Reporte_Incidentes.docx       (7 pp, Art. 73 LOPCYMAT)

Versión 3.0 — _common.py mapea automáticamente "RR.HH." → "DIRECTORA GERENTE" (Esnatlim Simoza).
Márgenes 1.8 cm para protocolos vet (más contenido técnico).
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _common import (
    setup_a4_portrait, add_membrete, add_doc_title, add_chapter, add_section,
    add_article, add_para, add_bullet, add_sanciones_table, add_signature_block,
    add_reception_letter, add_footer, add_hr,
    set_cell_bg, set_cell_borders, set_cell_margins, write_cell,
    TEAL_DARK, TEAL_HDR_BG, GRAY_ALT, AMBER_BG, SLATE_BG, WHITE, BLACK,
    GRAY_TEXT, GRAY_MUTED, RED_CRIT, GREEN_OK,
    EMPRESA, RIF_EMP, MARCA, DOMICILIO_EMP,
    REP_LEGAL_NOMBRE, REP_LEGAL_CARGO, REP_LEGAL_CI,
    DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI,
)

OUT_DIR_05 = os.path.join(os.path.dirname(os.path.abspath(__file__)), "05_SEGURIDAD_LABORAL")
OUT_DIR_08 = os.path.join(os.path.dirname(os.path.abspath(__file__)), "08_POLITICAS_INTERNAS")
os.makedirs(OUT_DIR_05, exist_ok=True)
os.makedirs(OUT_DIR_08, exist_ok=True)

FOOTER_BASE = "Protocolo Veterinario  ·  v3.0"

# ============================================================
# Helpers locales (tablas con encabezado teal y filas alternas)
# ============================================================
def add_data_table(doc, headers, rows, col_widths_cm=None, header_size=9, body_size=9,
                   center_cols=None):
    center_cols = center_cols or []
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        if col_widths_cm:
            c.width = Cm(col_widths_cm[ci])
        write_cell(c, h, size=header_size, bold=True,
                   color=RGBColor(0xFF, 0xFF, 0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            if col_widths_cm:
                c.width = Cm(col_widths_cm[ci])
            align = WD_ALIGN_PARAGRAPH.CENTER if ci in center_cols else (
                WD_ALIGN_PARAGRAPH.LEFT if ci > 0 else WD_ALIGN_PARAGRAPH.CENTER)
            color_txt = BLACK
            bold = False
            if isinstance(val, str):
                upper = val.upper()
                if "CRÍTICO" in upper or "OBLIGATORIO" in upper:
                    color_txt = RED_CRIT
                    bold = True
            write_cell(c, str(val), size=body_size, bold=bold, color=color_txt, bg=bg,
                       align=align)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl

def add_alert_box(doc, lines, bg=AMBER_BG, border="F59E0B"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(17.4)
    c = tbl.rows[0].cells[0]
    c.width = Cm(17.4)
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    for i, ln in enumerate(lines):
        if i > 0:
            br = p.add_run()
            br.add_break()
        if ln.startswith("**"):
            r = p.add_run(ln.strip("*"))
            _style_loc(r, size=9, bold=True, color=RED_CRIT)
        elif ln.startswith(">>"):
            r = p.add_run(ln.lstrip(">"))
            _style_loc(r, size=9, bold=True, color=GREEN_OK)
        else:
            r = p.add_run(ln)
            _style_loc(r, size=9, color=BLACK)
    set_cell_bg(c, bg)
    set_cell_borders(c, color=border, sz="8")
    set_cell_margins(c, top=120, bottom=120, left=160, right=160)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

def _style_loc(run, size=10, bold=False, color=None, font="Calibri", italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_page_break(doc):
    doc.add_page_break()


# ============================================================
# DOCUMENTO 1: CARTILLA DE BIOSEGURIDAD VETERINARIA
# ============================================================
def gen_cartilla_bioseguridad():
    doc = Document()
    setup_a4_portrait(doc, margins_cm=1.8)
    add_membrete(doc, "CARTILLA DE BIOSEGURIDAD", "Norma NT-01-2008", "v3.0  ·  SST")
    add_doc_title(doc, "CARTILLA DE BIOSEGURIDAD VETERINARIA")

    add_para(doc,
        f"La presente Cartilla establece los lineamientos de bioseguridad aplicables a "
        f"todo el personal de {EMPRESA} (marca {MARCA}), conforme a la Norma Técnica "
        f"NT-01-2008 del Instituto Nacional de Prevención, Salud y Seguridad Laborales "
        f"(INPSASEL), la LOPCYMAT, la Ley de Ejercicio de la Medicina Veterinaria y los "
        f"lineamientos de la Organización Mundial de Sanidad Animal (OMS/OIE). Su "
        f"cumplimiento es OBLIGATORIO para proteger la salud del personal, de los "
        f"pacientes y del público usuario.",
        size=10, space_after=6)

    # ===== CAPÍTULO I — PRINCIPIOS DE BIOSEGURIDAD =====
    add_chapter(doc, "I", "PRINCIPIOS DE BIOSEGURIDAD")
    add_article(doc, 1, "Universalidad",
        "Todo paciente y todo fluido biológico deben considerarse potencialmente "
        "infecciosos. Las medidas de bioseguridad se aplican a todos los procedimientos "
        "con animales, independientemente de su estado de salud aparente.")
    add_article(doc, 2, "Uso de barreras",
        "El uso de elementos de protección personal (EPP) es obligatorio. La selección del "
        "EPP depende del procedimiento y del riesgo evaluado. El EPP NO sustituye las "
        "medidas de control de ingeniería ni las prácticas seguras.")
    add_article(doc, 3, "Medios de eliminación",
        "Todo material contaminado debe eliminarse conforme a la Norma COVENIN 2747-93 "
        "sobre Manejo de Residuos en Establecimientos de Salud. Se prohíbe la "
        "manipulación manual de residuos cortopunzantes.")
    add_article(doc, 4, "Reducción de riesgos",
        "Toda exposición a agentes biológicos, físicos o químicos debe minimizarse mediante "
        "controles de ingeniería, prácticas seguras y EPP, en este orden jerárquico.")

    # ===== CAPÍTULO II — VÍAS DE TRANSMISIÓN =====
    add_chapter(doc, "II", "VÍAS DE TRANSMISIÓN DE ENFERMEDADES ZOONÓTICAS")
    add_article(doc, 5, "Clasificación",
        "Las principales vías de transmisión de zoonosis en el medio veterinario son:")
    add_bullet(doc, "Contacto directo con el animal, sus fluidos o secreciones (mordedura, arañazo, lamido de mucosas).", bold_lead="Directa:  ")
    add_bullet(doc, "Contacto con fómites (instrumental, camas, jaulas, superficies) contaminados.", bold_lead="Indirecta:  ")
    add_bullet(doc, "Inhalación de aerosoles generados por tos, estornudos o procedimientos (ecografía, limpieza).", bold_lead="Aérea:  ")
    add_bullet(doc, "Ingesta accidental (manos contaminadas llevadas a la boca) o por vectores (garrapatas, mosquitos, flebótomos).", bold_lead="Oral/Vectorial:  ")
    add_bullet(doc, "Percutánea por pinchazos con agujas, bisturíes, garras, dientes o mordeduras.", bold_lead="Percutánea:  ")

    add_article(doc, "6", "Factores que aumentan el riesgo",
        "Se consideran factores agravantes del riesgo de transmisión:")
    add_bullet(doc, "Trabajador(a) sin esquema completo de vacunación (antirrábica, hepatitis B, tétanos).")
    add_bullet(doc, "Inmunosupresión: embarazo, lactancia, VIH, diabetes, tratamiento inmunosupresor, trasplantes.")
    add_bullet(doc, "Heridas abiertas o dermatitis en manos y antebrazos sin protección.")
    add_bullet(doc, "Ausencia o uso incorrecto de EPP.")
    add_bullet(doc, "Procedimientos sin sedación en animales agresivos o estresados.")
    add_bullet(doc, "Manipulación de animales sin historia clínica o sin esquema vacunal conocido.")
    add_bullet(doc, "Mala higiene de manos entre pacientes y entre procedimientos.")
    add_bullet(doc, "Acumulación de residuos biomédicos sin descarte oportuno.")

    add_article(doc, "7", "Medidas generales de bloqueo de transmisión",
        "Para interrumpir la cadena de transmisión se aplican las siguientes medidas:")
    add_bullet(doc, "Lavado de manos en los 5 momentos OMS (ver Capítulo V).")
    add_bullet(doc, "Uso de EPP específico por procedimiento (ver Capítulo IV).")
    add_bullet(doc, "Desinfección y esterilización de instrumental entre pacientes (ver Capítulo VIII).")
    add_bullet(doc, "Aislamiento de animales con sospecha de zoonosis en área designada.")
    add_bullet(doc, "Señalización visual de áreas con riesgo biológico.")
    add_bullet(doc, "Manejo correcto de residuos según COVENIN 2747-93 (ver Capítulo VII).")
    add_bullet(doc, "Vacunación del personal y vigilancia médica periódica (NT-02-2008).")
    add_bullet(doc, "Reporte inmediato de cualquier accidente o exposición (Art. 73 LOPCYMAT).")

    # ===== CAPÍTULO II.5 — NIVELES DE BIOSEGURIDAD =====
    add_chapter(doc, "II.b", "NIVELES DE BIOSEGURIDAD APLICABLES")
    add_article(doc, "7b", "Clasificación de niveles",
        "Adaptando los niveles de bioseguridad del CDC y OMS al medio veterinario, se "
        "establecen los siguientes cuatro niveles operativos en LA EMPRESA:")
    nb_headers = ["Nivel", "Tipo de agente", "Áreas", "Prácticas mínimas"]
    nb_rows = [
        ["NB-1", "Agentes no patógenos / baja riesgo",
         "Recepción, tienda, áreas comunes",
         "Higiene de manos, EPP básico, desinfección de superficies"],
        ["NB-2", "Agentes de riesgo moderado (zoonosis comunes)",
         "Consulta, hospitalización, peluquería",
         "Bata, guantes, señalética, contenedor cortopunzantes, EPP específico"],
        ["NB-3", "Agentes de alto riesgo (rabia, brucelosis, leptospirosis)",
         "Quirófano, sala de necropsias, aislamiento",
         "Doble guante, N95, gafas, bata impermeable, control de aerosoles"],
        ["NB-4", "Agentes letales / emergentes",
         "NO disponible en LA EMPRESA — derivar a centro especializado",
         "Prohibido manipular; reportar al MPPS y canalizar a RAV"],
    ]
    add_data_table(doc, nb_headers, nb_rows,
                   col_widths_cm=[1.2, 4.8, 4.4, 7.0], header_size=8, body_size=8, center_cols=[0])

    # ===== CAPÍTULO III — ZOONOSIS EN VENEZUELA =====
    add_chapter(doc, "III", "ZOONOSIS RELEVANTES EN VENEZUELA")
    add_article(doc, 6, "Enfermedades endémicas",
        "El personal debe conocer las principales zoonosis presentes en Venezuela y sus "
        "medidas de prevención:")
    zoonosis_headers = ["Zoonosis", "Agente", "Reservorio", "Vía", "Prevención"]
    zoonosis_rows = [
        ["Rabia", "Virus (Lisavirus)", "Perros, gatos, quirópticos", "Mordedura, saliva", "Vacunación pre-exposición 3 dosis + PEP"],
        ["Leptospirosis", "Bacteria (Leptospira)", "Roedores, perros", "Contacto con orina", "EPP, control de roedores, doxiciclina profiláctica"],
        ["Toxoplasmosis", "Protozoo (T. gondii)", "Felinos", "Heces, ingestión", "Higiene manos, guantes al limpiar arenero"],
        ["Brucelosis", "Bacteria (Brucella)", "Bovinos, caprinos, caninos", "Contacto fluidos", "EPP en abortos, vacunación animales"],
        ["Tiña (Dermatofitosis)", "Hongos (Microsporum)", "Perros, gatos", "Contacto piel", "Guantes, lavado manos, desinfección"],
        ["Sarna sarcóptica", "Ácaro (Sarcoptes)", "Perros, zorros", "Contacto piel", "Guantes, bata, lavado ropa"],
        ["Campilobacteriosis", "Bacteria (C. jejuni)", "Aves, perros, gatos", "Vía fecal-oral", "Higiene, EPP, no comer en áreas clínicas"],
        ["Salmonelosis", "Bacteria (Salmonella)", "Reptiles, aves, perros", "Vía fecal-oral", "Higiene, EPP, desinfección"],
        ["Toxocariasis", "Nemátodo (T. canis)", "Perros, gatos", "Ingesta de huevos", "Higiene, desparasitación, EPP"],
    ]
    add_data_table(doc, zoonosis_headers, zoonosis_rows,
                   col_widths_cm=[3.0, 3.2, 3.5, 2.8, 4.8], header_size=8, body_size=8)

    add_article(doc, "7c", "Fichas clínicas de zoonosis prioritarias",
        "Las siguientes fichas resumen signos en humanos y acciones inmediatas:")
    add_section(doc, "Rabia")
    add_bullet(doc, "Incubación: 1-3 meses (puede llegar a 1 año).", bold_lead="Periodo:  ")
    add_bullet(doc, "Parestesia en sitio de mordedura, fiebre, hidrofobia, aerofobia, delirio, parálisis, muerte.", bold_lead="Signos humanos:  ")
    add_bullet(doc, "PEP inmediato según Tipo 1-4 OMS. Lavado 15 min + vacuna + IGAR. Vigilar animal mordedor 10 días.", bold_lead="Acción:  ")
    add_section(doc, "Leptospirosis")
    add_bullet(doc, "Incubación: 5-14 días.", bold_lead="Periodo:  ")
    add_bullet(doc, "Fiebre bifásica, cefalea, mialgia (pantorrillas), ictericia, insuficiencia renal (síndrome de Weil).", bold_lead="Signos humanos:  ")
    add_bullet(doc, "Serología IgM, doxiciclina 100 mg/12h × 7 días. Hospitalizar si ictericia o falla renal.", bold_lead="Acción:  ")
    add_section(doc, "Toxoplasmosis")
    add_bullet(doc, "Incubación: 5-18 días.", bold_lead="Periodo:  ")
    add_bullet(doc, "Adenopatías cervicales, fiebre, fatiga, mialgia. Peligroso en embarazadas (transmisión transplacentaria) e inmunodeprimidos.", bold_lead="Signos humanos:  ")
    add_bullet(doc, "Serología IgG/IgM, espiramicina si embarazo, sulfadiazina + pirimetamina si inmunodeprimido.", bold_lead="Acción:  ")
    add_section(doc, "Brucelosis")
    add_bullet(doc, "Incubación: 1-3 semanas.", bold_lead="Periodo:  ")
    add_bullet(doc, "Fiebre ondulante, sudores profusos, artralgias, astenia, hepatosplenomegalia.", bold_lead="Signos humanos:  ")
    add_bullet(doc, "Serología (Rose Bengal, Wright), rifampicina 900 mg/día + doxiciclina 100 mg/12h × 6 semanas.", bold_lead="Acción:  ")
    add_section(doc, "Tiña / Dermatofitosis")
    add_bullet(doc, "Incubación: 4-10 días.", bold_lead="Periodo:  ")
    add_bullet(doc, "Lesiones anulares eritematoescamosas pruriginosas en piel expuesta (manos, antebrazos, cuello).", bold_lead="Signos humanos:  ")
    add_bullet(doc, "Examen micológico. Terapéutica tópica (clotrimazol) o sistémica (terbinafina, itraconazol).", bold_lead="Acción:  ")
    add_section(doc, "Sarna sarcóptica")
    add_bullet(doc, "Incubación: 1-2 semanas (primer contagio); 1-3 días si reinfección.", bold_lead="Periodo:  ")
    add_bullet(doc, "Prurito intenso nocturno, surcos acarinos, pápulas en espacios interdigitales, muñecas, cintura.", bold_lead="Signos humanos:  ")
    add_bullet(doc, "Permetrina 5% tópica, ivermectina oral. Lavar ropa de cama a 60°C. Tratar animales.", bold_lead="Acción:  ")

    # ===== CAPÍTULO IV — EPP POR PROCEDIMIENTO =====
    add_chapter(doc, "IV", "EQUIPOS DE PROTECCIÓN PERSONAL POR PROCEDIMIENTO")
    add_article(doc, 7, "Selección del EPP",
        "Cada procedimiento requiere un EPP específico. La siguiente tabla establece el EPP "
        "obligatorio según el tipo de procedimiento:")
    epp_headers = ["Procedimiento", "EPP obligatorio"]
    epp_rows = [
        ["Consulta clínica general", "Bata, guantes de examen, mascarilla quirúrgica (si tos/estornudo animal)"],
        ["Vacunación y aplicación de inyectables", "Bata, guantes de examen, contenedor cortopunzantes"],
        ["Toma de muestras (sangre, heces, orina)", "Bata, guantes de examen, mascarilla quirúrgica, gafas"],
        ["Cirugía mayor (laparotomía, OVH)", "Bata quirúrgica estéril, guantes estériles, mascarilla, gafas, gorro, campo estéril"],
        ["Cirugía menor (sutura herida, drenaje)", "Bata, guantes estériles, mascarilla quirúrgica, gafas"],
        ["Odontología veterinaria", "Bata, guantes, mascarilla con visera, gafas, gorro"],
        ["Eutanasia y necropsia", "Bata impermeable, doble guante, mascarilla N95, gafas selladas, gorro, botas"],
        ["Peluquería canina", "Delantal impermeable, guantes de nitrilo, mascarilla quirúrgica, gafas (si salpicaduras)"],
        ["Limpieza de jaulas y áreas", "Guantes de goma largos, mascarilla, botas, delantal impermeable"],
        ["Manejo de animal sospechoso de rabia", "Bata impermeable, doble guante, mascarilla N95, gafas, gorro, botas — NUNCA solo"],
    ]
    add_data_table(doc, epp_headers, epp_rows,
                   col_widths_cm=[5.5, 11.8], header_size=9, body_size=9)

    # ===== CAPÍTULO V — 5 MOMENTOS OMS PARA HIGIENE DE MANOS =====
    add_chapter(doc, "V", "CINCO MOMENTOS OMS PARA LA HIGIENE DE MANOS")
    add_article(doc, 8, "Momentos OMS",
        "Adaptado al medio veterinario, la higiene de manos debe realizarse en los "
        "siguientes cinco momentos:")
    add_bullet(doc, "Antes de tocar al paciente animal.", bold_lead="1.  ")
    add_bullet(doc, "Antes de procedimientos limpios/asépticos (inyección, curación, cirugía).", bold_lead="2.  ")
    add_bullet(doc, "Después de riesgo de exposición a fluidos corporales.", bold_lead="3.  ")
    add_bullet(doc, "Después de tocar al paciente animal.", bold_lead="4.  ")
    add_bullet(doc, "Después del contacto con el entorno del paciente (jaula, camilla, instrumental).", bold_lead="5.  ")
    add_para(doc,
        "La higiene se realiza con agua y jabón (40-60 seg) o con solución alcoholada (20-30 "
        "seg). La solución alcoholada NO sustituye el lavado cuando hay suciedad visible, "
        "exposición a esporas (Clostridium) o contacto con fluidos potencialmente "
        "infecciosos.",
        size=9, space_before=2, italic=True)

    # ===== CAPÍTULO VI — CORTOPUNZANTES =====
    add_chapter(doc, "VI", "MANEJO DE CORTOPUNZANTES")
    add_article(doc, 9, "Descarte inmediato",
        "Toda aguja, bisturí, hoja, ampolla rota o material cortopunzante debe descartarse "
        "INMEDIATAMENTE después de su uso en el contenedor rígido rojo con símbolo de "
        "bioriesgo. Se prohíbe: reencapuchar agujas, doblarlas, romperlas, retirarlas con "
        "la mano o desecharlas en bolsas comunes.")
    add_article(doc, 10, "Accidente cortopunzante",
        "En caso de accidente cortopunzante: lavar inmediatamente con agua y jabón, "
        "sangrar la herida, aplicar antiséptico (alcohol 70% o povidona), cubrir con "
        "apósito estéril y reportar al supervisor en un máximo de DOS (2) horas conforme "
        "al artículo 73 de la LOPCYMAT y al Protocolo de Mordeduras/Zoonosis.")

    # ===== CAPÍTULO VII — RESIDUOS BIOMÉDICOS =====
    add_chapter(doc, "VII", "MANEJO DE RESIDUOS — NORMA COVENIN 2747-93")
    add_article(doc, 11, "Clasificación por color",
        "Los residuos se separan por tipo en bolsas del color correspondiente:")
    residuos_headers = ["Color", "Tipo de residuo", "Ejemplos"]
    residuos_rows = [
        ["ROJO", "Biológicos / infecciosos", "Gasas con sangre, tejidos, órganos, cultivos, vacunas vivas"],
        ["AMARILLO", "Cortopunzantes", "Agujas, bisturíes, ampollas, lancetas (contenedor rígido)"],
        ["NEGRO", "Comunes / ordinarios", "Papel, plástico no contaminado, restos de alimentos"],
        ["VERDE", "Reciclables", "Cartón limpio, plástico de envases de medicamentos vacíos"],
        ["GRIS / ESPECIAL", "Químicos y farmacéuticos", "Medicamentos vencidos, reactivos, anestésicos residuales"],
    ]
    add_data_table(doc, residuos_headers, residuos_rows,
                   col_widths_cm=[2.8, 4.5, 10.0], header_size=9, body_size=9, center_cols=[0])
    add_para(doc,
        "Las bolsas se llenan hasta 80% de su capacidad, se cierran y rotulan con fecha y "
        "responsable. El depósito temporal interno no excederá 72 horas. La disposición "
        "final se realiza mediante empresa certificada por el MINEC.",
        size=9, space_before=2, italic=True)

    # ===== CAPÍTULO VIII — DESINFECCIÓN Y ESTERILIZACIÓN =====
    add_chapter(doc, "VIII", "DESINFECCIÓN Y ESTERILIZACIÓN")
    add_article(doc, 12, "Niveles de desinfección",
        "Los desinfectantes se utilizan según el nivel requerido:")
    add_bullet(doc, "Glutaraldehído al 2% (20 min), peróxido de hidrógeno 6% (30 min), ácido peracético.", bold_lead="Alto nivel (críticos - instrumental quirúrgico):  ")
    add_bullet(doc, "Compuestos de amonio cuaternario, alcohol 70%, hipoclorito 1000 ppm.", bold_lead="Nivel medio (semicríticos - endoscopios, termómetros):  ")
    add_bullet(doc, "Detergente común, hipoclorito 200 ppm.", bold_lead="Bajo nivel (no críticos - superficies, pisos):  ")
    add_article(doc, 13, "Esterilización",
        "El instrumental quirúrgico se esteriliza preferentemente en autoclave (121°C, 15 "
        "psi, 20-30 min). Como alternativa se usa óxido de etileno o esterilización por "
        "radiación UV (limitada a superficies). Cada ciclo se controla con indicadores "
        "químicos y biológicos.")
    add_article(doc, "13b", "Diccionario de desinfectantes",
        "La siguiente tabla resume los desinfectantes de uso habitual en LA EMPRESA:")
    desinf_headers = ["Desinfectante", "Concentración", "Tiempo", "Uso / Notas"]
    desinf_rows = [
        ["Alcohol etílico / isopropílico", "70%", "1-3 min", "Piel, instrumental pequeño. Volátil, inflamable."],
        ["Hipoclorito de sodio", "200-1000 ppm", "10-30 min", "Superficies, pisos, derrames de sangre. Inactivado por materia orgánica."],
        ["Glutaraldehído", "2%", "20-45 min", "Alto nivel. Irritante vías respiratorias. Usar en cámara cerrada."],
        ["Amonio cuaternario", "0.5-2%", "5-10 min", "Superficies no metálicas. Incompatible con jabones aniónicos."],
        ["Peróxido de hidrógeno", "3-6%", "30 min", "Alto nivel. Material oxidante. Acción contra esporas."],
        ["Ácido peracético", "0.2-0.35%", "10-15 min", "Alto nivel. Esterilizante en frio para endoscopios."],
        ["Povidona yodada", "10%", "1-5 min", "Antiséptico piel y mucosas. Baja actividad contra esporas."],
        ["Clorhexidina", "2-4%", "1-3 min", "Antiséptico piel. Persistente, no irritante."],
        ["UV-C", "254 nm", "30-60 min", "Aire y superficies. No usar en presencia de personal (daño ocular)."],
    ]
    add_data_table(doc, desinf_headers, desinf_rows,
                   col_widths_cm=[3.5, 2.5, 2.0, 9.3], header_size=8, body_size=8)
    add_para(doc,
        "Advertencia: NUNCA mezclar desinfectantes (especialmente hipoclorito con amoníaco "
        "o ácidos) — genera gases tóxicos. Conservar en envases originales etiquetados, "
        "fuera del alcance de animales y fuera de alimentos.",
        size=9, italic=True, color=RED_CRIT)

    # ===== CAPÍTULO IX — PREVENCIÓN DE MORDEDURAS =====
    add_chapter(doc, "IX", "PREVENCIÓN DE MORDEDURAS Y ARañAZOS")
    add_article(doc, 14, "Principios",
        "Toda manipulación de animal se realiza con sujeción adecuada (bozal, toalla, "
        "campana, lazo) y, cuando sea necesario, con sedación. Se prohíbe: realizar "
        "procedimientos en animales no sedados que muestren agresividad, trabajar solo con "
        "animales peligrosos, y no reportar incidentes.")
    add_article(doc, 15, "Protección",
        "El personal usará guantes de cuero largos para manejo de animales agresivos o "
        "no conocidos. La sujeción química (sedación) se prefiere siempre que el estado "
        "clínico del animal lo permita.")

    # ===== CAPÍTULO X — PROFILAXIS VETERINARIA =====
    add_chapter(doc, "X", "PROFILAXIS DEL PERSONAL VETERINARIO")
    add_article(doc, 16, "Vacunación obligatoria",
        "Todo el personal veterinario y auxiliar debe mantener esquema completo de:")
    add_bullet(doc, "Antirrábica pre-exposición: 3 dosis (días 0, 7, 21) + titulación cada 2 años.", bold_lead="Rabia:  ")
    add_bullet(doc, "Esquema completo al día (refuerzo cada 10 años).", bold_lead="Tétanos:  ")
    add_bullet(doc, "3 dosis (esquema 0-1-6 meses) + titulación.", bold_lead="Hepatitis B:  ")
    add_bullet(doc, "Triple vírica (sarampión, paperas, rubéola), influenza estacional, COVID-19.", bold_lead="Otras:  ")
    add_article(doc, 17, "Exámenes médicos",
        "Conforme a la NT-02-2008, el personal se somete a examen médico pre-empleo, "
        "periódico anual y post-ocupacional, con énfasis en serologías para leptospirosis, "
        "toxoplasmosis, brucelosis y hepatitis B.")

    # ===== CAPÍTULO XI — ACCIDENTES DE TRABAJO =====
    add_chapter(doc, "XI", "ATENCIÓN ANTE ACCIDENTES DE TRABAJO")
    add_article(doc, 18, "Procedimiento inmediato",
        "Ante cualquier accidente ocupacional (mordedura, cortopunzante, salpicadura en "
        "mucosas, exposición química):")
    add_bullet(doc, "Lavar la zona afectada con agua abundante (15 minutos en mucosas, cortopunzantes o químicos).")
    add_bullet(doc, "Aplicar antiséptico y cubrir la herida.")
    add_bullet(doc, "Reportar al supervisor en un máximo de 2 horas (Art. 73 LOPCYMAT).")
    add_bullet(doc, "Acudir al centro asistencial designado por el IVSS-PMSSO.")
    add_bullet(doc, "Completar formulario de reporte de incidente (Anexo del Procedimiento de Reporte).")
    add_bullet(doc, "Seguir esquema PEP (post-exposición) si aplica: rabia, hepatitis B, VIH, tétanos.")
    add_article(doc, "18b", "Algoritmo de respuesta ante accidentes",
        "El siguiente algoritmo resume la respuesta inmediata:")
    add_alert_box(doc, [
        "**ALGORITMO ANTE EXPOSICIÓN OCUPACIONAL",
        "1. PARAR la actividad inmediatamente.",
        "2. LAVAR la zona afectada con agua y jabón durante 15 minutos (mucosas: irrigar con suero fisiológico).",
        "3. ANTISÉPTICO: alcohol 70% o povidona yodada en piel lesionada.",
        "4. CUBRIR con apósito estéril. NO suturar mordeduras.",
        "5. NOTIFICAR al supervisor directo en ≤30 minutos.",
        "6. REPORTAR formalmente en formulario de incidente (≤2 horas Art. 73 LOPCYMAT).",
        "7. CLASIFICAR la exposición (Tipo 1-4 OMS / categoría del material).",
        "8. CENTRO ASISTENCIAL del IVSS-PMSSO para evaluación y PEP.",
        "9. PEP según corresponda: vacuna antirrábica + IGAR (Tipo 3-4), vacuna Hepatitis B + IGHB, antitetánica, profilaxis VIH (dentro de 72 h).",
        "10. SEGUIMIENTO serológico a las 6 semanas, 3 y 6 meses para VIH, HBV y HCV.",
        "11. INVESTIGACIÓN de causa raíz y plan de acción por el Comité SST.",
        "12. CIERRE del expediente y archivo por 10 años (Art. 183 LOTTT)."
    ])
    add_article(doc, "18c", "Plazos legales de reporte",
        "Los plazos máximos de reporte conforme a la LOPCYMAT son:")
    add_bullet(doc, "Accidente grave: 2 horas ante el IVSS y la Inspectoría de INPSASEL.", bold_lead="Accidente grave:  ")
    add_bullet(doc, "Accidente común: 1 día hábil siguiente.", bold_lead="Accidente común:  ")
    add_bullet(doc, "Accidente in itinere: 1 día hábil al IVSS.", bold_lead="In itinere:  ")
    add_bullet(doc, "Enfermedad ocupacional: 5 días hábiles ante INPSASEL.", bold_lead="Enfermedad ocupacional:  ")
    add_bullet(doc, "Accidente mortal: reporte inmediato al IVSS, INPSASEL y CICPC.", bold_lead="Mortal:  ")

    # ===== CAPÍTULO XII — EMBARAZO Y LACTANCIA =====
    add_chapter(doc, "XII", "PROTECCIÓN AL PERSONAL EMBARAZADA Y EN LACTANCIA")
    add_article(doc, 19, "Restricciones",
        "Conforme al artículo 78 de la LOPCYMAT y a la Ley Especial para la Dignificación "
        "de la Mujer Trabajadora, las trabajadoras embarazadas o en lactancia NO "
        "desempeñarán labores que impliquen:")
    add_bullet(doc, "Exposición a radiaciones ionizantes (rayos X veterinarios).")
    add_bullet(doc, "Manejo de anestésicos volátiles (isoflurano, sevoflurano) sin sistema de gas scavenging.")
    add_bullet(doc, "Manejo de quimioterápicos veterinarios.")
    add_bullet(doc, "Manejo de animales con zoonosis de elevada riesgo (brucelosis, leptospirosis, toxoplasmosis).")
    add_bullet(doc, "Trabajo en altura, manipulación de cargas pesadas o posturas forzadas.")
    add_article(doc, 20, "Reasignación de funciones",
        "LA EMPRESA reasignará a la trabajadora embarazada a labores compatibles con su "
        "estado, sin reducción salarial, conforme al artículo 78 de la LOPCYMAT.")
    add_article(doc, "20b", "Derechos de la trabajadora embarazada",
        "Además de la reasignación, la trabajadora embarazada tiene derecho a:")
    add_bullet(doc, "Descanso prenatal de seis (6) semanas antes del parto (Ley para la Protección de la Familia, la Maternidad y la Paternidad).")
    add_bullet(doc, "Descanso postnatal de veinte (20) semanas después del parto.")
    add_bullet(doc, "Dos (2) descansos de treinta (30) minutos diarios para lactancia durante el primer año del bebé.")
    add_bullet(doc, "Prohibición de despido durante el embarazo y hasta un (1) año post-parto (Art. 79 LOTTT).")
    add_bullet(doc, "Atención médica integral a través del Seguro Social Obligatorio.")
    add_article(doc, "20c", "Glosario de términos clave",
        "Para facilitar la comprensión de la presente Cartilla, se define a continuación "
        "el glosario de términos técnicos:")
    glos_headers = ["Término", "Definición"]
    glos_rows = [
        ["Bioseguridad", "Conjunto de medidas preventivas para proteger la salud del personal y del entorno frente a agentes biológicos."],
        ["EPP", "Equipo de Protección Personal: dispositivos individuales que protegen al trabajador de riesgos específicos."],
        ["Fómite", "Objeto inanimado que puede transportar agentes infecciosos (instrumental, ropa, superficies)."],
        ["Zoonosis", "Enfermedad transmisible de animales a humanos y viceversa."],
        ["PEP", "Profilaxis Post-Exposición: tratamiento preventivo posterior a una exposición (ej. rabia)."],
        ["IGAR", "Inmunoglobulina Antirrábica: anticuerpos específicos contra el virus de la rabia."],
        ["PEP Essen", "Esquema de vacunación antirrábica en días 0, 3, 7, 14 y 28."],
        ["Cortopunzante", "Objeto capaz de cortar o perforar la piel (agujas, bisturíes, vidrios)."],
        ["NB-1 a NB-4", "Niveles de Bioseguridad del 1 al 4, según el riesgo del agente."],
        ["INHRR", "Instituto Nacional de Higiene 'Rafael Rangel' — referencia nacional para diagnóstico de rabia."],
        ["IVSS-PMSSO", "Programa de Medicina del Seguro Social Obligatorio — atención al trabajador."],
        ["INPSASEL", "Instituto Nacional de Prevención, Salud y Seguridad Laborales."],
        ["NT-01-2008", "Norma Técnica para la selección y uso de EPP en Venezuela."],
        ["NT-02-2008", "Norma Técnica para la Vigilancia de la Salud de los trabajadores."],
        ["COVENIN 2747-93", "Norma venezolana para el manejo de residuos en establecimientos de salud."],
        ["Comité SST", "Comité paritario de Seguridad y Salud Laboral, regulado por el Art. 46 LOPCYMAT."],
        ["Casi-accidente", "Suceso que pudo causar daño pero no lo hizo (near miss)."],
    ]
    add_data_table(doc, glos_headers, glos_rows,
                   col_widths_cm=[3.5, 13.8], header_size=9, body_size=9)

    # ===== ANEXO: Directorio de emergencias =====
    add_section(doc, "ANEXO — DIRECTORIO DE EMERGENCIAS Y CENTROS DE REFERENCIA")
    dir_headers = ["Tipo", "Contacto", "Teléfono / Dirección"]
    dir_rows = [
        ["Emergencia médica", "Bomberos / 911", "911 — cobertura nacional"],
        ["Centro antirrábico", "Distrito Sanitario N° 2 — Altos Mirandinos", "Los Teques, Miranda — consultar directorio MPPS"],
        ["Instituto Nacional de Higiene", "INHRR 'Rafael Rangel'", "Caracas — para envío de muestras de rabia"],
        ["IVSS", "Centro Asistencial designado", "Consultar PMSSO — Orden de Atención"],
        ["INPSASEL", "Inspectoría del Trabajo — Miranda", "Los Teques, Estado Miranda"],
        ["CICPC", "Denuncia por robo o sustancias", "911 o delegación local"],
        ["SENAC", "Reporte de sustancias controladas", "Caracas — vía portal del MPPP"],
        ["Coordinación SST — LA EMPRESA", "Directora Gerente", "Esnatlim Elena Simoza — interno Gerencia"],
        ["Comité SST", "Buzón de denuncias", "Recepción — LA EMPRESA"],
    ]
    add_data_table(doc, dir_headers, dir_rows,
                   col_widths_cm=[4.0, 5.5, 7.8], header_size=8, body_size=8)

    # ===== TABLA RESUMEN FINAL =====
    add_section(doc, "TABLA RESUMEN — EPP OBLIGATORIO GENERAL")
    add_para(doc, "El siguiente cuadro resume el EPP de uso obligatorio en el centro de trabajo:",
        size=9, italic=True, space_after=4)
    res_headers = ["Cargo / Área", "EPP Básico", "EPP Especial"]
    res_rows = [
        ["Médico Veterinario", "Bata, guantes, mascarilla quirúrgica", "Bata estéril, gafas, gorro, N95 (necropsia)"],
        ["Auxiliar Veterinario", "Bata, guantes, mascarilla quirúrgica", "Delantal impermeable, botas (limpieza jaulas)"],
        ["Dog Groomer", "Delantal impermeable, guantes de nitrilo", "Mascarilla N95 (asma/alergias), botas antideslizantes"],
        ["Encargado de Tienda", "Guantes (carga), calzado cerrado", "Mascarilla (productos químicos)"],
        ["Personal de Limpieza", "Guantes de goma, botas, delantal", "Mascarilla con filtro (desinfectantes fuertes)"],
    ]
    add_data_table(doc, res_headers, res_rows,
                   col_widths_cm=[3.5, 6.0, 7.8], header_size=9, body_size=9, center_cols=[0])

    # ===== Carta de recepción =====
    add_reception_letter(doc, "Cartilla de Bioseguridad Veterinaria")

    add_footer(doc.sections[0], FOOTER_BASE)
    out_path = os.path.join(OUT_DIR_05, "Cartilla_Bioseguridad_Veterinaria.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# DOCUMENTO 2: PROTOCOLO DE SUSTANCIAS CONTROLADAS
# ============================================================
def gen_sustancias_controladas():
    doc = Document()
    setup_a4_portrait(doc, margins_cm=1.8)
    add_membrete(doc, "PROTOCOLO SUSTANCIAS CONTROLADAS", "SENAC / Ley Orgánica Drogas", "v3.0  ·  SST")
    add_doc_title(doc, "PROTOCOLO DE MANEJO DE SUSTANCIAS CONTROLADAS")

    # ===== CAPÍTULO I — ÁMBITO Y MARCO LEGAL =====
    add_chapter(doc, "I", "ÁMBITO Y MARCO LEGAL")
    add_article(doc, 1, "Objeto",
        "Regular la adquisición, almacenamiento, prescripción, dispensación, "
        "administración, descarte y reporte de sustancias controladas utilizadas en la "
        "actividad veterinaria de LA EMPRESA.")
    add_article(doc, 2, "Sustancias incluidas",
        "Se consideran sustancias controladas: anestésicos generales (ketamina, "
        "propofol, tiopental), opioides (morfina, tramadol, fentanilo, buprenorfina, "
        "metadona), benzodiacepinas (diazepam, midazolam), agonistas alfa-2 (medetomidina, "
        "xilacina) y otros psicótropos según clasificación del SENAC y la Ley Orgánica de "
        "Drogas.")
    add_article(doc, 3, "Marco legal",
        "Ley Orgánica de Drogas (2010), Resoluciones del SENAC (Servicio Nacional "
        "Autónomo de Sustancias Químicas, Precursoras y Estupefacientes), Convención "
        "Única de Estupefacientes de 1961, Convención de Viena de 1988, Ley de Ejercicio "
        "de la Medicina Veterinaria, Código Penal y Norma COVENIN aplicable.")
    add_article(doc, "3b", "Fichas técnicas de sustancias controladas",
        "Las siguientes fichas resumen las características, indicaciones y "
        "precauciones de cada sustancia controlada utilizada en LA EMPRESA:")
    sub_headers = ["Sustancia", "Clase / Lista", "Indicación principal", "Antagonista"]
    sub_rows = [
        ["Ketamina", "Lista IV SENAC", "Anestésico disociativo IV/IM en perros, gatos, roedores, aves", "No específico (sintomático)"],
        ["Propofol", "No controlado (vigilar)", "Inducción anestésica IV en perros y gatos", "No específico (sintomático)"],
        ["Tiopental sódico", "Lista III SENAC", "Inducción anestésica IV (uso limitado)", "No específico (sintomático)"],
        ["Morfina", "Lista I SENAC — OPIOIDE", "Analgesia mayor perioperatoria", "Naloxona 0.04 mg/kg IV/IM"],
        ["Fentanilo", "Lista I SENAC — OPIOIDE", "Analgesia mayor, anestesia balanceada", "Naloxona 0.04 mg/kg IV/IM"],
        ["Metadona", "Lista I SENAC — OPIOIDE", "Analgesia mayor, programa sustitución", "Naloxona 0.04 mg/kg IV/IM"],
        ["Buprenorfina", "Lista I SENAC — OPIOIDE", "Analgesia moderada, postoperatorio", "Naloxona (parcial), Naltrexona"],
        ["Tramadol", "Lista IV SENAC — OPIOIDE", "Analgesia moderada, dolor crónico", "Naloxona (parcial)"],
        ["Diazepam", "Lista IV SENAC", "Sedación, anticonvulsivante, relajante muscular", "Flumazenilo 0.02 mg/kg IV"],
        ["Midazolam", "Lista IV SENAC", "Sedación, premedicación, inducción", "Flumazenilo 0.02 mg/kg IV"],
        ["Medetomidina", "No controlado (vigilar)", "Sedación alfa-2, premedicación", "Atipamezol 5× dosis mg/kg IM"],
        ["Xilacina", "No controlado (vigilar)", "Sedación alfa-2 en equinos, rumiantes, caninos", "Atipamezol o Yohimbina"],
    ]
    add_data_table(doc, sub_headers, sub_rows,
                   col_widths_cm=[3.0, 3.5, 7.5, 4.3], header_size=8, body_size=8, center_cols=[1])

    # ===== CAPÍTULO II — INVENTARIO Y LIBRO FOLIADO =====
    add_chapter(doc, "II", "INVENTARIO Y LIBRO FOLIADO")
    add_article(doc, 4, "Libro foliado obligatorio",
        "LA EMPRESA mantiene un libro foliado, habilitado y sellado por el SENAC, donde "
        "se registra TODO movimiento de sustancias controladas. Cada asiento contendrá:")
    add_data_table(doc,
        ["N°", "Campo a registrar"],
        [
            ["1", "Fecha del movimiento (día, mes, año, hora)"],
            ["2", "Tipo de movimiento: INGRESO, DISPENSACIÓN, ADMINISTRACIÓN, DESCARTE, AJUSTE, AUDITORÍA"],
            ["3", "Nombre genérico y comercial del producto"],
            ["4", "Forma farmacéutica y concentración"],
            ["5", "Lote del producto"],
            ["6", "Fecha de vencimiento"],
            ["7", "Cantidad (en unidades o mL)"],
            ["8", "Saldo anterior / Saldo nuevo"],
            ["9", "Nombre del paciente (especie, raza, nombre)"],
            ["10", "Nombre del propietario y C.I."],
            ["11", "Médico veterinario responsable y C.M.V.V."],
            ["12", "Firma del profesional y firma de quien recibe/descarta"],
        ],
        col_widths_cm=[1.2, 16.1], header_size=9, body_size=9, center_cols=[0])
    add_article(doc, 5, "Conciliación",
        "El libro foliado se concilia mensualmente con el inventario físico. Cualquier "
        "discrepancia se investiga y documenta en acta firmada por el responsable.")

    # ===== CAPÍTULO III — ALMACENAMIENTO Y DOBLE CERRADURA =====
    add_chapter(doc, "III", "ALMACENAMIENTO Y DOBLE CERRADURA")
    add_alert_box(doc, [
        "**MEDIDA DE SEGURIDAD OBLIGATORIA — DOBLE CERRADURA PARA OPIOIDES",
        f"Los opioides (morfina, fentanilo, metadona, buprenorfina) y otras sustancias de "
        f"mayor riesgo de abuso se almacenan en caja fuerte o gabinete de seguridad con "
        f"DOS (2) cerraduras independientes. Las llaves se custodian por separado: una "
        f"con el médico veterinario de turno y otra con la Dirección. El acceso requiere "
        f"presencia simultánea de dos personas autorizadas y registro en libro foliado."
    ])
    add_article(doc, 6, "Características del almacén",
        "El área de almacenamiento cumple con: ventilación, temperatura controlada (15-25°C), "
        "protección contra luz directa, separación de sustancias incompatibles, señalética "
        "de seguridad, extintor ABC, acceso restringido con registro de entradas.")
    add_article(doc, 7, "Inventario mínimo y máximo",
        "Se maneja inventario mínimo de una (1) semana de consumo promedio y máximo de un "
        "(1) mes. La adquisición se realiza con proveedor autorizado por el SENAC, con "
        "factura y guía de remisión que se conservan por cinco (5) años.")

    # ===== CAPÍTULO IV — CADENA DE CUSTODIA =====
    add_chapter(doc, "IV", "CADENA DE CUSTODIA — 5 ETAPAS")
    add_article(doc, 8, "Etapas",
        "La cadena de custodia es el registro ininterrumpido de todos los movimientos de "
        "una sustancia controlada desde su adquisición hasta su administración o descarte. "
        "Las cinco etapas son:")
    etapas_headers = ["#", "Etapa", "Responsable", "Documento"]
    etapas_rows = [
        ["1", "Adquisición", "Directora Gerente + MV", "Orden de compra, factura, guía remisión"],
        ["2", "Recepción y verificación", "Médico Veterinario encargado", "Acta de recepción, registro libro foliado"],
        ["3", "Almacenamiento", "Médico Veterinario + Director", "Inventario, libro foliado (INGRESO)"],
        ["4", "Prescripción y dispensación", "Médico Veterinario tratante", "Receta médica, fórmula magistral, libro foliado (DISPENSACIÓN)"],
        ["5", "Administración o descarte", "MV + Auxiliar (doble chequeo)", "Registro clínico, libro foliado (ADMIN/DESCARTE), acta de descarte"],
    ]
    add_data_table(doc, etapas_headers, etapas_rows,
                   col_widths_cm=[0.8, 4.2, 4.2, 8.1], header_size=9, body_size=9, center_cols=[0])
    add_article(doc, 9, "Doble chequeo de opioides",
        "La administración de opioides requiere DOBLE CHEQUEO por dos profesionales: el "
        "médico veterinario que prescribe y otro profesional que verifica dosis, vía y "
        "paciente. Ambos firman el registro clínico y el libro foliado.")

    # ===== CAPÍTULO V — PRESCRIPCIÓN Y ADMINISTRACIÓN =====
    add_chapter(doc, "V", "PRESCRIPCIÓN Y ADMINISTRACIÓN")
    add_article(doc, 10, "Receta médica",
        "La receta de sustancias controladas debe contener: fecha, datos del MV (nombre, "
        "C.I., C.M.V.V., firma y sello), datos del propietario y del paciente, nombre del "
        "producto (genérico), concentración, dosis, vía, frecuencia, duración y "
        "cantidad dispensada. La receta se conserva por cinco (5) años.")
    add_article(doc, 11, "Restricciones",
        "Se prohíbe: prescribir sustancias controladas sin fundamento clínico documentado, "
        "auto-prescripción, dispensación a personal no autorizado, fraccionamiento no "
        "justificado, y uso veterinario de productos no registrados ante el SENAC.")
    add_article(doc, "11b", "Doble chequeo obligatorio",
        "La administración de opioides y anestésicos generales requiere doble chequeo "
        "por DOS (2) profesionales calificados: el médico veterinario que prescribe y "
        "otro profesional que verifica en voz alta: (i) nombre del paciente, (ii) "
        "nombre del producto, (iii) concentración, (iv) dosis calculada, (v) vía de "
        "administración, y (vi) hora de administración. Ambos firman el registro clínico "
        "y el libro foliado. Solo en caso de guardia nocturna con un único profesional, "
        "el doble chequeo puede ser sustituido por verificación con checklist firmado, "
        "informando a la Dirección en el siguiente día hábil.")
    add_article(doc, "11c", "Manejo de sobredosis accidental",
        "En caso de sobredosis accidental de sustancias controladas en el paciente animal, "
        "se aplican las siguientes acciones inmediatas:")
    sobredosis_headers = ["Sustancia", "Signos en el paciente", "Antagonista / Acción"]
    sobredosis_rows = [
        ["Morfina, Fentanilo, Metadona, Tramadol, Buprenorfina",
         "Depresión respiratoria, miosis, bradicardia, sedación profunda",
         "Naloxona 0.04 mg/kg IV/IM (repetir cada 15 min). Oxígeno suplementario. Ventilación asistida."],
        ["Ketamina",
         "Salivación, espasmos musculares, convulsiones, taquicardia",
         "Diazepam 0.5 mg/kg IV. Soporte cardiovascular. Líquidos IV."],
        ["Propofol",
         "Apnea, hipotensión, cianosis",
         "Oxígeno y ventilación asistida. Detener administración. Soporte cardiovascular."],
        ["Diazepam, Midazolam",
         "Sedación profunda, ataxia, depresión respiratoria leve",
         "Flumazenilo 0.02 mg/kg IV (uso hospitalario). Soporte respiratorio."],
        ["Medetomidina, Xilacina",
         "Bradicardia severa, hipotensión, hipotermia, sedación profunda",
         "Atipamezol 5× la dosis de medetomidina IM. Anticolinérgicos solo si bradicardia severa."],
    ]
    add_data_table(doc, sobredosis_headers, sobredosis_rows,
                   col_widths_cm=[4.0, 5.5, 8.3], header_size=8, body_size=8)
    add_para(doc,
        "IMPORTANTE: Todo evento de sobredosis accidental debe registrarse en la historia "
        "clínica, notificarse a la Directora Gerente y, si cause daño al paciente o al "
        "personal, se completa el Formulario de Reporte de Incidente conforme al "
        "Procedimiento de Reporte de Incidentes (Capítulo III).",
        size=9, italic=True, color=RED_CRIT, space_before=2)

    # ===== CAPÍTULO VI — DESCARTE Y ELIMINACIÓN =====
    add_chapter(doc, "VI", "DESCARTE Y ELIMINACIÓN")
    add_article(doc, 12, "Procedimiento de descarte",
        "El descarte de sobrantes, productos vencidos o recetas no retiradas se realiza "
        "mediante: (i) acta de descarte firmada por dos testigos, (ii) registro en libro "
        "foliado, (iii) neutralización química o incineración por empresa certificada. "
        "Se prohíbe el descarte por alcantarillado o basura común.")
    add_article(doc, "12b", "Casos frecuentes de descarte",
        "Se consideran sujetos a descarte: ampolletas sobrantes no utilizadas en 24 horas, "
        "frascos de opioides vencidos, jeringas con restos visibles de anestésico, "
        "restos de infusión de fentanilo en bolsas de suero, productos dañados por "
        "exposición térmica, y recetas no retiradas por el cliente después de 30 días. "
        "En todos los casos el acta de descarte incluye: fecha, producto, lote, cantidad, "
        "causa del descarte, método de destrucción, y firmas del MV responsable, de un "
        "testigo y de la Directora Gerente.")

    # ===== CAPÍTULO VII — REPORTES Y AUDITORÍA =====
    add_chapter(doc, "VII", "REPORTES AL SENAC Y AUDITORÍA INTERNA")
    add_article(doc, 13, "Reporte trimestral",
        "LA EMPRESA presenta al SENAC el informe trimestral de movimientos de sustancias "
        "controladas, conforme a la normativa vigente. El reporte lo elabora el médico "
        "veterinario responsable y lo firma la Directora Gerente.")
    add_article(doc, 14, "Pérdida o robo",
        "En caso de pérdida, robo o sustracción de sustancias controladas, se debe: (i) "
        "denunciar ante el CICPC en un máximo de 24 horas, (ii) notificar al SENAC en 24 "
        "horas, (iii) realizar investigación interna, (iv) registrar acta en libro foliado.")
    add_article(doc, 15, "Auditoría interna mensual",
        "Mensualmente se realiza auditoría física del inventario, conciliación con libro "
        "foliado, verificación de fechas de vencimiento y revisión de actas de descarte. "
        "El acta de auditoría la firman el MV responsable y la Directora Gerente.")

    # ===== CAPÍTULO VIII — SANCIONES =====
    add_chapter(doc, "VIII", "SANCIONES")
    add_article(doc, 16, "Régimen sancionatorio",
        "El incumplimiento del presente Protocolo se sanciona conforme al Reglamento "
        "Interno de Trabajo y a la legislación aplicable:")
    add_sanciones_table(doc,
        ["Falta", "Tipo", "Sanción"],
        [
            ["Olvido de registro en libro foliado", "LEVE", "Amonestación verbal"],
            ["Error en dosis sin daño al paciente", "GRAVE", "Amonestación escrita + reentrenamiento"],
            ["Acceso no autorizado al almacén", "GRAVE", "Amonestación escrita + investigación"],
            ["Sustraer sustancia controlada", "MUY GRAVE", "DESPIDO JUSTIFICADO + denuncia CICPC"],
            ["Auto-administración o desvío", "MUY GRAVE", "DESPIDO JUSTIFICADO + denuncia penal"],
            ["Falsificar registro o receta", "MUY GRAVE", "DESPIDO JUSTIFICADO + denuncia penal"],
        ],
        col_widths_cm=[6.5, 2.5, 7.6])

    # ===== CAPÍTULO IX — CAPACITACIÓN DEL PERSONAL =====
    add_chapter(doc, "IX", "CAPACITACIÓN DEL PERSONAL")
    add_article(doc, "16b", "Programa de capacitación",
        "Todo el personal con acceso a sustancias controladas recibirá capacitación "
        "obligatoria sobre el presente Protocolo, con una duración mínima de cuatro (4) "
        "horas al ingreso y refrescos anuales de dos (2) horas. La capacitación cubre:")
    add_bullet(doc, "Marco legal: Ley Orgánica de Drogas, Resoluciones SENAC, Convenciones internacionales.")
    add_bullet(doc, "Clasificación de sustancias controladas y Listas del SENAC.")
    add_bullet(doc, "Manejo del libro foliado: cómo y cuándo registrar movimientos.")
    add_bullet(doc, "Almacenamiento con doble cerradura y control de llaves.")
    add_bullet(doc, "Cadena de custodia de 5 etapas: responsabilidades por actor.")
    add_bullet(doc, "Doble chequeo de opioides y anestésicos: protocolo verbal.")
    add_bullet(doc, "Manejo de sobredosis accidental y antagonistas disponibles.")
    add_bullet(doc, "Procedimiento de descarte y acta de destrucción.")
    add_bullet(doc, "Reporte al SENAC: plazos y formularios.")
    add_bullet(doc, "Actuación ante pérdida, robo o sustracción.")
    add_article(doc, "16c", "Evaluación y constancia",
        "Al finalizar la capacitación, el personal completa una evaluación escrita con "
        "diez (10) preguntas; se requiere una puntuación mínima de ochenta por ciento "
        "(80%) para considerar aprobada. La constancia de capacitación se archiva en el "
        "expediente del trabajador(a) por cinco (5) años.")
    add_article(doc, "16d", "Inducción a nuevo personal",
        "El nuevo personal que tenga acceso a sustancias controladas recibe una inducción "
        "específica de dos (2) horas el primer día de trabajo, antes de cualquier "
        "manipulación. La inducción es responsabilidad del MV encargado y queda "
        "documentada en el Formato de Inducción SST.")

    # ===== ANEXO A: Formulario dispensación =====
    add_section(doc, "ANEXO A — FORMULARIO DE DISPENSACIÓN DE SUSTANCIAS CONTROLADAS")
    add_para(doc, "Este formulario se completa por triplicado (copia clínica, copia libro foliado, copia SENAC).",
        size=9, italic=True, space_after=4)
    add_data_table(doc,
        ["Campo", "Información"],
        [
            ["Fecha y hora", "____ / ____ / ______  ·  ____:____"],
            ["Producto (genérico y comercial)", "_______________________________________________"],
            ["Concentración / forma farmacéutica", "_______________________________________________"],
            ["Lote", "__________________  ·  Vence: ____ / ____ / ______"],
            ["Cantidad dispensada (unidades/mL)", "_______________________________________________"],
            ["Saldo anterior", "____________  ·  Saldo nuevo: ____________"],
            ["Paciente (especie, raza, nombre)", "_______________________________________________"],
            ["Propietario / C.I.", "_______________________________________________"],
            ["Médico Veterinario / C.M.V.V.", "_______________________________________________"],
            ["Firma MV", "_______________________________________________"],
            ["Firma Auxiliar (doble chequeo)", "_______________________________________________"],
        ],
        col_widths_cm=[5.0, 12.3], header_size=9, body_size=9)

    # ===== Carta de recepción =====
    add_reception_letter(doc, "Protocolo de Manejo de Sustancias Controladas")

    add_footer(doc.sections[0], FOOTER_BASE)
    out_path = os.path.join(OUT_DIR_08, "Protocolo_Sustancias_Controladas.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# DOCUMENTO 3: PROTOCOLO DE MORDEDURAS Y ZOONOSIS
# ============================================================
def gen_mordeduras_zoonosis():
    doc = Document()
    setup_a4_portrait(doc, margins_cm=1.8)
    add_membrete(doc, "PROTOCOLO MORDEDURAS", "OMS / OIE — Profilaxis PEP", "v3.0  ·  SST")
    add_doc_title(doc, "PROTOCOLO DE MORDEDURAS Y ZOONOSIS POST-EXPOSICIÓN")

    # ===== CAPÍTULO I — OBJETIVO Y ÁMBITO =====
    add_chapter(doc, "I", "OBJETIVO Y ÁMBITO")
    add_article(doc, 1, "Objetivo",
        "Establecer el procedimiento de atención inmediata ante mordeduras, arañazos y "
        "exposición a fluidos de animales en el personal de LA EMPRESA, conforme a los "
        "lineamientos de la Organización Mundial de la Salud (OMS), la Organización "
        "Mundial de Sanidad Animal (OIE), el Ministerio del Poder Popular para la Salud y "
        "la LOPCYMAT.")
    add_article(doc, 2, "Ámbito",
        "Aplica a todo el personal, en especial al médico veterinario, auxiliar "
        "veterinario y dog groomer, así como a cualquier trabajador(a) que sufra exposición "
        "ocupacional a mordedura, arañazo o fluido de animal.")

    # ===== CAPÍTULO II — CLASIFICACIÓN OMS/OIE =====
    add_chapter(doc, "II", "CLASIFICACIÓN DE LA EXPOSICIÓN — OMS/OIE")
    add_article(doc, 3, "Categorías",
        "La OMS clasifica la exposición a material potencialmente rábico en cuatro tipos:")
    clasif_headers = ["Tipo", "Descripción", "Acción inmediata"]
    clasif_rows = [
        ["TIPO 1", "Tocar, alimentar o lamido de piel intacta (sin lesión)", "Lavado de manos. No requiere PEP. Vigilar animal 10 días."],
        ["TIPO 2", "Arañazo superficial sin sangrado, lamido de piel lesionada", "Lavado 15 min con jabón. PEP según estado animal. Vigilar 10 días."],
        ["TIPO 3", "Mordedura única o múltiples con sangrado, lamido de mucosa", "Lavado 15 min + antiséptico. PEP COMPLETO (vacuna + IGAR)."],
        ["TIPO 4", "Exposición grave: mordedura cabeza/cuello/mano, animal rábico confirmado", "Lavado 15 min + IGAR + vacuna regimen Essen completo + serología."],
    ]
    add_data_table(doc, clasif_headers, clasif_rows,
                   col_widths_cm=[1.6, 7.4, 8.3], header_size=9, body_size=9, center_cols=[0])

    # ===== CAPÍTULO III — PROCEDIMIENTO INMEDIATO (8 PASOS) =====
    add_chapter(doc, "III", "PROCEDIMIENTO INMEDIATO — 8 PASOS")
    add_article(doc, 4, "Acción inmediata",
        "Ante cualquier mordedura, arañazo o salpicadura de fluido en mucosa o piel "
        "lesionada, se aplican los siguientes OCHO (8) pasos:")
    add_bullet(doc, "Lavar inmediatamente la herida con agua y jabón durante 15 minutos, sangrando la herida si es posible.", bold_lead="PASO 1:  ")
    add_bullet(doc, "Aplicar antiséptico: alcohol 70% o povidona yodada.", bold_lead="PASO 2:  ")
    add_bullet(doc, "Cubrir la herida con apósito estéril limpio. NO suturar en caso de mordedura (riesgo de infección).", bold_lead="PASO 3:  ")
    add_bullet(doc, "Notificar al supervisor en un máximo de 2 horas (Art. 73 LOPCYMAT).", bold_lead="PASO 4:  ")
    add_bullet(doc, "Completar formulario de reporte de exposición (Anexo B).", bold_lead="PASO 5:  ")
    add_bullet(doc, "Clasificar la exposición (Tipo 1-4) según Capítulo II.", bold_lead="PASO 6:  ")
    add_bullet(doc, "Acudir al centro antirrábico o al IVSS-PMSSO para evaluación y PEP.", bold_lead="PASO 7:  ")
    add_bullet(doc, "Investigar al animal mordedor (vacunación, procedencia, comportamiento) y completar esquema PEP si corresponde.", bold_lead="PASO 8:  ")

    # ===== CAPÍTULO IV — MANEJO DEL ANIMAL MORDEDOR =====
    add_chapter(doc, "IV", "MANEJO DEL ANIMAL MORDEDOR")
    add_article(doc, 5, "Casos posibles",
        "El manejo del animal mordedor depende de su estatus sanitario y procedencia:")
    add_bullet(doc, "Si está vacunado y sin signos: observación veterinaria durante 10 días. Si continúa sano, no se indica PEP. Si enferma o muere, se inicia PEP completo y se envía cabeza al INHRR.", bold_lead="Animal doméstico vacunado:  ")
    add_bullet(doc, "Si NO está vacunado o vacuna dudosa: iniciar PEP inmediatamente y observar 10 días. Si enferma o muere, enviar cabeza al INHRR.", bold_lead="Animal doméstico no vacunado:  ")
    add_bullet(doc, "Capturar si es posible y observar 10 días. Si no se captura, iniciar PEP completo inmediatamente.", bold_lead="Animal callejero:  ")
    add_bullet(doc, "En caso de muerte o sacrificio del animal sospechoso: enviar la cabeza refrigerada (NO congelada) al Instituto Nacional de Higiene Rafael Rangel (INHRR) para prueba de inmunofluorescencia.", bold_lead="Animal muerto o sacrificado:  ")
    add_article(doc, "5b", "Signos de alerta en el animal mordedor",
        "El personal debe identificar signos compatibles con rabia en el animal mordedor:")
    add_bullet(doc, "Cambio brusco de conducta (agresividad inusual o excesiva afectividad).")
    add_bullet(doc, "Parálisis progresiva: mandíbula caída, dificultad para tragar, salivación abundante.")
    add_bullet(doc, "Fotofobia, hidrofobia, incoordinación, convulsiones.")
    add_bullet(doc, "Mordedura sin provocación, destrucción de objetos, vagabundeo.")
    add_bullet(doc, "Muerte en un plazo de 5 a 10 días desde la aparición de signos.")
    add_para(doc,
        "Ante cualquiera de estos signos, el animal se considera sospechoso de rabia. "
        "Se notifica inmediatamente al MPPS / Servicio Autónomo de Salud Animal (SASA) y "
        "se inicia PEP completo en el personal expuesto.",
        size=9, italic=True, color=RED_CRIT)
    add_article(doc, "5c", "Sacrificio humanitario y envío de muestra",
        "Si el animal muere o se sacrifica por indicación del MV o de la autoridad "
        "sanitaria, se procede así:")
    add_bullet(doc, "Sacrificio humanitario por MV con método compatible con bienestar animal (pentobarbital IV).")
    add_bullet(doc, "Decapitación aséptica con EPP completo (doble guante, N95, gafas, bata impermeable).")
    add_bullet(doc, "Colocar la cabeza en bolsa estéril, rotular con fecha, especie, número de caso.")
    add_bullet(doc, "Refrigerar a 4°C (NO congelar). Enviar al INHRR en un máximo de 24 horas.")
    add_bullet(doc, "Acompañar con ficha epidemiológica del caso y del personal expuesto.")
    add_bullet(doc, "Mantener cuarentena del área hasta recibir resultado negativo.")

    # ===== CAPÍTULO V — ESQUEMA PEP ESSEN =====
    add_chapter(doc, "V", "ESQUEMA PEP — ESSEN (OMS)")
    add_alert_box(doc, [
        "**ESQUEMA PEP ESSEN — DÍAS 0, 3, 7, 14, 28 + IGAR",
        "La profilaxis post-exposición (PEP) antirrábica conforme al esquema Essen de la OMS:",
        ">> DÍA 0: 1ra dosis + IGAR (Inmunoglobulina Antirrábica) si Tipo 3 o 4 (20 UI/kg, mitad perilesional, mitad IM)",
        ">> DÍA 3: 2da dosis",
        ">> DÍA 7: 3ra dosis",
        ">> DÍA 14: 4ta dosis",
        ">> DÍA 28: 5ta dosis (algunos esquemas OMS omiten esta última)",
        "Vía intradérmica (ID) o intramuscular (IM) en deltoides (adultos) o cara anterolateral del muslo (niños).",
        "Tétanos: refuerzo si >5 años sin vacuna. Antibioticoterapia: amoxicilina/ácido clavulánico 7-10 días.",
    ])
    add_article(doc, 6, "Esquema de vacunación",
        "El esquema Essen se aplica en los días 0, 3, 7, 14 y 28, con vacuna antirrábica "
        "de cultivo celular (Verorab, Rabipur). En exposiciones Tipo 3 y 4 se adiciona "
        "IGAR (inmunoglobulina antirrábica) el día 0.")
    add_article(doc, "6b", "Esquemas alternativos aprobados por la OMS",
        "Como alternativa al esquema Essen clásico de 5 dosis, la OMS ha aprobado los "
        "siguientes esquemas abreviados:")
    alt_headers = ["Esquema", "Días", "Notas"]
    alt_rows = [
        ["Essen clásico IM", "0, 3, 7, 14, 28 (5 visitas)", "1 dosis IM por visita. Adultos en deltoides; niños en muslo anterolateral."],
        ["Essen abreviado ID", "0, 3, 7 (3 visitas)", "2 sitios ID por visita. Aprobado OMS 2018. Menor cantidad de vacuna."],
        ["IPC (Institut Pasteur de Cambodia) ID", "0, 3, 7, 28 (4 visitas)", "2 sitios ID en cada visita. Eficacia equivalente al Essen clásico."],
        ["2-sites ID (Thai Red Cross)", "0, 3, 7, 28 (4 visitas)", "2 sitios ID. Reduce a 4 visitas."],
    ]
    add_data_table(doc, alt_headers, alt_rows,
                   col_widths_cm=[4.0, 4.0, 9.3], header_size=8, body_size=8)
    add_article(doc, "6c", "IGAR — Inmunoglobulina Antirrábica",
        "La IGAR se aplica únicamente en exposiciones Tipo 3 y 4, en el DÍA 0 (y preferiblemente "
        "antes de las 7 horas post-exposición). Dosis: 20 UI/kg. La mitad se infiltra "
        "alrededor y dentro de la herida (perilesional), la otra mitad se aplica "
        "intramuscular en glúteo o muslo. Si la herida es muy pequeña o en dedos, se "
        "infilttra lo posible perilesional y el resto IM. NO se aplica IGAR si el "
        "trabajador(a) ya recibió vacuna pre-exposición completa.")

    # ===== CAPÍTULO VI — ZOONOSIS POST-EXPOSICIÓN =====
    add_chapter(doc, "VI", "OTRAS ZOONOSIS POST-EXPOSICIÓN")
    add_article(doc, 7, "Vigilancia",
        "Además de rabia, el personal con mordedura debe vigilarse por otras zoonosis:")
    zoon_post_headers = ["Zoonosis", "Periodo incubación", "Signos alerta", "Acción"]
    zoon_post_rows = [
        ["Leptospirosis", "5-14 días", "Fiebre, cefalea, mialgia, ictericia", "Doxiciclina 100 mg/12h × 7 días"],
        ["Toxoplasmosis", "5-18 días", "Adenopatías, fiebre, fatiga", "Serología IgM, espiramicina si embarazo"],
        ["Brucelosis", "1-3 semanas", "Fiebre ondulante, sudores, artralgias", "Serología, rifampicina + doxiciclina"],
        ["Pasteurelosis", "24-72 horas", "Celulitis, absceso", "Amoxicilina/ácido clavulánico"],
        ["Capnocytophaga", "3-7 días", "Sepsis, especialmente inmunodeprimidos", "Hospitalización, IV ceftriaxona"],
        ["Tétanos", "3-21 días", "Trismo, rigidez, opistótonos", "Vacuna + inmunoglobulina antitetánica"],
    ]
    add_data_table(doc, zoon_post_headers, zoon_post_rows,
                   col_widths_cm=[3.0, 2.8, 5.5, 6.0], header_size=9, body_size=9)

    # ===== CAPÍTULO VII — REGISTRO E INVESTIGACIÓN =====
    add_chapter(doc, "VII", "REGISTRO E INVESTIGACIÓN DE CAUSA RAÍZ")
    add_article(doc, 8, "Investigación",
        "Todo incidente de mordedura se investiga para identificar la causa raíz: falla de "
        "sujeción, ausencia de sedación, falta de EPP, defecto de procedimiento, presencia "
        "de terceros, etc. La investigación la realiza el Comité SST con apoyo del MV "
        "tratante y genera un plan de acción documentado.")
    add_article(doc, 9, "Reposo y cobertura",
        "El trabajador(a) accidentado tiene derecho a reposo médico según orden del "
        "centro asistencial del IVSS-PMSSO. La atención médica, medicamentos e "
        "indemnización por accidente de trabajo se cubren conforme al Seguro Social "
        "Obligatorio y a la LOPCYMAT (arts. 56, 73, 130 y siguientes).")

    # ===== CAPÍTULO VIII — PROHIBICIÓN DE REPRESALIAS =====
    add_chapter(doc, "VIII", "PROHIBICIÓN DE REPRESALIAS")
    add_article(doc, 10, "Protección al denunciante",
        "Conforme al artículo 26 de la LOPCYMAT, queda prohibida cualquier represalia "
        "contra el trabajador(a) que reporte una mordedura o exposición. El reporte oportuno "
        "es un derecho y un deber, no constituye falta disciplinaria.")

    # ===== CAPÍTULO IX — DIRECTORIO Y CENTROS DE REFERENCIA =====
    add_chapter(doc, "IX", "DIRECTORIO Y CENTROS DE REFERENCIA")
    add_article(doc, "10b", "Centros antirrábicos y de referencia",
        "Ante cualquier exposición ocupacional, el personal será canalizado a los siguientes "
        "centros especializados:")
    centros_headers = ["Centro", "Función", "Ubicación"]
    centros_rows = [
        ["Distrito Sanitario N° 2 — Altos Mirandinos", "Aplicación de PEP antirrábica (vacuna + IGAR)", "Los Teques, Estado Miranda"],
        ["Instituto Nacional de Higiene 'Rafael Rangel' (INHRR)", "Diagnóstico de rabia por inmunofluorescencia", "Caracas — envío de muestras"],
        ["Centro asistencial del IVSS-PMSSO", "Atención de lesiones, antibioticoterapia, reposo", "Designado por el PMSSO del trabajador"],
        ["Hospital 'Domingo Luciani' (El Llanito)", "Atención de heridas graves, cirugía plástica", "Caracas — referencia nacional"],
        ["Hospital 'Pérez de León' (Petare)", "Atención de mordeduras graves en cara y manos", "Caracas — referencia"],
        ["CICPC — División de Protección Animal", "Investigación de animales agresores sin dueño", "Sede CICPC Miranda"],
        ["SASA — Servicio Autónomo de Salud Animal", "Reporte de rabia animal, vigilancia epidemiológica", "MPPS — sede Caracas"],
        ["Coordinación SST — LA EMPRESA", "Reporte interno, acompañamiento al trabajador", "Internos de la empresa"],
    ]
    add_data_table(doc, centros_headers, centros_rows,
                   col_widths_cm=[5.5, 6.5, 5.3], header_size=8, body_size=8)
    add_article(doc, "10c", "Documentación que debe portar el trabajador(a)",
        "Al acudir al centro de referencia, el trabajador(a) debe portar:")
    add_bullet(doc, "Cédula de identidad y carnet del IVSS / PMSSO.")
    add_bullet(doc, "Constancia de trabajo emitida por LA EMPRESA (con cargo y antigüedad).")
    add_bullet(doc, "Formulario de Reporte de Exposición (Anexo B) completado.")
    add_bullet(doc, "Carnet de vacunación antirrábica pre-exposición (si lo posee).")
    add_bullet(doc, "Información sobre el animal mordedor: especie, raza, vacunación, procedencia, estado actual.")

    # ===== ANEXO B: Formulario de reporte de exposición =====
    add_section(doc, "ANEXO B — FORMULARIO DE REPORTE DE EXPOSICIÓN")
    add_data_table(doc,
        ["Campo", "Información"],
        [
            ["Fecha y hora de la exposición", "____ / ____ / ______  ·  ____:____"],
            ["Lugar del incidente", "_______________________________________________"],
            ["Trabajador(a) expuesto(a) / C.I. / Cargo", "_______________________________________________"],
            ["Tipo de exposición (mordedura, arañazo, salpicadura mucosa)", "_______________________________________________"],
            ["Tipo OMS (1, 2, 3, 4)", "_______"],
            ["Localización anatómica de la herida", "_______________________________________________"],
            ["Descripción de la herida (profundidad, sangrado)", "_______________________________________________"],
            ["Animal mordedor (especie, raza, edad, sexo)", "_______________________________________________"],
            ["Procedencia (doméstico vacunado, no vacunado, callejero)", "_______________________________________________"],
            ["Estado sanitario del animal", "_______________________________________________"],
            ["Destino del animal (observación, sacrificio, escapó)", "_______________________________________________"],
            ["Lavado de la herida (duración, antiséptico)", "_______________________________________________"],
            ["Atención médica recibida (centro, médico)", "_______________________________________________"],
            ["PEP indicado (vacuna, IGAR, antibiótico, antitetánica)", "_______________________________________________"],
            ["Investigación causa raíz", "_______________________________________________"],
            ["Firma del trabajador(a)", "_______________________________________________"],
            ["Firma del supervisor", "_______________________________________________"],
            ["Firma del MV tratante", "_______________________________________________"],
        ],
        col_widths_cm=[6.0, 11.3], header_size=9, body_size=9)

    # ===== Carta de recepción =====
    add_reception_letter(doc, "Protocolo de Mordeduras y Zoonosis Post-Exposición")

    add_footer(doc.sections[0], FOOTER_BASE)
    out_path = os.path.join(OUT_DIR_08, "Protocolo_Mordeduras_Zoonosis.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# DOCUMENTO 4: PROCEDIMIENTO DE REPORTE DE INCIDENTES
# ============================================================
def gen_reporte_incidentes():
    doc = Document()
    setup_a4_portrait(doc, margins_cm=1.8)
    add_membrete(doc, "PROCEDIMIENTO DE REPORTE", "Art. 73 LOPCYMAT", "v3.0  ·  SST")
    add_doc_title(doc, "PROCEDIMIENTO DE REPORTE DE INCIDENTES Y ACCIDENTES DE TRABAJO")

    # ===== CAPÍTULO I — OBJETIVO Y ALCANCE =====
    add_chapter(doc, "I", "OBJETIVO Y ALCANCE")
    add_article(doc, 1, "Objetivo",
        "Establecer el procedimiento uniforme para el reporte, investigación y "
        "seguimiento de los incidentes y accidentes de trabajo ocurridos en las "
        "instalaciones de LA EMPRESA o en Misión de trabajo, conforme al artículo 73 de "
        "la LOPCYMAT y la Norma Técnica NT-01-2008.")
    add_article(doc, 2, "Alcance",
        "Aplica a todo el personal de LA EMPRESA, sea fijo, contratado, eventual, de "
        "prueba o pasante, así como a visitantes, proveedores y contratistas presentes en "
        "el centro de trabajo.")

    # ===== CAPÍTULO II — DEFINICIONES =====
    add_chapter(doc, "II", "DEFINICIONES")
    add_bullet(doc, "Suceso acaecido en el curso del trabajo que NO causó lesión pero que pudo haberla causado.", bold_lead="Casi-accidente (near miss):  ")
    add_bullet(doc, "Suceso acaecido en el curso del trabajo que causó lesión leve sin incapacidad.", bold_lead="Incidente:  ")
    add_bullet(doc, "Suceso acaecido en el curso del trabajo que causó lesión con incapacidad temporal o permanente, o muerte.", bold_lead="Accidente de trabajo:  ")
    add_bullet(doc, "Accidente ocurrido en el trayecto entre el domicilio del trabajador(a) y el centro de trabajo, y viceversa (Art. 69 LOPCYMAT).", bold_lead="Accidente in itinere:  ")
    add_bullet(doc, "Enfermedad contraída por exposición a factores de riesgo en el trabajo (Art. 70 LOPCYMAT y Lista Oficial de Enfermedades Ocupacionales).", bold_lead="Enfermedad ocupacional:  ")
    add_bullet(doc, "Accidente que causa la muerte del trabajador(a), o lesiones que ponen en peligro la vida, o amputación, o pérdida de un sentido, o invalidez permanente (Art. 73 LOPCYMAT).", bold_lead="Accidente grave:  ")
    add_article(doc, "2b", "Lista de enfermedades ocupacionales relevantes",
        "Conforme al artículo 70 de la LOPCYMAT y la Lista Oficial de Enfermedades "
        "Ocupacionales, en LA EMPRESA se vigila especialmente:")
    enf_headers = ["CIE-10", "Enfermedad", "Causa / Agente", "Cargo expuesto"]
    enf_rows = [
        ["A82", "Rabia", "Virus rábico (mordedura)", "MV, Aux Vet, Groomer"],
        ["A27", "Leptospirosis", "Leptospira (orina animales)", "MV, Aux Vet, Groomer"],
        ["B58", "Toxoplasmosis", "Toxoplasma gondii (felinos)", "MV, Aux Vet"],
        ["A23", "Brucelosis", "Brucella (fluidos, abortos)", "MV, Aux Vet"],
        ["B35", "Dermatofitosis (tiña)", "Hongos dermatofitos", "MV, Aux Vet, Groomer"],
        ["B86", "Sarna sarcóptica", "Sarcoptes scabiei", "MV, Aux Vet, Groomer"],
        ["L24", "Dermatitis irritativa de contacto", "Desinfectantes, champús", "Aux Vet, Groomer, Encargado Tienda"],
        ["J45", "Asma ocupacional", "Aerosoles, pelos, caspa", "MV, Aux Vet, Groomer"],
        ["M75", "Trastornos musculoesqueléticos", "Posturas forzadas, movimientos repetitivos", "Todos"],
        ["Z57", "Exposición a radiaciones", "Rayos X veterinarios", "MV"],
        ["F43", "Trastornos por estrés laboral", "Carga mental, guardias", "MV, Gerente"],
        ["A36", "Tétanos", "Clostridium tetani (heridas)", "Todos"],
    ]
    add_data_table(doc, enf_headers, enf_rows,
                   col_widths_cm=[1.4, 4.0, 6.5, 5.3], header_size=8, body_size=8, center_cols=[0])

    # ===== CAPÍTULO III — OBLIGATORIEDAD =====
    add_chapter(doc, "III", "OBLIGATORIEDAD DEL REPORTE")
    add_alert_box(doc, [
        "**OBLIGACIÓN LEGAL — ARTÍCULO 73 LOPCYMAT",
        f"Todo accidente de trabajo debe ser reportado formalmente por LA EMPRESA ante el "
        f"IVSS y ante INPSASEL dentro de las DOS (2) HORAS siguientes al evento cuando se "
        f"trate de accidente grave; y dentro del primer día hábil siguiente para los "
        f"demás accidentes. El incumplimiento genera responsabilidad administrativa, "
        f"civil y penal para LA EMPRESA y para los directivos responsables."
    ])
    add_article(doc, "2c", "Responsabilidades específicas",
        "Las responsabilidades en materia de reporte se distribuyen así:")
    resp_headers = ["Actor", "Responsabilidad principal"]
    resp_rows = [
        ["Trabajador(a)", "Reportar inmediatamente cualquier incidente o condición insegura a su supervisor."],
        ["Supervisor directo", "Recibir el reporte, completar el formulario y notificar a la Gerencia y al Comité SST."],
        ["Médico Veterinario (según caso)", "Brindar atención primaria, evaluar gravedad, indicar traslado."],
        ["Comité SST", "Investigar causa raíz, elaborar plan de acción, dar seguimiento al cierre."],
        ["Directora Gerente", "Reportar formalmente al IVSS e INPSASEL dentro de los plazos legales; firmar reportes."],
        ["Representante Legal", "Firmar reportes oficiales y asumir representación legal ante autoridades."],
        ["Delegado(s) de Prevención", "Participar en la investigación, velar por derechos del trabajador(a)."],
    ]
    add_data_table(doc, resp_headers, resp_rows,
                   col_widths_cm=[4.5, 12.8], header_size=9, body_size=9)

    # ===== CAPÍTULO IV — CANALES DE REPORTE =====
    add_chapter(doc, "IV", "CANALES DE REPORTE")
    add_article(doc, 3, "Canales disponibles",
        "El trabajador(a) puede reportar el incidente a través de los siguientes canales, "
        "en orden de preferencia:")
    add_bullet(doc, "Verbal inmediata + escrita en formulario (Anexo C) ante el supervisor directo.", bold_lead="Canal 1:  ")
    add_bullet(doc, "Comité de Seguridad y Salud Laboral (SST).", bold_lead="Canal 2:  ")
    add_bullet(doc, "Directora Gerente o Representante Legal.", bold_lead="Canal 3:  ")
    add_bullet(doc, "Buzón físico de denuncia ubicado en recepción.", bold_lead="Canal 4:  ")
    add_article(doc, 4, "Categorías de incidentes",
        "Los incidentes se clasifican en las siguientes categorías para fines de reporte:")
    add_data_table(doc,
        ["#", "Categoría", "Ejemplos"],
        [
            ["1", "Accidente leve", "Arañazo sin incapacidad, pequeño corte sin sutura"],
            ["2", "Accidente con incapacidad temporal", "Mordedura que requiere reposo 3 días"],
            ["3", "Accidente grave", "Mordedura en cara/mano con riesgo funcional"],
            ["4", "Accidente mortal", "Fallecimiento del trabajador(a)"],
            ["5", "Casi-accidente", "Casi-mordedura, caída sin lesión"],
            ["6", "Enfermedad ocupacional", "Dermatitis por químicos, leptospirosis"],
            ["7", "Accidente in itinere", "Caída en ruta al trabajo"],
            ["8", "Incidente material / daño", "Daño a equipo, derrame químico"],
        ],
        col_widths_cm=[0.8, 4.5, 12.0], header_size=9, body_size=9, center_cols=[0])

    # ===== CAPÍTULO V — PROCEDIMIENTO DE REPORTE =====
    add_chapter(doc, "V", "PROCEDIMIENTO DE REPORTE — 9 PASOS")
    add_article(doc, 5, "Pasos",
        "El procedimiento de reporte se compone de nueve (9) pasos obligatorios:")
    add_bullet(doc, "Atender inmediatamente al trabajador(a): primeros auxilios, traslado a centro asistencial del IVSS-PMSSO si es necesario.", bold_lead="PASO 1:  ")
    add_bullet(doc, "Notificar verbalmente al supervisor directo en un máximo de 30 minutos.", bold_lead="PASO 2:  ")
    add_bullet(doc, "Completar el Formulario de Reporte de Incidente (Anexo C) dentro de las 2 horas.", bold_lead="PASO 3:  ")
    add_bullet(doc, "Reportar formalmente al IVSS e INPSASEL (accidente grave: 2 horas; otros: 1 día hábil).", bold_lead="PASO 4:  ")
    add_bullet(doc, "Preservar la escena del accidente para la investigación (fotos, testigos).", bold_lead="PASO 5:  ")
    add_bullet(doc, "Realizar investigación de causa raíz (5 por qué, análisis de barreras).", bold_lead="PASO 6:  ")
    add_bullet(doc, "Elaborar plan de acción correctivo y preventivo con responsables y plazos.", bold_lead="PASO 7:  ")
    add_bullet(doc, "Hacer seguimiento al plan de acción hasta su cierre.", bold_lead="PASO 8:  ")
    add_bullet(doc, "Cerrar el caso y archivar el expediente por 10 años (Art. 183 LOTTT).", bold_lead="PASO 9:  ")

    # ===== CAPÍTULO VI — PLAZOS LEGALES DE REPORTE =====
    add_chapter(doc, "VI", "PLAZOS LEGALES DE REPORTE")
    add_article(doc, 6, "Tabla de plazos",
        "Los plazos máximos de reporte conforme a la LOPCYMAT y la LOTTT son:")
    plazos_headers = ["Tipo de evento", "Ante quién", "Plazo máximo"]
    plazos_rows = [
        ["Accidente grave (Art. 73 LOPCYMAT)", "IVSS + INPSASEL", "2 HORAS"],
        ["Accidente común con incapacidad", "IVSS + INPSASEL", "1 día hábil"],
        ["Accidente in itinere", "IVSS", "1 día hábil"],
        ["Enfermedad ocupacional", "INPSASEL", "5 días hábiles"],
        ["Accidente mortal", "IVSS + INPSASEL + CICPC", "Inmediato"],
        ["Incidente (casi-accidente)", "Registro interno + Comité SST", "1 día hábil"],
    ]
    add_data_table(doc, plazos_headers, plazos_rows,
                   col_widths_cm=[6.5, 5.2, 5.6], header_size=9, body_size=9, center_cols=[2])

    # ===== CAPÍTULO VII — COMITÉ SST =====
    add_chapter(doc, "VII", "COMITÉ DE SEGURIDAD Y SALUD LABORAL")
    add_article(doc, 7, "Integración",
        "Conforme al artículo 46 de la LOPCYMAT, LA EMPRESA mantiene un Comité de "
        "Seguridad y Salud Laboral paritario, integrado por representantes del empleador "
        "y de los trabajadores, elegidos conforme al Reglamento de la LOPCYMAT. El Comité "
        "se reúne mensualmente y participa en la investigación de todo accidente grave.")
    add_article(doc, 8, "Funciones del Comité",
        "Son funciones del Comité SST: vigilar el cumplimiento del Programa SST, "
        "investigar accidentes y enfermedades ocupacionales, promover la capacitación, "
        "evaluar el Programa de Vigilancia Médica y recabar información sobre condiciones "
        "inseguras.")

    # ===== CAPÍTULO VIII — CONFIDENCIALIDAD Y NO REPRESALIAS =====
    add_chapter(doc, "VIII", "CONFIDENCIALIDAD Y PROHIBICIÓN DE REPRESALIAS")
    add_article(doc, 9, "Confidencialidad",
        "Los datos personales del trabajador(a) accidentado se tratan conforme a la LOPDP "
        "y se comparten exclusivamente con fines de prevención, investigación, atención "
        "médica y reporte legal. El expediente se conserva por 10 años.")
    add_article(doc, 10, "Prohibición de represalias",
        "Queda terminantemente prohibida cualquier represalia contra el trabajador(a) que "
        "reporte un incidente, accidente o condición insegura, conforme al artículo 26 de "
        "la LOPCYMAT. El reporte de buena fe constituye un derecho y un deber.")

    # ===== CAPÍTULO IX — INDICADORES, ESTADÍSTICAS Y MEJORA CONTINUA =====
    add_chapter(doc, "IX", "INDICADORES Y MEJORA CONTINUA")
    add_article(doc, "10b", "Indicadores de gestión SST",
        "LA EMPRESA calcula y monitorea los siguientes indicadores mensualmente:")
    ind_headers = ["Indicador", "Fórmula", "Meta"]
    ind_rows = [
        ["Índice de Frecuencia (IF)", "N° accidentes con incapacidad × 200.000 / horas trabajadas", "< 5"],
        ["Índice de Severidad (IS)", "Días perdidos × 200.000 / horas trabajadas", "< 100"],
        ["Tasa de incidencia", "N° accidentes × 100 / N° trabajadores", "< 10%"],
        ["Tasa de reporte de casi-accidentes", "N° casi-accidentes / N° accidentes", "> 10 (cultura proactiva)"],
        ["Cumplimiento de capacitación SST", "Horas capacitación ejecutadas / planificadas × 100", "≥ 95%"],
        ["Cierre de planes de acción", "Acciones cerradas / acciones planificadas × 100", "≥ 90%"],
        ["Ausentismo por salud", "Días perdidos por salud / días laborables × 100", "< 3%"],
    ]
    add_data_table(doc, ind_headers, ind_rows,
                   col_widths_cm=[5.5, 7.5, 4.3], header_size=8, body_size=8)
    add_article(doc, "10c", "Revisión del Programa SST",
        "El Programa de Seguridad y Salud en el Trabajo se revisa anualmente, conforme "
        "al artículo 89 de la LOPCYMAT, con base en los indicadores calculados, los "
        "reportes de incidentes y accidentes, las auditorías internas y externas, y las "
        "recomendaciones del Comité SST. La revisión da lugar a un Plan Anual de "
        "Prevención con objetivos, metas, responsables y plazos.")
    add_article(doc, "10d", "Capacitación continua",
        "Todo el personal recibe capacitación anual mínima de dieciséis (16) horas en "
        "materia de SST, conforme al artículo 56 de la LOPCYMAT. Las capacitaciones "
        "incluyen: bioseguridad veterinaria, manejo de sustancias controladas, "
        "protocolo de mordeduras, RCP, prevención de incendios, manejo de extintores, "
        "y revisión del presente Procedimiento de Reporte de Incidentes.")
    add_article(doc, "10e", "Auditorías externas",
        "LA EMPRESA se somete a auditorías externas periódicas por parte del INPSASEL, "
        "del IVSS y, cuando aplique, del SENAC. Los hallazgos se documentan y dan lugar "
        "a planes de acción con responsables y plazos, cuyo seguimiento hace el Comité SST.")

    # ===== ANEXO C: Formulario de reporte de incidente =====
    add_section(doc, "ANEXO C — FORMULARIO DE REPORTE DE INCIDENTE")
    add_para(doc, "Este formulario se completa por duplicado: copia para expediente SST y copia para el trabajador(a).",
        size=9, italic=True, space_after=4)
    add_para(doc, "SECCIÓN 1 — DATOS DEL TRABAJADOR(A)", size=9, bold=True, color=TEAL_DARK, space_after=2)
    add_data_table(doc,
        ["Campo", "Información"],
        [
            ["Nombre completo", "_______________________________________________"],
            ["C.I. / Cargo / Departamento", "_______________________________________________"],
            ["Antigüedad / Fecha de ingreso", "_______________________________________________"],
            ["Edad / Sexo", "_______________________________________________"],
        ],
        col_widths_cm=[5.0, 12.3], header_size=9, body_size=9)
    add_para(doc, "SECCIÓN 2 — DATOS DEL INCIDENTE", size=9, bold=True, color=TEAL_DARK, space_after=2)
    add_data_table(doc,
        ["Campo", "Información"],
        [
            ["Fecha y hora del incidente", "____ / ____ / ______  ·  ____:____"],
            ["Lugar exacto del incidente", "_______________________________________________"],
            ["Tipo de incidente (leve, grave, mortal, in itinere, enfermedad)", "_______________________________________________"],
            ["Categoría (ver Cap. IV)", "_______________________________________________"],
            ["Testigos (nombres y C.I.)", "_______________________________________________"],
        ],
        col_widths_cm=[5.0, 12.3], header_size=9, body_size=9)
    add_para(doc, "SECCIÓN 3 — DESCRIPCIÓN DEL INCIDENTE", size=9, bold=True, color=TEAL_DARK, space_after=2)
    add_data_table(doc,
        ["Campo", "Información"],
        [
            ["Descripción detallada", "_______________________________________________"],
            ["Tarea que realizaba el trabajador(a)", "_______________________________________________"],
            ["Equipo / herramienta / sustancia involucrada", "_______________________________________________"],
            ["Factores contribuyentes (condiciones inseguras)", "_______________________________________________"],
            ["Actos inseguros detectados", "_______________________________________________"],
        ],
        col_widths_cm=[5.0, 12.3], header_size=9, body_size=9)
    add_para(doc, "SECCIÓN 4 — LESIONES Y ATENCIÓN MÉDICA", size=9, bold=True, color=TEAL_DARK, space_after=2)
    add_data_table(doc,
        ["Campo", "Información"],
        [
            ["Tipo de lesión / parte del cuerpo", "_______________________________________________"],
            ["Gravedad", "_______________________________________________"],
            ["Centro asistencial que atendió", "_______________________________________________"],
            ["Tratamiento / reposo indicado (días)", "_______________________________________________"],
            ["Reportado al IVSS / INPSASEL (fecha y hora)", "_______________________________________________"],
        ],
        col_widths_cm=[5.0, 12.3], header_size=9, body_size=9)
    add_para(doc, "SECCIÓN 5 — INVESTIGACIÓN DE CAUSA RAÍZ", size=9, bold=True, color=TEAL_DARK, space_after=2)
    add_data_table(doc,
        ["Campo", "Información"],
        [
            ["Investigador designado", "_______________________________________________"],
            ["Causa inmediata", "_______________________________________________"],
            ["Causa básica (organizacional / humano)", "_______________________________________________"],
            ["Causa raíz (5 por qué)", "_______________________________________________"],
            ["Medidas correctivas propuestas", "_______________________________________________"],
            ["Responsables / Plazos", "_______________________________________________"],
        ],
        col_widths_cm=[5.0, 12.3], header_size=9, body_size=9)
    add_para(doc, "SECCIÓN 6 — FIRMAS", size=9, bold=True, color=TEAL_DARK, space_after=2)
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "TESTIGO", "DIRECTORA GERENTE"])

    # ===== Carta de recepción =====
    add_reception_letter(doc, "Procedimiento de Reporte de Incidentes y Accidentes de Trabajo")

    add_footer(doc.sections[0], FOOTER_BASE)
    out_path = os.path.join(OUT_DIR_08, "Procedimiento_Reporte_Incidentes.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# Main
# ============================================================
def main():
    print("=" * 64)
    print("gen_protocolos_vet.py — Generando 4 protocolos veterinarios ALIKA PETS")
    print("=" * 64)
    paths = []
    for fn in (gen_cartilla_bioseguridad, gen_sustancias_controladas,
               gen_mordeduras_zoonosis, gen_reporte_incidentes):
        try:
            p = fn()
            paths.append(p)
            print(f"  ✓ {os.path.basename(p)}  →  {os.path.getsize(p)//1024} KB")
        except Exception as e:
            print(f"  ✗ Error en {fn.__name__}: {e}")
            raise
    print("=" * 64)
    print(f"Total: {len(paths)} documentos generados")
    return paths


if __name__ == "__main__":
    main()
