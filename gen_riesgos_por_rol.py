"""
gen_riesgos_por_rol.py — Genera 5 Notificaciones de Riesgos por cargo ALIKA PETS.

Documentos generados en /home/z/my-project/output/05_SEGURIDAD_LABORAL/:
  1. Notificacion_Riesgos_Gerente.docx
  2. Notificacion_Riesgos_Encargado_Tienda.docx
  3. Notificacion_Riesgos_Medico_Veterinario.docx   (10 riesgos)
  4. Notificacion_Riesgos_Auxiliar_Veterinario.docx (9 riesgos)
  5. Notificacion_Riesgos_Dog_Groomer.docx          (10 riesgos)

Metodología NT-01-2008 (Prob × Cons):
  Probabilidad: Baja(1) / Media(2) / Alta(3)
  Consecuencia: Ligeramente dañino(1) / Dañino(2) / Muy dañino(3)
  Nivel:  1=T (Trivial) · 2=TO (Tolerable) · 3,4=MO (Moderado) · 6=IM (Importante) · 9=IN (Intolerable)
Colores: T/TO → verde · MO → ámbar · IM → naranja · IN → rojo

Cada notificación contiene:
  - Encabezado (membrete + título)
  - Identificación del trabajador y cargo
  - Descripción del cargo y tareas
  - Matriz de identificación de peligros (8-10 riesgos)
  - Matriz visual 4×4 prob×cons con códigos coloreados
  - Tabla de EPP obligatorio
  - Medidas de prevención
  - Procedimiento ante emergencias
  - Declaración del trabajador conforme Art. 56 LOPCYMAT
  - Firma: Trabajador + DIRECTORA GERENTE (Esnatlim)

Versión 3.0 — _common.py mapea automáticamente "RR.HH." → "DIRECTORA GERENTE".
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _common import (
    setup_a4_portrait, add_membrete, add_doc_title, add_chapter, add_section,
    add_article, add_para, add_bullet, add_sanciones_table, add_signature_block,
    add_reception_letter, add_footer, add_hr,
    set_cell_bg, set_cell_borders, set_cell_margins, write_cell,
    TEAL_DARK, TEAL_HDR_BG, GRAY_ALT, AMBER_BG, SLATE_BG, WHITE, BLACK,
    GRAY_TEXT, GRAY_MUTED, RED_CRIT, GREEN_OK,
    EMPRESA, RIF_EMP, MARCA, DOMICILIO_EMP,
    REP_LEGAL_NOMBRE, REP_LEGAL_CARGO, REP_LEGAL_CI,
    DIRECTORA_NOMBRE, DIRECTORA_CARGO, DIRECTORA_CI,
)

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "05_SEGURIDAD_LABORAL")
os.makedirs(OUT_DIR, exist_ok=True)
FOOTER_BASE = "Notificación de Riesgos  ·  Art. 56 LOPCYMAT  ·  v3.0"

# ============================================================
# Constantes de la metodología NT-01-2008
# ============================================================
# Probabilidad
PROB = {"Baja": 1, "Media": 2, "Alta": 3}
# Consecuencia
CONS = {"Ligeramente dañino": 1, "Dañino": 2, "Muy dañino": 3}

# Códigos y colores por nivel
LEVEL_INFO = {
    "T":  {"code": "T",  "name": "TRIVIAL",     "bg": "DCFCE7", "color": RGBColor(0x04, 0x78, 0x57)},
    "TO": {"code": "TO", "name": "TOLERABLE",   "bg": "F0FDF4", "color": RGBColor(0x04, 0x78, 0x57)},
    "MO": {"code": "MO", "name": "MODERADO",    "bg": "FEF3C7", "color": RGBColor(0xB4, 0x53, 0x09)},
    "IM": {"code": "IM", "name": "IMPORTANTE",  "bg": "FED7AA", "color": RGBColor(0xC2, 0x41, 0x0C)},
    "IN": {"code": "IN", "name": "INTOLERABLE", "bg": "FEE2E2", "color": RGBColor(0xB9, 0x1C, 0x1C)},
}

def calc_level(prob_val, cons_val):
    """Calcula el código de nivel (T/TO/MO/IM/IN) según prob×cons."""
    score = prob_val * cons_val
    if score <= 1:   # 1
        return "T"
    elif score <= 4: # 2,3,4 → TOLERABLE 2-3, MODERADO 4
        if score == 2:
            return "TO"
        elif score == 3:
            return "MO"
        else:  # 4
            return "MO"
    elif score == 6:
        return "IM"
    elif score == 9:
        return "IN"
    else:
        return "TO"


# ============================================================
# Helpers
# ============================================================
def add_data_table(doc, headers, rows, col_widths_cm=None, header_size=9, body_size=9,
                   center_cols=None, level_col_idx=None):
    """Tabla con encabezado teal y filas alternas. Si level_col_idx se especifica,
    colorea esa columna según el código T/TO/MO/IM/IN."""
    center_cols = center_cols or []
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            tbl.columns[i].width = Cm(w)
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        if col_widths_cm:
            c.width = Cm(col_widths_cm[ci])
        write_cell(c, h, size=header_size, bold=True,
                   color=RGBColor(0xFF, 0xFF, 0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    for ri, row in enumerate(rows, start=1):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            if col_widths_cm:
                c.width = Cm(col_widths_cm[ci])
            # Color especial para columna de nivel
            if level_col_idx is not None and ci == level_col_idx and isinstance(val, str) and val in LEVEL_INFO:
                info = LEVEL_INFO[val]
                display = f"{info['code']}  {info['name']}"
                write_cell(c, display, size=body_size, bold=True, color=info["color"],
                           align=WD_ALIGN_PARAGRAPH.CENTER, bg=info["bg"])
            else:
                align = WD_ALIGN_PARAGRAPH.CENTER if ci in center_cols else (
                    WD_ALIGN_PARAGRAPH.LEFT if ci > 0 else WD_ALIGN_PARAGRAPH.CENTER)
                write_cell(c, str(val), size=body_size, color=BLACK, bg=bg, align=align)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl


def add_risk_matrix(doc):
    """Construye la matriz visual 4×4 (header + 3 filas × header + 3 columnas).
    Columnas = Consecuencia (LD/D/MD), Filas = Probabilidad (Alta/Media/Baja)."""
    add_section(doc, "MATRIZ VISUAL DE EVALUACIÓN DE RIESGOS (NT-01-2008)")
    add_para(doc,
        "Probabilidad: Baja(1) / Media(2) / Alta(3).  "
        "Consecuencia: Ligeramente dañino(1) / Dañino(2) / Muy dañino(3).  "
        "Nivel = Probabilidad × Consecuencia.",
        size=9, italic=True, space_after=4)

    # Tabla 4x4 (header + 3 columnas/rows)
    tbl = doc.add_table(rows=4, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cw = [4.5, 4.0, 4.0, 4.0]
    for i, w in enumerate(cw):
        tbl.columns[i].width = Cm(w)

    # Esquina superior izquierda
    write_cell(tbl.rows[0].cells[0], "Probabilidad ↓ / Consecuencia →",
               size=8, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
    tbl.rows[0].cells[0].width = Cm(cw[0])
    # Encabezado de columnas (consecuencias)
    cons_list = [("Ligeramente dañino (1)", 1),
                 ("Dañino (2)", 2),
                 ("Muy dañino (3)", 3)]
    for ci, (label, _) in enumerate(cons_list, start=1):
        write_cell(tbl.rows[0].cells[ci], label,
                   size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=TEAL_HDR_BG)
        tbl.rows[0].cells[ci].width = Cm(cw[ci])

    # Filas: Alta (3), Media (2), Baja (1) — de arriba a abajo
    prob_list = [("Alta (3)", 3), ("Media (2)", 2), ("Baja (1)", 1)]
    for ri, (plabel, pval) in enumerate(prob_list, start=1):
        # Etiqueta de fila (probabilidad)
        write_cell(tbl.rows[ri].cells[0], plabel,
                   size=9, bold=True, color=BLACK,
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=SLATE_BG)
        tbl.rows[ri].cells[0].width = Cm(cw[0])
        # Celdas de nivel
        for ci, (_, cval) in enumerate(cons_list, start=1):
            level_code = calc_level(pval, cval)
            info = LEVEL_INFO[level_code]
            txt = f"{info['code']}\n{info['name']}\n({pval}×{cval}={pval*cval})"
            write_cell(tbl.rows[ri].cells[ci], txt,
                       size=8, bold=True, color=info["color"],
                       align=WD_ALIGN_PARAGRAPH.CENTER, bg=info["bg"])
            tbl.rows[ri].cells[ci].width = Cm(cw[ci])
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

    # Leyenda de niveles
    add_para(doc, "Leyenda de niveles:", size=9, bold=True, space_after=2)
    legend_tbl = doc.add_table(rows=1, cols=5)
    legend_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    legend_tbl.autofit = False
    for i in range(5):
        legend_tbl.columns[i].width = Cm(3.3)
    legend_codes = ["T", "TO", "MO", "IM", "IN"]
    for i, code in enumerate(legend_codes):
        info = LEVEL_INFO[code]
        c = legend_tbl.rows[0].cells[i]
        write_cell(c, f"{info['code']}  ·  {info['name']}",
                   size=8, bold=True, color=info["color"],
                   align=WD_ALIGN_PARAGRAPH.CENTER, bg=info["bg"])
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def add_epp_table(doc, epp_rows):
    """Tabla de EPP obligatorio: prendas + uso."""
    add_section(doc, "EQUIPOS DE PROTECCIÓN PERSONAL (EPP) OBLIGATORIO")
    add_data_table(doc,
        ["EPP", "Uso / Procedimiento"],
        epp_rows,
        col_widths_cm=[5.0, 12.3], header_size=9, body_size=9)


def add_worker_identification(doc, cargo):
    """Tabla de identificación del trabajador."""
    add_section(doc, "IDENTIFICACIÓN DEL TRABAJADOR(A)")
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.0)
    tbl.columns[1].width = Cm(13.3)
    labels = [
        ("Nombre y Apellido:", "_______________________________________________"),
        ("Cédula de Identidad:", "V-____________________________________________"),
        ("Cargo:", cargo),
        ("Fecha de notificación:", "____ / ____ / ________"),
    ]
    for i, (lab, val) in enumerate(labels):
        write_cell(tbl.rows[i].cells[0], lab, size=9, bold=True, color=TEAL_DARK,
                   align=WD_ALIGN_PARAGRAPH.LEFT, bg=SLATE_BG)
        write_cell(tbl.rows[i].cells[1], val, size=9, color=BLACK,
                   align=WD_ALIGN_PARAGRAPH.LEFT, bg=WHITE)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def add_declaration_block(doc):
    """Declaración del trabajador conforme al Art. 56 LOPCYMAT."""
    add_section(doc, "DECLARACIÓN DEL TRABAJADOR(A) — Art. 56 LOPCYMAT")
    add_para(doc,
        "Yo, _____________________________________________, titular de la cédula de "
        "identidad N° V-___________________, en mi condición de trabajador(a) de "
        f"{EMPRESA} (marca {MARCA}), declaro haber sido notificado(a) por la empresa "
        "sobre los riesgos laborales asociados a mi cargo, así como sobre las medidas "
        "de prevención, control y protección establecidas, conforme al artículo 56 de la "
        "Ley Orgánica de Prevención, Condiciones y Medio Ambiente de Trabajo (LOPCYMAT).",
        size=10, space_after=4)
    add_bullet(doc, "He recibido información clara y suficiente sobre los riesgos a los que estaré expuesto(a) en el desempeño de mis funciones.")
    add_bullet(doc, "Conozco y me comprometo a utilizar correctamente los Equipos de Protección Personal (EPP) entregados por la empresa.")
    add_bullet(doc, "Conozco los procedimientos de emergencia y de reporte de incidentes establecidos.")
    add_bullet(doc, "Me comprometo a participar en las actividades de capacitación y en el Programa de Seguridad y Salud Laboral.")
    add_bullet(doc, "Entiendo que tengo derecho a ser reubicado(a) en caso de embarazo, lactancia o condición de salud que así lo requiera (Art. 78 LOPCYMAT).")
    add_para(doc,
        "En fe de lo cual firmo la presente notificación en la ciudad de Los Teques, "
        "Estado Miranda, a los ____ días del mes de ________________ de ________.",
        size=10, space_before=4, space_after=8)
    # Firma: Trabajador + Directora Gerente (mapeo automático)
    add_signature_block(doc, ["EL(LA) TRABAJADOR(A)", "DIRECTORA GERENTE"])


def add_emergency_procedure(doc, specific_items):
    """Procedimiento ante emergencias."""
    add_section(doc, "PROCEDIMIENTO ANTE EMERGENCIAS")
    add_para(doc,
        "Ante cualquier accidente de trabajo o emergencia, el trabajador(a) debe "
        "aplicar el siguiente procedimiento:",
        size=10, space_after=2)
    add_bullet(doc, "Prestar primeros auxilios inmediatos según el tipo de lesión.", bold_lead="1.  ")
    add_bullet(doc, "Notificar al supervisor directo en un máximo de 30 minutos.", bold_lead="2.  ")
    add_bullet(doc, "Reportar formalmente en el Formulario de Incidente en un máximo de 2 horas (Art. 73 LOPCYMAT).", bold_lead="3.  ")
    add_bullet(doc, "Acudir al centro asistencial del IVSS-PMSSO designado.", bold_lead="4.  ")
    add_bullet(doc, "Completar el reporte al IVSS/INPSASEL según el plazo legal.", bold_lead="5.  ")
    add_para(doc, "Acciones específicas para este cargo:", size=9, bold=True, space_after=2)
    for item in specific_items:
        add_bullet(doc, item)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def build_notification(doc, *, cargo, descripción, tareas, riesgos, epp_rows,
                       medidas, emergencias, footer_label):
    """Construye una Notificación de Riesgos con todos los elementos comunes.
    `riesgos` es una lista de tuplas: (puesto_riesgo, prob_str, cons_str, código_nivel)
    """
    setup_a4_portrait(doc, margins_cm=2.0)
    add_membrete(doc, "NOTIFICACIÓN DE RIESGOS", "Art. 56 LOPCYMAT", "v3.0  ·  SST")
    add_doc_title(doc, f"NOTIFICACIÓN DE RIESGOS — {cargo.upper()}")

    # Introducción legal
    add_para(doc,
        f"Conforme al artículo 56 de la LOPCYMAT, {EMPRESA} (marca {MARCA}) notifica "
        f"formalmente al trabajador(a) los riesgos laborales asociados al cargo de "
        f"«{cargo}», así como las medidas de prevención, control y atención en caso de "
        f"accidente de trabajo o enfermedad ocupacional.",
        size=10, space_after=6)

    # Identificación del trabajador
    add_worker_identification(doc, cargo)

    # Descripción del cargo y tareas
    add_chapter(doc, "I", "DESCRIPCIÓN DEL CARGO Y TAREAS")
    add_article(doc, 1, "Descripción del cargo",
        descripción)
    add_article(doc, 2, "Tareas principales",
        "Las tareas principales del cargo incluyen:")
    for t in tareas:
        add_bullet(doc, t)

    # Matriz de identificación de peligros
    add_chapter(doc, "II", "MATRIZ DE IDENTIFICACIÓN DE PELIGROS Y EVALUACIÓN")
    add_para(doc,
        "Se identifican los siguientes peligros asociados al cargo y se evalúa el nivel "
        "de riesgo según la metodología NT-01-2008:",
        size=10, space_after=4)
    headers = ["#", "Peligro / Riesgo", "Probabilidad", "Consecuencia", "Nivel"]
    rows = []
    for i, (peligro, prob, cons, code) in enumerate(riesgos, start=1):
        rows.append([str(i), peligro, prob, cons, code])
    add_data_table(doc, headers, rows,
                   col_widths_cm=[0.8, 8.2, 2.5, 2.5, 3.3],
                   header_size=9, body_size=9,
                   center_cols=[0, 2, 3], level_col_idx=4)

    # Matriz visual 4×4
    add_chapter(doc, "III", "MATRIZ VISUAL DE EVALUACIÓN")
    add_risk_matrix(doc)

    # EPP obligatorio
    add_chapter(doc, "IV", "EQUIPOS DE PROTECCIÓN PERSONAL")
    add_epp_table(doc, epp_rows)

    # Medidas de prevención
    add_chapter(doc, "V", "MEDIDAS DE PREVENCIÓN Y CONTROL")
    for m in medidas:
        add_bullet(doc, m)

    # Procedimiento ante emergencias
    add_chapter(doc, "VI", "PROCEDIMIENTO ANTE EMERGENCIAS")
    add_emergency_procedure(doc, emergencias)

    # Declaración del trabajador + firma
    add_chapter(doc, "VII", "DECLARACIÓN DEL TRABAJADOR(A) Y FIRMA")
    add_declaration_block(doc)

    # Carta de recepción
    add_reception_letter(doc, f"Notificación de Riesgos — {cargo}")

    # Footer
    add_footer(doc.sections[0], footer_label)


# ============================================================
# RIESGOS POR CARGO — Datos
# ============================================================

# --- GERENTE ---
RIESGOS_GERENTE = [
    ("Estrés laboral y carga mental", "Media", "Dañino", "MO"),
    ("Trastornos musculoesqueléticos (postura sedente)", "Baja", "Ligeramente dañino", "T"),
    ("Fatiga visual (pantallas)", "Media", "Ligeramente dañino", "TO"),
    ("Sobreesfuerzo al movilizar carga ocasional", "Baja", "Dañino", "TO"),
    ("Riesgo eléctrico (equipos oficina)", "Baja", "Dañino", "TO"),
    ("Riesgo ergonómico por movilidad entre áreas", "Media", "Ligeramente dañino", "TO"),
    ("Exposición eventual a fluidos/biológicos (visita áreas)", "Baja", "Dañino", "TO"),
    ("Riesgo psicosocial por conflicto interpersonal", "Media", "Dañino", "MO"),
]

# --- ENCARGADO DE TIENDA ---
RIESGOS_ENC_TIENDA = [
    ("Sobreesfuerzo por movilización de cargas (sacos alimento, jaulas)", "Alta", "Dañino", "IM"),
    ("Trastornos musculoesqueléticos (bipedación prolongada)", "Alta", "Dañino", "IM"),
    ("Caídas al mismo nivel (pisos mojados, mercancía)", "Media", "Dañino", "MO"),
    ("Riesgo eléctrico (iluminación, refrigeradores)", "Baja", "Dañino", "TO"),
    ("Exposición a productos químicos (limpiadores, desinfectantes)", "Media", "Ligeramente dañino", "TO"),
    ("Riesgo psicosocial por atención al público / asaltos", "Media", "Muy dañino", "IM"),
    ("Mordeduras / arañazos por animales pequeños en venta", "Baja", "Dañino", "TO"),
    ("Fatiga visual y postural por cajero / POS", "Alta", "Ligeramente dañino", "MO"),
    ("Cortes con cúter / apertura de cajas", "Media", "Ligeramente dañino", "TO"),
]

# --- MÉDICO VETERINARIO ---
RIESGOS_VET = [
    ("Mordeduras / arañazos por pacientes", "Alta", "Muy dañino", "IN"),
    ("Exposición a fluidos biológicos (sangre, secreciones)", "Alta", "Dañino", "IM"),
    ("Zoonosis (rabia, leptospirosis, toxoplasmosis, brucelosis)", "Media", "Muy dañino", "IM"),
    ("Accidentes cortopunzantes (agujas, bisturíes)", "Alta", "Dañino", "IM"),
    ("Exposición a anestésicos volátiles (isoflurano, sevoflurano)", "Media", "Dañino", "MO"),
    ("Radiaciones ionizantes (rayos X veterinarios)", "Media", "Muy dañino", "IM"),
    ("Manejo de sustancias controladas (opioides, ketamina)", "Media", "Muy dañino", "IM"),
    ("Trastornos musculoesqueléticos (posturas forzadas en cirugía)", "Alta", "Dañino", "IM"),
    ("Estrés laboral por emergencias y guardias", "Media", "Dañino", "MO"),
    ("Riesgo eléctrico (equipos médicos, lámparas quirúrgicas)", "Baja", "Muy dañino", "MO"),
]

# --- AUXILIAR VETERINARIO ---
RIESGOS_AUX_VET = [
    ("Mordeduras / arañazos durante sujeción y manejo de animales", "Alta", "Dañino", "IM"),
    ("Exposición a fluidos biológicos (limpieza de jaulas, curaciones)", "Alta", "Dañino", "IM"),
    ("Zoonosis (tiña, sarna, leptospirosis, salmonelosis)", "Media", "Dañino", "MO"),
    ("Accidentes cortopunzantes (agujas, vidrio, instrumental)", "Alta", "Dañino", "IM"),
    ("Exposición a desinfectantes (glutaraldehído, amonio cuaternario)", "Media", "Dañino", "MO"),
    ("Trastornos musculoesqueléticos (movilización pacientes, bipedación)", "Alta", "Dañino", "IM"),
    ("Caídas al mismo nivel (pisos mojados, fluidos derramados)", "Media", "Dañino", "MO"),
    ("Exposición a anestésicos (asistencia en quirófano)", "Baja", "Dañino", "TO"),
    ("Riesgo eléctrico (limpieza de equipos, áreas húmedas)", "Baja", "Muy dañino", "MO"),
]

# --- DOG GROOMER ---
RIESGOS_GROOMER = [
    ("Mordeduras / arañazos por mascotas en baño y peluquería", "Alta", "Dañino", "IM"),
    ("Zoonosis cutáneas (tiña, sarna, pulgas, garrapatas)", "Alta", "Dañino", "IM"),
    ("Exposición a champús, jabones y productos químicos", "Media", "Ligeramente dañino", "TO"),
    ("Cortes con tijeras, maquinillas y cuchillas", "Media", "Dañino", "MO"),
    ("Riesgo eléctrico en zona húmeda (secadores, bañeras eléctricas)", "Media", "Muy dañino", "IM"),
    ("Trastornos musculoesqueléticos (posturas forzadas, bipedación)", "Alta", "Dañino", "IM"),
    ("Exposición a ruido (secadores, maquinillas, ladridos)", "Alta", "Ligeramente dañino", "MO"),
    ("Asma / alergias respiratorias (pelos, caspa, aerosoles)", "Media", "Dañino", "MO"),
    ("Dermatitis de contacto (champús, desinfectantes)", "Media", "Dañino", "MO"),
    ("Caídas al mismo nivel (superficies mojadas, resbaladizas)", "Media", "Dañino", "MO"),
]


# ============================================================
# Generadores por cargo
# ============================================================
def gen_gerente():
    doc = Document()
    build_notification(
        doc,
        cargo="Gerente",
        descripción=(
            "El/La Gerente es responsable de la planificación, coordinación y supervisión "
            "general de las operaciones de LA EMPRESA (clínica veterinaria, tienda y "
            "peluquería canina), así como del cumplimiento de los objetivos comerciales, "
            "legales y administrativos. Realiza tareas administrativas, atención al "
            "cliente, manejo de personal, control de inventario y reportes contables."
        ),
        tareas=[
            "Supervisión del personal y cumplimiento de horarios.",
            "Gestión administrativa: nómina, contratos, declaraciones IVSS/FAOV/INCES/ISLR.",
            "Atención al cliente y resolución de quejas.",
            "Control de inventario y compras.",
            "Manejo de caja y conciliación diaria.",
            "Reportes contables y reunion con la Directora Gerente.",
            "Movilización entre áreas (clínica, tienda, peluquería, depósito).",
        ],
        riesgos=RIESGOS_GERENTE,
        epp_rows=[
            ("Calzado cerrado antideslizante", "Uso permanente en jornada laboral."),
            ("Guantes de nitrilo", "Para manejo ocasional de mercancía o productos químicos."),
            ("Mascarilla quirúrgica", "Visitas a áreas clínicas o de hospitalización."),
        ],
        medidas=[
            "Realizar pausas activas cada 90 minutos para reducir el sedentarismo y la fatiga visual.",
            "Mantener postura ergonómica frente al computador: monitor a la altura de los ojos, silla con respaldo lumbar.",
            "Capacitación en manejo del estrés, gestión del tiempo y resolución de conflictos.",
            "Cumplir el Programa de Vigilancia Médica (NT-02-2008): examen periódico anual.",
            "Mantener iluminación adecuada y evitar reflejos en pantallas.",
            "Revisar instalaciones eléctricas periódicamente (responsable de mantenimiento).",
            "Capacitación en prevención de asaltos y manejo de situaciones de estrés.",
        ],
        emergencias=[
            "En caso de asalto: NO oponer resistencia, colaborar y activar botón de pánico o llamar al 911 una vez fuera de peligro.",
            "En caso de malestar físico o estrés agudo: notificar a la Directora Gerente y acudir al centro asistencial.",
            "En caso de accidente en tránsito (in itinere): reportar en menos de 1 día hábil al IVSS.",
        ],
        footer_label=f"{FOOTER_BASE}  ·  Cargo: Gerente",
    )
    out_path = os.path.join(OUT_DIR, "Notificacion_Riesgos_Gerente.docx")
    doc.save(out_path)
    return out_path


def gen_encargado_tienda():
    doc = Document()
    build_notification(
        doc,
        cargo="Encargado de Tienda",
        descripción=(
            "El/La Encargado(a) de Tienda es responsable de la atención al cliente, "
            "venta de productos para mascotas, manejo de caja, reposición de inventario, "
            "limpieza del área de venta y supervisión de mercancía. Pasa la mayor parte "
            "de la jornada de pie, atendiendo al público y movilizando cargas medias."
        ),
        tareas=[
            "Atención al cliente, asesoría de productos y ventas.",
            "Operación de caja registradora / POS.",
            "Reposición de góndolas y exhibición de productos.",
            "Movilización de sacos de alimento (hasta 20 kg), jaulas y cajas.",
            "Limpieza del área de venta y baños.",
            "Control de inventario y reporte de mermas.",
            "Apertura y cierre de tienda.",
        ],
        riesgos=RIESGOS_ENC_TIENDA,
        epp_rows=[
            ("Calzado cerrado antideslizante", "OBLIGATORIO. Uso permanente."),
            ("Guantes de nitrilo / goma", "Para manipulación de productos químicos y limpieza."),
            ("Mascarilla quirúrgica", "Para limpieza con productos químicos fuertes."),
            ("Faja lumbar", "Para movilización de cargas mayores a 10 kg."),
            ("Botas de seguridad (opcional)", "Para descarga de mercancía pesada."),
        ],
        medidas=[
            "Levantamiento de cargas con técnica correcta: flexionar rodillas, espalda recta, carga pegada al cuerpo.",
            "No levantar cargas mayores a 25 kg sin ayuda. Usar carretilla para sacos de alimento.",
            "Rotación de tareas para evitar bipedación prolongada. Pausas activas cada 2 horas.",
            "Mantener pisos secos y limpios. Colocar señalización 'piso mojado' durante limpieza.",
            "Capacitación en atención al cliente y prevención de asaltos.",
            "No operar equipos eléctricos con manos mojadas.",
            "Ventilación adecuada al usar productos de limpieza.",
        ],
        emergencias=[
            "En caso de corte: lavar, aplicar antiséptico, cubrir y reportar.",
            "En caso de caída: evaluar movilidad, no intentar levantar al lesionado si hay dolor lumbar.",
            "En caso de asalto: colaborar, no oponer resistencia, reportar a la Gerencia y al CICPC.",
            "En caso de mareo por productos químicos: salir a zona ventilada y reportar.",
        ],
        footer_label=f"{FOOTER_BASE}  ·  Cargo: Encargado de Tienda",
    )
    out_path = os.path.join(OUT_DIR, "Notificacion_Riesgos_Encargado_Tienda.docx")
    doc.save(out_path)
    return out_path


def gen_medico_veterinario():
    doc = Document()
    build_notification(
        doc,
        cargo="Médico Veterinario",
        descripción=(
            "El/La Médico(a) Veterinario(a) es responsable de la atención clínica, "
            "diagnóstica, terapéutica y quirúrgica de los animales pacientes. Realiza "
            "consultas, cirugías, hospitalización, imagenología, manejo de sustancias "
            "controladas, atención de emergencias y guardias. Está expuesto a riesgos "
            "biológicos, físicos, químicos y psicosociales de elevada magnitud."
        ),
        tareas=[
            "Consulta clínica y diagnóstico de animales (perros, gatos, aves, exóticos).",
            "Procedimientos quirúrgicos (OVH, laparotomía, ortopedia, dental).",
            "Toma de muestras (sangre, heces, biopsias) y laboratorio.",
            "Aplicación de vacunas, anestésicos y tratamientos.",
            "Imagenología: radiografías y ecografías.",
            "Manejo y prescripción de sustancias controladas (opioides, ketamina).",
            "Hospitalización, monitoreo y atención de emergencias.",
            "Eutanasia y necropsias cuando aplique.",
            "Registro en historia clínica y libro foliado.",
            "Atención de guardias y emergencias fuera de horario.",
        ],
        riesgos=RIESGOS_VET,
        epp_rows=[
            ("Bata quirúrgica / blanca", "Uso permanente durante jornada."),
            ("Guantes de examen (nitrilo)", "OBLIGATORIO en todo contacto con animales y fluidos."),
            ("Guantes estériles", "Para procedimientos quirúrgicos y asépticos."),
            ("Mascarilla quirúrgica", "Para atención de pacientes y procedimientos."),
            ("Mascarilla N95", "Para necropsias, sospecha de zoonosis respiratoria, quimioterapia."),
            ("Gafas de protección / visera", "Para cirugías, odontología, necropsias, salpicaduras."),
            ("Gorro quirúrgico", "Para cirugías y procedimientos con riesgo de aerosoles."),
            ("Delantal plomado + dosímetro", "OBLIGATORIO en radiografías. Verificación trimestral."),
            ("Contenedor cortopunzantes", "Disponible en cada mesa de procedimiento."),
            ("Botas de seguridad", "Para necropsias y limpieza de jaulas."),
            ("Sistema de gas scavenging", "OBLIGATORIO para anestesia inhalatoria."),
        ],
        medidas=[
            "Esquema completo de vacunación: antirrábica pre-exposición (3 dosis + titulación cada 2 años), Hepatitis B, tétanos, influenza.",
            "Examen médico periódico anual con serologías (leptospirosis, toxoplasmosis, brucelosis, hepatitis B) conforme NT-02-2008.",
            "Cumplir protocolo de Bioseguridad Veterinaria (5 momentos OMS de higiene de manos, EPP por procedimiento).",
            "Uso de doble guante en necropsias y sospecha de zoonosis.",
            "No reencapuchar agujas. Descarte inmediato en contenedor rígido.",
            "Manejo de sustancias controladas con libro foliado y doble chequeo de opioides.",
            "Sujeción adecuada de animales: bozal, toalla, sedación. Nunca realizar procedimiento solo con animal agresivo.",
            "Mantener vigente la inscripción en el CMVV y certificación de vacunación antirrábica.",
            "Capacitación periódica en RCP, manejo de emergencias y protocolo PEP Essen.",
            "Pausas activas entre cirugías para reducir fatiga postural.",
            "Embarazadas: reubicación en tareas sin exposición a radiaciones, anestésicos ni quimioterápicos (Art. 78 LOPCYMAT).",
        ],
        emergencias=[
            "Mordedura: lavar 15 min con agua y jabón, antiséptico, NO suturar. Reportar en 2 horas. Clasificar Tipo 1-4 OMS.",
            "Accidente cortopunzante: sangrar, lavar, antiséptico, cubrir, reportar. Evaluar PEP (hepatitis B, VIH, rabia).",
            "Salpicadura mucosa: lavar con suero fisiológico 15 min, reportar, evaluar PEP.",
            "Exposición a anestésicos: ventilar el área, salir si hay mareo, reportar incidente.",
            "Sobredosis accidental de anestésico en paciente: seguir protocolo de antagonistas (naloxona, atipamezol) y notificar a la Directora Gerente.",
            "Embarazada con exposición accidental: acudir inmediatamente al ginecólogo y al IVSS-PMSSO.",
        ],
        footer_label=f"{FOOTER_BASE}  ·  Cargo: Médico Veterinario",
    )
    out_path = os.path.join(OUT_DIR, "Notificacion_Riesgos_Medico_Veterinario.docx")
    doc.save(out_path)
    return out_path


def gen_auxiliar_veterinario():
    doc = Document()
    build_notification(
        doc,
        cargo="Auxiliar Veterinario",
        descripción=(
            "El/La Auxiliar Veterinario(a) apoya al Médico Veterinario en consultas, "
            "cirugías, hospitalización, limpieza y desinfección de áreas clínicas, manejo "
            "y sujeción de animales, preparación de instrumental y atención básica a "
            "pacientes. Está expuesto a riesgos biológicos, químicos, ergonómicos y "
            "eléctricos."
        ),
        tareas=[
            "Sujeción y manejo de animales durante procedimientos.",
            "Limpieza y desinfección de jaulas, camillas, quirófano y hospitalización.",
            "Preparación y esterilización de instrumental quirúrgico.",
            "Asistencia en cirugías y procedimientos.",
            "Administración de medicamentos bajo indicación del MV.",
            "Toma de muestras (sangre, heces, orina).",
            "Alimentación, hidratación y paseo de animales hospitalizados.",
            "Manejo de residuos biomédicos (COVENIN 2747-93).",
            "Registro en historia clínica y libro foliado bajo supervisión.",
        ],
        riesgos=RIESGOS_AUX_VET,
        epp_rows=[
            ("Bata blanca / impermeable", "Uso permanente durante jornada."),
            ("Guantes de nitrilo / goma largos", "OBLIGATORIO en limpieza, manejo de fluidos y residuos."),
            ("Guantes de examen", "Para procedimientos con animales."),
            ("Mascarilla quirúrgica", "Para atención de pacientes y limpieza con aerosoles."),
            ("Mascarilla N95", "Para desinfección con glutaraldehído, manejo de animales con zoonosis respiratoria."),
            ("Gafas de protección", "Para salpicaduras, limpieza con químicos fuertes."),
            ("Delantal impermeable", "Para limpieza de jaulas y áreas húmedas."),
            ("Botas de goma antideslizantes", "Para limpieza de pisos y manejo de residuos."),
            ("Contenedor cortopunzantes", "Acceso inmediato en cada área de procedimiento."),
        ],
        medidas=[
            "Vacunación: antirrábica pre-exposición, tétanos, Hepatitis B (esquema completo).",
            "Examen médico periódico anual con serologías conforme NT-02-2008.",
            "Capacitación en 5 momentos OMS de higiene de manos y protocolo de bioseguridad.",
            "Uso de EPP completo en cada procedimiento. NO omitir guantes ni mascarilla.",
            "Técnica correcta de sujeción animal: bozal, toalla, sedación cuando sea necesario.",
            "No recoger vidrio roto ni cortopunzantes con las manos. Usar pinzas y contenedor rígido.",
            "Rotación de tareas para evitar bipedación prolongada y movimientos repetitivos.",
            "Ventilación adecuada al usar desinfectantes. No mezclar productos químicos.",
            "Mantener pisos secos. Señalizar 'piso mojado' durante limpieza.",
            "Pausas activas cada 2 horas para reducir fatiga musculoesquelética.",
            "Embarazadas: reubicación en tareas sin contacto con zoonosis de riesgo, anestésicos ni químicos fuertes.",
        ],
        emergencias=[
            "Mordedura: lavar 15 min con agua y jabón, antiséptico, NO suturar. Reportar en 2 horas. Evaluar PEP Essen si Tipo 3 o 4.",
            "Accidente cortopunzante: sangrar, lavar, antiséptico, cubrir, reportar. Evaluar PEP para hepatitis B, VIH, rabia.",
            "Salpicadura en mucosa: lavar 15 min con suero fisiológico, reportar y evaluar PEP.",
            "Dermatitis por químicos: suspender uso, lavar zona afectada, reportar.",
            "Caída: evaluar movilidad, no mover al lesionado si hay dolor lumbar, llamar al supervisor.",
            "Sobredosis accidental de anestésico en paciente: notificar inmediatamente al MV tratante.",
        ],
        footer_label=f"{FOOTER_BASE}  ·  Cargo: Auxiliar Veterinario",
    )
    out_path = os.path.join(OUT_DIR, "Notificacion_Riesgos_Auxiliar_Veterinario.docx")
    doc.save(out_path)
    return out_path


def gen_dog_groomer():
    doc = Document()
    build_notification(
        doc,
        cargo="Dog Groomer (Peluquero Canino)",
        descripción=(
            "El/La Dog Groomer es responsable del baño, peluquería, corte, secado y "
            "estética de perros y gatos. Realiza manejo y sujeción de animales, uso de "
            "herramientas cortantes, secadores eléctricos y productos químicos, en un "
            "ambiente húmedo con exposición a ruido, pelos y aerosoles."
        ),
        tareas=[
            "Recepción, evaluación y sujeción de la mascota.",
            "Baño con champús, acondicionadores y desparasitantes tópicos.",
            "Corte de pelo, despeluche, rasurado y arreglo según raza.",
            "Corte y limado de uñas.",
            "Limpieza de oídos y glándulas anales.",
            "Secado con secador eléctrico y toallas.",
            "Limpieza y desinfección de bañeras, mesas y herramientas.",
            "Manejo de mascotas con problemas de conducta.",
        ],
        riesgos=RIESGOS_GROOMER,
        epp_rows=[
            ("Delantal impermeable", "Uso permanente durante jornada."),
            ("Guantes de nitrilo largos", "OBLIGATORIO en baño y manejo de productos químicos."),
            ("Guantes de corte (KEVLAR)", "Para uso de tijeras y maquinillas con mascotas agresivas."),
            ("Mascarilla N95 / quirúrgica", "Para reducir inhalación de pelos, caspa y aerosoles. N95 si asma/alergias."),
            ("Gafas de protección", "Para salpicaduras de champús y productos químicos."),
            ("Botas de goma antideslizantes", "OBLIGATORIO en zona húmeda. Suela antideslizante."),
            ("Protectores auditivos (tapones)", "Para exposición prolongada a secadores y maquinillas."),
            ("Gorro / redecilla", "Para reducir inhalación de pelos."),
            ("Contenedor para cortopunzantes (cuchillas)", "Disponible en cada mesa de peluquería."),
        ],
        medidas=[
            "Vacunación: antirrábica pre-exposición (3 dosis), tétanos al día.",
            "Examen médico periódico anual con énfasis en piel, vías respiratorias y audiometría.",
            "Técnica correcta de sujeción animal: bozal, lazo, toalla. Sedación por MV si mascota agresiva.",
            "Mantener pisos secos. Usar pisos antideslizantes en zona de baño.",
            "Inspección visual de cables eléctricos antes de usar secadores. NO usar con manos mojadas.",
            "Ventilación adecuada para reducir concentración de aerosoles, pelos y olores.",
            "Rotación de tareas para evitar movimientos repetitivos y posturas forzadas.",
            "Pausas activas cada 90 minutos. Estiramientos de cuello, hombros y muñecas.",
            "Capacitación en manejo de mascotas agresivas y primeros auxilios.",
            "Capacitación en higiene de manos y desinfección de herramientas entre mascotas.",
            "Embarazadas: reubicación en tareas sin manejo de zoonosis cutáneas ni químicos fuertes.",
        ],
        emergencias=[
            "Mordedura: lavar 15 min con agua y jabón, antiséptico, NO suturar. Reportar en 2 horas. Clasificar Tipo 1-4 OMS.",
            "Corte con tijera o cuchilla: presionar, lavar, antiséptico, cubrir. Si sangrado abundante o profundo, acudir al IVSS-PMSSO.",
            "Contacto de champú con ojos: lavar con suero fisiológico 15 min, evaluación médica si persiste irritación.",
            "Dermatitis: suspender uso del producto causante, lavar zona, reportar.",
            "Shock eléctrico: desconectar fuente, NO tocar al lesionado con manos mojadas, usar material no conductor, llamar 911.",
            "Asma o crisis alérgica: salir a zona ventilada, usar inhalador si está indicado, reportar.",
            "Caída en zona húmeda: evaluar movilidad, no mover si hay dolor lumbar.",
        ],
        footer_label=f"{FOOTER_BASE}  ·  Cargo: Dog Groomer",
    )
    out_path = os.path.join(OUT_DIR, "Notificacion_Riesgos_Dog_Groomer.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# Main
# ============================================================
def main():
    print("=" * 64)
    print("gen_riesgos_por_rol.py — Generando 5 Notificaciones de Riesgos ALIKA PETS")
    print("=" * 64)
    paths = []
    for fn in (gen_gerente, gen_encargado_tienda, gen_medico_veterinario,
               gen_auxiliar_veterinario, gen_dog_groomer):
        try:
            p = fn()
            paths.append(p)
            print(f"  ✓ {os.path.basename(p)}  →  {os.path.getsize(p)//1024} KB")
        except Exception as e:
            print(f"  ✗ Error en {fn.__name__}: {e}")
            raise
    print("=" * 64)
    print(f"Total: {len(paths)} notificaciones generadas en {OUT_DIR}")
    return paths


if __name__ == "__main__":
    main()
