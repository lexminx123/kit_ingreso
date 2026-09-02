"""
gen_solicitud_v3.py — Solicitud de Empleo COMPACTA (2 páginas) + firma Directora Gerente
Versión corregida: layout más eficiente, menos espacio desperdiciado.
Directora Gerente: ESNATLIM ELENA SIMOZA, C.I. V-17.976.287
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# Paleta y constantes
# ============================================================
TEAL_DARK   = RGBColor(0x0F, 0x76, 0x6E)
TEAL_HDR_BG = "0F766E"
GRAY_ALT    = "F3F4F6"
GRAY_LIGHT  = "F9FAFB"
SLATE_BG    = "E2E8F0"
WHITE       = "FFFFFF"
BLACK       = RGBColor(0x11, 0x18, 0x27)
GRAY_TEXT   = RGBColor(0x55, 0x65, 0x75)
GRAY_MUTED  = RGBColor(0x90, 0x98, 0xA8)

# ============================================================
# Utilidades (versión compacta)
# ============================================================
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="CBD5E1", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)

def set_cell_margins_compact(cell, top=20, bottom=20, left=80, right=80):
    """Márgenes de celda MUY compactos (versión reducida)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, v in (('top',top),('bottom',bottom),('left',left),('right',right)):
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tc_pr.append(tcMar)

def style_run(run, size=9, bold=False, color=None, font="Calibri", italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)

def write_field_cell(cell, label, value="", label_w=None, bg_label=SLATE_BG):
    """Celda con label + valor en una sola celda (más compacto que 2 celdas separadas)."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if label:
        rl = p.add_run(label + " ")
        style_run(rl, size=8, bold=True, color=TEAL_DARK)
    if value:
        rv = p.add_run(value)
        style_run(rv, size=9, color=BLACK)
    else:
        rv = p.add_run("_" * 20)
        style_run(rv, size=9, color=GRAY_MUTED)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(cell)
    set_cell_margins_compact(cell)

def write_label_cell(cell, label, bg=SLATE_BG):
    """Celda solo con label (encabezado)."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(label)
    style_run(r, size=8, bold=True, color=TEAL_DARK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(cell, bg)
    set_cell_borders(cell)
    set_cell_margins_compact(cell)

def write_value_cell(cell, value="", bg=WHITE):
    """Celda solo con valor (campo para llenar)."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if value:
        r = p.add_run(value)
        style_run(r, size=9, color=BLACK)
    else:
        r = p.add_run(" ")
        style_run(r, size=9, color=GRAY_MUTED)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(cell, bg)
    set_cell_borders(cell)
    set_cell_margins_compact(cell)

def add_hr(paragraph, color="0F766E", sz="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    p_pr.append(pBdr)

def add_para(doc, text, size=9, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=2, italic=False, font="Calibri", line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color, font=font, italic=italic)
    return p

def add_section_title(doc, num, title):
    """Título de sección compacto: '1. DATOS PERSONALES' en una sola línea."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run(f"{num}. ")
    style_run(r1, size=9, bold=True, color=TEAL_DARK)
    r2 = p.add_run(title)
    style_run(r2, size=9, bold=True, color=TEAL_DARK)

# ============================================================
# Construcción de tablas compactas (4 columnas: label/val/label/val)
# ============================================================
def build_compact_table(doc, rows, col_widths_cm=(3.5, 5.0, 3.5, 5.0)):
    """
    rows = lista de tuplas (label1, val1, label2, val2)
    Si label2 es None, la fila tiene solo 2 columnas (la 3a y 4a se fusionan).
    """
    n_rows = len(rows)
    tbl = doc.add_table(rows=n_rows, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate(col_widths_cm):
        tbl.columns[i].width = Cm(w)
    for ri, row in enumerate(rows):
        l1, v1, l2, v2 = row
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        # celda label 1
        c = tbl.rows[ri].cells[0]
        c.width = Cm(col_widths_cm[0])
        write_label_cell(c, l1, bg=SLATE_BG)
        # celda valor 1
        c = tbl.rows[ri].cells[1]
        c.width = Cm(col_widths_cm[1])
        write_value_cell(c, v1, bg=bg)
        if l2 is not None:
            # celda label 2
            c = tbl.rows[ri].cells[2]
            c.width = Cm(col_widths_cm[2])
            write_label_cell(c, l2, bg=SLATE_BG)
            # celda valor 2
            c = tbl.rows[ri].cells[3]
            c.width = Cm(col_widths_cm[3])
            write_value_cell(c, v2, bg=bg)
        else:
            # fusionar celdas 2 y 3 para que el valor ocupe más espacio
            merged = tbl.rows[ri].cells[2].merge(tbl.rows[ri].cells[3])
            merged.width = Cm(col_widths_cm[2] + col_widths_cm[3])
            write_value_cell(merged, v1 if v1 else "", bg=bg)
            # limpiar celda 1 ya que su valor va en la fusionada
            tbl.rows[ri].cells[1].text = ""
            write_value_cell(tbl.rows[ri].cells[1], "", bg=bg)
    # evitar corte de fila
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit')
        cs.set(qn('w:val'), 'true')
        trPr.append(cs)
    return tbl

def build_references_table(doc, headers, n_rows=2, col_widths_cm=(4.5, 3.5, 4.5, 4.5)):
    """Tabla de referencias con encabezado teal y filas para llenar."""
    tbl = doc.add_table(rows=1+n_rows, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, w in enumerate(col_widths_cm):
        tbl.columns[i].width = Cm(w)
    # header
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.width = Cm(col_widths_cm[ci])
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(h)
        style_run(r, size=8, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(c, TEAL_HDR_BG)
        set_cell_borders(c)
        set_cell_margins_compact(c)
    # filas vacías
    for ri in range(1, 1+n_rows):
        bg = GRAY_ALT if ri % 2 == 0 else WHITE
        for ci in range(len(headers)):
            c = tbl.rows[ri].cells[ci]
            c.width = Cm(col_widths_cm[ci])
            write_value_cell(c, "", bg=bg)
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit')
        cs.set(qn('w:val'), 'true')
        trPr.append(cs)
    return tbl

# ============================================================
# Firma de Directora Gerente
# ============================================================
def add_signature_block_directora(doc):
    """Bloque de firmas: Solicitante + Directora Gerente (Esnatlim Elena Simoza)."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(8.5)
    tbl.columns[1].width = Cm(8.5)
    # fila 1: labels
    c = tbl.rows[0].cells[0]
    write_label_cell(c, "EL(LA) SOLICITANTE", bg=SLATE_BG)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = tbl.rows[0].cells[1]
    write_label_cell(c, "DIRECTORA GERENTE", bg=SLATE_BG)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # fila 2: líneas de firma
    for ci, (nombre, ci_num, cargo) in enumerate([
        ("", "", "Firma del Solicitante"),
        ("ESNATLIM ELENA SIMOZA", "V-17.976.287", "Directora Gerente"),
    ]):
        c = tbl.rows[1].cells[ci]
        c.width = Cm(8.5)
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        if ci == 0:
            r = p.add_run("______________________________")
            style_run(r, size=9)
        else:
            r = p.add_run("______________________________")
            style_run(r, size=9)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        if nombre:
            r2 = p2.add_run(nombre)
            style_run(r2, size=8, bold=True, color=BLACK)
            p3 = c.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_after = Pt(0)
            p3.paragraph_format.line_spacing = 1.0
            r3 = p3.add_run(f"C.I. {ci_num}")
            style_run(r3, size=8, color=GRAY_TEXT)
            p4 = c.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_after = Pt(0)
            p4.paragraph_format.line_spacing = 1.0
            r4 = p4.add_run(cargo)
            style_run(r4, size=8, italic=True, color=GRAY_TEXT)
        else:
            r2 = p2.add_run("Firma · C.I.: ________________")
            style_run(r2, size=8, italic=True, color=GRAY_TEXT)
        p5 = c.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.paragraph_format.space_before = Pt(2)
        p5.paragraph_format.space_after = Pt(0)
        p5.paragraph_format.line_spacing = 1.0
        r5 = p5.add_run("Fecha: ____ / ____ / ________")
        style_run(r5, size=8, italic=True, color=GRAY_TEXT)
        set_cell_borders(c)
        set_cell_margins_compact(c)

# ============================================================
# Documento principal
# ============================================================
def gen_solicitud_compacta():
    out = "/home/z/my-project/output/01_Solicitud_de_Empleo/Solicitud_de_Empleo_Ficha_Ingreso.docx"
    os.makedirs(os.path.dirname(out), exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(9)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(2)

    # ===== Membrete compacto =====
    hdr = doc.add_table(rows=1, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.autofit = False
    hdr.columns[0].width = Cm(13)
    hdr.columns[1].width = Cm(5)
    for c in hdr.rows[0].cells:
        set_cell_borders(c, color="FFFFFF", sz="0")
        set_cell_margins_compact(c)

    left = hdr.rows[0].cells[0]
    left.text = ""
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("ALIKA PETS")
    style_run(r, size=18, bold=True, color=TEAL_DARK)
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    r2 = p2.add_run("Grupo Caval 1003, C.A.  ·  RIF: J501662533")
    style_run(r2, size=9, color=BLACK)
    p3 = left.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.line_spacing = 1.0
    r3 = p3.add_run("Av. Francisco de Miranda, Local N° 1, Los Teques, Miranda  ·  Zona Postal 1201")
    style_run(r3, size=8, color=GRAY_TEXT)

    right = hdr.rows[0].cells[1]
    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    rp.paragraph_format.line_spacing = 1.0
    rr = rp.add_run("SOLICITUD DE EMPLEO")
    style_run(rr, size=10, bold=True, color=TEAL_DARK)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_before = Pt(0)
    rp2.paragraph_format.space_after = Pt(0)
    rp2.paragraph_format.line_spacing = 1.0
    rr2 = rp2.add_run("Ficha de Ingreso v3.0")
    style_run(rr2, size=8, italic=True, color=GRAY_TEXT)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after = Pt(4)
    add_hr(sep, color="0F766E", sz="10")

    # ===== Declaración inicial compacta =====
    add_para(doc,
        "Yo, ______________________________________, titular de la cédula de identidad N° V-______________, "
        "solicito empleo en GRUPO CAVAL 1003, C.A. (ALIKA PETS) y declaro que la información aquí "
        "proporcionada es veraz, completa y exacta. Acepto que cualquier falsificación u omisión podrá "
        "ser causal de rechazo de la solicitud o de despido justificado si se descubriere con posterioridad.",
        size=9, space_after=4)

    # ===== Sección 1: Datos Personales =====
    add_section_title(doc, "1", "DATOS PERSONALES")
    build_compact_table(doc, [
        ("Apellidos:", "", "Nombres:", ""),
        ("Cédula de Identidad:", "", "Lugar de expedición:", ""),
        ("Fecha de nacimiento:", "", "Lugar de nacimiento:", ""),
        ("Nacionalidad:", "", "Estado civil:", ""),
        ("Sexo:", "", "Edad:", ""),
        ("Estatura (cm):", "", "Peso (kg):", ""),
        ("Tipo de sangre:", "", "¿Posee discapacidad?:", ""),
        ("N° de hijos:", "", "Edades de hijos:", ""),
        ("Dirección de habitación:", "", None, None),
        ("¿Desde cuándo reside ahí?:", "", None, None),
    ])

    # ===== Sección 2: Datos de Contacto =====
    add_section_title(doc, "2", "DATOS DE CONTACTO")
    build_compact_table(doc, [
        ("Teléfono móvil:", "", "Teléfono fijo:", ""),
        ("Correo electrónico:", "", "¿Posee WhatsApp?:", ""),
        ("Contacto de emergencia:", "", "Parentesco:", ""),
        ("Teléfono de emergencia:", "", "Dirección de emergencia:", ""),
    ])

    # ===== Sección 3: Datos Laborales =====
    add_section_title(doc, "3", "DATOS LABORALES (A COMPLETAR POR RR.HH.)")
    build_compact_table(doc, [
        ("Cargo al que aplica:", "", "Departamento / Área:", ""),
        ("Fecha de ingreso:", "", "Salario mensual ofrecido:", ""),
        ("Tipo de contrato:", "", "Fecha de inicio:", ""),
        ("Fecha de culminación:", "", "Supervisor inmediato:", ""),
        ("Horario de trabajo:", "", "Días de descanso:", ""),
    ])

    # ===== Sección 4: Datos Bancarios =====
    add_section_title(doc, "4", "DATOS BANCARIOS (para pago de nómina)")
    build_compact_table(doc, [
        ("Banco:", "", "Tipo de cuenta:", ""),
        ("N° de cuenta:", "", "N° de tarjeta de débito:", ""),
        ("Cuenta email / pago móvil:", "", "Teléfono pago móvil:", ""),
    ])

    # ===== Sección 5: Carga Familiar =====
    add_section_title(doc, "5", "CARGA FAMILIAR (para prestaciones)")
    build_compact_table(doc, [
        ("Total personas a cargo:", "", "Esposo(a) (nombre):", ""),
        ("C.I. Esposo(a):", "", "Hijo 1 (nombre / edad):", ""),
        ("Hijo 2 (nombre / edad):", "", "Hijo 3 (nombre / edad):", ""),
        ("Padre (nombre / edad):", "", "Madre (nombre / edad):", ""),
    ])

    # ===== Sección 6: Documentos Consignados =====
    add_section_title(doc, "6", "DOCUMENTOS CONSIGNADOS (marque con X)")
    build_compact_table(doc, [
        ("☐ Cédula de identidad (copia):", "", "☐ RIF (copia):", ""),
        ("☐ Constancia IVSS (14-100):", "", "☐ Constancia FAOV:", ""),
        ("☐ Constancia INCES:", "", "☐ Constancia estudios / títulos:", ""),
        ("☐ Certificado de salud:", "", "☐ Constancia trabajo anterior:", ""),
        ("☐ Referencias personales (2):", "", "☐ Referencia laboral (1):", ""),
        ("☐ Constancia vacunación antirrábica:", "", "☐ Foto carnet (2):", ""),
    ])

    # ===== Sección 7: Referencias Personales =====
    add_section_title(doc, "7", "REFERENCIAS PERSONALES")
    build_references_table(doc,
        ["Nombre y Apellido", "Parentesco", "Teléfono", "Dirección"],
        n_rows=2,
        col_widths_cm=(4.5, 3.5, 4.0, 5.0))

    # ===== Sección 8: Referencias Laborales =====
    add_section_title(doc, "8", "REFERENCIAS LABORALES")
    build_references_table(doc,
        ["Empresa anterior", "Cargo desempeñado", "Teléfono", "Motivo de egreso"],
        n_rows=2,
        col_widths_cm=(4.5, 3.5, 3.0, 6.0))

    # ===== Declaración Jurada =====
    add_section_title(doc, "9", "DECLARACIÓN JURADA")
    add_para(doc,
        "Declaro bajo fe de juramento que toda la información aquí contenida es veraz y completa. "
        "Autorizo a GRUPO CAVAL 1003, C.A. a verificar los datos suministrados, así como a consultar "
        "mis referencias personales y laborales. Acepto que la falsificación u omisión de información "
        "podrá ser causal de rechazo de la solicitud o de despido justificado conforme al artículo 79 "
        "de la LOTTT.",
        size=9, space_after=6)

    add_para(doc,
        "En la ciudad de Los Teques, Estado Miranda, a los ____ días del mes de "
        "________________ de ________.",
        size=9, space_after=8)

    # ===== Firmas =====
    add_signature_block_directora(doc)

    # ===== Footer =====
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.line_spacing = 1.0
    fr = fp.add_run("ALIKA PETS  ·  Grupo Caval 1003, C.A.  ·  RIF J501662533  ·  Solicitud de Empleo v3.0  ·  Página ")
    style_run(fr, size=7, color=GRAY_MUTED)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r2 = fp.add_run()
    r2._r.append(fldChar1)
    r2._r.append(instrText)
    r2._r.append(fldChar2)
    style_run(r2, size=7, color=GRAY_MUTED)

    doc.save(out)
    return out

# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    out = gen_solicitud_compacta()
    size_kb = os.path.getsize(out) / 1024
    print(f"✓ {out}  ({size_kb:.1f} KB)")
